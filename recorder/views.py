import csv
import datetime
import io
import json
import os
import requests
import signal
import re
import subprocess
import sys
import threading
import time
import uuid
from concurrent.futures import Future
from typing import Any

from PIL import Image, ImageDraw

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.auth import logout as auth_logout
from django.core import signing
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.core.paginator import Paginator
from django.http import Http404, HttpResponse, JsonResponse
from django.shortcuts import render, redirect
from django.urls import reverse
from django.db import connection, close_old_connections, IntegrityError
from django.db import transaction
from django.contrib import messages
from django.db.models import Count, Sum, Case, When, IntegerField, Value, Q
from django.db.models.functions import TruncDate
from django.utils import timezone as _tz

from .models import Step, Locator, DataEntry, RunResult, SessionMeta, Recording, RemoteExecution, RemoteTarget
from .replay import replay_session
try:
    from .playwright_replay import replay_session as playwright_replay_session
except Exception:
    playwright_replay_session = None  # type: ignore[assignment]
from workflow_agent import (
    WorkflowGenerationError,
    WorkflowNotFoundError,
    create_test_case_from_workflow,
)

_MAIN_PY = os.path.join(os.path.dirname(os.path.dirname(__file__)), "main.py")
_PLAYWRIGHT_PY = os.path.join(os.path.dirname(os.path.dirname(__file__)), "playwright_recorder.py")
_RECORDING_LOGS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "logs")

# ---------------------------------------------------------------------------
# Global bounded thread-pool for replay workers.
# Size is read from app_config at first use; reload Django to change it.
# All replay jobs — single and bulk — share this pool so one bulk submission
# cannot starve individual replays launched from the UI.
# ---------------------------------------------------------------------------
_REPLAY_SEMAPHORE: threading.Semaphore | None = None
_REPLAY_SEMAPHORE_LOCK = threading.Lock()


def _get_replay_semaphore() -> threading.Semaphore:
    """Return (and lazily create) the concurrency-limiting semaphore."""
    global _REPLAY_SEMAPHORE
    if _REPLAY_SEMAPHORE is not None:
        return _REPLAY_SEMAPHORE
    with _REPLAY_SEMAPHORE_LOCK:
        if _REPLAY_SEMAPHORE is not None:
            return _REPLAY_SEMAPHORE
        try:
            from django.db import connection as _c
            with _c.cursor() as _cur:
                _cur.execute(
                    "SELECT value FROM app_config WHERE key = 'replay.max_parallel_replays'"
                )
                _row = _cur.fetchone()
                _max = max(1, int(_row[0])) if _row else 8
        except Exception:
            _max = getattr(settings, "REPLAY_MAX_WORKERS", 8)
        _REPLAY_SEMAPHORE = threading.Semaphore(_max)
        return _REPLAY_SEMAPHORE


def _submit_replay_job(fn) -> Future:
    """Run fn on a daemon thread with semaphore-based concurrency control.

    Returns a concurrent.futures.Future for API compatibility with callers
    that inspect .result() / .done() / .exception().
    """
    future: Future = Future()
    sem = _get_replay_semaphore()

    def _wrapper():
        sem.acquire()
        try:
            if not future.set_running_or_notify_cancel():
                return
            try:
                result = fn()
                future.set_result(result)
            except BaseException as exc:
                future.set_exception(exc)
        finally:
            sem.release()

    t = threading.Thread(target=_wrapper, daemon=True)
    t.start()
    return future


# In-memory replay job registry  {str(run_id): job_dict}
_REPLAY_JOBS: dict[str, dict] = {}
# Maps (record_id, username) → latest run_id so we can find the active job per user+session
_SESSION_TO_RUN: dict[tuple, str] = {}
# Global RLock — protects all compound read-modify-write ops on _REPLAY_JOBS/_SESSION_TO_RUN
_JOBS_LOCK = threading.RLock()

# In-memory recording state (populated by start/continue_recording, cleared on stop).
# Allows hotkey endpoints to stop/pause/resume without a Django session.
# {record_id: {"pid": int, "folder": str, "user_id": int, "paused": bool}}
_ACTIVE_RECORDING: dict[str, dict] = {}
_ACTIVE_RECORDING_LOCK = threading.Lock()
# Schema-ensure flag — runs once per process
_SCHEMA_ENSURED = False
_AI_DATABANK_SCHEMA_ENSURED = False
_AI_WORKFLOW_SCHEMA_ENSURED = False
# Track last remote execution timestamp for page auto-refresh
_LAST_REMOTE_ACTION = 0.0
_FT_TRAINING_LOCK = threading.Lock()
_FT_TRAINING_PROCESS: subprocess.Popen | None = None
_FT_TRAINING_LOG_HANDLE = None
_FT_TRAINING_STATE: dict[str, Any] = {
    "status": "idle",
    "is_active": False,
    "pid": None,
    "started_at": None,
    "finished_at": None,
    "returncode": None,
    "log_path": None,
    "dataset_path": None,
    "trainer_script_path": None,
    "modelfile_path": None,
    "import_script_path": None,
    "gguf_path": None,
    "base_model": None,
    "ft_model_name": None,
    "last_error": None,
    "command": None,
}

# ── Edit-lock registry ────────────────────────────────────────────────────
# {record_id: {"user": str, "acquired_at": float}}
# A lock expires after _EDIT_LOCK_TTL seconds without a heartbeat.
_EDITING_SESSIONS: dict[str, dict] = {}
_EDITING_LOCK = threading.Lock()
_EDIT_LOCK_TTL = 30.0  # seconds


def _is_pid_alive(pid: int) -> bool:
    """Return True if a process with *pid* is still running.

    os.kill(pid, 0) is the POSIX idiom but on Windows it can raise SystemError
    for certain process states (zombie / access-denied).  We use psutil when
    available (most reliable), fall back to a ctypes OpenProcess check on
    Windows, and finally try the os.kill approach.
    """
    try:
        import psutil
        return psutil.pid_exists(pid)
    except ImportError:
        pass

    if sys.platform == "win32":
        import ctypes
        SYNCHRONIZE = 0x00100000
        handle = ctypes.windll.kernel32.OpenProcess(SYNCHRONIZE, False, pid)
        if handle == 0:
            return False
        ctypes.windll.kernel32.CloseHandle(handle)
        return True

    # POSIX fallback
    try:
        os.kill(pid, 0)
        return True
    except (OSError, ProcessLookupError, SystemError):
        return False


def _kill_browser_on_port(port: int) -> None:
    """Kill the browser process listening on *port* (e.g. Chrome RDP port).

    Uses netstat to find the PID of the process bound to the port, then
    kills it with /F /T so the entire process tree (child renderers) is gone.
    No-op on non-Windows or if nothing is listening.
    """
    if sys.platform != "win32" or port <= 0:
        return
    try:
        result = subprocess.run(
            ["netstat", "-ano"],
            capture_output=True, text=True, timeout=6,
        )
        for line in result.stdout.splitlines():
            if f":{port}" in line and "LISTENING" in line:
                parts = line.split()
                if parts:
                    _pid = parts[-1]
                    if _pid.isdigit() and int(_pid) > 0:
                        subprocess.call(
                            ["taskkill", "/F", "/T", "/PID", _pid],
                            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                        )
                        break
    except Exception:
        pass
# ---------------------------------------------------------------------------

@csrf_exempt
def logout_view(request):
    """Accept GET or POST without CSRF — logout is safe to exempt."""
    auth_logout(request)
    return redirect("login")


def _build_monitor_token(user: User) -> str:
    signer = signing.TimestampSigner(salt="local-monitor-auth")
    return signer.sign(str(user.pk))


def _finetune_workspace_dir() -> str:
    return os.path.join(str(settings.BASE_DIR), "llm_workflow_assistant", "finetune")


def _finetune_uploads_dir() -> str:
    return os.path.join(str(settings.BASE_DIR), "llm_workflow_assistant", "uploads")


def _latest_uploaded_finetune_dataset() -> str | None:
    uploads_dir = _finetune_uploads_dir()
    if not os.path.isdir(uploads_dir):
        return None
    candidates = [
        os.path.join(uploads_dir, name)
        for name in os.listdir(uploads_dir)
        if name.lower().endswith(".jsonl")
    ]
    if not candidates:
        return None
    return max(candidates, key=os.path.getmtime)


def _read_text_tail(path: str | None, max_chars: int = 16000) -> str:
    if not path or not os.path.exists(path):
        return ""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            data = fh.read()
        return data[-max_chars:]
    except Exception:
        return ""


def _serialize_training_state() -> dict[str, Any]:
    global _FT_TRAINING_PROCESS, _FT_TRAINING_LOG_HANDLE, _FT_TRAINING_STATE
    with _FT_TRAINING_LOCK:
        proc = _FT_TRAINING_PROCESS
        if proc is not None:
            returncode = proc.poll()
            if returncode is not None:
                _FT_TRAINING_STATE["status"] = "completed" if returncode == 0 else "failed"
                _FT_TRAINING_STATE["is_active"] = False
                _FT_TRAINING_STATE["finished_at"] = datetime.datetime.now(datetime.timezone.utc).isoformat()
                _FT_TRAINING_STATE["returncode"] = int(returncode)
                if returncode != 0 and not _FT_TRAINING_STATE.get("last_error"):
                    _FT_TRAINING_STATE["last_error"] = f"Training exited with code {returncode}."
                _FT_TRAINING_PROCESS = None
                if _FT_TRAINING_LOG_HANDLE is not None:
                    try:
                        _FT_TRAINING_LOG_HANDLE.close()
                    except Exception:
                        pass
                    _FT_TRAINING_LOG_HANDLE = None

        payload = dict(_FT_TRAINING_STATE)
        payload["log_tail"] = _read_text_tail(payload.get("log_path"))
        return payload


def _start_finetune_training(*, base_model: str, ft_model_name: str) -> dict[str, Any]:
    global _FT_TRAINING_PROCESS, _FT_TRAINING_LOG_HANDLE, _FT_TRAINING_STATE
    with _FT_TRAINING_LOCK:
        if _FT_TRAINING_PROCESS is not None and _FT_TRAINING_PROCESS.poll() is None:
            raise RuntimeError("Training is already running.")

        latest_dataset = _latest_uploaded_finetune_dataset()
        workspace_dataset = os.path.join(_finetune_workspace_dir(), "fine_tune.jsonl")
        if latest_dataset is None and not os.path.exists(workspace_dataset):
            raise RuntimeError("Upload a fine-tune dataset before starting training.")

        from llm_workflow_assistant.finetune_service import get_training_environment_status, prepare_finetune_workspace

        env_status = get_training_environment_status()
        if not env_status.get("supported"):
            _FT_TRAINING_STATE = {
                **_FT_TRAINING_STATE,
                "status": "failed",
                "is_active": False,
                "pid": None,
                "finished_at": datetime.datetime.now(datetime.timezone.utc).isoformat(),
                "returncode": None,
                "base_model": base_model,
                "ft_model_name": ft_model_name,
                "last_error": str(env_status.get("message") or "Training environment is not supported."),
            }
            raise RuntimeError(_FT_TRAINING_STATE["last_error"])

        workspace = prepare_finetune_workspace(
            dataset_path=latest_dataset or workspace_dataset,
            base_dir=str(settings.BASE_DIR),
            base_model=base_model,
            ft_model_name=ft_model_name,
        )

        log_path = os.path.join(workspace["workspace_dir"], "training.log")
        if _FT_TRAINING_LOG_HANDLE is not None:
            try:
                _FT_TRAINING_LOG_HANDLE.close()
            except Exception:
                pass
            _FT_TRAINING_LOG_HANDLE = None

        log_handle = open(log_path, "w", encoding="utf-8", newline="\n")
        command = [
            sys.executable,
            workspace["trainer_script_path"],
            "--dataset",
            workspace["dataset_copy_path"],
            "--base-model",
            base_model,
            "--ft-model-name",
            ft_model_name,
        ]
        creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0) if sys.platform == "win32" else 0
        process = subprocess.Popen(
            command,
            cwd=workspace["workspace_dir"],
            stdout=log_handle,
            stderr=subprocess.STDOUT,
            creationflags=creationflags,
        )

        _FT_TRAINING_LOG_HANDLE = log_handle
        _FT_TRAINING_PROCESS = process
        _FT_TRAINING_STATE = {
            "status": "running",
            "is_active": True,
            "pid": process.pid,
            "started_at": datetime.datetime.now(datetime.timezone.utc).isoformat(),
            "finished_at": None,
            "returncode": None,
            "log_path": log_path,
            "dataset_path": workspace["dataset_copy_path"],
            "trainer_script_path": workspace["trainer_script_path"],
            "modelfile_path": workspace["modelfile_path"],
            "import_script_path": workspace["import_script_path"],
            "gguf_path": workspace["gguf_path"],
            "base_model": base_model,
            "ft_model_name": ft_model_name,
            "last_error": None,
            "command": command,
        }
    return _serialize_training_state()


def _resolve_monitor_user(request):
    if getattr(request, "user", None) and request.user.is_authenticated:
        return request.user

    token = (
        request.GET.get("monitor_token")
        or request.POST.get("monitor_token")
        or request.headers.get("X-Monitor-Token")
    )
    if not token:
        return None

    try:
        signer = signing.TimestampSigner(salt="local-monitor-auth")
        user_pk = int(signer.unsign(token, max_age=60 * 60 * 12))
    except Exception:
        return None
    return User.objects.filter(pk=user_pk, is_active=True).first()


def _get_user_tenant_id(user: User | None) -> str | None:
    if user is None:
        return None
    try:
        tenant_id = user.profile.tenant_id
    except Exception:
        return None
    return str(tenant_id) if tenant_id else None


@csrf_exempt
def local_monitor_login(request):
    """Local-only helper for the detached monitor browser profile."""
    remote_addr = (request.META.get("REMOTE_ADDR") or "").strip()
    if remote_addr not in {"127.0.0.1", "::1", "localhost"}:
        raise Http404()

    user = _resolve_monitor_user(request)
    if user is None:
        user = User.objects.filter(is_active=True, is_superuser=True).order_by("id").first()
        if user is None:
            user = User.objects.filter(is_active=True).order_by("id").first()
        if user is None:
            return redirect("login")

    minimized = "1" if request.GET.get("minimized") == "1" else "0"
    token = _build_monitor_token(user)
    return redirect(f"/active_runs/?minimized={minimized}&monitor_token={token}")


# ---------------------------------------------------------------------------
# Dashboard helpers
# ---------------------------------------------------------------------------

def _get_active_users():
    """Return a list of dicts describing currently logged-in users.
    Flags is_recording / is_running are set from live process / job state.
    Safe to call from any view; returns [] on any error.
    """
    from django.contrib.sessions.models import Session
    from django.contrib.auth import get_user_model
    import datetime as _dt
    _User = get_user_model()
    try:
        now = _dt.datetime.now(_dt.timezone.utc)
        active_sessions_qs = Session.objects.filter(expire_date__gte=now)
        user_ids      = set()
        recording_pks = set()
        for _s in active_sessions_qs:
            _decoded = _s.get_decoded()
            _uid = _decoded.get("_auth_user_id")
            if _uid:
                user_ids.add(int(_uid))
            _pid = _decoded.get("recording_pid")
            if _pid and _uid:
                if _is_pid_alive(int(_pid)):
                    recording_pks.add(int(_uid))
        running_usernames = {
            _job.get("runner", "")
            for _job in _REPLAY_JOBS.values()
            if _job.get("status") in ("running", "paused") and _job.get("runner")
        }
        users = []
        for _u in (
            _User.objects.filter(pk__in=user_ids, is_active=True)
            .values("pk", "username", "first_name", "last_name", "is_superuser", "last_login")
            .order_by("username")
        ):
            _u["is_recording"] = _u["pk"] in recording_pks
            _u["is_running"]   = _u["username"] in running_usernames
            users.append(_u)
        return users
    except Exception:
        return []


# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------

@login_required
def dashboard(request):
    with connection.cursor() as cur:
        # Total sessions
        cur.execute("SELECT COUNT(DISTINCT record_id) FROM steps;")
        total_sessions = cur.fetchone()[0] or 0

        # Total steps recorded
        cur.execute("SELECT COUNT(*) FROM steps;")
        total_steps = cur.fetchone()[0] or 0

        # Total replay runs
        cur.execute("SELECT COUNT(DISTINCT run_id) FROM run_table;")
        total_runs = cur.fetchone()[0] or 0

        # Count unique sessions by outcome of their LATEST run:
        #   passed  = latest run has no failures and no not_executed steps
        #   failed  = latest run has at least one failure
        #   skipped = latest run has no failures but at least one not_executed
        cur.execute("""
            WITH latest_runs AS (
                SELECT DISTINCT ON (record_id) record_id, run_id
                FROM run_table
                ORDER BY record_id, created_at DESC
            ),
            session_outcome AS (
                SELECT
                    lr.record_id,
                    SUM(CASE WHEN r.status = 'fail'         THEN 1 ELSE 0 END) AS fails,
                    SUM(CASE WHEN r.status = 'not_executed' THEN 1 ELSE 0 END) AS not_execs
                FROM run_table r
                INNER JOIN latest_runs lr ON lr.run_id = r.run_id
                GROUP BY lr.record_id
            )
            SELECT
                COUNT(CASE WHEN fails = 0 AND not_execs = 0 THEN 1 END) AS total_pass,
                COUNT(CASE WHEN fails > 0                   THEN 1 END) AS total_fail,
                COUNT(CASE WHEN fails = 0 AND not_execs > 0 THEN 1 END) AS total_not_exec
            FROM session_outcome;
        """)
        row = cur.fetchone()
        total_pass, total_fail, total_not_exec = (row[0] or 0, row[1] or 0, row[2] or 0)

    # ── Active (in-flight) runs from the in-memory job registry ──────────────
    active_runs = []
    for _run_id, _job in list(_REPLAY_JOBS.items()):
        if _job.get("status") not in ("running", "paused"):
            continue
        _sid = _job.get("record_id", "")
        try:
            _rname = SessionMeta.objects.get(record_id=_sid).record_name or _sid
        except Exception:
            _rname = _sid
        with _job["lock"]:
            _results = list(_job["results"])
        _passes  = sum(1 for r in _results if r.get("ok"))
        _fails   = sum(1 for r in _results if not r.get("ok"))
        _done    = len(_results)
        active_runs.append({
            "run_id":      _run_id,
            "record_id":  _sid,
            "record_name": _rname,
            "status":      _job["status"],
            "total":       _job["total"],
            "done":        _done,
            "passes":      _passes,
            "fails":       _fails,
            "started_at":  _job.get("started_at"),
            "runner":      _job.get("runner", ""),
        })

    # ── Active logged-in users (non-expired sessions) ─────────────────────
    active_users = _get_active_users()

    return render(request, "recorder/dashboard.html", {
        "total_sessions": total_sessions,
        "total_steps": total_steps,
        "total_runs": total_runs,
        "total_pass": total_pass,
        "total_fail": total_fail,
        "total_not_exec": total_not_exec,
        "active_runs":     active_runs,
        "active_users":    active_users,
    })


@login_required
def burndown_api(request):
    """GET /api/dashboard/burndown/
    Returns per-day pass/fail/not_executed step counts for the burndown chart.
    Optional params:
      folder  – parent_folder_id UUID (empty = all folders)
      days    – look-back window in days (default 30, max 365)
    """
    folder = request.GET.get("folder", "").strip()
    try:
        days = max(1, min(int(request.GET.get("days", 30)), 365))
    except (ValueError, TypeError):
        days = 30

    since = datetime.datetime.now() - datetime.timedelta(days=days)

    # Build base queryset from RunResult (run_table)
    qs = RunResult.objects.filter(run_date__gte=since)

    if folder:
        qs = qs.filter(parent_folder_id=folder)

    # Distinct parent folders: one row per parent_folder_id (latest run_date)
    with connection.cursor() as _cur:
        _cur.execute("""
            SELECT row_id, record_id, folder_id, name, latest_run_date
            FROM (
                SELECT DISTINCT ON (rt.parent_folder_id)
                    rt.id                      AS row_id,
                    rt.record_id::text         AS record_id,
                    rt.parent_folder_id::text  AS folder_id,
                    pf.parent_folder           AS name,
                    rt.run_date                AS latest_run_date
                FROM run_table rt
                JOIN parent_folders pf ON pf.parent_folder_id = rt.parent_folder_id
                WHERE rt.parent_folder_id IS NOT NULL
                ORDER BY rt.parent_folder_id, rt.run_date DESC NULLS LAST
            ) sub
            ORDER BY latest_run_date DESC NULLS LAST
        """)
        folders = [
            {
                "row_id":    r[0],
                "record_id": r[1],
                "id":        r[2],
                "name":      r[3],
                "run_date":  r[4].isoformat() if r[4] else None,
            }
            for r in _cur.fetchall()
        ]

    # Per-day aggregation
    daily = (
        qs
        .annotate(day=TruncDate("run_date"))
        .values("day")
        .annotate(
            passes=Sum(Case(When(status=RunResult.STATUS_PASS,         then=Value(1)), default=Value(0), output_field=IntegerField())),
            fails=Sum(Case(When(status=RunResult.STATUS_FAIL,         then=Value(1)), default=Value(0), output_field=IntegerField())),
            skips=Sum(Case(When(status=RunResult.STATUS_NOT_EXECUTED, then=Value(1)), default=Value(0), output_field=IntegerField())),
        )
        .order_by("day")
    )

    return JsonResponse({
        "labels":  [str(r["day"]) for r in daily],
        "pass":    [r["passes"] for r in daily],
        "fail":    [r["fails"]  for r in daily],
        "skip":    [r["skips"]  for r in daily],
        "folders": folders,
    })


@login_required
def active_users_api(request):
    """GET /api/active-users/  — returns live list of logged-in users as JSON."""
    users = []
    for u in _get_active_users():
        users.append({
            "username":     u["username"],
            "first_name":   u["first_name"],
            "last_name":    u["last_name"],
            "is_superuser": u["is_superuser"],
            "is_recording": u["is_recording"],
            "is_running":   u["is_running"],
        })
    return JsonResponse({"users": users})


@login_required
def dashboard_stats_api(request):
    """GET /api/dashboard/stats/  — live pass/fail/not_exec counts (latest run per session).
    Optional param: folder (empty = all folders) — now a parent folder name.
    """
    folder = request.GET.get("folder", "").strip()

    if folder:
        folder_cond = "r.parent_folder_id = %s::uuid"
        cond_params: list = [folder]
    else:
        folder_cond = "1=1"
        cond_params = []

    with connection.cursor() as cur:
        # Folders dropdown — one row per parent_folder_id: id, record_id, and latest run_date
        cur.execute("""
            SELECT row_id, record_id, folder_id, name, latest_run_date
            FROM (
                SELECT DISTINCT ON (rt.parent_folder_id)
                    rt.id                      AS row_id,
                    rt.record_id::text         AS record_id,
                    rt.parent_folder_id::text  AS folder_id,
                    pf.parent_folder           AS name,
                    rt.run_date                AS latest_run_date
                FROM run_table rt
                JOIN parent_folders pf ON pf.parent_folder_id = rt.parent_folder_id
                WHERE rt.parent_folder_id IS NOT NULL
                ORDER BY rt.parent_folder_id, rt.run_date DESC NULLS LAST
            ) sub
            ORDER BY latest_run_date DESC NULLS LAST
        """)
        folders = [
            {
                "row_id":    r[0],
                "record_id": r[1],
                "id":        r[2],
                "name":      r[3],
                "run_date":  r[4].isoformat() if r[4] else None,
            }
            for r in cur.fetchall()
        ]

        cur.execute(f"""
            WITH latest_runs AS (
                SELECT DISTINCT ON (r.record_id) r.record_id, r.run_id
                FROM run_table r
                WHERE {folder_cond}
                ORDER BY r.record_id, r.created_at DESC
            ),
            session_outcome AS (
                SELECT
                    lr.record_id,
                    SUM(CASE WHEN r.status = 'fail'         THEN 1 ELSE 0 END) AS fails,
                    SUM(CASE WHEN r.status = 'not_executed' THEN 1 ELSE 0 END) AS not_execs
                FROM run_table r
                INNER JOIN latest_runs lr ON lr.run_id = r.run_id
                GROUP BY lr.record_id
            )
            SELECT
                COUNT(CASE WHEN fails = 0 AND not_execs = 0 THEN 1 END) AS total_pass,
                COUNT(CASE WHEN fails > 0                   THEN 1 END) AS total_fail,
                COUNT(CASE WHEN fails = 0 AND not_execs > 0 THEN 1 END) AS total_not_exec
            FROM session_outcome;
        """, cond_params)
        row = cur.fetchone()
    total_pass, total_fail, total_not_exec = (row[0] or 0, row[1] or 0, row[2] or 0)
    return JsonResponse({
        "pass":    total_pass,
        "fail":    total_fail,
        "not_run": total_not_exec,
        "folders": folders,
    })


# ---------------------------------------------------------------------------
# History
# ---------------------------------------------------------------------------

_RECORDINGS_ALIASES_H = {"baseline", ""}


@login_required
def history(request):
    """Full run history with date filtering and folder tabs."""
    if get_config("features.history_enabled", "true") != "true":
        messages.warning(request, "The History page is currently disabled.")
        return redirect("dashboard")
    date_from_str = request.GET.get("date_from", "").strip()
    date_to_str   = request.GET.get("date_to",   "").strip()
    active_folder = request.GET.get("folder",    "").strip()

    # Parse dates (best-effort)
    date_from = None
    date_to   = None
    try:
        if date_from_str:
            date_from = datetime.date.fromisoformat(date_from_str)
    except ValueError:
        pass
    try:
        if date_to_str:
            date_to = datetime.date.fromisoformat(date_to_str)
    except ValueError:
        pass

    # Build WHERE clause for date filter
    where_parts = []
    params: list = []
    if date_from:
        where_parts.append("MIN(DATE(r.run_date)) >= %s")
        params.append(date_from)
    if date_to:
        where_parts.append("MIN(DATE(r.run_date)) <= %s")
        params.append(date_to)
    having_clause = ("HAVING " + " AND ".join(where_parts)) if where_parts else ""

    sql = f"""
        SELECT
            r.run_id,
            COALESCE(m.record_name, '')          AS record_name,
            r.record_id,
            MIN(r.created_at)                    AS run_at,
            MIN(r.run_date)                      AS run_date,
            SUM(CASE WHEN r.status='pass'         THEN 1 ELSE 0 END) AS passes,
            SUM(CASE WHEN r.status='fail'         THEN 1 ELSE 0 END) AS failures,
            SUM(CASE WHEN r.status='not_executed' THEN 1 ELSE 0 END) AS skipped,
            MAX(r.runner) AS runner,
            COALESCE(
                NULLIF(TRIM(MIN(r.folder_name)), ''),
                NULLIF(TRIM((
                    SELECT MIN(s.folder_name) FROM steps s
                    WHERE s.record_id = r.record_id
                      AND s.folder_name IS NOT NULL AND s.folder_name <> ''
                )), ''),
                'Baseline'
            ) AS folder_name
        FROM run_table r
        LEFT JOIN session_meta m ON m.record_id = r.record_id
        GROUP BY r.run_id, m.record_name, r.record_id
        {having_clause}
        ORDER BY MIN(r.run_date) DESC NULLS LAST, MIN(r.created_at) DESC, r.run_id DESC;
    """

    with connection.cursor() as cur:
        cur.execute(sql, params)
        cols = [c.name for c in cur.description]
        rows_flat = [dict(zip(cols, row)) for row in cur.fetchall()]

    # Normalise folder names and build folder tree
    recordings_folder_label = _get_recordings_folder_label()
    folder_map: dict[str, list] = {}
    unfiled_runs: list = []
    for _r in rows_flat:
        _raw = (_r.get("folder_name") or "").strip()
        if _is_recordings_folder_name(_raw):
            _r["folder_name"] = recordings_folder_label
            folder_map.setdefault(recordings_folder_label, []).append(_r)
        elif _raw:
            folder_map.setdefault(_raw, []).append(_r)
        else:
            unfiled_runs.append(_r)

    rows_flat = [
        run for run in rows_flat
        if _is_project_path_visible_to_user(request.user, (run.get("folder_name") or "").strip())
    ]
    folder_map = {
        folder_name: runs
        for folder_name, runs in folder_map.items()
        if _is_project_path_visible_to_user(request.user, folder_name)
    }

    total_runs = len(rows_flat)
    folder_roots = _build_folder_tree(folder_map, [], {})
    # Promote recordings root to front
    recordings_children: list = []
    remaining_roots: list = []
    recordings_prefix = recordings_folder_label + "/"
    for _root in folder_roots:
        if _root["path"].startswith(recordings_prefix):
            recordings_children.append(_root)
        else:
            remaining_roots.append(_root)
    folder_roots = remaining_roots
    if folder_map.get(recordings_folder_label) or recordings_children:
        recs_root = _build_special_folder_group(recordings_folder_label, folder_map.get(recordings_folder_label, []))
        recs_root["children"] = recordings_children
        recs_root["total_sessions"] = (
            len(folder_map.get(recordings_folder_label, []))
            + sum(c["total_sessions"] for c in recordings_children)
        )
        folder_roots.insert(0, recs_root)
    if unfiled_runs:
        folder_roots.append(_build_special_folder_group("", unfiled_runs))

    # Per-user visible actions (for delete_folder button etc.)
    _pref_raw = get_user_pref(
        request.user.id,
        "projects.visible_actions",
        ",".join(sorted(_ALL_PROJ_ACTIONS)),
    ) if request.user.is_authenticated else ""
    visible_actions = _resolve_project_visible_actions(_pref_raw)

    return render(request, "recorder/history.html", {
        "folder_roots":    folder_roots,
        "total_runs":      total_runs,
        "date_from":       date_from_str,
        "date_to":         date_to_str,
        "visible_actions": visible_actions,
    })


@login_required
def clear_history(request):
    """DELETE run_table rows for selected run_ids (if provided) or whole folder."""
    if request.method != "POST":
        return redirect("history")

    folder_name = request.POST.get("folder_name", "").strip()
    raw_ids     = request.POST.getlist("run_ids")

    import uuid as _uuid
    valid_ids = []
    for v in raw_ids:
        try:
            valid_ids.append(str(_uuid.UUID(v.strip())))
        except (ValueError, AttributeError):
            pass

    with connection.cursor() as cur:
        if valid_ids and folder_name:
            # Delete specific run_ids scoped to the given folder_name
            placeholders = ",".join(["%s"] * len(valid_ids))
            if _is_recordings_folder_name(folder_name):
                _ch_aliases = list(_recordings_sql_aliases())
                _ch_ph = ",".join(["%s"] * len(_ch_aliases))
                cur.execute(
                    f"DELETE FROM run_table WHERE run_id IN ({placeholders})"
                    f" AND LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_ch_ph})",
                    valid_ids + _ch_aliases,
                )
            else:
                cur.execute(
                    f"DELETE FROM run_table WHERE run_id IN ({placeholders})"
                    f" AND TRIM(folder_name) = %s",
                    valid_ids + [folder_name],
                )
        elif valid_ids:
            placeholders = ",".join(["%s"] * len(valid_ids))
            cur.execute(f"DELETE FROM run_table WHERE run_id IN ({placeholders})", valid_ids)
        elif folder_name:
            if _is_recordings_folder_name(folder_name):
                _ch_aliases = list(_recordings_sql_aliases())
                _ch_ph = ",".join(["%s"] * len(_ch_aliases))
                cur.execute(
                    f"DELETE FROM run_table WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_ch_ph})",
                    _ch_aliases,
                )
            else:
                cur.execute(
                    "DELETE FROM run_table WHERE TRIM(folder_name) = %s",
                    [folder_name],
                )
        else:
            return redirect("history")

    return redirect("/history/")


@login_required
def delete_runs(request):
    """DELETE specific run_table rows by run_id (POST only, run_ids[] list)."""
    if request.method != "POST":
        return redirect("history")
    raw_ids = request.POST.getlist("run_ids")
    # validate each value looks like a UUID
    import uuid as _uuid
    valid_ids = []
    for v in raw_ids:
        try:
            valid_ids.append(str(_uuid.UUID(v.strip())))
        except (ValueError, AttributeError):
            pass
    if valid_ids:
        placeholders = ",".join(["%s"] * len(valid_ids))
        with connection.cursor() as cur:
            cur.execute(f"DELETE FROM run_table WHERE run_id IN ({placeholders})", valid_ids)
    count = len(valid_ids)
    messages.success(request, f"{count} run{'s' if count != 1 else ''} deleted.")
    # preserve date filter / folder params
    next_url = request.POST.get("next", "/history/")
    return redirect(next_url)


@login_required
def run_stats(request):
    """GET: return stats for given run_ids[] — used by delete confirmation modal."""
    import uuid as _uuid
    raw_ids = request.GET.getlist("run_ids")
    valid_ids = []
    for v in raw_ids:
        try:
            valid_ids.append(str(_uuid.UUID(v.strip())))
        except (ValueError, AttributeError):
            pass
    if not valid_ids:
        return JsonResponse({"filenames": [], "steps": 0, "locators": 0, "data_entries": 0})

    ph = ",".join(["%s"] * len(valid_ids))
    with connection.cursor() as cur:
        # filenames (distinct record_names for the selected runs)
        cur.execute(f"""
            SELECT DISTINCT COALESCE(NULLIF(TRIM(m.record_name), ''), r.record_id::text)
            FROM run_table r
            LEFT JOIN session_meta m ON m.record_id = r.record_id
            WHERE r.run_id IN ({ph})
            ORDER BY 1
        """, valid_ids)
        filenames = [row[0] for row in cur.fetchall()]

        # step rows
        cur.execute(f"SELECT COUNT(*) FROM run_table r WHERE r.run_id IN ({ph})", valid_ids)
        steps = cur.fetchone()[0] or 0

        # locators
        cur.execute(f"""
            SELECT COUNT(DISTINCT l.id) FROM locators l
            WHERE l.record_id IN (
                SELECT DISTINCT r.record_id FROM run_table r WHERE r.run_id IN ({ph})
            )
        """, valid_ids)
        locators = cur.fetchone()[0] or 0

        # data entries
        cur.execute(f"""
            SELECT COUNT(DISTINCT d.id) FROM data d
            WHERE d.record_id IN (
                SELECT DISTINCT r.record_id FROM run_table r WHERE r.run_id IN ({ph})
            )
        """, valid_ids)
        data_entries = cur.fetchone()[0] or 0

    return JsonResponse({
        "filenames":    filenames,
        "steps":        steps,
        "locators":     locators,
        "data_entries": data_entries,
    })


@login_required
def history_delete_folder(request):
    """POST: delete run_table rows for a folder and its sub-folders only.

    Does NOT touch steps, locators, data, session_meta, or the folder
    registry — so the folder and its sessions remain intact on /projects/.
    """
    if request.method != "POST":
        return redirect("history")
    folder = _normalize_folder_path(request.POST.get("folder_name", ""))
    if not folder:
        messages.warning(request, "No folder specified.")
        return redirect("history")

    is_recordings = _is_recordings_folder_name(folder)
    if is_recordings:
        aliases = list(_recordings_sql_aliases())
        ph = ",".join(["%s"] * len(aliases))
        with connection.cursor() as cur:
            cur.execute(
                f"DELETE FROM run_table WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({ph})",
                aliases,
            )
    else:
        db_folder = folder
        like_pattern = db_folder + "/%"
        with connection.cursor() as cur:
            cur.execute(
                "DELETE FROM run_table "
                "WHERE TRIM(COALESCE(folder_name, '')) = %s "
                "   OR TRIM(COALESCE(folder_name, '')) LIKE %s",
                [db_folder, like_pattern],
            )

    messages.success(request, f'Run history for "{folder}" has been deleted.')
    next_url = request.POST.get("next", "").strip()
    return redirect(next_url if next_url.startswith("/") else "history")


@login_required
def history_clear_stats(request):
    """Return JSON stats for the Clear History confirmation modal."""
    folder_name = request.GET.get("folder_name", "").strip()
    if not folder_name:
        return JsonResponse({"error": "no folder"}, status=400)

    is_recordings = _is_recordings_folder_name(folder_name)
    _hcs_aliases = list(_recordings_sql_aliases())
    _hcs_ph = ",".join(["%s"] * len(_hcs_aliases))

    # Condition scoped to run_table alias 'r'
    rt_cond = (
        f"LOWER(TRIM(COALESCE(r.folder_name, ''))) IN ({_hcs_ph})"
        if is_recordings
        else "TRIM(r.folder_name) = %s"
    )
    # Same condition for locators/data aliases
    loc_cond = (
        f"LOWER(TRIM(COALESCE(l.folder_name, ''))) IN ({_hcs_ph})"
        if is_recordings
        else "TRIM(l.folder_name) = %s"
    )
    dat_cond = (
        f"LOWER(TRIM(COALESCE(d.folder_name, ''))) IN ({_hcs_ph})"
        if is_recordings
        else "TRIM(d.folder_name) = %s"
    )
    p1 = _hcs_aliases if is_recordings else [folder_name]

    # If specific run_ids are provided, scope all stats to those runs only
    import uuid as _uuid
    raw_ids = request.GET.getlist("run_ids")
    valid_ids = []
    for v in raw_ids:
        try:
            valid_ids.append(str(_uuid.UUID(v.strip())))
        except (ValueError, AttributeError):
            pass

    with connection.cursor() as cur:
        if valid_ids:
            id_ph = ",".join(["%s"] * len(valid_ids))
            # file count = distinct (run_id, folder_name) among selected
            cur.execute(f"""
                SELECT COUNT(*) FROM (
                    SELECT DISTINCT r.run_id, r.folder_name
                    FROM run_table r
                    WHERE r.run_id IN ({id_ph})
                ) sub
            """, valid_ids)
            file_count = cur.fetchone()[0] or 0

            # steps = total rows for selected run_ids
            cur.execute(f"""
                SELECT COALESCE(SUM(step_rows), 0) FROM (
                    SELECT COUNT(*) AS step_rows
                    FROM run_table r
                    WHERE r.run_id IN ({id_ph})
                    GROUP BY r.run_id, r.folder_name
                ) sub
            """, valid_ids)
            step_count = cur.fetchone()[0] or 0

            cur.execute(f"""
                SELECT COUNT(DISTINCT l.id) FROM locators l
                WHERE l.record_id IN (
                    SELECT DISTINCT r.record_id FROM run_table r WHERE r.run_id IN ({id_ph})
                )
            """, valid_ids)
            locator_count = cur.fetchone()[0] or 0

            cur.execute(f"""
                SELECT COUNT(DISTINCT d.id) FROM data d
                WHERE d.record_id IN (
                    SELECT DISTINCT r.record_id FROM run_table r WHERE r.run_id IN ({id_ph})
                )
            """, valid_ids)
            data_count = cur.fetchone()[0] or 0
        else:
            # distinct (run_id, folder_name) pairs = number of files in this folder
            cur.execute(f"""
                SELECT COUNT(*) FROM (
                    SELECT DISTINCT r.run_id, r.folder_name
                    FROM run_table r
                    WHERE {rt_cond}
                ) sub
            """, p1)
            file_count = cur.fetchone()[0] or 0

            # sum of rows per distinct (run_id, folder_name) group = total step records
            cur.execute(f"""
                SELECT COALESCE(SUM(step_rows), 0) FROM (
                    SELECT COUNT(*) AS step_rows
                    FROM run_table r
                    WHERE {rt_cond}
                    GROUP BY r.run_id, r.folder_name
                ) sub
            """, p1)
            step_count = cur.fetchone()[0] or 0

            # locators scoped by folder_name directly
            cur.execute(f"""
                SELECT COUNT(DISTINCT l.id)
                FROM locators l
                WHERE {loc_cond}
                  AND l.record_id IN (
                      SELECT DISTINCT r.record_id FROM run_table r WHERE {rt_cond}
                  )
            """, p1 + p1)
            locator_count = cur.fetchone()[0] or 0

            # data entries scoped by folder_name directly
            cur.execute(f"""
                SELECT COUNT(DISTINCT d.id)
                FROM data d
                WHERE {dat_cond}
                  AND d.record_id IN (
                      SELECT DISTINCT r.record_id FROM run_table r WHERE {rt_cond}
                  )
            """, p1 + p1)
            data_count = cur.fetchone()[0] or 0

    return JsonResponse({
        "file_count":    file_count,
        "step_count":    step_count,
        "locator_count": locator_count,
        "data_count":    data_count,
    })


# ---------------------------------------------------------------------------

@login_required
def sessions_list(request):
    """Show all recorded sessions grouped by folder."""
    global _SCHEMA_ENSURED
    if not _SCHEMA_ENSURED:
        try:
            _ensure_schema_columns()
        except Exception:
            pass
        _SCHEMA_ENSURED = True

    recordings_folder_label = _get_recordings_folder_label()

    with connection.cursor() as cursor:
        cursor.execute("""
            SELECT
                s.record_id,
                COUNT(DISTINCT s.step_no)                AS total_steps,
                COUNT(DISTINCT l.id)                     AS total_locators,
                COUNT(DISTINCT s.data_id)                AS total_data,
                MIN(s.created_at)    AS started_at,
                MAX(s.created_at)    AS last_at,
                MIN(s.page_url)      AS first_url,
                MIN(COALESCE(s.file_order, 1)) AS file_order,
                COALESCE(m.record_name, '') AS record_name,
                m.recorder,
                COALESCE(s.folder_name, 'Unfiled') AS folder,
                BOOL_OR(s.headless_state) AS headless_state,
                MIN(s.pos_x)         AS pos_x,
                MIN(s.pos_y)         AS pos_y,
                BOOL_OR(s.is_primary) AS is_primary,
                MIN(s.locator_rank)  AS locator_rank,
                MIN(s.author)        AS author,
                MIN(s.last_updated_by) AS last_updated_by,
                MIN(s.parent_record_id::text) AS parent_record_id,
                MIN(s.sub_record_id::text)    AS sub_record_id,
                MIN(s.end_record::text)       AS end_record,
                BOOL_OR(s.is_baseline) AS is_baseline,
                COALESCE(m.engine, 'selenium') AS engine
            FROM steps s
            LEFT JOIN locators l ON l.record_id = s.record_id
            LEFT JOIN session_meta m ON m.record_id = s.record_id
            GROUP BY s.record_id, m.record_name, m.recorder, s.folder_name, m.engine
            ORDER BY folder ASC, file_order ASC, started_at ASC;
        """)
        cols = [c.name for c in cursor.description]
        sessions = [dict(zip(cols, row)) for row in cursor.fetchall()]

    # ── Enrich sessions with latest-run outcome ───────────────────────────
    if sessions:
        sids = [s["record_id"] for s in sessions]
        ph   = ",".join(["%s"] * len(sids))
        with connection.cursor() as cur:
            cur.execute(f"""
                WITH latest AS (
                    SELECT DISTINCT ON (record_id) record_id, run_id
                    FROM run_table
                    WHERE record_id IN ({ph})
                    ORDER BY record_id, created_at DESC
                ),
                expected_steps AS (
                    SELECT
                        l.record_id,
                        CASE
                            WHEN EXISTS (SELECT 1 FROM steps s WHERE s.record_id = l.record_id) THEN
                                (SELECT COUNT(DISTINCT s2.step_no) FROM steps s2 WHERE s2.record_id = l.record_id)
                            ELSE
                                (SELECT COUNT(DISTINCT rec.step_no) FROM recordings rec WHERE rec.record_id = l.record_id)
                        END AS expected_total
                    FROM latest l
                )
                SELECT
                    l.record_id,
                    SUM(CASE WHEN r.status = 'fail'         THEN 1 ELSE 0 END) AS fails,
                    SUM(CASE WHEN r.status = 'not_executed' THEN 1 ELSE 0 END) AS not_execs,
                    COUNT(DISTINCT r.step_no) AS actual_total,
                    COALESCE(MAX(es.expected_total), 0) AS expected_total
                FROM run_table r
                INNER JOIN latest l ON l.run_id = r.run_id
                LEFT JOIN expected_steps es ON es.record_id = l.record_id
                GROUP BY l.record_id
            """, sids)
            outcome_map = {}
            for sid, fails, not_execs, actual_total, expected_total in cur.fetchall():
                if fails > 0:
                    status = "fail"
                elif not_execs > 0 or actual_total < expected_total:
                    status = "not_completed"
                else:
                    status = "pass"
                outcome_map[str(sid)] = status
        for s in sessions:
            s["run_status"] = outcome_map.get(str(s["record_id"]), None)

    folder_map: dict[str, list[dict]] = {}
    unfoldered: list[dict] = []
    recordings_sessions: list[dict] = []
    for s in sessions:
        folder_name = (s["folder"] or "").strip()
        if _is_recordings_folder_name(folder_name):
            s["folder"] = recordings_folder_label
            recordings_sessions.append(s)
        elif folder_name:
            folder_map.setdefault(folder_name, []).append(s)
        else:
            unfoldered.append(s)

    sessions = [
        s for s in sessions
        if _is_project_path_visible_to_user(request.user, (s.get("folder") or "").strip())
    ]
    folder_map = {
        folder_name: items
        for folder_name, items in folder_map.items()
        if _is_project_path_visible_to_user(request.user, folder_name)
    }

    registered_folder_rows = _list_registered_project_folders()
    registered_folders = [
        row["folder_name"]
        for row in registered_folder_rows
        if _is_project_path_visible_to_user(request.user, row["folder_name"])
    ]
    folder_order_map = {
        row["folder_name"]: row.get("folder_order")
        for row in registered_folder_rows
    }
    for root_name, root_order in _list_parent_folder_order_map().items():
        if _is_project_path_visible_to_user(request.user, root_name):
            folder_order_map[root_name] = root_order
    for sub_name, sub_order in _list_sub_folder_order_map().items():
        if _is_project_path_visible_to_user(request.user, sub_name):
            folder_order_map[sub_name] = sub_order
    all_folders = []
    if recordings_sessions:
        all_folders.append(recordings_folder_label)
    for folder_name in sorted(
        set(list(folder_map.keys()) + registered_folders),
        key=lambda value: _project_folder_sort_key(value, folder_order_map),
    ):
        all_folders.append(folder_name)

    folder_roots = _build_folder_tree(folder_map, registered_folders, folder_order_map)
    _folder_meta_map = _get_project_folder_metadata_map(list(folder_map.keys()) + registered_folders)
    _annotate_project_tree_metadata(folder_roots, request.user, _folder_meta_map)
    recordings_children: list[dict] = []
    remaining_roots: list[dict] = []
    recordings_prefix = recordings_folder_label + "/"
    for root in folder_roots:
        if root["path"].startswith(recordings_prefix):
            recordings_children.append(root)
        else:
            remaining_roots.append(root)
    folder_roots = remaining_roots
    if recordings_sessions or recordings_children:
        recordings_root = _build_special_folder_group(recordings_folder_label, recordings_sessions)
        recordings_root["children"] = recordings_children
        recordings_root["total_sessions"] = len(recordings_sessions) + sum(child["total_sessions"] for child in recordings_children)
        folder_roots.insert(0, recordings_root)
    if unfoldered:
        folder_roots.append(_build_special_folder_group("", unfoldered))

    recording_pid  = request.session.get("recording_pid")
    recording_url  = request.session.get("recording_url", "")
    recording_name = request.session.get("recording_name", "")
    recording_id   = str(request.session.get("recording_id", ""))
    is_recording = False
    if recording_pid:
        if _is_pid_alive(recording_pid):
            is_recording = True
        else:
            request.session.pop("recording_pid", None)
            request.session.pop("recording_url", None)
            request.session.pop("recording_name", None)
            recording_id = ""

    # Per-user visible actions preference
    _pref_raw = get_user_pref(
        request.user.id,
        "projects.visible_actions",
        ",".join(sorted(_ALL_PROJ_ACTIONS)),
    )
    visible_actions = _resolve_project_visible_actions(_pref_raw)

    _headless_default = get_config("replay.headless_default", "false") == "true"
    _execution_mode = (get_config("replay.execution_mode") or "parallel").strip().lower()

    # ── Cross-module search data (API endpoints + DB test cases) ───────────
    from api_testcases.models import ApiModule, TestCase as ApiTestCase
    from db_testcases.models import TestCase as DbTestCase

    api_modules_qs = ApiModule.objects.order_by('name').values('id', 'name')
    all_api_modules = []
    for mod in api_modules_qs:
        module_name = (mod.get('name') or '').strip()
        encoded_name = requests.utils.quote(module_name, safe='')
        all_api_modules.append({
            'id': mod.get('id'),
            'name': module_name,
            'name_lower': module_name.lower(),
            'testcases_url': f"/api-lab/testcases/?module={encoded_name}" if module_name else '/api-lab/testcases/',
        })

    api_testcases_qs = ApiTestCase.objects.filter(is_active=True).order_by('module', 'name')
    api_endpoints_list = []
    for tc in api_testcases_qs:
        module_name = (tc.module or '').strip()
        encoded_module_name = requests.utils.quote(module_name, safe='') if module_name else ''
        api_endpoints_list.append({
            'id': tc.id,
            'name': tc.name,
            'method': tc.http_method,
            'endpoint': tc.endpoint,
            'module_name': module_name,
            'module_id': None,
            'testcases_url': f"/api-lab/testcases/?module={encoded_module_name}" if encoded_module_name else '/api-lab/testcases/',
        })

    db_testcases_qs = DbTestCase.objects.filter(is_active=True).order_by('sort_order', 'name')
    db_testcases_list = []
    for tc in db_testcases_qs:
        latest = tc.executions.order_by('-executed_at').first()
        db_testcases_list.append({
            'id': tc.id,
            'name': tc.name,
            'test_type': tc.get_test_type_display(),
            'connection': str(tc.connection),
            'folder': str(tc.project_folder) if tc.project_folder else '',
            'status': latest.status if latest else None,
            'executed_at': latest.executed_at if latest else None,
        })

    return render(request, "recorder/sessions.html", {
        "sessions": sessions,
        "folder_roots": folder_roots,
        "all_folders": all_folders,
        "is_recording": is_recording,
        "recording_url": recording_url,
        "recording_name": recording_name,
        "recording_id": recording_id,
        "visible_actions": visible_actions,
        "visible_actions_json": json.dumps(sorted(visible_actions)),
        "all_proj_actions": sorted(_ALL_PROJ_ACTIONS),
        "headless_default": _headless_default,
        "execution_mode": _execution_mode,
        "all_api_modules": all_api_modules,
        "api_endpoints": api_endpoints_list,
        "db_testcases": db_testcases_list,
    })


def _redirect_to_ai_page(request, page_name: str):
    params = request.GET.copy()
    params.pop("embedded", None)
    params.pop("tab", None)
    query = params.urlencode()
    target = redirect(page_name)
    if query:
        target["Location"] = f"{target['Location']}?{query}"
    return target


def _build_ai_databank_context(request, *, embedded_page: bool = False) -> dict[str, Any]:
    query = (request.GET.get("q") or "").strip()
    selected_page_name = (request.GET.get("page_name") or "").strip()
    page_number = request.GET.get("page") or 1

    filters: list[str] = []
    params: list[Any] = []
    if query:
        like = f"%{query}%"
        filters.append("(page_url ILIKE %s OR page_name ILIKE %s OR element_type ILIKE %s OR locator_property::text ILIKE %s)")
        params.extend([like, like, like, like])
    if selected_page_name:
        filters.append("COALESCE(NULLIF(TRIM(page_name), ''), 'Untitled page') = %s")
        params.append(selected_page_name)

    where_sql = f"WHERE {' AND '.join(filters)}" if filters else ""

    with connection.cursor() as cur:
        available_page_names = _get_ai_databank_page_name_options(cur)
        cur.execute(f"SELECT id, page_url, page_name, element_type, locator_property, created_at, updated_at FROM ai_databank {where_sql} ORDER BY updated_at DESC, created_at DESC, id DESC", params)
        rows = cur.fetchall()

    items = []
    for row in rows:
        locator_property = _normalize_ai_databank_locator_property(row[4])
        ordered = _build_ai_databank_ordered_locators(locator_property)
        primary = ordered[0] if ordered else None
        highlight = _build_ai_databank_highlight(locator_property)
        page_name = (row[2] or "").strip() or "Untitled page"
        items.append({
            "id": row[0],
            "page_url": row[1],
            "page_name": page_name,
            "element_type": row[3],
            "locator_property": locator_property,
            "created_at": row[5],
            "updated_at": row[6],
            "primary": primary,
            "available_locators": ordered,
            "tag_name": locator_property.get("tag_name"),
            "text": locator_property.get("text"),
            "highlight": highlight,
        })

    paginator = Paginator(items, 25)
    page_obj = paginator.get_page(page_number)
    return {
        "page_obj": page_obj,
        "query": query,
        "selected_page_name": selected_page_name,
        "available_page_names": available_page_names,
        "total_rows": len(items),
        "current_page_name": "AI" if embedded_page else "AI Databank",
        "current_page_url": request.build_absolute_uri(),
        "ai_databank_tab": "objects",
        "active_ai_section": "databank",
        "embedded_page": embedded_page,
    }


def _build_ai_databank_flow_context(request, *, embedded_page: bool = False) -> dict[str, Any]:
    with connection.cursor() as cur:
        page_cards = _get_ai_databank_flow_cards(cur)
        workflow_names = _get_ai_workflow_name_options(cur)
        workflow_source_sessions = _get_ai_workflow_source_sessions(cur)

    return {
        "page_cards": page_cards,
        "workflow_names": workflow_names,
        "workflow_source_sessions": workflow_source_sessions,
        "total_pages": len(page_cards),
        "current_page_name": "AI" if embedded_page else "AI Databank Flow",
        "current_page_url": request.build_absolute_uri(),
        "ai_databank_tab": "flow",
        "active_ai_section": "flow",
        "embedded_page": embedded_page,
    }


def _build_ai_locator_stats_context(request, *, embedded_page: bool = False) -> dict[str, Any]:
    query = (request.GET.get("q") or "").strip()
    selected_record_name = (request.GET.get("record_name") or "").strip()
    selected_strategy = (request.GET.get("strategy") or "").strip()
    selected_preset = (request.GET.get("preset") or "").strip().lower()
    date_from_str = (request.GET.get("date_from") or "").strip()
    date_to_str = (request.GET.get("date_to") or "").strip()
    page_number = request.GET.get("page") or 1
    failed_page_number = request.GET.get("failed_page") or 1
    date_to = None
    today = _tz.now().date()

    if not date_from_str and not date_to_str:
        if selected_preset == "today":
            date_from = today
            date_to = today
            date_from_str = today.isoformat()
            date_to_str = today.isoformat()
        elif selected_preset == "7d":
            date_to = today
            date_from = today - datetime.timedelta(days=6)
            date_from_str = date_from.isoformat()
            date_to_str = date_to.isoformat()
        elif selected_preset == "this_month":
            date_to = today
            date_from = today.replace(day=1)
            date_from_str = date_from.isoformat()
            date_to_str = date_to.isoformat()
        else:
            selected_preset = "30d"
            date_to = today
            date_from = today - datetime.timedelta(days=29)
            date_from_str = date_from.isoformat()
            date_to_str = date_to.isoformat()

    try:
        if date_from is None and date_from_str:
            date_from = datetime.date.fromisoformat(date_from_str)
    except ValueError:
        date_from = None
        date_from_str = ""
    try:
        if date_to is None and date_to_str:
            date_to = datetime.date.fromisoformat(date_to_str)
    except ValueError:
        date_to = None
        date_to_str = ""

    if date_from_str and date_to_str and not selected_preset:
        if date_from == today and date_to == today:
            selected_preset = "today"
        elif date_to == today and date_from == today - datetime.timedelta(days=6):
            selected_preset = "7d"
        elif date_to == today and date_from == today - datetime.timedelta(days=29):
            selected_preset = "30d"
        elif date_from == today.replace(day=1) and date_to == today:
            selected_preset = "this_month"

    filters: list[str] = []
    params: list[Any] = []
    if query:
        like = f"%{query}%"
        filters.append(
            "(" 
            "COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text) ILIKE %s OR "
            "ls.record_id::text ILIKE %s OR "
            "COALESCE(ls.strategy, '') ILIKE %s OR "
            "COALESCE(ls.action, '') ILIKE %s OR "
            "COALESCE(ls.page_url, '') ILIKE %s OR "
            "COALESCE(ls.locator, '') ILIKE %s"
            ")"
        )
        params.extend([like, like, like, like, like, like])
    if selected_record_name:
        filters.append("COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text) = %s")
        params.append(selected_record_name)
    if selected_strategy:
        filters.append("COALESCE(NULLIF(TRIM(ls.strategy), ''), 'unknown') = %s")
        params.append(selected_strategy)
    if date_from:
        filters.append("DATE(ls.created_at) >= %s")
        params.append(date_from)
    if date_to:
        filters.append("DATE(ls.created_at) <= %s")
        params.append(date_to)

    where_sql = f"WHERE {' AND '.join(filters)}" if filters else ""

    failed_filters: list[str] = ["r.status = 'fail'"]
    failed_params: list[Any] = []
    if query:
        like = f"%{query}%"
        failed_filters.append(
            "("
            "COALESCE(NULLIF(TRIM(m.record_name), ''), r.record_id::text) ILIKE %s OR "
            "r.record_id::text ILIKE %s OR "
            "COALESCE(NULLIF(TRIM(l.strategy), ''), 'unknown') ILIKE %s OR "
            "COALESCE(r.action, '') ILIKE %s OR "
            "COALESCE(r.page_url, '') ILIKE %s OR "
            "COALESCE(l.locator, '') ILIKE %s OR "
            "COALESCE(r.message, '') ILIKE %s"
            ")"
        )
        failed_params.extend([like, like, like, like, like, like, like])
    if selected_record_name:
        failed_filters.append("COALESCE(NULLIF(TRIM(m.record_name), ''), r.record_id::text) = %s")
        failed_params.append(selected_record_name)
    if selected_strategy:
        failed_filters.append("COALESCE(NULLIF(TRIM(l.strategy), ''), 'unknown') = %s")
        failed_params.append(selected_strategy)
    if date_from:
        failed_filters.append("DATE(COALESCE(r.run_date, r.created_at)) >= %s")
        failed_params.append(date_from)
    if date_to:
        failed_filters.append("DATE(COALESCE(r.run_date, r.created_at)) <= %s")
        failed_params.append(date_to)
    failed_where_sql = f"WHERE {' AND '.join(failed_filters)}"

    with connection.cursor() as cur:
        cur.execute(
            """
            SELECT DISTINCT COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text) AS record_name
            FROM locators_stat ls
            LEFT JOIN session_meta m ON m.record_id = ls.record_id
            ORDER BY 1
            """
        )
        available_record_names = [row[0] for row in cur.fetchall() if row[0]]

        cur.execute(
            """
            SELECT DISTINCT COALESCE(NULLIF(TRIM(strategy), ''), 'unknown') AS strategy
            FROM locators_stat
            ORDER BY 1
            """
        )
        available_strategies = [row[0] for row in cur.fetchall() if row[0]]

        cur.execute(
            f"""
            SELECT
                ls.record_id,
                COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text) AS record_name,
                COUNT(*) AS total_hits,
                COUNT(DISTINCT ls.run_id) AS total_runs,
                COUNT(DISTINCT ls.step_no) AS unique_steps,
                MAX(ls.created_at) AS last_used
            FROM locators_stat ls
            LEFT JOIN session_meta m ON m.record_id = ls.record_id
            {where_sql}
            GROUP BY ls.record_id, COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text)
            ORDER BY total_hits DESC, last_used DESC NULLS LAST, record_name ASC
            """,
            params,
        )
        summary_rows = cur.fetchall()

        cur.execute(
            f"""
            SELECT
                COALESCE(NULLIF(TRIM(ls.strategy), ''), 'unknown') AS strategy,
                COUNT(*) AS total_hits
            FROM locators_stat ls
            LEFT JOIN session_meta m ON m.record_id = ls.record_id
            {where_sql}
            GROUP BY COALESCE(NULLIF(TRIM(ls.strategy), ''), 'unknown')
            ORDER BY total_hits DESC, strategy ASC
            """,
            params,
        )
        strategy_rows = cur.fetchall()

        cur.execute(
            f"""
            SELECT
                COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text) AS record_name,
                COALESCE(NULLIF(TRIM(ls.strategy), ''), 'unknown') AS strategy,
                COUNT(*) AS total_hits
            FROM locators_stat ls
            LEFT JOIN session_meta m ON m.record_id = ls.record_id
            {where_sql}
            GROUP BY COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text),
                     COALESCE(NULLIF(TRIM(ls.strategy), ''), 'unknown')
            ORDER BY COUNT(*) DESC, record_name ASC, strategy ASC
            """,
            params,
        )
        record_strategy_rows = cur.fetchall()

        cur.execute(
            f"""
            SELECT
                ls.run_id,
                ls.record_id,
                COALESCE(NULLIF(TRIM(m.record_name), ''), ls.record_id::text) AS record_name,
                ls.step_no,
                COALESCE(ls.action, '') AS action,
                COALESCE(NULLIF(TRIM(ls.strategy), ''), 'unknown') AS strategy,
                COALESCE(ls.locator, '') AS locator,
                ls.locator_rank,
                ls.is_primary,
                COALESCE(ls.page_url, '') AS page_url,
                COALESCE(ls.runner, '') AS runner,
                ls.created_at
            FROM locators_stat ls
            LEFT JOIN session_meta m ON m.record_id = ls.record_id
            {where_sql}
            ORDER BY ls.run_id ASC NULLS LAST, ls.step_no ASC, ls.created_at DESC NULLS LAST
            """,
            params,
        )
        detail_rows = cur.fetchall()

        cur.execute(
            f"""
            SELECT
                r.run_id,
                r.record_id,
                COALESCE(NULLIF(TRIM(m.record_name), ''), r.record_id::text) AS record_name,
                r.step_no,
                COALESCE(r.action, '') AS action,
                COALESCE(NULLIF(TRIM(l.strategy), ''), 'unknown') AS strategy,
                COALESCE(l.locator, '') AS locator,
                l.locator_rank,
                l.is_primary,
                COALESCE(r.page_url, '') AS page_url,
                COALESCE(r.runner, '') AS runner,
                COALESCE(r.message, '') AS message,
                COALESCE(r.run_date, r.created_at) AS run_date
            FROM run_table r
            LEFT JOIN locators l ON l.id = r.locator_id
            LEFT JOIN session_meta m ON m.record_id = r.record_id
            {failed_where_sql}
            ORDER BY r.run_id ASC NULLS LAST, r.step_no ASC, COALESCE(r.run_date, r.created_at) DESC NULLS LAST
            """,
            failed_params,
        )
        failed_rows = cur.fetchall()

    summary_items = [
        {
            "record_id": row[0],
            "record_name": row[1],
            "total_hits": row[2],
            "total_runs": row[3],
            "unique_steps": row[4],
            "last_used": row[5],
        }
        for row in summary_rows
    ]
    strategy_items = [
        {"strategy": row[0], "total_hits": row[1]}
        for row in strategy_rows
    ]
    chart_legend: list[str] = []
    chart_records_map: dict[str, dict[str, Any]] = {}
    for record_name, strategy, total_hits in record_strategy_rows:
        if strategy not in chart_legend:
            chart_legend.append(strategy)
        record_bucket = chart_records_map.setdefault(
            record_name,
            {"record_name": record_name, "total_hits": 0, "segments": {}},
        )
        record_bucket["total_hits"] += total_hits
        record_bucket["segments"][strategy] = total_hits

    chart_palette = [
        "#178472",
        "#1f6fb6",
        "#d97706",
        "#7c3aed",
        "#dc3545",
        "#0ea5e9",
        "#475569",
        "#22c55e",
    ]
    chart_legend_items = [
        {"strategy": strategy, "color": chart_palette[index % len(chart_palette)]}
        for index, strategy in enumerate(chart_legend)
    ]
    chart_color_map = {item["strategy"]: item["color"] for item in chart_legend_items}
    chart_records = sorted(
        chart_records_map.values(),
        key=lambda item: (-item["total_hits"], item["record_name"].lower()),
    )

    run_overlay_map: dict[str, dict[str, int]] = {}
    if chart_records_map:
        record_ids = [item["record_id"] for item in summary_items]
        placeholders = ",".join(["%s"] * len(record_ids))
        run_filters: list[str] = [f"r.record_id IN ({placeholders})"]
        run_params: list[Any] = list(record_ids)
        if selected_record_name:
            run_filters.append("COALESCE(NULLIF(TRIM(m.record_name), ''), r.record_id::text) = %s")
            run_params.append(selected_record_name)
        if date_from:
            run_filters.append("DATE(COALESCE(r.run_date, r.created_at)) >= %s")
            run_params.append(date_from)
        if date_to:
            run_filters.append("DATE(COALESCE(r.run_date, r.created_at)) <= %s")
            run_params.append(date_to)
        run_where_sql = f"WHERE {' AND '.join(run_filters)}" if run_filters else ""
        with connection.cursor() as cur:
            cur.execute(
                f"""
                SELECT
                    record_name,
                    SUM(CASE WHEN run_outcome = 'pass' THEN 1 ELSE 0 END) AS pass_runs,
                    SUM(CASE WHEN run_outcome = 'fail' THEN 1 ELSE 0 END) AS fail_runs,
                    SUM(CASE WHEN run_outcome = 'partial' THEN 1 ELSE 0 END) AS partial_runs
                FROM (
                    SELECT
                        r.run_id,
                        COALESCE(NULLIF(TRIM(m.record_name), ''), r.record_id::text) AS record_name,
                        CASE
                            WHEN SUM(CASE WHEN r.status = 'fail' THEN 1 ELSE 0 END) > 0 THEN 'fail'
                            WHEN SUM(CASE WHEN r.status = 'not_executed' THEN 1 ELSE 0 END) > 0 THEN 'partial'
                            ELSE 'pass'
                        END AS run_outcome
                    FROM run_table r
                    LEFT JOIN session_meta m ON m.record_id = r.record_id
                    {run_where_sql}
                    GROUP BY r.run_id, COALESCE(NULLIF(TRIM(m.record_name), ''), r.record_id::text)
                ) run_summary
                GROUP BY record_name
                """,
                run_params,
            )
            run_overlay_map = {
                row[0]: {
                    "pass_runs": row[1] or 0,
                    "fail_runs": row[2] or 0,
                    "partial_runs": row[3] or 0,
                }
                for row in cur.fetchall()
            }

    if not selected_record_name:
        chart_records = chart_records[:8]
    for item in chart_records:
        total_hits = item["total_hits"] or 1
        ordered_segments = []
        for legend_item in chart_legend_items:
            strategy = legend_item["strategy"]
            hits = item["segments"].get(strategy, 0)
            if hits <= 0:
                continue
            ordered_segments.append({
                "strategy": strategy,
                "hits": hits,
                "pct": round((hits / total_hits) * 100, 2),
                "color": chart_color_map[strategy],
            })
        item["segments"] = ordered_segments
        outcome = run_overlay_map.get(item["record_name"], {"pass_runs": 0, "fail_runs": 0, "partial_runs": 0})
        total_runs = outcome["pass_runs"] + outcome["fail_runs"] + outcome["partial_runs"]
        overlay_segments = []
        if total_runs > 0:
            for label, key, color in (
                ("Pass", "pass_runs", "#198754"),
                ("Fail", "fail_runs", "#dc3545"),
                ("Not Completed", "partial_runs", "#d97706"),
            ):
                count = outcome[key]
                if count <= 0:
                    continue
                overlay_segments.append({
                    "label": label,
                    "count": count,
                    "pct": round((count / total_runs) * 100, 2),
                    "color": color,
                })
        item["run_overlay"] = {
            "total_runs": total_runs,
            "segments": overlay_segments,
            **outcome,
        }

    detail_items = [
        {
            "run_id": row[0],
            "record_id": row[1],
            "record_name": row[2],
            "step_no": row[3],
            "action": row[4],
            "strategy": row[5],
            "locator": row[6],
            "locator_rank": row[7],
            "is_primary": row[8],
            "page_url": row[9],
            "runner": row[10],
            "created_at": row[11],
        }
        for row in detail_rows
    ]

    paginator = Paginator(detail_items, 50)
    page_obj = paginator.get_page(page_number)

    failed_items = [
        {
            "run_id": row[0],
            "record_id": row[1],
            "record_name": row[2],
            "step_no": row[3],
            "action": row[4],
            "strategy": row[5],
            "locator": row[6],
            "locator_rank": row[7],
            "is_primary": row[8],
            "page_url": row[9],
            "runner": row[10],
            "message": row[11],
            "run_date": row[12],
        }
        for row in failed_rows
    ]
    failed_paginator = Paginator(failed_items, 50)
    failed_page_obj = failed_paginator.get_page(failed_page_number)
    total_hits = sum(item["total_hits"] for item in summary_items)
    top_strategy = strategy_items[0]["strategy"] if strategy_items else "-"
    last_used = max((item["created_at"] for item in detail_items if item["created_at"]), default=None)

    return {
        "page_obj": page_obj,
        "failed_page_obj": failed_page_obj,
        "query": query,
        "summary_items": summary_items,
        "strategy_items": strategy_items,
        "chart_records": chart_records,
        "chart_legend_items": chart_legend_items,
        "total_hits": total_hits,
        "total_records": len(summary_items),
        "top_strategy": top_strategy,
        "last_used": last_used,
        "current_page_name": "AI" if embedded_page else "Locator Stats",
        "current_page_url": request.build_absolute_uri(),
        "ai_databank_tab": "stats",
        "active_ai_section": "stats",
        "embedded_page": embedded_page,
    }


def _build_workflow_assistant_context(request, *, embedded_page: bool = False) -> dict[str, Any]:
    _ensure_ai_workflow_schema()
    rag_status = {}
    try:
        from llm_workflow_assistant.rag_service import get_rag_status

        rag_status = get_rag_status()
    except Exception:
        rag_status = {}

    ollama_api = (get_config("ai.ollama_api", "http://localhost:11434/api") or "http://localhost:11434/api").strip()
    ollama_model = (get_config("ai.ollama_model", "llama3.2:3b") or "llama3.2:3b").strip()
    ft_model = (get_config("ai.ft_model", "llama3-finetuned") or "llama3-finetuned").strip()
    with connection.cursor() as cur:
        workflow_names = _get_ai_workflow_name_options(cur)

    return {
        "ollama_api": ollama_api,
        "ollama_model": ollama_model,
        "ft_model": ft_model,
        "workflow_names": workflow_names,
        "rag_status": rag_status,
        "ft_status": {},
        "models": [],
        "current_page_name": "AI" if embedded_page else "AI RAG",
        "current_page_url": request.build_absolute_uri(),
        "active_ai_section": "workflow",
        "embedded_page": embedded_page,
    }


@login_required
def ai_hub(request):
    tab = (request.GET.get("tab") or "databank").strip().lower()
    if tab not in {"databank", "flow", "stats", "workflow"}:
        tab = "databank"
    page_name = {
        "databank": "ai_databank",
        "flow": "ai_databank_flow",
        "stats": "ai_locator_stats",
        "workflow": "workflow_assistant",
    }[tab]
    return _redirect_to_ai_page(request, page_name)


@login_required
def ai_databank(request):
    _ensure_ai_databank_schema()
    _ensure_ai_workflow_schema()
    embedded_page = (request.GET.get("embedded") or "").strip() == "1"
    ctx = _build_ai_databank_context(request, embedded_page=embedded_page)
    # Merge flow context so the Flow tab renders in-page
    flow_ctx = _build_ai_databank_flow_context(request, embedded_page=embedded_page)
    ctx.update({
        "page_cards": flow_ctx.get("page_cards", []),
        "workflow_names": flow_ctx.get("workflow_names", []),
        "workflow_source_sessions": flow_ctx.get("workflow_source_sessions", []),
        "total_pages": flow_ctx.get("total_pages", 0),
    })
    # Check if user navigated to the flow tab
    if request.GET.get("tab") == "flow":
        ctx["active_inner_tab"] = "flow"
    else:
        ctx["active_inner_tab"] = "databank"
    return render(request, "recorder/ai_databank.html", ctx)


@login_required
def ai_databank_flow(request):
    """Redirect to the unified AI Databank page with the Flow tab active."""
    url = reverse("ai_databank") + "?tab=flow"
    embedded = (request.GET.get("embedded") or "").strip()
    if embedded == "1":
        url += "&embedded=1"
    return redirect(url)


@login_required
def ai_locator_stats(request):
    _ensure_ai_databank_schema()
    embedded_page = (request.GET.get("embedded") or "").strip() == "1"
    return render(request, "recorder/ai_locator_stats.html", _build_ai_locator_stats_context(request, embedded_page=embedded_page))


@login_required
def workflow_assistant(request):
    _ensure_ai_databank_schema()
    _ensure_ai_workflow_schema()
    embedded_page = (request.GET.get("embedded") or "").strip() == "1"
    return render(request, "recorder/workflow_assistant.html", _build_workflow_assistant_context(request, embedded_page=embedded_page))


@csrf_exempt
@login_required
def workflow_assistant_query(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required."}, status=405)

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid JSON payload."}, status=400)

    message = str(body.get("message") or "").strip()
    ollama_api = str(body.get("ollama_api") or get_config("ai.ollama_api", "http://localhost:11434/api") or "http://localhost:11434/api").strip()
    ollama_model = str(body.get("model") or get_config("ai.ollama_model", "llama3.2:3b") or "llama3.2:3b").strip()

    if not message:
        return JsonResponse({"ok": False, "error": "Message is required."}, status=400)

    try:
        from llm_workflow_assistant.rag_service import query_workflow_assistant

        result = query_workflow_assistant(
            message,
            tenant_id=_get_user_tenant_id(request.user),
            ollama_api=ollama_api,
            model=ollama_model,
        )
    except requests.RequestException as exc:
        return JsonResponse({"ok": False, "error": f"Ollama request failed: {exc}"}, status=502)
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({
        "ok": True,
        "answer": result.get("answer", ""),
        "model": result.get("model", ollama_model),
        "ollama_api": result.get("ollama_api", ollama_api),
        "context": result.get("context", {}),
        "timings": result.get("timings", {}),
    })


@csrf_exempt
@login_required
def workflow_assistant_rebuild(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required."}, status=405)

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid JSON payload."}, status=400)

    ollama_api = str(body.get("ollama_api") or get_config("ai.ollama_api", "http://localhost:11434/api") or "http://localhost:11434/api").strip()

    try:
        from llm_workflow_assistant.rag_service import queue_rag_rebuild 

        result = queue_rag_rebuild(
            ollama_api=ollama_api,
            timeout=300,
        )
    except requests.RequestException as exc:
        return JsonResponse({"ok": False, "error": f"Ollama request failed: {exc}"}, status=502)
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({
        "ok": True,
        "sync": result,
    })


@login_required
def workflow_assistant_status(request):
    if request.method != "GET":
        return JsonResponse({"ok": False, "error": "GET required."}, status=405)

    try:
        from llm_workflow_assistant.rag_service import get_rag_status

        result = get_rag_status()
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({
        "ok": True,
        "sync": result,
    })


@csrf_exempt
@login_required
def workflow_assistant_search(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required."}, status=405)

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid JSON payload."}, status=400)

    query = str(body.get("query") or "").strip()
    if not query:
        return JsonResponse({"ok": False, "error": "Query is required."}, status=400)

    ollama_api = str(body.get("ollama_api") or get_config("ai.ollama_api", "http://localhost:11434/api") or "http://localhost:11434/api").strip()
    top_k = int(body.get("top_k") or 10)

    try:
        from llm_workflow_assistant.finetune_service import search_rag_documents

        result = search_rag_documents(query=query, ollama_api=ollama_api, top_k=top_k)
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({"ok": True, **result})


def _normalize_generate_prompt(prompt: str) -> str:
    text = (prompt or "").strip()
    if not text:
        return ""

    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(
        r"^(?:generate|create|build)(?:\s+(?:a|the))?(?:\s+(?:full|complete))?(?:\s+(?:robot|test))?(?:\s+(?:framework))?(?:\s+(?:file|script|suite))?(?:\s+for)?\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"^(?:test\s+case|test\s+script|script|session|record)\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    return text.strip().strip('"').strip("'")


def _find_generate_record_matches(prompt: str) -> list[dict[str, Any]]:
    raw = (prompt or "").strip()
    if not raw:
        return []

    uuid_match = re.search(
        r"\b[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}\b",
        raw,
    )
    target = uuid_match.group(0) if uuid_match else _normalize_generate_prompt(raw)
    if not target:
        return []

    with connection.cursor() as cur:
        if uuid_match:
            cur.execute(
                """
                SELECT
                    record_id::text,
                    COALESCE(NULLIF(TRIM(record_name), ''), record_id::text) AS record_name,
                    COALESCE(NULLIF(TRIM(folder_name), ''), '') AS folder_name,
                    created_at
                FROM session_meta
                WHERE record_id = %s::uuid
                ORDER BY created_at DESC
                LIMIT 5
                """,
                [target],
            )
        else:
            cur.execute(
                """
                SELECT
                    record_id::text,
                    COALESCE(NULLIF(TRIM(record_name), ''), record_id::text) AS record_name,
                    COALESCE(NULLIF(TRIM(folder_name), ''), '') AS folder_name,
                    created_at
                FROM session_meta
                WHERE record_name ILIKE %s
                ORDER BY
                    CASE WHEN LOWER(record_name) = LOWER(%s) THEN 0 ELSE 1 END,
                    created_at DESC
                LIMIT 10
                """,
                [f"%{target}%", target],
            )
        rows = cur.fetchall()

    matches = []
    for record_id, record_name, folder_name, created_at in rows:
        matches.append(
            {
                "record_id": record_id,
                "record_name": record_name,
                "folder_name": folder_name or "",
                "created_at": created_at,
            }
        )
    return matches


def _safe_json_value(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (dict, list)):
        return value
    if isinstance(value, str):
        try:
            return json.loads(value)
        except Exception:
            return value
    return value


def _step_identity_locators(step_obj: Any, all_locators: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Augment rendered identity rows with human-readable checkbox/radio ids and labels."""
    rendered = [dict(item) for item in (all_locators or [])]
    raw_event = getattr(step_obj, "raw_event", None)
    if not isinstance(raw_event, dict):
        return rendered

    input_type = str(raw_event.get("inputType") or raw_event.get("type") or "").strip().lower()
    if input_type not in {"checkbox", "radio"}:
        return rendered

    info = raw_event.get("selenium_info") if isinstance(raw_event.get("selenium_info"), dict) else {}
    pw_info = raw_event.get("playwright_info") if isinstance(raw_event.get("playwright_info"), dict) else {}
    attrs = {}
    for source in (info, pw_info):
        if isinstance(source.get("attributes"), dict):
            attrs.update(source["attributes"])

    raw_id = str(raw_event.get("id") or info.get("id") or pw_info.get("id") or attrs.get("id") or "").strip()
    text = str(
        raw_event.get("text")
        or (raw_event.get("locators") or {}).get("text")
        or info.get("labelText")
        or info.get("accessibleName")
        or pw_info.get("labelText")
        or pw_info.get("accessibleName")
        or attrs.get("aria-label")
        or ""
    ).strip()
    label = str(
        (raw_event.get("locators") or {}).get("label")
        or info.get("labelText")
        or info.get("accessibleName")
        or pw_info.get("labelText")
        or pw_info.get("accessibleName")
        or attrs.get("aria-label")
        or text
        or ""
    ).strip()
    label_for = f"for={raw_id}" if raw_id else label

    seen_strategies = {str(item.get("strategy") or "").strip().lower() for item in rendered}
    for item in rendered:
        strategy = str(item.get("strategy") or "").strip().lower()
        if strategy == "id" and raw_id:
            item["display_locator"] = raw_id
        elif strategy == "label" and label_for:
            item["display_locator"] = label_for
        elif strategy == "text" and text:
            item["display_locator"] = text

    if raw_id and "id" not in seen_strategies:
        rendered.append({
            "strategy": "id",
            "locator": raw_id,
            "display_locator": raw_id,
            "is_primary": False,
            "rank": 2,
        })
    if label and "label" not in seen_strategies:
        rendered.append({
            "strategy": "label",
            "locator": label,
            "display_locator": label_for,
            "is_primary": False,
            "rank": 10,
        })
    if text and "text" not in seen_strategies:
        rendered.append({
            "strategy": "text",
            "locator": text,
            "display_locator": text,
            "is_primary": False,
            "rank": 11,
        })

    return sorted(rendered, key=lambda item: (int(item.get("rank") or 99), str(item.get("strategy") or "")))


def _step_strategy_options(step_obj: Any, all_locators: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Build selectable locator strategies, including synthetic checkbox/radio label rows."""
    options: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in all_locators or []:
        strategy = str(item.get("strategy") or "").strip()
        if not strategy or strategy in seen:
            continue
        seen.add(strategy)
        options.append({
            "strategy": strategy,
            "is_primary": bool(item.get("is_primary")),
            "locator": item.get("locator") or "",
            "locator_id": item.get("locator_id"),
            "rank": int(item.get("rank") or 99),
        })
    return options


def _synthetic_step_locator(step_obj: Any, strategy: str) -> tuple[str | None, int | None]:
    """Return a locator value for strategies shown in the UI but not yet stored in locators."""
    raw_event = getattr(step_obj, "raw_event", None)
    if not isinstance(raw_event, dict):
        return (None, None)

    strategy = str(strategy or "").strip().lower()
    input_type = str(raw_event.get("inputType") or raw_event.get("type") or "").strip().lower()
    info = raw_event.get("selenium_info") if isinstance(raw_event.get("selenium_info"), dict) else {}
    pw_info = raw_event.get("playwright_info") if isinstance(raw_event.get("playwright_info"), dict) else {}
    attrs = {}
    for source in (info, pw_info):
        if isinstance(source.get("attributes"), dict):
            attrs.update(source["attributes"])

    raw_id = str(raw_event.get("id") or info.get("id") or pw_info.get("id") or attrs.get("id") or "").strip()
    if strategy == "label" and input_type in {"checkbox", "radio"} and raw_id:
        return (f"for={raw_id}", 10)
    return (None, None)


def _load_record_bundle(record_id: str) -> dict[str, Any]:
    with connection.cursor() as cur:
        cur.execute(
            """
            SELECT
                m.record_id::text,
                COALESCE(NULLIF(TRIM(m.record_name), ''), m.record_id::text) AS record_name,
                COALESCE(NULLIF(TRIM(m.folder_name), ''), '') AS folder_name,
                m.created_at
            FROM session_meta m
            WHERE m.record_id = %s::uuid
            LIMIT 1
            """,
            [record_id],
        )
        meta_row = cur.fetchone()
        if not meta_row:
            raise ValueError("Record not found in session_meta.")

        cur.execute(
            """
            SELECT
                s.step_no,
                COALESCE(s.action, ''),
                COALESCE(s.page_url, ''),
                COALESCE(s.element_tag, ''),
                COALESCE(s.field_name, ''),
                COALESCE(s.field_value, ''),
                COALESCE(s.steps_description, ''),
                COALESCE(s.page_title, ''),
                COALESCE(s.strategy, ''),
                COALESCE(s.locator, ''),
                s.locator_id,
                s.data_id,
                s.raw_event,
                s.locators_raw,
                COALESCE(s.validation, ''),
                d.id,
                COALESCE(d.field_name, ''),
                COALESCE(d.value, ''),
                l.id,
                COALESCE(l.strategy, ''),
                COALESCE(l.locator, ''),
                COALESCE(l.locator_rank, 0),
                COALESCE(l.is_primary, FALSE)
            FROM steps s
            LEFT JOIN data d ON d.id = s.data_id
            LEFT JOIN locators l ON l.id = s.locator_id
            WHERE s.record_id = %s::uuid
            ORDER BY s.step_no ASC, s.id ASC
            """,
            [record_id],
        )
        step_rows = cur.fetchall()

    meta = {
        "record_id": meta_row[0],
        "record_name": meta_row[1],
        "folder_name": meta_row[2],
        "created_at": meta_row[3],
    }
    steps = []
    for row in step_rows:
        steps.append(
            {
                "step_no": row[0],
                "action": row[1],
                "page_url": row[2],
                "element_tag": row[3],
                "field_name": row[4],
                "field_value": row[5],
                "steps_description": row[6],
                "page_title": row[7],
                "step_strategy": row[8],
                "step_locator": row[9],
                "locator_id": row[10],
                "data_id": row[11],
                "raw_event": _safe_json_value(row[12]),
                "locators_raw": _safe_json_value(row[13]),
                "validation": row[14],
                "data_field_name": row[16],
                "data_value": row[17],
                "locator_strategy": row[19],
                "locator_value": row[20],
                "locator_rank": row[21],
                "locator_is_primary": row[22],
            }
        )
    return {"meta": meta, "steps": steps}


def _robot_safe_name(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 _-]+", " ", (value or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned or fallback


def _robot_cell(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r", " ").replace("\n", " ")
    return re.sub(r"\s+", " ", text).strip()


def _robot_key_token(value: str) -> str:
    normalized = _robot_cell(value).strip()
    if not normalized:
        return ""

    aliases = {
        "enter": "ENTER",
        "return": "ENTER",
        "ctrl": "CONTROL",
        "control": "CONTROL",
        "cmd": "COMMAND",
        "command": "COMMAND",
        "meta": "META",
        "win": "META",
        "windows": "META",
        "option": "ALT",
        "alt": "ALT",
        "shift": "SHIFT",
        "tab": "TAB",
        "esc": "ESCAPE",
        "escape": "ESCAPE",
        "space": "SPACE",
        "spacebar": "SPACE",
        "backspace": "BACKSPACE",
        "delete": "DELETE",
        "del": "DELETE",
        "home": "HOME",
        "end": "END",
        "pageup": "PAGE_UP",
        "page up": "PAGE_UP",
        "pagedown": "PAGE_DOWN",
        "page down": "PAGE_DOWN",
        "arrowup": "ARROW_UP",
        "arrow up": "ARROW_UP",
        "up": "ARROW_UP",
        "arrowdown": "ARROW_DOWN",
        "arrow down": "ARROW_DOWN",
        "down": "ARROW_DOWN",
        "arrowleft": "ARROW_LEFT",
        "arrow left": "ARROW_LEFT",
        "left": "ARROW_LEFT",
        "arrowright": "ARROW_RIGHT",
        "arrow right": "ARROW_RIGHT",
        "right": "ARROW_RIGHT",
    }

    direct = aliases.get(normalized.lower())
    if direct:
        return direct

    if len(normalized) == 1:
        return normalized

    if re.fullmatch(r"f\d{1,2}", normalized, re.IGNORECASE):
        return normalized.upper()

    chord_parts = [part for part in re.split(r"\s*(?:\+|-)\s*", normalized) if part]
    if len(chord_parts) > 1:
        mapped_parts = [_robot_key_token(part) for part in chord_parts]
        mapped_parts = [part for part in mapped_parts if part]
        if mapped_parts:
            return "+".join(mapped_parts)

    return normalized.upper().replace(" ", "_")


def _robot_input_type(step: dict[str, Any]) -> str:
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    return _robot_cell(raw_event.get("inputType")).lower()


def _robot_truthy_state(value: Any) -> bool | None:
    normalized = _robot_cell(value).lower()
    if normalized in {"true", "1", "yes", "on", "checked", "selected"}:
        return True
    if normalized in {"false", "0", "no", "off", "unchecked", "unselected", ""}:
        return False
    return None


def _robot_toggle_lines(locator: str, input_type: str, value: Any) -> list[str]:
    desired_state = _robot_truthy_state(value)
    label = "checkbox" if input_type == "checkbox" else "radio"

    if input_type == "checkbox":
        if desired_state is True:
            return [f"    Select Checkbox    {locator}"]
        if desired_state is False:
            return [f"    Unselect Checkbox    {locator}"]
        return [
            f"    Click Element    {locator}",
            f"    Comment    Recorded {label} state was not captured explicitly; verify the final checked state.",
        ]

    return [
        f"    Click Element    {locator}",
        f"    Comment    Recorded {label} selection executed as a click.",
    ]


def _robot_select_lines(locator: str, value: str, raw_event: dict[str, Any]) -> list[str]:
    visible_text = _robot_cell(raw_event.get("text"))
    option_value = _robot_cell(raw_event.get("value"))

    if visible_text:
        return [f"    Select From List By Label    {locator}    {visible_text}"]
    if value:
        return [f"    Select From List By Label    {locator}    {value}"]
    if option_value:
        return [f"    Select From List By Value    {locator}    {option_value}"]
    return [f"    Wait Until Element Is Visible    {locator}"]


def _preferred_robot_locator(step: dict[str, Any]) -> str:
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    locators_raw = step.get("locators_raw") if isinstance(step.get("locators_raw"), dict) else {}

    for source in (locators_raw, raw_event.get("locators") if isinstance(raw_event.get("locators"), dict) else {}):
        for strategy in ("xpath", "id", "name", "css", "linkText", "partialLinkText"):
            locator_value = source.get(strategy) if isinstance(source, dict) else None
            if isinstance(locator_value, str) and locator_value.strip():
                return f"{strategy}:{locator_value.strip()}"

    strategy = _robot_cell(step.get("locator_strategy") or step.get("step_strategy"))
    locator = _robot_cell(step.get("locator_value") or step.get("step_locator"))
    if strategy and locator:
        if ":" in locator:
            return locator
        return f"{strategy}:{locator}"

    field_name = _robot_cell(step.get("data_field_name") or step.get("field_name"))
    if field_name:
        return f"name:{field_name}"
    return ""


def _robot_step_lines(step: dict[str, Any]) -> list[str]:
    step_no = step.get("step_no")
    action = _robot_cell(step.get("action")).lower()
    page_url = _robot_cell(step.get("page_url"))
    locator = _preferred_robot_locator(step)
    value = _robot_cell(step.get("data_value") or step.get("field_value"))
    field_name = _robot_cell(step.get("data_field_name") or step.get("field_name"))
    description = _robot_cell(step.get("steps_description"))
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    key_name = _robot_key_token(raw_event.get("key") if isinstance(raw_event, dict) else "")
    element_tag = _robot_cell(step.get("element_tag")).lower()
    input_type = _robot_input_type(step)

    lines: list[str] = []

    if action in {"navigate", "open", "goto"} and page_url:
        lines.append(f"    Go To    {page_url}")
        return lines
    if action == "click" and locator:
        if input_type in {"checkbox", "radio"}:
            lines.extend(_robot_toggle_lines(locator, input_type, value))
            return lines
        lines.append(f"    Click Element    {locator}")
        return lines
    if action == "dblclick" and locator:
        lines.append(f"    Double Click Element    {locator}")
        return lines
    if action == "contextmenu" and locator:
        lines.append(f"    Open Context Menu    {locator}")
        return lines
    if action in {"input", "change"} and locator:
        if element_tag == "select":
            lines.extend(_robot_select_lines(locator, value, raw_event))
        elif input_type in {"checkbox", "radio"}:
            lines.extend(_robot_toggle_lines(locator, input_type, value))
        elif value:
            lines.append(f"    Input Text    {locator}    {value}")
        else:
            lines.append(f"    Wait Until Element Is Visible    {locator}")
        return lines
    if action == "keydown" and key_name:
        if key_name == "TAB":
            lines.append("    No Operation")
            return lines
        target = locator or "NONE"
        lines.append(f"    Press Keys    {target}    {key_name}")
        return lines
    if action == "submit" and locator:
        lines.append(f"    Submit Form    {locator}")
        return lines
    if action == "scroll":
        raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
        delta_x = int(raw_event.get("delta_x") or 0)
        delta_y = int(raw_event.get("delta_y") or 0)
        lines.append(f"    Execute Javascript    window.scrollBy({delta_x}, {delta_y})")
        return lines
    if action == "navigate_back":
        lines.append("    Go Back")
        return lines
    if action == "navigate_forward":
        lines.append("    Execute Javascript    window.history.forward()")
        return lines
    if action == "navigate_unknown":
        lines.append("    No Operation")
        return lines
    lines.append("    No Operation")
    return lines


def _build_robot_script(record_bundle: dict[str, Any]) -> str:
    meta = record_bundle["meta"]
    steps = record_bundle["steps"]
    record_name = _robot_safe_name(meta.get("record_name", ""), f"Session {meta['record_id'][:8]}")
    test_case_name = _robot_safe_name(meta.get("record_name", ""), f"Generated Session {meta['record_id'][:8]}")
    folder_name = _robot_cell(meta.get("folder_name"))
    first_url = next((_robot_cell(step.get("page_url")) for step in steps if _robot_cell(step.get("page_url"))), "about:blank")

    # Resolve chromedriver path from configuration; fall back to first file found
    webdriver_filename = (get_config("chrome.webdriver_path") or "").strip()
    if webdriver_filename:
        _chrome_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "webdrivers", "chrome")
        chromedriver_path = os.path.join(_chrome_dir, webdriver_filename).replace("\\", "/")
    else:
        try:
            _chrome_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "webdrivers", "chrome")
            candidates = sorted(f for f in os.listdir(_chrome_dir) if f.lower().endswith(".exe"))
            chromedriver_path = os.path.join(_chrome_dir, candidates[0]).replace("\\", "/") if candidates else "C:/web__automation/webdrivers/chrome/chromedriver.exe"
        except Exception:
            chromedriver_path = "C:/web__automation/webdrivers/chrome/chromedriver.exe"

    # Build DB URL for the LocatorStatListener — read from Django settings
    # which already resolves env vars / .env / settings.ini via python-decouple.
    _db = settings.DATABASES["default"]
    _db_user = _db.get("USER") or "postgres"
    _db_pass = _db.get("PASSWORD") or ""
    _db_host = _db.get("HOST") or "localhost"
    _db_port = str(_db.get("PORT") or "5432")
    _db_name = _db.get("NAME") or "automation_db"
    # Encode password special chars so the URL stays valid
    _db_pass_enc = _db_pass.replace("@", "%40").replace(":", "%3A")
    db_url = f"postgresql://{_db_user}:{_db_pass_enc}@{_db_host}:{_db_port}/{_db_name}"

    # Absolute path to the listener file (project root)
    _base_dir = os.path.dirname(os.path.dirname(__file__))
    _listener_path = os.path.join(_base_dir, "LocatorStatListener.py").replace("\\", "/")

    # Align variable values to column 20 (2-space minimum separator)
    def _var(name: str, value: str) -> str:
        token = f"${{{name}}}"
        return token + " " * max(2, 20 - len(token)) + value

    lines = [
        "*** Settings ***",
        "Library    SeleniumLibrary",
        f"Library    {_listener_path}",
        f"...        ${{RECORD_ID}}    record_name=${{RECORD_NAME}}",
        f"...        folder_name=${{PROJECT_FOLDER}}    db_url=${{DB_URL}}",
        "Suite Setup    Open Recorded Application",
        "Suite Teardown    Close All Browsers",
        "Test Teardown    Capture Page Screenshot",
        "",
        "*** Variables ***",
        _var("BROWSER", "Chrome"),
        _var("START_URL", first_url),
        _var("RECORD_ID", meta["record_id"]),
        _var("RECORD_NAME", record_name),
        _var("PROJECT_FOLDER", folder_name or "(none)"),
        _var("DEFAULT_TIMEOUT", "10s"),
        _var("DB_URL", db_url),
        "",
        "*** Test Cases ***",
        test_case_name,
        f"    [Documentation]    Generated from session_meta, steps, data, and locators for record_id ${{RECORD_ID}}.",
        "    Set Test Documentation    Generated from PostgreSQL source tables.",
        "    Execute Recorded Workflow",
        "",
        "*** Keywords ***",
        "Open Recorded Application",
        f"    Create Webdriver    Chrome    executable_path={chromedriver_path}",
        "    Go To    ${START_URL}",
        "    Maximize Browser Window",
        "    Set Selenium Timeout    ${DEFAULT_TIMEOUT}",
        "",
        "Execute Recorded Workflow",
    ]

    for step in steps:
        lines.extend(_robot_step_lines(step))

    return "\n".join(lines).strip() + "\n"


# ══════════════════════════════════════════════════════════════════════════════
# Playwright Python Code Generator (playwright codegen style)
# ══════════════════════════════════════════════════════════════════════════════

def _pw_locator(step: dict[str, Any]) -> str:
    """Build a Playwright-style locator string from step data."""
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    locators_raw = step.get("locators_raw") if isinstance(step.get("locators_raw"), dict) else {}
    pw_info = _pw_get_playwright_info(step)
    attrs = pw_info.get("attributes") if isinstance(pw_info.get("attributes"), dict) else {}

    # Merge locator sources
    locs: dict[str, str] = {}
    for source in (locators_raw, raw_event.get("locators") if isinstance(raw_event.get("locators"), dict) else {}):
        if isinstance(source, dict):
            for k, v in source.items():
                if isinstance(v, str) and v.strip() and k not in locs:
                    locs[k] = v.strip()

    for attr_name, strategy in {
        "id": "id",
        "name": "name",
        "placeholder": "placeholder",
        "role": "role",
        "title": "title",
        "alt": "alt",
        "href": "href",
        "aria-label": "ariaLabel",
        "data-testid": "dataTestId",
        "data-test-id": "dataTestId",
        "type": "type",
    }.items():
        attr_val = attrs.get(attr_name)
        if isinstance(attr_val, str) and attr_val.strip() and strategy not in locs:
            locs[strategy] = attr_val.strip()

    semantic_locator = _pw_semantic_locator(step)
    if semantic_locator:
        return semantic_locator

    # Playwright prefers role/test-id/label/text locators over CSS/XPath
    # 1. data-testid
    if locs.get("dataTestId"):
        return f'page.get_by_test_id("{_pw_escape(locs["dataTestId"])}")'
    # 2. aria-label / role
    if locs.get("ariaLabel"):
        role = locs.get("role", "").lower()
        label = locs["ariaLabel"]
        if role and role in ("button", "link", "textbox", "checkbox", "radio", "combobox", "heading", "tab", "menuitem", "option", "listbox"):
            return f'page.get_by_role("{role}", name="{_pw_escape(label)}")'
        return f'page.get_by_label("{_pw_escape(label)}")'
    # 3. placeholder
    if locs.get("placeholder"):
        return f'page.get_by_placeholder("{_pw_escape(locs["placeholder"])}")'
    # 4. link text (for <a>)
    element_tag = (step.get("element_tag") or "").lower()
    if element_tag == "a" and locs.get("linkText"):
        return f'page.get_by_role("link", name="{_pw_escape(locs["linkText"])}")'
    # 5. text content (button/link/heading)
    if locs.get("text") and element_tag in ("button", "a", "h1", "h2", "h3", "h4", "h5", "h6", "span", "label"):
        text = locs["text"]
        if element_tag == "button":
            return f'page.get_by_role("button", name="{_pw_escape(text)}")'
        if element_tag == "a":
            return f'page.get_by_role("link", name="{_pw_escape(text)}")'
        if element_tag.startswith("h"):
            return f'page.get_by_role("heading", name="{_pw_escape(text)}")'
        return f'page.get_by_text("{_pw_escape(text)}")'
    # 6. id
    if locs.get("id"):
        return f'page.locator("#{_pw_css_escape(_pw_normalize_id_value(locs["id"]))}")'
    # 7. name attribute
    if locs.get("name"):
        if _pw_looks_like_full_selector(locs["name"]):
            return f'page.locator("{_pw_escape(locs["name"])}")'
        return f'page.locator("[name=\\"{_pw_escape(locs["name"])}\\"]")'
    # 8. CSS selector
    if locs.get("css"):
        return f'page.locator("{_pw_escape(locs["css"])}")'
    # 9. XPath
    if locs.get("xpath"):
        return f'page.locator("xpath={_pw_escape(locs["xpath"])}")'
    # 10. Fallback from step-level data
    strategy = (step.get("locator_strategy") or step.get("step_strategy") or "").strip()
    locator_val = (step.get("locator_value") or step.get("step_locator") or "").strip()
    if strategy and locator_val:
        if strategy == "id":
            return f'page.locator("#{_pw_css_escape(_pw_normalize_id_value(locator_val))}")'
        if strategy == "name":
            if _pw_looks_like_full_selector(locator_val):
                return f'page.locator("{_pw_escape(locator_val)}")'
            return f'page.locator("[name=\\"{_pw_escape(locator_val)}\\"]")'
        if strategy == "xpath":
            return f'page.locator("xpath={_pw_escape(locator_val)}")'
        if strategy == "css":
            return f'page.locator("{_pw_escape(locator_val)}")'
        return f'page.locator("{_pw_escape(locator_val)}")'
    # 11. field_name fallback
    field_name = (step.get("data_field_name") or step.get("field_name") or "").strip()
    if field_name:
        return f'page.locator("[name=\\"{_pw_escape(field_name)}\\"]")'
    return ""


def _pw_escape(value: str) -> str:
    """Escape a string for use inside Python string literals."""
    return value.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n").replace("\r", "")


def _pw_css_escape(value: str) -> str:
    """Escape an ID for use in a CSS selector."""
    return re.sub(r'([^a-zA-Z0-9_-])', r'\\\1', value)


def _pw_looks_like_full_selector(value: str) -> bool:
    """Return True when the stored value is already a CSS/XPath selector."""
    v = (value or "").strip()
    if not v:
        return False
    if v.startswith(("//", "./", "(//", "xpath=")):
        return True
    return bool(re.search(r'\[|>|~|\+|^#|^\.\w', v))


def _pw_normalize_id_value(value: str) -> str:
    """Normalize recorded id values that may already include a leading #."""
    return (value or "").strip().lstrip("#")


def _pw_humanize_identifier(value: str) -> str:
    """Turn ids/names like 'user_name' or 'ctlBranchCode' into readable labels."""
    raw = (value or "").strip()
    if not raw:
        return ""
    tail = raw.split(":")[-1]
    tail = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", tail)
    tail = re.sub(r"[_\-]+", " ", tail)
    tail = re.sub(r"\s+", " ", tail).strip()
    if not tail:
        return ""
    return " ".join(part.capitalize() if not part.isupper() else part for part in tail.split())


def _pw_guess_label_from_identifier(value: str) -> str:
    """Use identifier-based labels only for short, human-looking names."""
    raw = (value or "").strip()
    if not raw or ":" in raw or raw.startswith("_") or len(raw) > 40:
        return ""
    return _pw_humanize_identifier(raw)


def _pw_get_playwright_info(step: dict[str, Any]) -> dict[str, Any]:
    """Return Playwright-specific recorder metadata when available."""
    pw_info = step.get("raw_event_playwright")
    if isinstance(pw_info, dict) and pw_info:
        return pw_info
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    pw_info = raw_event.get("playwright_info")
    return pw_info if isinstance(pw_info, dict) else {}


def _pw_role_name(role: str, name: str, *, exact: bool = False) -> str:
    if exact:
        return f'page.get_by_role("{_pw_escape(role)}", name="{_pw_escape(name)}", exact=True)'
    return f'page.get_by_role("{_pw_escape(role)}", name="{_pw_escape(name)}")'


def _pw_semantic_locator(step: dict[str, Any]) -> str:
    """Build a semantic Playwright locator when enough metadata is available."""
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    pw_info = _pw_get_playwright_info(step)
    attrs = pw_info.get("attributes") if isinstance(pw_info.get("attributes"), dict) else {}

    tag = str(step.get("element_tag") or raw_event.get("tag") or pw_info.get("tagName") or "").strip().lower()
    input_type = str(raw_event.get("inputType") or raw_event.get("type") or attrs.get("type") or "").strip().lower()
    text = str(raw_event.get("text") or "").strip()
    value = str(raw_event.get("value") or pw_info.get("value") or "").strip()
    placeholder = str((raw_event.get("locators") or {}).get("placeholder") or attrs.get("placeholder") or "").strip()
    aria_label = str((raw_event.get("locators") or {}).get("ariaLabel") or attrs.get("aria-label") or "").strip()
    accessible_name = str(pw_info.get("accessibleName") or pw_info.get("labelText") or aria_label).strip()
    raw_name = str(raw_event.get("name") or attrs.get("name") or "").strip()
    raw_id = str(raw_event.get("id") or attrs.get("id") or "").strip()
    fallback_name = accessible_name or _pw_guess_label_from_identifier(raw_name) or _pw_guess_label_from_identifier(raw_id)

    if tag == "button" or input_type in ("submit", "button", "reset"):
        button_name = text or value or fallback_name
        if button_name:
            return _pw_role_name("button", button_name)

    if tag == "a":
        link_name = text or fallback_name or value
        if link_name:
            return _pw_role_name("link", link_name)

    if tag == "select":
        combo_name = fallback_name or placeholder
        if combo_name:
            return _pw_role_name("combobox", combo_name)

    if tag in ("input", "textarea") and input_type not in ("checkbox", "radio", "submit", "button", "reset", "hidden"):
        if fallback_name:
            return _pw_role_name("textbox", fallback_name)
        if placeholder:
            return f'page.get_by_placeholder("{_pw_escape(placeholder)}")'

    if input_type == "checkbox":
        check_name = fallback_name or text
        if check_name:
            return _pw_role_name("checkbox", check_name)

    if input_type == "radio":
        radio_name = fallback_name or text
        if radio_name:
            return _pw_role_name("radio", radio_name)

    if text:
        role_hint = str(attrs.get("role") or "").strip().lower()
        if role_hint in {"cell", "gridcell"} or tag in {"td", "th"}:
            return _pw_role_name("cell", text, exact=True)
        if tag in {"span", "label", "div"}:
            return f'page.get_by_text("{_pw_escape(text)}")'

    return ""


def _pw_key_name(value: str) -> str:
    """Map captured key names to Playwright key identifiers."""
    if not value:
        return ""
    normalized = value.strip()
    aliases = {
        "enter": "Enter", "return": "Enter",
        "tab": "Tab",
        "escape": "Escape", "esc": "Escape",
        "backspace": "Backspace",
        "delete": "Delete", "del": "Delete",
        "arrowup": "ArrowUp", "up": "ArrowUp",
        "arrowdown": "ArrowDown", "down": "ArrowDown",
        "arrowleft": "ArrowLeft", "left": "ArrowLeft",
        "arrowright": "ArrowRight", "right": "ArrowRight",
        "home": "Home", "end": "End",
        "pageup": "PageUp", "pagedown": "PageDown",
        "space": " ", "spacebar": " ",
        "ctrl": "Control", "control": "Control",
        "alt": "Alt", "option": "Alt",
        "shift": "Shift",
        "meta": "Meta", "cmd": "Meta", "command": "Meta", "win": "Meta",
    }
    mapped = aliases.get(normalized.lower())
    if mapped:
        return mapped
    if re.fullmatch(r"f\d{1,2}", normalized, re.IGNORECASE):
        return normalized.capitalize()
    if len(normalized) == 1:
        return normalized
    # chord: Ctrl+A => Control+a
    parts = [p for p in re.split(r"\s*\+\s*", normalized) if p]
    if len(parts) > 1:
        return "+".join(_pw_key_name(p) for p in parts)
    return normalized


def _pw_step_lines(step: dict[str, Any], locator_name: str) -> list[str]:
    """Generate one Playwright step block using inline Playwright expectations."""
    action = (step.get("action") or "").strip().lower()
    page_url = (step.get("page_url") or "").strip()
    locator = _pw_locator(step)
    value = (step.get("data_value") or step.get("field_value") or "").strip()
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    element_tag = (step.get("element_tag") or "").lower()
    input_type = (raw_event.get("inputType") or "").lower() if isinstance(raw_event, dict) else ""
    key = raw_event.get("key", "") if isinstance(raw_event, dict) else ""
    description = _generator_step_description(step)
    timeout_ms = _generator_step_timeout_millis(step)

    if action in ("navigate", "open", "goto") and page_url:
        return [
            f"    # {description}",
            f'    page.goto("{_pw_escape(page_url)}")',
            f"    wait_for_page_ready(page, timeout={timeout_ms})",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
        ]

    if action == "click" and locator:
        if input_type == "checkbox":
            if value.lower() in ("true", "1", "yes", "on", "checked"):
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator}",
                    f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                    f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                    f"    {locator_name}.check()",
                ]
            elif value.lower() in ("false", "0", "no", "off", "unchecked"):
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator}",
                    f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                    f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                    f"    {locator_name}.uncheck()",
                ]
            return [
                f"    # {description}",
                f"    {locator_name} = {locator}",
                f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                f"    expect({locator_name}).to_be_enabled(timeout={timeout_ms})",
                f"    {locator_name}.click()",
            ]
        return [
            f"    # {description}",
            f"    {locator_name} = {locator}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f"    expect({locator_name}).to_be_enabled(timeout={timeout_ms})",
            f"    {locator_name}.click()",
        ]

    if action == "dblclick" and locator:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f"    expect({locator_name}).to_be_enabled(timeout={timeout_ms})",
            f"    {locator_name}.dblclick()",
        ]

    if action == "contextmenu" and locator:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f"    expect({locator_name}).to_be_enabled(timeout={timeout_ms})",
            f"    {locator_name}.click(button=\"right\")",
        ]

    if action in ("input", "change") and locator:
        if element_tag == "select":
            visible_text = (raw_event.get("text") or "").strip() if isinstance(raw_event, dict) else ""
            option_value = (raw_event.get("value") or "").strip() if isinstance(raw_event, dict) else ""
            if visible_text:
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator}",
                    f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                    f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                    f'    {locator_name}.select_option(label="{_pw_escape(visible_text)}")',
                ]
            if value:
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator}",
                    f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                    f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                    f'    {locator_name}.select_option(label="{_pw_escape(value)}")',
                ]
            if option_value:
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator}",
                    f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                    f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                    f'    {locator_name}.select_option(value="{_pw_escape(option_value)}")',
                ]
            return [
                f"    # {description}",
                f"    {locator_name} = {locator}",
                f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                f"    expect({locator_name}).to_be_enabled(timeout={timeout_ms})",
                f"    {locator_name}.click()",
            ]
        if input_type in ("checkbox", "radio"):
            if value.lower() in ("true", "1", "yes", "on", "checked"):
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator}",
                    f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                    f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                    f"    {locator_name}.check()",
                ]
            elif value.lower() in ("false", "0", "no", "off", "unchecked"):
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator}",
                    f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                    f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                    f"    {locator_name}.uncheck()",
                ]
            return [
                f"    # {description}",
                f"    {locator_name} = {locator}",
                f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                f"    expect({locator_name}).to_be_enabled(timeout={timeout_ms})",
                f"    {locator_name}.click()",
            ]
        if value:
            return [
                f"    # {description}",
                f"    {locator_name} = {locator}",
                f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                f'    {locator_name}.fill("{_pw_escape(value)}")',
            ]
        return [
            f"    # {description}",
            f"    {locator_name} = {locator}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f"    expect({locator_name}).to_be_enabled(timeout={timeout_ms})",
            f"    {locator_name}.click()",
        ]

    if action == "keydown" and key:
        pw_key = _pw_key_name(key)
        if locator:
            return [
                f"    # {description}",
                f"    {locator_name} = {locator}",
                f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
                f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
                f'    {locator_name}.press("{_pw_escape(pw_key)}")',
            ]
        return [
            f"    # {description}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f'    page.keyboard.press("{_pw_escape(pw_key)}")',
        ]

    if action == "submit" and locator:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
            f'    {locator_name}.press("Enter")',
        ]

    if action == "scroll":
        delta_x = int(raw_event.get("delta_x") or 0) if isinstance(raw_event, dict) else 0
        delta_y = int(raw_event.get("delta_y") or 0) if isinstance(raw_event, dict) else 0
        return [
            f"    # {description}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f"    page.mouse.wheel({delta_x}, {delta_y})",
        ]

    if action == "navigate_back":
        return [f"    # {description}", "    page.go_back()", f"    wait_for_page_ready(page, timeout={timeout_ms})", f"    wait_for_overlay_gone(page, timeout={timeout_ms})"]

    if action == "navigate_forward":
        return [f"    # {description}", "    page.go_forward()", f"    wait_for_page_ready(page, timeout={timeout_ms})", f"    wait_for_overlay_gone(page, timeout={timeout_ms})"]

    if action == "hover" and locator:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator}",
            f"    wait_for_overlay_gone(page, timeout={timeout_ms})",
            f"    expect({locator_name}).to_be_visible(timeout={timeout_ms})",
            f"    {locator_name}.hover()",
        ]

    return [f"    # Step {step.get('step_no')}: unsupported action '{action}'", "    pass"]


def _pw_step_line(step: dict[str, Any]) -> str:
    """Backward-compatible single-string wrapper for per-step playwright_code updates."""
    locator_name = _selenium_var_stem(step, f"step_{step.get('step_no') or '1'}")
    lines = _pw_step_lines(step, locator_name)
    return "\n".join(line[4:] if line.startswith("    ") else line for line in lines)


def _populate_playwright_code(record_id: str):
    """Generate and store playwright_code for every step in a session."""
    sid = str(record_id)
    steps = list(Step.objects.filter(record_id=sid).order_by("step_no"))
    source_table = "steps"
    if not steps:
        steps = list(Recording.objects.filter(record_id=sid).order_by("step_no"))
        source_table = "recordings"
    if not steps:
        return

    # Build step dicts for _pw_step_line
    locator_ids = [s.locator_id for s in steps if s.locator_id]
    data_ids = [s.data_id for s in steps if s.data_id]
    locators_map = {loc.id: loc for loc in Locator.objects.filter(id__in=locator_ids)}
    data_map = {de.id: de for de in DataEntry.objects.filter(id__in=data_ids)}

    updates = []
    for step in steps:
        loc = locators_map.get(step.locator_id)
        data = data_map.get(step.data_id)
        step_dict = {
            "step_no": step.step_no,
            "action": step.action,
            "page_url": step.page_url,
            "element_tag": step.element_tag,
            "raw_event": step.raw_event if isinstance(step.raw_event, dict) else {},
            "raw_event_playwright": getattr(step, "raw_event_playwright", None),
            "locator_strategy": loc.strategy if loc else None,
            "locator_value": loc.locator if loc else None,
            "data_value": data.value if data else None,
            "field_value": getattr(step, "field_value", None) or (data.value if data else None),
        }
        pw_line = _pw_step_line(step_dict)
        # Strip leading whitespace (the 4-space indent)
        pw_code = pw_line.strip() if pw_line else ""
        updates.append((pw_code, sid, step.step_no))

    if updates:
        with connection.cursor() as cur:
            cur.executemany(
                f"UPDATE {source_table} SET playwright_code = %s WHERE record_id = %s AND step_no = %s",
                updates,
            )


# ── Selenium WebDriver Python Code Generator ─────────────────────────────────

def _selenium_looks_like_full_selector(value: str) -> bool:
    """Return True when the stored value is already a CSS/XPath selector."""
    v = (value or "").strip()
    if not v:
        return False
    if v.startswith(("//", "./", "(//")):
        return True
    return bool(re.search(r'\[|>|~|\+|^#|^\.\w', v))


def _selenium_xpath_literal(value: str) -> str:
    """Escape a string for safe XPath literal usage."""
    if "'" not in value:
        return f"'{value}'"
    if '"' not in value:
        return f'"{value}"'
    parts = value.split("'")
    return "concat(" + ", \"'\", ".join(_selenium_xpath_literal(part) for part in parts) + ")"


def _selenium_get_info(step: dict[str, Any]) -> dict[str, Any]:
    """Return Selenium-specific recorder metadata when available."""
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    info = raw_event.get("selenium_info")
    return info if isinstance(info, dict) else {}


def _selenium_semantic_locator(step: dict[str, Any]) -> str:
    """Build a semantic Selenium locator when enough metadata is available."""
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    info = _selenium_get_info(step)
    attrs = info.get("attributes") if isinstance(info.get("attributes"), dict) else {}

    tag = str(step.get("element_tag") or raw_event.get("tag") or info.get("tagName") or "").strip().lower()
    input_type = str(raw_event.get("inputType") or raw_event.get("type") or attrs.get("type") or "").strip().lower()
    text = str(raw_event.get("text") or "").strip()
    value = str(raw_event.get("value") or info.get("value") or "").strip()
    placeholder = str((raw_event.get("locators") or {}).get("placeholder") or attrs.get("placeholder") or "").strip()
    accessible_name = str(info.get("accessibleName") or info.get("labelText") or attrs.get("aria-label") or "").strip()
    raw_name = str(raw_event.get("name") or attrs.get("name") or "").strip()
    raw_id = str(raw_event.get("id") or attrs.get("id") or "").strip()
    fallback_name = accessible_name or _pw_guess_label_from_identifier(raw_name) or _pw_guess_label_from_identifier(raw_id)

    def _find_xpath(xpath: str) -> str:
        return f'driver.find_element(By.XPATH, "{_pw_escape(xpath)}")'

    if tag == "button" or input_type in ("submit", "button", "reset"):
        button_name = text or value or fallback_name
        if button_name:
            lit = _selenium_xpath_literal(button_name)
            return _find_xpath(
                f"(//button[normalize-space()={lit}] | //input[(@type='submit' or @type='button' or @type='reset') and @value={lit}] | //*[@role='button' and normalize-space()={lit}])[1]"
            )

    if tag == "a":
        link_name = text or fallback_name or value
        if link_name:
            lit = _selenium_xpath_literal(link_name)
            return _find_xpath(f"(//a[normalize-space()={lit}])[1]")

    if tag == "select":
        combo_name = fallback_name or placeholder
        if combo_name:
            lit = _selenium_xpath_literal(combo_name)
            return _find_xpath(
                f"(//*[@aria-label={lit}] | //*[@id=(//label[normalize-space()={lit}]/@for)] | //label[normalize-space()={lit}]//*[self::select])[1]"
            )

    if tag in ("input", "textarea") and input_type not in ("checkbox", "radio", "submit", "button", "reset", "hidden"):
        if fallback_name:
            lit = _selenium_xpath_literal(fallback_name)
            return _find_xpath(
                f"(//*[@aria-label={lit}] | //*[@id=(//label[normalize-space()={lit}]/@for)] | //label[normalize-space()={lit}]//*[self::input or self::textarea])[1]"
            )
        if placeholder:
            return f'driver.find_element(By.CSS_SELECTOR, "[placeholder=\\"{_pw_escape(placeholder)}\\"]")'

    if input_type == "checkbox":
        check_name = fallback_name or text
        if check_name:
            lit = _selenium_xpath_literal(check_name)
            return _find_xpath(
                f"(//*[@aria-label={lit} and self::input[@type='checkbox']] | //*[@id=(//label[normalize-space()={lit}]/@for)] | //label[normalize-space()={lit}]//*[self::input[@type='checkbox']])[1]"
            )

    if input_type == "radio":
        radio_name = fallback_name or text
        if radio_name:
            lit = _selenium_xpath_literal(radio_name)
            return _find_xpath(
                f"(//*[@aria-label={lit} and self::input[@type='radio']] | //*[@id=(//label[normalize-space()={lit}]/@for)] | //label[normalize-space()={lit}]//*[self::input[@type='radio']])[1]"
            )

    if text:
        role_hint = str(attrs.get("role") or "").strip().lower()
        lit = _selenium_xpath_literal(text)
        if role_hint in {"cell", "gridcell"} or tag in {"td", "th"}:
            return _find_xpath(f"(//*[@role='cell' and normalize-space()={lit}] | //td[normalize-space()={lit}] | //th[normalize-space()={lit}])[1]")
        if tag in {"span", "label", "div"}:
            return _find_xpath(f"//*[normalize-space()={lit}][1]")

    return ""

def _selenium_locator(step: dict[str, Any]) -> str:
    """Build a Selenium WebDriver find_element expression from step data.

    Priority:
    1. The user-selected primary locator from the ``locators`` table
       (``locator_strategy`` / ``locator_value``).  This reflects whatever
       the user chose in the Identity column of the steps UI.
    2. Fall back to raw-event locator data (id → name → data-testid →
       aria-label → placeholder → css → xpath).
    3. Field name as last resort.
    """
    # ── 1. Primary locator (user-selected) ────────────────────────────────
    strategy    = (step.get("locator_strategy") or step.get("step_strategy") or "").strip().lower()
    locator_val = (step.get("locator_value")    or step.get("step_locator")  or "").strip()
    if _selenium_semantic_locator(step):
        return _selenium_semantic_locator(step)
    if strategy and locator_val:
        # Strip CSS-selector prefix characters that Selenium By.ID/By.NAME don't need
        _id_val   = locator_val.lstrip("#")
        _name_val = locator_val.lstrip("[").rstrip("]")
        if strategy not in {"xpath", "css", "text", "linkText", "partialLinkText", "class", "classname", "tagname"} and _selenium_looks_like_full_selector(locator_val):
            if locator_val.strip().startswith(("//", "./", "(//")):
                return f'driver.find_element(By.XPATH, "{_pw_escape(locator_val)}")'
            return f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locator_val)}")'
        _by_map = {
            "id":                f'driver.find_element(By.ID, "{_pw_escape(_id_val)}")',
            "name":              f'driver.find_element(By.NAME, "{_pw_escape(locator_val)}")',
            "css":               f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locator_val)}")',
            "css_selector":      f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locator_val)}")',
            "xpath":             f'driver.find_element(By.XPATH, "{_pw_escape(locator_val)}")',
            "link_text":         f'driver.find_element(By.LINK_TEXT, "{_pw_escape(locator_val)}")',
            "partial_link_text": f'driver.find_element(By.PARTIAL_LINK_TEXT, "{_pw_escape(locator_val)}")',
            "class_name":        f'driver.find_element(By.CLASS_NAME, "{_pw_escape(locator_val)}")',
            "tag_name":          f'driver.find_element(By.TAG_NAME, "{_pw_escape(locator_val)}")',
        }
        return _by_map.get(strategy, f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locator_val)}")')

    # ── 2. Fall back to raw-event locator data ─────────────────────────────
    raw_event    = step.get("raw_event")    if isinstance(step.get("raw_event"), dict) else {}
    locators_raw = step.get("locators_raw") if isinstance(step.get("locators_raw"), dict) else {}
    info = _selenium_get_info(step)
    attrs = info.get("attributes") if isinstance(info.get("attributes"), dict) else {}

    locs: dict[str, str] = {}
    for source in (locators_raw, raw_event.get("locators") if isinstance(raw_event.get("locators"), dict) else {}):
        if isinstance(source, dict):
            for k, v in source.items():
                if isinstance(v, str) and v.strip() and k not in locs:
                    locs[k] = v.strip()

    for attr_name, strategy_name in {
        "id": "id",
        "name": "name",
        "placeholder": "placeholder",
        "role": "role",
        "title": "title",
        "alt": "alt",
        "href": "href",
        "aria-label": "ariaLabel",
        "data-testid": "dataTestId",
        "data-test-id": "dataTestId",
        "type": "type",
        "value": "value",
    }.items():
        attr_val = attrs.get(attr_name)
        if isinstance(attr_val, str) and attr_val.strip() and strategy_name not in locs:
            locs[strategy_name] = attr_val.strip()

    if locs.get("id"):
        return f'driver.find_element(By.ID, "{_pw_escape(locs["id"].lstrip("#"))}")'
    if locs.get("name"):
        if _selenium_looks_like_full_selector(locs["name"]):
            return f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locs["name"])}")'
        return f'driver.find_element(By.NAME, "{_pw_escape(locs["name"])}")'
    if locs.get("dataTestId"):
        if _selenium_looks_like_full_selector(locs["dataTestId"]):
            return f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locs["dataTestId"])}")'
        return f'driver.find_element(By.CSS_SELECTOR, "[data-testid=\\"{_pw_escape(locs["dataTestId"])}\\"]")'
    if locs.get("ariaLabel"):
        if _selenium_looks_like_full_selector(locs["ariaLabel"]):
            return f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locs["ariaLabel"])}")'
        return f'driver.find_element(By.CSS_SELECTOR, "[aria-label=\\"{_pw_escape(locs["ariaLabel"])}\\"]")'
    if locs.get("placeholder"):
        if _selenium_looks_like_full_selector(locs["placeholder"]):
            return f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locs["placeholder"])}")'
        return f'driver.find_element(By.CSS_SELECTOR, "[placeholder=\\"{_pw_escape(locs["placeholder"])}\\"]")'
    if locs.get("css"):
        return f'driver.find_element(By.CSS_SELECTOR, "{_pw_escape(locs["css"])}")'
    if locs.get("xpath"):
        return f'driver.find_element(By.XPATH, "{_pw_escape(locs["xpath"])}")'

    # ── 3. Field name last resort ──────────────────────────────────────────
    field_name = (step.get("data_field_name") or step.get("field_name") or "").strip()
    if field_name:
        return f'driver.find_element(By.NAME, "{_pw_escape(field_name)}")'
    return ""


def _selenium_key_name(value: str) -> str:
    """Map a captured key name to a selenium Keys.XXX expression or quoted string."""
    if not value:
        return '""'
    normalized = value.strip().lower()
    _aliases = {
        "enter": "ENTER", "return": "ENTER",
        "tab": "TAB",
        "escape": "ESCAPE", "esc": "ESCAPE",
        "backspace": "BACK_SPACE",
        "delete": "DELETE", "del": "DELETE",
        "arrowup": "ARROW_UP",    "up":    "ARROW_UP",
        "arrowdown": "ARROW_DOWN","down":  "ARROW_DOWN",
        "arrowleft": "ARROW_LEFT","left":  "ARROW_LEFT",
        "arrowright": "ARROW_RIGHT","right":"ARROW_RIGHT",
        "home": "HOME", "end": "END",
        "pageup": "PAGE_UP", "pagedown": "PAGE_DOWN",
        "space": "SPACE", "spacebar": "SPACE",
        "ctrl": "CONTROL", "control": "CONTROL",
        "alt": "ALT",  "option": "ALT",
        "shift": "SHIFT",
        "meta": "META", "cmd": "META", "command": "META", "win": "META",
        **{f"f{n}": f"F{n}" for n in range(1, 13)},
    }
    mapped = _aliases.get(normalized)
    if mapped:
        return f"Keys.{mapped}"
    if len(value.strip()) == 1:
        return f'"{_pw_escape(value.strip())}"'
    return f'"{_pw_escape(value.strip())}"'


def _selenium_locator_tuple(step: dict[str, Any]) -> str:
    """Convert the chosen Selenium locator into a reusable ``(By.*, value)`` tuple."""
    expr = (_selenium_locator(step) or "").strip()
    match = re.match(r'^driver\.find_element\((By\.[A-Z_]+),\s+"((?:[^"\\]|\\.)*)"\)$', expr)
    if not match:
        return ""
    return f'({match.group(1)}, "{match.group(2)}")'


def _generator_var_slug(value: Any) -> str:
    text = str(value or "").strip().lower()
    if not text:
        return ""
    text = re.sub(r"\bsign\s+in\b", "signin", text)
    text = re.sub(r"\blog\s+in\b", "login", text)
    text = re.sub(r"\blog\s+out\b", "logout", text)
    text = re.sub(r"\bsign\s+out\b", "signout", text)
    text = re.sub(r"[^0-9a-z]+", "_", text).strip("_")
    if text and text[0].isdigit():
        text = f"step_{text}"
    return text


def _generator_var_suffix(tag: str, input_type: str, role: str) -> str:
    role = (role or "").strip().lower()
    tag = (tag or "").strip().lower()
    input_type = (input_type or "").strip().lower()
    if role == "button" or tag == "button" or input_type in {"submit", "button", "reset"}:
        return "btn"
    if role == "link" or tag == "a":
        return "link"
    if role in {"checkbox"} or input_type == "checkbox":
        return "checkbox"
    if role in {"radio"} or input_type == "radio":
        return "radio"
    if role in {"combobox", "listbox"} or tag == "select":
        return "select"
    return ""


def _generator_semantic_var_name(step: dict[str, Any]) -> str:
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    sel_info = _selenium_get_info(step)
    pw_info = _pw_get_playwright_info(step)
    attrs: dict[str, Any] = {}
    for source in (sel_info, pw_info):
        if isinstance(source.get("attributes"), dict):
            attrs.update(source["attributes"])

    tag = str(step.get("element_tag") or raw_event.get("tag") or sel_info.get("tagName") or pw_info.get("tagName") or "").strip().lower()
    input_type = str(raw_event.get("inputType") or raw_event.get("type") or attrs.get("type") or "").strip().lower()
    role = str(attrs.get("role") or raw_event.get("role") or "").strip().lower()
    text = str(raw_event.get("text") or "").strip()
    value = str(raw_event.get("value") or sel_info.get("value") or pw_info.get("value") or "").strip()
    accessible_name = str(
        sel_info.get("accessibleName")
        or sel_info.get("labelText")
        or pw_info.get("accessibleName")
        or pw_info.get("labelText")
        or attrs.get("aria-label")
        or ""
    ).strip()
    placeholder = str((raw_event.get("locators") or {}).get("placeholder") or attrs.get("placeholder") or "").strip()
    raw_name = str(raw_event.get("name") or attrs.get("name") or "").strip()
    raw_id = str(raw_event.get("id") or attrs.get("id") or "").strip()
    guessed_name = _pw_guess_label_from_identifier(raw_name) or _pw_guess_label_from_identifier(raw_id)

    if role == "button" or tag == "button" or input_type in {"submit", "button", "reset"}:
        label = text or value or accessible_name or guessed_name
    elif role == "link" or tag == "a":
        label = text or accessible_name or guessed_name or value
    elif role in {"combobox", "listbox"} or tag == "select":
        label = accessible_name or guessed_name or placeholder or text
    elif role in {"checkbox", "radio"} or input_type in {"checkbox", "radio"}:
        label = accessible_name or text or guessed_name
    else:
        label = accessible_name or guessed_name or placeholder or text

    stem = _generator_var_slug(label)
    suffix = _generator_var_suffix(tag, input_type, role)
    if stem and suffix and not stem.endswith(f"_{suffix}"):
        return f"{stem}_{suffix}"
    return stem


def _generator_step_timeout_seconds(step: dict[str, Any], default: float = 15.0) -> float:
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    try:
        timeout_s = float(raw_event.get("_recorded_step_delay_s") or default)
    except (TypeError, ValueError):
        timeout_s = default
    return max(0.0, timeout_s)


def _generator_step_timeout_literal(step: dict[str, Any], default: float = 15.0) -> str:
    timeout_s = _generator_step_timeout_seconds(step, default=default)
    if float(timeout_s).is_integer():
        return str(int(timeout_s))
    return f"{timeout_s:.3f}".rstrip("0").rstrip(".")


def _generator_step_timeout_millis(step: dict[str, Any], default: float = 15.0) -> int:
    timeout_s = _generator_step_timeout_seconds(step, default=default)
    return max(0, int(round(timeout_s * 1000)))


def _generator_step_description(step: dict[str, Any]) -> str:
    action = (step.get("action") or "").strip().lower()
    return str(step.get("steps_description") or f"Step {step.get('step_no')}: {action}").strip()


def _selenium_var_stem(step: dict[str, Any], fallback: str) -> str:
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    semantic_name = _generator_semantic_var_name(step)
    if semantic_name:
        return semantic_name
    candidates = [
        step.get("data_field_name"),
        step.get("field_name"),
        raw_event.get("accessibleName") if isinstance(raw_event, dict) else "",
        raw_event.get("text") if isinstance(raw_event, dict) else "",
        step.get("steps_description"),
        step.get("action"),
    ]
    for candidate in candidates:
        text = _generator_var_slug(candidate)
        if text:
            return text
    return fallback


def _selenium_step_lines(step: dict[str, Any], locator_name: str, element_name: str) -> list[str]:
    """Generate one Selenium step block using explicit waits and locator tuples."""
    action      = (step.get("action")   or "").strip().lower()
    page_url    = (step.get("page_url") or "").strip()
    locator_tpl = _selenium_locator_tuple(step)
    value       = (step.get("data_value") or step.get("field_value") or "").strip()
    raw_event   = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    element_tag = (step.get("element_tag") or "").lower()
    input_type  = (raw_event.get("inputType") or "").lower() if isinstance(raw_event, dict) else ""
    key         = (raw_event.get("key") or "") if isinstance(raw_event, dict) else ""
    description = _generator_step_description(step)
    timeout_s   = _generator_step_timeout_literal(step)

    if action in ("navigate", "open", "goto") and page_url:
        return [
            f"    # {description}",
            f'    driver.get("{_pw_escape(page_url)}")',
            f"    WebDriverWait(driver, {timeout_s}).until(",
            '        lambda d: d.execute_script("return document.readyState") == "complete"',
            "    )",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
        ]

    if action == "click" and locator_tpl:
        if input_type == "checkbox":
            want = value.lower() in ("true", "1", "yes", "on", "checked")
            return [
                f"    # {description}",
                f"    {locator_name} = {locator_tpl}",
                f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
                f"    if {element_name}.is_selected() != {str(want).title()}:",
                f"        {element_name}.click()",
            ]
        return [
            f"    # {description}",
            f"    {locator_name} = {locator_tpl}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    {element_name} = wait_for_clickable(driver, {locator_name}, timeout={timeout_s})",
            f"    {element_name}.click()",
        ]

    if action == "dblclick" and locator_tpl:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator_tpl}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    {element_name} = wait_for_clickable(driver, {locator_name}, timeout={timeout_s})",
            f"    ActionChains(driver).double_click({element_name}).perform()",
        ]

    if action == "contextmenu" and locator_tpl:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator_tpl}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    {element_name} = wait_for_clickable(driver, {locator_name}, timeout={timeout_s})",
            f"    ActionChains(driver).context_click({element_name}).perform()",
        ]

    if action in ("input", "change") and locator_tpl:
        if element_tag == "select":
            visible_text = (raw_event.get("text")  or "").strip() if isinstance(raw_event, dict) else ""
            option_value = (raw_event.get("value") or "").strip() if isinstance(raw_event, dict) else ""
            if visible_text:
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator_tpl}",
                    f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                    f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
                    f'    Select({element_name}).select_by_visible_text("{_pw_escape(visible_text)}")',
                ]
            if value:
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator_tpl}",
                    f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                    f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
                    f'    Select({element_name}).select_by_visible_text("{_pw_escape(value)}")',
                ]
            if option_value:
                return [
                    f"    # {description}",
                    f"    {locator_name} = {locator_tpl}",
                    f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                    f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
                    f'    Select({element_name}).select_by_value("{_pw_escape(option_value)}")',
                ]
            return [
                f"    # {description}",
                f"    {locator_name} = {locator_tpl}",
                f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                f"    {element_name} = wait_for_clickable(driver, {locator_name}, timeout={timeout_s})",
                f"    {element_name}.click()",
            ]
        if input_type in ("checkbox", "radio"):
            want = value.lower() in ("true", "1", "yes", "on", "checked")
            return [
                f"    # {description}",
                f"    {locator_name} = {locator_tpl}",
                f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
                f"    if {element_name}.is_selected() != {str(want).title()}:",
                f"        {element_name}.click()",
            ]
        if value:
            return [
                f"    # {description}",
                f"    {locator_name} = {locator_tpl}",
                f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
                f"    {element_name}.clear()",
                f'    {element_name}.send_keys("{_pw_escape(value)}")',
            ]
        return [
            f"    # {description}",
            f"    {locator_name} = {locator_tpl}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    {element_name} = wait_for_clickable(driver, {locator_name}, timeout={timeout_s})",
            f"    {element_name}.click()",
        ]

    if action == "keydown" and key:
        key_expr = _selenium_key_name(key)
        if locator_tpl:
            return [
                f"    # {description}",
                f"    {locator_name} = {locator_tpl}",
                f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
                f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
                f"    {element_name}.send_keys({key_expr})",
            ]
        return [
            f"    # {description}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    ActionChains(driver).send_keys({key_expr}).perform()",
        ]

    if action == "submit" and locator_tpl:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator_tpl}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
            f"    {element_name}.submit()",
        ]

    if action == "scroll":
        delta_x = int(raw_event.get("delta_x") or 0) if isinstance(raw_event, dict) else 0
        delta_y = int(raw_event.get("delta_y") or 0) if isinstance(raw_event, dict) else 0
        return [
            f"    # {description}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    driver.execute_script('window.scrollBy({delta_x}, {delta_y})')",
        ]

    if action == "navigate_back":
        return [f"    # {description}", "    driver.back()", f"    wait_for_overlay_gone(driver, timeout={timeout_s})"]

    if action == "navigate_forward":
        return [f"    # {description}", "    driver.forward()", f"    wait_for_overlay_gone(driver, timeout={timeout_s})"]

    if action == "hover" and locator_tpl:
        return [
            f"    # {description}",
            f"    {locator_name} = {locator_tpl}",
            f"    wait_for_overlay_gone(driver, timeout={timeout_s})",
            f"    {element_name} = wait_for_visible(driver, {locator_name}, timeout={timeout_s})",
            f"    ActionChains(driver).move_to_element({element_name}).perform()",
        ]

    return [f"    # Step {step.get('step_no')}: unsupported action '{action}'", "    pass"]


def _build_selenium_script(record_bundle: dict[str, Any]) -> str:
    """Generate a complete Selenium WebDriver Python script from a record bundle."""
    meta        = record_bundle["meta"]
    steps       = record_bundle["steps"]
    record_name = meta.get("record_name", "Untitled")
    first_url   = next(
        (s.get("page_url") for s in steps if (s.get("page_url") or "").strip()),
        "about:blank",
    )
    first_nav_step = next(
        (s for s in steps if (s.get("page_url") or "").strip()),
        {},
    )
    first_nav_timeout_s = _generator_step_timeout_literal(first_nav_step)

    lines = [
        '"""',
        "Enhanced Selenium WebDriver Python script with explicit waits",
        '"""',
        "import time",
        "import traceback",
        "import re",
        "import shutil",
        "from datetime import datetime",
        "from html import escape",
        "from pathlib import Path",
        "from selenium import webdriver",
        "from selenium.common.exceptions import WebDriverException",
        "from selenium.webdriver.common.by import By",
        "from selenium.webdriver.common.keys import Keys",
        "from selenium.webdriver.common.action_chains import ActionChains",
        "from selenium.webdriver.chrome.service import Service",
        "from selenium.webdriver.support.ui import Select, WebDriverWait",
        "from selenium.webdriver.support import expected_conditions as EC",
        "",
        "",
        "try:",
        "    from webdriver_manager.chrome import ChromeDriverManager",
        "except Exception:",
        "    ChromeDriverManager = None",
        "",
        "STEP_RESULTS = []",
        "CURRENT_STEP = None",
        "CURRENT_STEP_STARTED = None",
        "ACTIVE_DRIVER = None",
        "ACTIVE_ARTIFACT_DIR = None",
        "FAILURE_SCREENSHOT = ''",
        "",
        "",
        "def build_output_paths():",
        "    script_path = Path(__file__).resolve() if '__file__' in globals() else Path.cwd() / 'generated_selenium.py'",
        "    script_dir = script_path.parent",
        "    script_stem = script_path.stem",
        "    artifact_dir = script_dir / f'{script_stem}_artifacts'",
        "    artifact_dir.mkdir(parents=True, exist_ok=True)",
        "    report_path = script_dir / f'{script_stem}_report.html'",
        "    return report_path, artifact_dir",
        "",
        "",
        "def begin_step(step_no, description):",
        "    global CURRENT_STEP, CURRENT_STEP_STARTED",
        "    CURRENT_STEP = {'step_no': step_no, 'description': description}",
        "    CURRENT_STEP_STARTED = time.perf_counter()",
        "",
        "",
        "def save_step_screenshot(driver, artifact_dir, step_no):",
        "    if driver is None or artifact_dir is None or step_no in (None, ''):",
        "        return ''",
        "    screenshot_path = artifact_dir / f'step_{int(step_no):04d}_{datetime.now().strftime(\"%Y%m%d_%H%M%S\")}.png'",
        "    driver.save_screenshot(str(screenshot_path))",
        "    return screenshot_path.relative_to(artifact_dir.parent).as_posix()",
        "",
        "",
        "def finish_step(status, error_message=''):",
        "    global CURRENT_STEP, CURRENT_STEP_STARTED",
        "    if CURRENT_STEP is None:",
        "        return",
        "    duration_ms = int((time.perf_counter() - CURRENT_STEP_STARTED) * 1000) if CURRENT_STEP_STARTED is not None else 0",
        "    screenshot_rel_path = ''",
        "    try:",
        "        screenshot_rel_path = save_step_screenshot(ACTIVE_DRIVER, ACTIVE_ARTIFACT_DIR, CURRENT_STEP.get('step_no'))",
        "    except Exception:",
        "        screenshot_rel_path = ''",
        "    STEP_RESULTS.append({",
        "        'step_no': CURRENT_STEP.get('step_no', ''),",
        "        'description': CURRENT_STEP.get('description', ''),",
        "        'status': status,",
        "        'duration_ms': duration_ms,",
        "        'error': error_message,",
        "        'screenshot': screenshot_rel_path,",
        "    })",
        "    CURRENT_STEP = None",
        "    CURRENT_STEP_STARTED = None",
        "",
        "",
        "def save_failure_screenshot(driver, artifact_dir):",
        "    screenshot_path = artifact_dir / f'failure_{datetime.now().strftime(\"%Y%m%d_%H%M%S\")}.png'",
        "    driver.save_screenshot(str(screenshot_path))",
        "    return screenshot_path.relative_to(artifact_dir.parent).as_posix()",
        "",
        "",
        "def write_execution_report(report_path, report_title, started_at, finished_at, overall_status, step_results, error_message='', screenshot_rel_path=''):",
        "    duration_ms = int((finished_at - started_at).total_seconds() * 1000)",
        "    pass_count = sum(1 for i in step_results if i.get('status') == 'PASS')",
        "    fail_count = sum(1 for i in step_results if i.get('status') != 'PASS')",
        "    html_lines = [",
        "        '<!doctype html>',",
        "        '<html lang=\"en\"><head><meta charset=\"utf-8\"><meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">',",
        "        f'<title>{escape(report_title)}</title>',",
        "        '<style>',",
        "        ':root{--bg:#f8fafc;--card:#ffffff;--border:#e2e8f0;--text:#1e293b;--muted:#64748b;--accent:#6366f1;--pass-bg:#ecfdf5;--pass:#059669;--fail-bg:#fef2f2;--fail:#dc2626;--shadow:0 4px 6px -1px rgb(0 0 0/0.1),0 2px 4px -2px rgb(0 0 0/0.1);}',",
        "        '*{box-sizing:border-box;margin:0;padding:0;}',",
        "        'body{font-family:Inter,Segoe UI,system-ui,sans-serif;background:var(--bg);color:var(--text);padding:32px;line-height:1.6;}',",
        "        '.container{max-width:1400px;margin:0 auto;}',",
        "        '.header{background:linear-gradient(135deg,#6366f1 0%,#8b5cf6 100%);border-radius:16px;padding:32px 40px;color:#fff;margin-bottom:24px;box-shadow:var(--shadow);}',",
        "        '.header h1{font-size:1.75rem;font-weight:700;margin-bottom:8px;}',",
        "        '.header p{opacity:0.9;font-size:0.9rem;}',",
        "        '.summary{display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:16px;margin-bottom:24px;}',",
        "        '.card{background:var(--card);border:1px solid var(--border);border-radius:12px;padding:20px;box-shadow:var(--shadow);}',",
        "        '.card .label{font-size:0.8rem;text-transform:uppercase;letter-spacing:0.05em;color:var(--muted);font-weight:600;}',",
        "        '.card .value{font-size:1.5rem;font-weight:700;margin-top:4px;}',",
        "        '.card.pass-card .value{color:var(--pass);}',",
        "        '.card.fail-card .value{color:var(--fail);}',",
        "        '.badge{display:inline-block;padding:4px 12px;border-radius:20px;font-size:0.8rem;font-weight:600;}',",
        "        '.badge-pass{background:var(--pass-bg);color:var(--pass);}',",
        "        '.badge-fail{background:var(--fail-bg);color:var(--fail);}',",
        "        'table{width:100%;border-collapse:separate;border-spacing:0;background:var(--card);border-radius:12px;overflow:hidden;box-shadow:var(--shadow);margin-bottom:24px;}',",
        "        'thead{background:linear-gradient(135deg,#f1f5f9,#e2e8f0);}',",
        "        'th{padding:14px 16px;font-size:0.8rem;text-transform:uppercase;letter-spacing:0.05em;color:var(--muted);font-weight:700;text-align:left;}',",
        "        'td{padding:12px 16px;border-top:1px solid var(--border);vertical-align:top;font-size:0.9rem;}',",
        "        'tr:hover td{background:#f8fafc;}',",
        "        '.step-no{font-weight:700;color:var(--accent);min-width:50px;}',",
        "        '.screenshot-thumb{cursor:pointer;border-radius:8px;border:2px solid var(--border);max-width:200px;max-height:120px;object-fit:cover;transition:transform 0.2s,box-shadow 0.2s;}',",
        "        '.screenshot-thumb:hover{transform:scale(1.05);box-shadow:0 8px 25px rgb(0 0 0/0.15);}',",
        "        '.error-box{background:#1e1b4b;color:#e0e7ff;padding:20px;border-radius:12px;font-family:JetBrains Mono,Consolas,monospace;font-size:0.85rem;overflow:auto;white-space:pre-wrap;margin-top:16px;box-shadow:var(--shadow);}',",
        "        '.lightbox{display:none;position:fixed;inset:0;background:rgba(0,0,0,0.85);z-index:9999;justify-content:center;align-items:center;padding:24px;cursor:zoom-out;}',",
        "        '.lightbox.active{display:flex;}',",
        "        '.lightbox img{max-width:95vw;max-height:92vh;border-radius:12px;box-shadow:0 20px 60px rgb(0 0 0/0.5);}',",
        "        '.lightbox-close{position:absolute;top:20px;right:28px;color:#fff;font-size:2rem;cursor:pointer;opacity:0.7;transition:opacity 0.2s;}',",
        "        '.lightbox-close:hover{opacity:1;}',",
        "        '</style>',",
        "        '</head><body>',",
        "        '<div class=\"container\">',",
        "        f'<div class=\"header\"><h1>{escape(report_title)}</h1>',",
        "        f'<p>{escape(started_at.isoformat(sep=\" \", timespec=\"seconds\"))} &mdash; {escape(finished_at.isoformat(sep=\" \", timespec=\"seconds\"))}</p></div>',",
        "        '<div class=\"summary\">',",
        "        f'<div class=\"card\"><div class=\"label\">Duration</div><div class=\"value\">{duration_ms} ms</div></div>',",
        "        f'<div class=\"card\"><div class=\"label\">Total Steps</div><div class=\"value\">{len(step_results)}</div></div>',",
        "        f'<div class=\"card pass-card\"><div class=\"label\">Passed</div><div class=\"value\">{pass_count}</div></div>',",
        "        f'<div class=\"card fail-card\"><div class=\"label\">Failed</div><div class=\"value\">{fail_count}</div></div>',",
        "        f'<div class=\"card\"><div class=\"label\">Status</div><div class=\"value\"><span class=\"badge badge-{\"pass\" if overall_status == \"PASS\" else \"fail\"}\">{escape(overall_status)}</span></div></div>',",
        "        '</div>',",
        "        '<table><thead><tr><th>Step</th><th>Description</th><th>Status</th><th>Duration</th><th>Error</th><th>Screenshot</th></tr></thead><tbody>',",
        "    ]",
        "    for item in step_results:",
        "        badge_class = 'badge-pass' if item.get('status') == 'PASS' else 'badge-fail'",
        "        screenshot_cell = ''",
        "        if item.get('screenshot'):",
        "            screenshot_href = escape(str(item.get('screenshot')))",
        "            screenshot_cell = f'<img src=\"{screenshot_href}\" alt=\"step {item.get(\"step_no\")}\" class=\"screenshot-thumb\" onclick=\"openLightbox(this.src)\">'",
        "        error_text = escape(str(item.get('error', ''))) if item.get('error') else ''",
        "        html_lines.append(",
        "            '<tr>'",
        "            f'<td class=\"step-no\">{escape(str(item.get(\"step_no\", \"\")))}</td>'",
        "            f'<td>{escape(str(item.get(\"description\", \"\")))}</td>'",
        "            f'<td><span class=\"badge {badge_class}\">{escape(str(item.get(\"status\", \"\")))}</span></td>'",
        "            f'<td>{escape(str(item.get(\"duration_ms\", \"\")))} ms</td>'",
        "            f'<td>{error_text}</td>'",
        "            f'<td>{screenshot_cell}</td>'",
        "            '</tr>'",
        "        )",
        "    html_lines.append('</tbody></table>')",
        "    if screenshot_rel_path:",
        "        html_lines.append(f'<div class=\"card\" style=\"margin-bottom:16px\"><div class=\"label\">Failure Screenshot</div><img src=\"{escape(screenshot_rel_path)}\" class=\"screenshot-thumb\" style=\"max-width:600px;max-height:400px;margin-top:8px\" onclick=\"openLightbox(this.src)\"></div>')",
        "    if error_message:",
        "        html_lines.append(f'<div class=\"card\"><div class=\"label\">Error Details</div><div class=\"error-box\">{escape(error_message)}</div></div>')",
        "    html_lines.append('</div>')",
        "    html_lines.append('<div class=\"lightbox\" id=\"lightbox\" onclick=\"closeLightbox()\"><span class=\"lightbox-close\">&times;</span><img id=\"lightbox-img\" src=\"\" alt=\"screenshot\"></div>')",
        "    html_lines.append('<script>')",
        "    html_lines.append('function openLightbox(src){var lb=document.getElementById(\"lightbox\");document.getElementById(\"lightbox-img\").src=src;lb.classList.add(\"active\");}')",
        "    html_lines.append('function closeLightbox(){document.getElementById(\"lightbox\").classList.remove(\"active\");}')",
        "    html_lines.append('document.addEventListener(\"keydown\",function(e){if(e.key===\"Escape\")closeLightbox();});')",
        "    html_lines.append('</script>')",
        "    html_lines.append('</body></html>')",
        "    report_path.write_text('\\n'.join(html_lines), encoding='utf-8')",
        "",
        "",
        "def resolve_local_chromedriver():",
        '    """Find a nearby ChromeDriver before falling back to auto-detection."""',
        "    script_dir = Path(__file__).resolve().parent if '__file__' in globals() else Path.cwd()",
        "    home_dir = Path.home()",
        "",
        "    def add_search_root(search_roots, seen_roots, candidate_root):",
        "        try:",
        "            normalized = str(candidate_root.resolve())",
        "        except Exception:",
        "            normalized = str(candidate_root)",
        "        if normalized not in seen_roots:",
        "            seen_roots.add(normalized)",
        "            search_roots.append(candidate_root)",
        "",
        "    search_roots = [",
        "    ]",
        "    seen_roots = set()",
        "    add_search_root(search_roots, seen_roots, script_dir)",
        "    add_search_root(search_roots, seen_roots, Path.cwd())",
        "    add_search_root(search_roots, seen_roots, script_dir / 'webdrivers' / 'chrome')",
        "    add_search_root(search_roots, seen_roots, Path.cwd() / 'webdrivers' / 'chrome')",
        "    add_search_root(search_roots, seen_roots, home_dir / 'webdrivers' / 'chrome')",
        "    add_search_root(search_roots, seen_roots, home_dir / 'Downloads' / 'webdrivers' / 'chrome')",
        "    add_search_root(search_roots, seen_roots, home_dir / 'Downloads')",
        "    add_search_root(search_roots, seen_roots, home_dir / 'Desktop' / 'webdrivers' / 'chrome')",
        "    add_search_root(search_roots, seen_roots, home_dir / 'Desktop')",
        "    for parent in [script_dir, *script_dir.parents, Path.cwd(), *Path.cwd().parents]:",
        "        candidate_root = parent / 'webdrivers' / 'chrome'",
        "        add_search_root(search_roots, seen_roots, candidate_root)",
        "    drive_roots = {Path(script_dir.anchor), Path.cwd().resolve().anchor and Path(Path.cwd().resolve().anchor), Path(home_dir.anchor)}",
        "    for drive_root in [root for root in drive_roots if root]:",
        "        add_search_root(search_roots, seen_roots, drive_root / 'webdrivers' / 'chrome')",
        "        try:",
        "            for child in drive_root.iterdir():",
        "                if child.is_dir():",
        "                    add_search_root(search_roots, seen_roots, child / 'webdrivers' / 'chrome')",
        "        except Exception:",
        "            pass",
        "",
        "    def driver_sort_key(candidate_path):",
        "        match = re.search(r'(\\d+)', candidate_path.name)",
        "        major = int(match.group(1)) if match else -1",
        "        return (major, candidate_path.name.lower())",
        "    patterns = ('chromedriver.exe', 'chromedriver*.exe', 'chromedriver', 'chromedriver*')",
        "    seen = set()",
        "    candidates = []",
        "    for root in search_roots:",
        "        if not root.exists():",
        "            continue",
        "        for pattern in patterns:",
        "            for candidate in root.glob(pattern):",
        "                candidate_str = str(candidate.resolve())",
        "                if candidate.is_file() and candidate_str not in seen:",
        "                    seen.add(candidate_str)",
        "                    candidates.append(candidate.resolve())",
        "    if candidates:",
        "        return str(sorted(candidates, key=driver_sort_key, reverse=True)[0])",
        "    return shutil.which('chromedriver') or ''",
        "",
        "",
        "def compact_error(exc):",
        '    """Collapse exception text into one line for diagnostics."""',
        "    return ' '.join(str(exc).split())",
        "",
        "",
        "def create_chrome_driver(options):",
        '    """Start ChromeDriver: nearby file -> Selenium Manager -> webdriver_manager -> PATH."""',
        "    local_driver = resolve_local_chromedriver()",
        "    attempted = []",
        "    if local_driver:",
        "        attempted.append(f'local:{local_driver}')",
        "        try:",
        "            return webdriver.Chrome(service=Service(local_driver), options=options)",
        "        except WebDriverException as exc:",
        "            attempted.append(f'local_error:{compact_error(exc)}')",
        "",
        "    try:",
        "        attempted.append('selenium-manager')",
        "        return webdriver.Chrome(options=options)",
        "    except WebDriverException as exc:",
        "        attempted.append(f'selenium_manager_error:{compact_error(exc)}')",
        "",
        "    if ChromeDriverManager is not None:",
        "        try:",
        "            attempted.append('webdriver-manager')",
        "            manager_path = ChromeDriverManager().install()",
        "            attempted.append(f'webdriver_manager_path:{manager_path}')",
        "            return webdriver.Chrome(service=Service(manager_path), options=options)",
        "        except Exception as exc:",
        "            attempted.append(f'webdriver_manager_error:{compact_error(exc)}')",
        "    else:",
        "        attempted.append('webdriver-manager-unavailable')",
        "",
        "    try:",
        "        attempted.append('path:chromedriver')",
        "        return webdriver.Chrome(service=Service('chromedriver'), options=options)",
        "    except WebDriverException as exc:",
        "        details = ' | '.join(attempted[-8:]) if attempted else 'no attempts recorded'",
        "        raise WebDriverException(",
        "            'Could not start ChromeDriver via nearby files, Selenium Manager, webdriver_manager, or PATH. '",
        "            'Place a matching chromedriver next to this script, in webdrivers/chrome, or install webdriver-manager. '",
        "            f'Attempts: {details}'",
        "        ) from exc",
        "",
        "",
        "def wait_for_visible(driver, locator, timeout=15):",
        '    """Wait until element is visible"""',
        "    return WebDriverWait(driver, timeout).until(",
        "        EC.visibility_of_element_located(locator)",
        "    )",
        "",
        "",
        "def wait_for_clickable(driver, locator, timeout=15):",
        '    """Wait until element is clickable"""',
        "    return WebDriverWait(driver, timeout).until(",
        "        EC.element_to_be_clickable(locator)",
        "    )",
        "",
        "",
        "def wait_for_overlay_gone(driver, timeout=15):",
        '    """Wait until the in-progress overlay is hidden or detached"""',
        "    overlay_locator = (By.XPATH, '//*[@id=\"inProgressPage\"]')",
        "    if not driver.find_elements(*overlay_locator):",
        "        return True",
        "    return WebDriverWait(driver, timeout).until(",
        "        EC.invisibility_of_element_located(overlay_locator)",
        "    )",
        "",
        "",
        "def run():",
        "    global ACTIVE_DRIVER",
        "    begin_step(0, 'Browser startup')",
        "    options = webdriver.ChromeOptions()",
        "    driver = create_chrome_driver(options)",
        "    ACTIVE_DRIVER = driver",
        "    driver.maximize_window()",
        "    finish_step('PASS')",
        "",
        "    # Keep small implicit wait (optional safety net)",
        "    driver.implicitly_wait(5)",
        "",
        f'    driver.get("{_pw_escape(first_url)}")',
        "",
        "    # Wait for page readiness (optional but useful)",
        f"    WebDriverWait(driver, {first_nav_timeout_s}).until(",
        '        lambda d: d.execute_script("return document.readyState") == "complete"',
        "    )",
        f"    wait_for_overlay_gone(driver, timeout={first_nav_timeout_s})",
        "",
    ]

    first_navigate_skipped = False
    used_names: dict[str, int] = {}
    for step in steps:
        action   = (step.get("action")   or "").strip().lower()
        page_url = (step.get("page_url") or "").strip()
        if action in ("navigate", "open", "goto") and page_url and not first_navigate_skipped:
            if page_url == first_url:
                first_navigate_skipped = True
                continue
        base_name = _selenium_var_stem(step, f"step_{step.get('step_no') or len(used_names) + 1}")
        used_names[base_name] = used_names.get(base_name, 0) + 1
        suffix = "" if used_names[base_name] == 1 else f"_{used_names[base_name]}"
        locator_name = f"{base_name}_locator{suffix}"
        element_name = f"{base_name}_element{suffix}"
        step_lines = _selenium_step_lines(step, locator_name, element_name)
        if step_lines:
            description = _generator_step_description(step)
            step_no = int(step.get("step_no") or 0)
            lines.append(f"    begin_step({step_no}, \"{_pw_escape(description)}\")")
            lines.extend(step_lines)
            lines.append("    finish_step('PASS')")
            lines.append("")

    lines.extend([
        "",
        "",
        "if __name__ == '__main__':",
        "    report_path, artifact_dir = build_output_paths()",
        "    ACTIVE_ARTIFACT_DIR = artifact_dir",
        f"    report_title = \"Selenium Execution Report - {_pw_escape(record_name)}\"",
        "    started_at = datetime.now()",
        "    overall_status = 'PASS'",
        "    error_message = ''",
        "    try:",
        "        run()",
        "    except Exception:",
        "        overall_status = 'FAIL'",
        "        error_message = traceback.format_exc()",
        "        finish_step('FAIL', error_message.strip())",
        "        if ACTIVE_DRIVER is not None:",
        "            try:",
        "                FAILURE_SCREENSHOT = save_failure_screenshot(ACTIVE_DRIVER, artifact_dir)",
        "            except Exception:",
        "                FAILURE_SCREENSHOT = ''",
        "        raise",
        "    finally:",
        "        finished_at = datetime.now()",
        "        write_execution_report(report_path, report_title, started_at, finished_at, overall_status, STEP_RESULTS, error_message, FAILURE_SCREENSHOT)",
        "        print(f'Execution report written to: {report_path}')",
        "        if ACTIVE_DRIVER is not None:",
        "            try:",
        "                ACTIVE_DRIVER.quit()",
        "            except Exception:",
        "                pass",
        "",
    ])

    return "\n".join(lines)


# ── End Selenium Code Generator ───────────────────────────────────────────────


def _build_playwright_script(record_bundle: dict[str, Any]) -> str:
    """Generate a complete Playwright Python script from a record bundle."""
    meta = record_bundle["meta"]
    steps = record_bundle["steps"]
    first_url = next((s.get("page_url") for s in steps if (s.get("page_url") or "").strip()), "about:blank")
    first_nav_step = next((s for s in steps if (s.get("page_url") or "").strip()), {})
    first_nav_timeout_ms = _generator_step_timeout_millis(first_nav_step)

    lines = [
        '"""',
        "Enhanced Playwright Python script with proper waiting",
        '"""',
        "import time",
        "import traceback",
        "from datetime import datetime",
        "from html import escape",
        "from pathlib import Path",
        "from playwright.sync_api import Playwright, sync_playwright, expect",
        "",
        "STEP_RESULTS = []",
        "CURRENT_STEP = None",
        "CURRENT_STEP_STARTED = None",
        "ACTIVE_BROWSER = None",
        "ACTIVE_CONTEXT = None",
        "ACTIVE_PAGE = None",
        "ACTIVE_ARTIFACT_DIR = None",
        "FAILURE_SCREENSHOT = ''",
        "",
        "",
        "def build_output_paths():",
        "    script_path = Path(__file__).resolve() if '__file__' in globals() else Path.cwd() / 'generated_playwright.py'",
        "    script_dir = script_path.parent",
        "    script_stem = script_path.stem",
        "    artifact_dir = script_dir / f'{script_stem}_artifacts'",
        "    artifact_dir.mkdir(parents=True, exist_ok=True)",
        "    report_path = script_dir / f'{script_stem}_report.html'",
        "    return report_path, artifact_dir",
        "",
        "",
        "def begin_step(step_no, description):",
        "    global CURRENT_STEP, CURRENT_STEP_STARTED",
        "    CURRENT_STEP = {'step_no': step_no, 'description': description}",
        "    CURRENT_STEP_STARTED = time.perf_counter()",
        "",
        "",
        "def save_step_screenshot(page, artifact_dir, step_no):",
        "    if page is None or artifact_dir is None or step_no in (None, ''):",
        "        return ''",
        "    screenshot_path = artifact_dir / f'step_{int(step_no):04d}_{datetime.now().strftime(\"%Y%m%d_%H%M%S\")}.png'",
        "    page.screenshot(path=str(screenshot_path), full_page=True)",
        "    return screenshot_path.relative_to(artifact_dir.parent).as_posix()",
        "",
        "",
        "def finish_step(status, error_message=''):",
        "    global CURRENT_STEP, CURRENT_STEP_STARTED",
        "    if CURRENT_STEP is None:",
        "        return",
        "    duration_ms = int((time.perf_counter() - CURRENT_STEP_STARTED) * 1000) if CURRENT_STEP_STARTED is not None else 0",
        "    screenshot_rel_path = ''",
        "    try:",
        "        screenshot_rel_path = save_step_screenshot(ACTIVE_PAGE, ACTIVE_ARTIFACT_DIR, CURRENT_STEP.get('step_no'))",
        "    except Exception:",
        "        screenshot_rel_path = ''",
        "    STEP_RESULTS.append({",
        "        'step_no': CURRENT_STEP.get('step_no', ''),",
        "        'description': CURRENT_STEP.get('description', ''),",
        "        'status': status,",
        "        'duration_ms': duration_ms,",
        "        'error': error_message,",
        "        'screenshot': screenshot_rel_path,",
        "    })",
        "    CURRENT_STEP = None",
        "    CURRENT_STEP_STARTED = None",
        "",
        "",
        "def save_failure_screenshot(page, artifact_dir):",
        "    screenshot_path = artifact_dir / f'failure_{datetime.now().strftime(\"%Y%m%d_%H%M%S\")}.png'",
        "    page.screenshot(path=str(screenshot_path), full_page=True)",
        "    return screenshot_path.relative_to(artifact_dir.parent).as_posix()",
        "",
        "",
        "def write_execution_report(report_path, report_title, started_at, finished_at, overall_status, step_results, error_message='', screenshot_rel_path=''):",
        "    duration_ms = int((finished_at - started_at).total_seconds() * 1000)",
        "    pass_count = sum(1 for i in step_results if i.get('status') == 'PASS')",
        "    fail_count = sum(1 for i in step_results if i.get('status') != 'PASS')",
        "    html_lines = [",
        "        '<!doctype html>',",
        "        '<html lang=\"en\"><head><meta charset=\"utf-8\"><meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">',",
        "        f'<title>{escape(report_title)}</title>',",
        "        '<style>',",
        "        ':root{--bg:#f8fafc;--card:#ffffff;--border:#e2e8f0;--text:#1e293b;--muted:#64748b;--accent:#6366f1;--pass-bg:#ecfdf5;--pass:#059669;--fail-bg:#fef2f2;--fail:#dc2626;--shadow:0 4px 6px -1px rgb(0 0 0/0.1),0 2px 4px -2px rgb(0 0 0/0.1);}',",
        "        '*{box-sizing:border-box;margin:0;padding:0;}',",
        "        'body{font-family:Inter,Segoe UI,system-ui,sans-serif;background:var(--bg);color:var(--text);padding:32px;line-height:1.6;}',",
        "        '.container{max-width:1400px;margin:0 auto;}',",
        "        '.header{background:linear-gradient(135deg,#6366f1 0%,#8b5cf6 100%);border-radius:16px;padding:32px 40px;color:#fff;margin-bottom:24px;box-shadow:var(--shadow);}',",
        "        '.header h1{font-size:1.75rem;font-weight:700;margin-bottom:8px;}',",
        "        '.header p{opacity:0.9;font-size:0.9rem;}',",
        "        '.summary{display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:16px;margin-bottom:24px;}',",
        "        '.card{background:var(--card);border:1px solid var(--border);border-radius:12px;padding:20px;box-shadow:var(--shadow);}',",
        "        '.card .label{font-size:0.8rem;text-transform:uppercase;letter-spacing:0.05em;color:var(--muted);font-weight:600;}',",
        "        '.card .value{font-size:1.5rem;font-weight:700;margin-top:4px;}',",
        "        '.card.pass-card .value{color:var(--pass);}',",
        "        '.card.fail-card .value{color:var(--fail);}',",
        "        '.badge{display:inline-block;padding:4px 12px;border-radius:20px;font-size:0.8rem;font-weight:600;}',",
        "        '.badge-pass{background:var(--pass-bg);color:var(--pass);}',",
        "        '.badge-fail{background:var(--fail-bg);color:var(--fail);}',",
        "        'table{width:100%;border-collapse:separate;border-spacing:0;background:var(--card);border-radius:12px;overflow:hidden;box-shadow:var(--shadow);margin-bottom:24px;}',",
        "        'thead{background:linear-gradient(135deg,#f1f5f9,#e2e8f0);}',",
        "        'th{padding:14px 16px;font-size:0.8rem;text-transform:uppercase;letter-spacing:0.05em;color:var(--muted);font-weight:700;text-align:left;}',",
        "        'td{padding:12px 16px;border-top:1px solid var(--border);vertical-align:top;font-size:0.9rem;}',",
        "        'tr:hover td{background:#f8fafc;}',",
        "        '.step-no{font-weight:700;color:var(--accent);min-width:50px;}',",
        "        '.screenshot-thumb{cursor:pointer;border-radius:8px;border:2px solid var(--border);max-width:200px;max-height:120px;object-fit:cover;transition:transform 0.2s,box-shadow 0.2s;}',",
        "        '.screenshot-thumb:hover{transform:scale(1.05);box-shadow:0 8px 25px rgb(0 0 0/0.15);}',",
        "        '.error-box{background:#1e1b4b;color:#e0e7ff;padding:20px;border-radius:12px;font-family:JetBrains Mono,Consolas,monospace;font-size:0.85rem;overflow:auto;white-space:pre-wrap;margin-top:16px;box-shadow:var(--shadow);}',",
        "        '.lightbox{display:none;position:fixed;inset:0;background:rgba(0,0,0,0.85);z-index:9999;justify-content:center;align-items:center;padding:24px;cursor:zoom-out;}',",
        "        '.lightbox.active{display:flex;}',",
        "        '.lightbox img{max-width:95vw;max-height:92vh;border-radius:12px;box-shadow:0 20px 60px rgb(0 0 0/0.5);}',",
        "        '.lightbox-close{position:absolute;top:20px;right:28px;color:#fff;font-size:2rem;cursor:pointer;opacity:0.7;transition:opacity 0.2s;}',",
        "        '.lightbox-close:hover{opacity:1;}',",
        "        '</style>',",
        "        '</head><body>',",
        "        '<div class=\"container\">',",
        "        f'<div class=\"header\"><h1>{escape(report_title)}</h1>',",
        "        f'<p>{escape(started_at.isoformat(sep=\" \", timespec=\"seconds\"))} &mdash; {escape(finished_at.isoformat(sep=\" \", timespec=\"seconds\"))}</p></div>',",
        "        '<div class=\"summary\">',",
        "        f'<div class=\"card\"><div class=\"label\">Duration</div><div class=\"value\">{duration_ms} ms</div></div>',",
        "        f'<div class=\"card\"><div class=\"label\">Total Steps</div><div class=\"value\">{len(step_results)}</div></div>',",
        "        f'<div class=\"card pass-card\"><div class=\"label\">Passed</div><div class=\"value\">{pass_count}</div></div>',",
        "        f'<div class=\"card fail-card\"><div class=\"label\">Failed</div><div class=\"value\">{fail_count}</div></div>',",
        "        f'<div class=\"card\"><div class=\"label\">Status</div><div class=\"value\"><span class=\"badge badge-{\"pass\" if overall_status == \"PASS\" else \"fail\"}\">{escape(overall_status)}</span></div></div>',",
        "        '</div>',",
        "        '<table><thead><tr><th>Step</th><th>Description</th><th>Status</th><th>Duration</th><th>Error</th><th>Screenshot</th></tr></thead><tbody>',",
        "    ]",
        "    for item in step_results:",
        "        badge_class = 'badge-pass' if item.get('status') == 'PASS' else 'badge-fail'",
        "        screenshot_cell = ''",
        "        if item.get('screenshot'):",
        "            screenshot_href = escape(str(item.get('screenshot')))",
        "            screenshot_cell = f'<img src=\"{screenshot_href}\" alt=\"step {item.get(\"step_no\")}\" class=\"screenshot-thumb\" onclick=\"openLightbox(this.src)\">'",
        "        error_text = escape(str(item.get('error', ''))) if item.get('error') else ''",
        "        html_lines.append(",
        "            '<tr>'",
        "            f'<td class=\"step-no\">{escape(str(item.get(\"step_no\", \"\")))}</td>'",
        "            f'<td>{escape(str(item.get(\"description\", \"\")))}</td>'",
        "            f'<td><span class=\"badge {badge_class}\">{escape(str(item.get(\"status\", \"\")))}</span></td>'",
        "            f'<td>{escape(str(item.get(\"duration_ms\", \"\")))} ms</td>'",
        "            f'<td>{error_text}</td>'",
        "            f'<td>{screenshot_cell}</td>'",
        "            '</tr>'",
        "        )",
        "    html_lines.append('</tbody></table>')",
        "    if screenshot_rel_path:",
        "        html_lines.append(f'<div class=\"card\" style=\"margin-bottom:16px\"><div class=\"label\">Failure Screenshot</div><img src=\"{escape(screenshot_rel_path)}\" class=\"screenshot-thumb\" style=\"max-width:600px;max-height:400px;margin-top:8px\" onclick=\"openLightbox(this.src)\"></div>')",
        "    if error_message:",
        "        html_lines.append(f'<div class=\"card\"><div class=\"label\">Error Details</div><div class=\"error-box\">{escape(error_message)}</div></div>')",
        "    html_lines.append('</div>')",
        "    html_lines.append('<div class=\"lightbox\" id=\"lightbox\" onclick=\"closeLightbox()\"><span class=\"lightbox-close\">&times;</span><img id=\"lightbox-img\" src=\"\" alt=\"screenshot\"></div>')",
        "    html_lines.append('<script>')",
        "    html_lines.append('function openLightbox(src){var lb=document.getElementById(\"lightbox\");document.getElementById(\"lightbox-img\").src=src;lb.classList.add(\"active\");}')",
        "    html_lines.append('function closeLightbox(){document.getElementById(\"lightbox\").classList.remove(\"active\");}')",
        "    html_lines.append('document.addEventListener(\"keydown\",function(e){if(e.key===\"Escape\")closeLightbox();});')",
        "    html_lines.append('</script>')",
        "    html_lines.append('</body></html>')",
        "    report_path.write_text('\\n'.join(html_lines), encoding='utf-8')",
        "",
        "",
        "def wait_for_page_ready(page, timeout=15000):",
        '    """Wait until page is fully loaded"""',
        '    page.wait_for_load_state("load", timeout=timeout)',
        "",
        "",
        "def wait_for_overlay_gone(page, timeout=15000):",
        '    """Wait until the in-progress overlay is hidden or detached"""',
        "    overlay = page.locator('xpath=//*[@id=\"inProgressPage\"]')",
        "    if overlay.count() == 0:",
        "        return",
        "    expect(overlay).to_be_hidden(timeout=timeout)",
        "",
        "",
        "def run(playwright: Playwright) -> None:",
        "    global ACTIVE_BROWSER, ACTIVE_CONTEXT, ACTIVE_PAGE",
        "    begin_step(0, 'Browser startup')",
        "    browser = playwright.chromium.launch(headless=False, args=['--start-maximized'])",
        "    ACTIVE_BROWSER = browser",
        "    context = browser.new_context(no_viewport=True)",
        "    ACTIVE_CONTEXT = context",
        "    page = context.new_page()",
        "    ACTIVE_PAGE = page",
        "    finish_step('PASS')",
        "",
        f'    page.goto("{_pw_escape(first_url)}")',
        "",
        "    # Wait for page readiness",
        f"    wait_for_page_ready(page, timeout={first_nav_timeout_ms})",
        f"    wait_for_overlay_gone(page, timeout={first_nav_timeout_ms})",
        "",
    ]

    first_navigate_skipped = False
    used_names: dict[str, int] = {}

    for step in steps:
        action = (step.get("action") or "").strip().lower()
        page_url = (step.get("page_url") or "").strip()

        if action in ("navigate", "open", "goto") and page_url and not first_navigate_skipped:
            if page_url == first_url:
                first_navigate_skipped = True
                continue

        base_name = _selenium_var_stem(step, f"step_{step.get('step_no') or len(used_names) + 1}")
        used_names[base_name] = used_names.get(base_name, 0) + 1
        suffix = "" if used_names[base_name] == 1 else f"_{used_names[base_name]}"
        locator_name = f"{base_name}{suffix}"
        step_lines = _pw_step_lines(step, locator_name)
        if step_lines:
            description = _generator_step_description(step)
            step_no = int(step.get("step_no") or 0)
            lines.append(f"    begin_step({step_no}, \"{_pw_escape(description)}\")")
            lines.extend(step_lines)
            lines.append("    finish_step('PASS')")
            lines.append("")

    lines.extend([
        "",
        "with sync_playwright() as playwright:",
        "    report_path, artifact_dir = build_output_paths()",
        "    ACTIVE_ARTIFACT_DIR = artifact_dir",
        f"    report_title = \"Playwright Execution Report - {_pw_escape(meta.get('record_name', 'Untitled'))}\"",
        "    started_at = datetime.now()",
        "    overall_status = 'PASS'",
        "    error_message = ''",
        "    try:",
        "        run(playwright)",
        "    except Exception:",
        "        overall_status = 'FAIL'",
        "        error_message = traceback.format_exc()",
        "        finish_step('FAIL', error_message.strip())",
        "        if ACTIVE_PAGE is not None:",
        "            try:",
        "                FAILURE_SCREENSHOT = save_failure_screenshot(ACTIVE_PAGE, artifact_dir)",
        "            except Exception:",
        "                FAILURE_SCREENSHOT = ''",
        "        raise",
        "    finally:",
        "        finished_at = datetime.now()",
        "        write_execution_report(report_path, report_title, started_at, finished_at, overall_status, STEP_RESULTS, error_message, FAILURE_SCREENSHOT)",
        "        print(f'Execution report written to: {report_path}')",
        "        if ACTIVE_CONTEXT is not None:",
        "            try:",
        "                ACTIVE_CONTEXT.close()",
        "            except Exception:",
        "                pass",
        "        if ACTIVE_BROWSER is not None:",
        "            try:",
        "                ACTIVE_BROWSER.close()",
        "            except Exception:",
        "                pass",
        "",
    ])

    return "\n".join(lines)


@csrf_exempt
@login_required
def workflow_assistant_generate(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required."}, status=405)

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid JSON payload."}, status=400)

    prompt = str(body.get("prompt") or "").strip()
    if not prompt:
        return JsonResponse({"ok": False, "error": "Prompt is required."}, status=400)

    try:
        started = time.perf_counter()
        matches = _find_generate_record_matches(prompt)
        if not matches:
            return JsonResponse(
                {
                    "ok": False,
                    "error": "No matching record was found in session_meta. Enter a record name or record_id from the recorded sessions list.",
                },
                status=404,
            )

        if len(matches) > 1:
            options = [
                f"{item['record_name']} ({item['folder_name'] or 'no project'})"
                for item in matches[:5]
            ]
            return JsonResponse(
                {
                    "ok": False,
                    "error": "Multiple records matched the prompt. Please use the exact record name or record_id.",
                    "matches": matches[:5],
                    "suggestions": options,
                },
                status=409,
            )

        bundle = _load_record_bundle(matches[0]["record_id"])
        if not bundle["steps"]:
            return JsonResponse(
                {
                    "ok": False,
                    "error": "The selected record exists in session_meta but has no steps in the steps table.",
                },
                status=404,
            )

        script = _build_robot_script(bundle)
        elapsed_ms = round((time.perf_counter() - started) * 1000, 1)
        result = {
            "script": script,
            "model": "postgresql-db-builder",
            "elapsed_ms": elapsed_ms,
            "record_id": bundle["meta"]["record_id"],
            "record_name": bundle["meta"]["record_name"],
            "step_count": len(bundle["steps"]),
            "folder_name": bundle["meta"]["folder_name"],
        }
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({"ok": True, **result})


@csrf_exempt
@login_required
@require_POST
def workflow_assistant_create_test_case(request):
    _ensure_ai_databank_schema()
    _ensure_ai_workflow_schema()

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid JSON payload."}, status=400)

    workflow_name = str(body.get("workflow_name") or "").strip()
    record_name = str(body.get("record_name") or "").strip()
    folder_name = str(body.get("folder_name") or "").strip()
    max_steps_per_page = body.get("max_steps_per_page", 8)
    use_llm_raw = body.get("use_llm", True)
    ollama_api = str(body.get("ollama_api") or get_config("ai.ollama_api", "http://localhost:11434/api") or "http://localhost:11434/api").strip()
    ollama_model = str(body.get("model") or get_config("ai.ollama_model", "llama3.2:3b") or "llama3.2:3b").strip()

    if isinstance(use_llm_raw, str):
        use_llm = use_llm_raw.strip().lower() not in {"0", "false", "no", "off"}
    else:
        use_llm = bool(use_llm_raw)

    try:
        max_steps_per_page = int(max_steps_per_page)
    except (TypeError, ValueError):
        return JsonResponse({"ok": False, "error": "max_steps_per_page must be an integer."}, status=400)

    try:
        result = create_test_case_from_workflow(
            workflow_name,
            record_name=record_name or None,
            folder_name=folder_name,
            author=(request.user.username or "system").strip() or "system",
            max_steps_per_page=max_steps_per_page,
            use_llm=use_llm,
            ollama_api=ollama_api,
            model=ollama_model,
        )
    except WorkflowNotFoundError as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=404)
    except WorkflowGenerationError as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=400)
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({"ok": True, **result})


@csrf_exempt
@login_required
def workflow_assistant_finetune(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required."}, status=405)

    is_multipart = request.content_type and request.content_type.startswith("multipart/form-data")
    if is_multipart:
        body = request.POST
    else:
        try:
            body = json.loads(request.body.decode("utf-8")) if request.body else {}
        except Exception:
            return JsonResponse({"ok": False, "error": "Invalid JSON payload."}, status=400)

    action = str(body.get("action") or "").strip()
    ollama_api = str(body.get("ollama_api") or get_config("ai.ollama_api", "http://localhost:11434/api") or "http://localhost:11434/api").strip()

    try:
        from llm_workflow_assistant.finetune_service import (
            build_finetune_dataset,
            create_finetuned_model,
            save_uploaded_finetune_dataset,
        )

        if action == "build_dataset":
            result = build_finetune_dataset()
            return JsonResponse({"ok": True, "action": "build_dataset", **result})

        elif action == "upload_dataset":
            uploaded_file = request.FILES.get("dataset")
            if uploaded_file is None:
                return JsonResponse({"ok": False, "error": "Choose a finetune.json or .jsonl file first."}, status=400)
            base_model = str(body.get("base_model") or get_config("ai.ollama_model", "llama3.2:3b") or "llama3.2:3b").strip()
            ft_model_name = str(body.get("ft_model_name") or get_config("ai.ft_model", "llama3-finetuned") or "llama3-finetuned").strip()
            result = save_uploaded_finetune_dataset(
                filename=uploaded_file.name,
                content=uploaded_file.read(),
                base_dir=str(settings.BASE_DIR),
                base_model=base_model,
                ft_model_name=ft_model_name,
            )
            return JsonResponse({"ok": True, "action": "upload_dataset", **result})

        elif action == "create_model":
            base_model = str(body.get("base_model") or get_config("ai.ollama_model", "llama3.2:3b") or "llama3.2:3b").strip()
            ft_model_name = str(body.get("ft_model_name") or get_config("ai.ft_model", "llama3-finetuned") or "llama3-finetuned").strip()
            result = create_finetuned_model(
                ollama_api=ollama_api,
                base_model=base_model,
                ft_model_name=ft_model_name,
            )
            return JsonResponse({"ok": True, "action": "create_model", **result})

        elif action == "start_training":
            base_model = str(body.get("base_model") or get_config("ai.ollama_model", "llama3.2:3b") or "llama3.2:3b").strip()
            ft_model_name = str(body.get("ft_model_name") or get_config("ai.ft_model", "llama3-finetuned") or "llama3-finetuned").strip()
            result = _start_finetune_training(base_model=base_model, ft_model_name=ft_model_name)
            return JsonResponse({"ok": True, "action": "start_training", "training_status": result})

        else:
            return JsonResponse({"ok": False, "error": "action must be 'build_dataset', 'upload_dataset', 'start_training', or 'create_model'."}, status=400)

    except requests.RequestException as exc:
        return JsonResponse({"ok": False, "error": f"Ollama request failed: {exc}"}, status=502)
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)


@login_required
def workflow_assistant_finetune_status(request):
    if request.method != "GET":
        return JsonResponse({"ok": False, "error": "GET required."}, status=405)

    ollama_api = str(request.GET.get("ollama_api") or get_config("ai.ollama_api", "http://localhost:11434/api") or "http://localhost:11434/api").strip()

    try:
        from llm_workflow_assistant.finetune_service import get_finetune_status, get_training_environment_status, list_ollama_models

        models = list_ollama_models(ollama_api=ollama_api)
        ft_status = get_finetune_status(ollama_api=ollama_api, _models=models)
        training_environment = get_training_environment_status()
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({"ok": True, "ft_status": ft_status, "models": models, "training_status": _serialize_training_state(), "training_environment": training_environment})


@login_required
def ai_databank_screenshot(request, row_id: int):
    _ensure_ai_databank_schema()
    with connection.cursor() as cur:
        cur.execute("SELECT screenshot_png, locator_property FROM ai_databank WHERE id = %s", [row_id])
        row = cur.fetchone()

    if not row or not row[0]:
        raise Http404()

    screenshot_bytes = bytes(row[0])
    locator_property = _normalize_ai_databank_locator_property(row[1])
    bounds = locator_property.get("bounds") or {}
    viewport = locator_property.get("viewport") or {}

    try:
        image = Image.open(io.BytesIO(screenshot_bytes)).convert("RGBA")
        viewport_width = float(viewport.get("width") or 0)
        viewport_height = float(viewport.get("height") or 0)
        left = max(0.0, float(bounds.get("left") or 0))
        top = max(0.0, float(bounds.get("top") or 0))
        width = max(0.0, float(bounds.get("width") or 0))
        height = max(0.0, float(bounds.get("height") or 0))

        if viewport_width <= 0 or viewport_height <= 0 or width <= 0 or height <= 0:
            return HttpResponse(screenshot_bytes, content_type="image/png")

        scale_x = image.width / viewport_width
        scale_y = image.height / viewport_height
        image_left = max(0, min(image.width, int(round(left * scale_x))))
        image_top = max(0, min(image.height, int(round(top * scale_y))))
        image_width = max(1, int(round(width * scale_x)))
        image_height = max(1, int(round(height * scale_y)))
        image_right = max(image_left + 1, min(image.width, image_left + image_width))
        image_bottom = max(image_top + 1, min(image.height, image_top + image_height))

        pad = max(16, int(max(image_right - image_left, image_bottom - image_top) * 0.35))
        crop_left = max(0, image_left - pad)
        crop_top = max(0, image_top - pad)
        crop_right = min(image.width, image_right + pad)
        crop_bottom = min(image.height, image_bottom + pad)

        cropped = image.crop((crop_left, crop_top, crop_right, crop_bottom))
        draw = ImageDraw.Draw(cropped)
        rect = (
            image_left - crop_left,
            image_top - crop_top,
            image_right - crop_left,
            image_bottom - crop_top,
        )
        draw.rectangle(rect, outline=(220, 53, 69, 255), width=4)

        out = io.BytesIO()
        cropped.save(out, format="PNG")
        return HttpResponse(out.getvalue(), content_type="image/png")
    except Exception:
        return HttpResponse(screenshot_bytes, content_type="image/png")


def _normalize_ai_databank_locator_property(raw_value: Any) -> dict[str, Any]:
    locator_property = raw_value or {}
    if isinstance(locator_property, str):
        try:
            locator_property = json.loads(locator_property)
        except Exception:
            locator_property = {}
    return locator_property if isinstance(locator_property, dict) else {}


def _get_ai_databank_page_name_options(cursor) -> list[str]:
    cursor.execute(
        """
        SELECT DISTINCT COALESCE(NULLIF(TRIM(page_name), ''), 'Untitled page') AS normalized_page_name
        FROM ai_databank
        ORDER BY normalized_page_name ASC
        """
    )
    return [row[0] for row in cursor.fetchall() if row[0]]


def _get_ai_databank_flow_cards(cursor) -> list[dict[str, Any]]:
    primary_locators = _get_primary_locator_lookup(cursor)
    cursor.execute(
        """
        SELECT
            id,
            COALESCE(NULLIF(TRIM(page_name), ''), 'Untitled page') AS normalized_page_name,
            NULLIF(TRIM(page_url), '') AS page_url,
            element_type,
            locator_property,
            updated_at
        FROM ai_databank
        ORDER BY normalized_page_name ASC, updated_at DESC, id DESC
        """
    )
    pages: dict[str, dict[str, Any]] = {}
    page_order: list[str] = []
    for row in cursor.fetchall():
        row_id = int(row[0])
        page_name = str(row[1] or "").strip() or "Untitled page"
        page_url = row[2] or ""
        element_type = str(row[3] or "element").strip() or "element"
        locator_property = _normalize_ai_databank_locator_property(row[4])
        updated_at = row[5]
        page_entry = pages.get(page_name)
        if page_entry is None:
            page_entry = {
                "page_name": page_name,
                "item_count": 0,
                "page_url": page_url,
                "last_seen_at": updated_at,
                "elements": [],
            }
            pages[page_name] = page_entry
            page_order.append(page_name)
        page_entry["item_count"] += 1
        if not page_entry.get("page_url") and page_url:
            page_entry["page_url"] = page_url
        if updated_at and (page_entry.get("last_seen_at") is None or updated_at > page_entry["last_seen_at"]):
            page_entry["last_seen_at"] = updated_at
        element_option = _build_ai_databank_flow_element_option(
            row_id,
            page_name,
            page_url,
            element_type,
            locator_property,
            primary_locators,
        )
        if element_option:
            # De-duplicate: skip if another row already claimed the same primary (strategy, locator) for this page
            locator_key = (element_option["primary_strategy"], element_option["primary_locator"])
            seen = page_entry.setdefault("_seen_locator_keys", set())
            if locator_key not in seen:
                seen.add(locator_key)
                if element_option["is_stable_primary"]:
                    page_entry["elements"].append(element_option)
                else:
                    page_entry.setdefault("_fallback_elements", []).append(element_option)

    cards: list[dict[str, Any]] = []
    for index, page_name in enumerate(page_order):
        page_entry = pages[page_name]
        elements = page_entry["elements"]
        if not elements:
            # No stable-primary elements found for this page — show one best fallback
            fallbacks = page_entry.get("_fallback_elements", [])
            if fallbacks:
                elements = [fallbacks[0]]
        cards.append({
            "id": f"page-card-{index + 1}",
            "page_name": page_entry["page_name"],
            "item_count": int(page_entry["item_count"] or 0),
            "page_url": page_entry.get("page_url") or "",
            "last_seen_at": page_entry["last_seen_at"].isoformat() if page_entry.get("last_seen_at") else None,
            "elements": elements,
        })
    return cards


def _get_primary_locator_lookup(cursor) -> dict[str, dict[str, set[str]]]:
    """
    Build a per-page-URL lookup of which CSS/xpath locator strings were marked
    as primary in actual recordings.

    Match chain:
      steps.locators_raw (JSONB) has the same locator keys as
      ai_databank.locator_property.locators, so we join on:
          steps.page_url  =  ai_databank.page_url
          steps.locators_raw->>'css'  =  locator_property->'locators'->>'css'
      (and equivalently for 'xpath').

    Returns:
      {page_url: {'css': {css_value, ...}, 'xpath': {xpath_value, ...}}}
    """
    lookup: dict[str, dict[str, set[str]]] = {}

    def _add(url: str, strategy: str, value: str) -> None:
        if not url or not value or value.strip().lower() in ("", "none", "null"):
            return
        entry = lookup.setdefault(url.strip(), {"css": set(), "xpath": set()})
        strat_key = strategy.strip().lower()
        if strat_key in ("css", "xpath"):
            entry.setdefault(strat_key, set()).add(value.strip())

    # Primary source: steps.locators_raw (JSONB) — both css and xpath columns
    cursor.execute(
        """
        SELECT DISTINCT
            NULLIF(TRIM(page_url), '') AS page_url,
            NULLIF(TRIM(locators_raw->>'css'), '') AS css_loc,
            NULLIF(TRIM(locators_raw->>'xpath'), '') AS xpath_loc
        FROM steps
        WHERE is_primary = TRUE
          AND locators_raw IS NOT NULL
          AND NULLIF(TRIM(page_url), '') IS NOT NULL
        """
    )
    for page_url, css_loc, xpath_loc in cursor.fetchall():
        _add(page_url, "css", css_loc or "")
        _add(page_url, "xpath", xpath_loc or "")

    # Secondary source: locators table joined to steps for page_url
    cursor.execute(
        """
        SELECT DISTINCT
            NULLIF(TRIM(s.page_url), '') AS page_url,
            TRIM(l.strategy),
            TRIM(l.locator)
        FROM locators l
        JOIN steps s ON s.record_id = l.record_id AND s.step_no = l.step_no
        WHERE l.is_primary = TRUE
          AND NULLIF(TRIM(l.strategy), '') IS NOT NULL
          AND NULLIF(TRIM(l.locator), '') IS NOT NULL
          AND NULLIF(TRIM(s.page_url), '') IS NOT NULL
        """
    )
    for page_url, strategy, locator in cursor.fetchall():
        _add(page_url or "", strategy or "", locator or "")

    return lookup


def _build_ai_databank_flow_element_option(
    row_id: int,
    page_name: str,
    page_url: str,
    element_type: str,
    locator_property: dict[str, Any],
    primary_locators: dict[str, dict[str, set[str]]],
) -> dict[str, Any] | None:
    # Elements are shown only when the recorder actually interacted with them
    # (cross-table match against locators.is_primary=TRUE + steps).
    # Non-recorded elements are only surfaced if they pass the interactive-type
    # filter AND have a fallback locator — they show as non-primary candidates.
    _INTERACTIVE_TYPES = frozenset({
        "button", "checkbox", "radio", "entry field", "dropdown", "link",
    })
    _STRUCTURAL_TAGS = frozenset({
        "footer", "header", "nav", "aside", "main", "section", "article",
        "html", "body", "head", "script", "style", "meta", "noscript", "p",
    })
    et = element_type.strip().lower()
    tag_name_raw = str(locator_property.get("tag_name") or "").strip().lower()
    tag_name = tag_name_raw

    # ── Build locator candidates ──────────────────────────────────────────
    ordered = _build_ai_databank_ordered_locators(locator_property)
    primary = None
    cross_table_match = False

    url_key = (page_url or "").strip()
    url_prim = primary_locators.get(url_key, {})
    prim_css_set = url_prim.get("css", set())
    prim_xpath_set = url_prim.get("xpath", set())

    locators_block = locator_property.get("locators") or {}
    elem_css = str(locators_block.get("css") or "").strip()
    elem_xpath = str(locators_block.get("xpath") or "").strip()

    # The scraper stores its own rank-1 locator in locator_property.primary.
    # Use it as an additional match candidate against the recording sets.
    scraped_primary = locator_property.get("primary") or {}
    sp_strat = str(scraped_primary.get("strategy") or "").strip().lower()
    sp_loc = str(scraped_primary.get("locator") or scraped_primary.get("prepared_locator") or "").strip()

    def _in_primary_sets(strategy: str, locator: str) -> bool:
        if not locator or locator.lower() in ("none", "null", ""):
            return False
        if strategy == "css":
            return locator in prim_css_set
        if strategy == "xpath":
            return locator in prim_xpath_set
        return False

    # Match via locators.css / locators.xpath, then via locator_property.primary
    if (
        _in_primary_sets("css", elem_css)
        or _in_primary_sets("xpath", elem_xpath)
        or _in_primary_sets(sp_strat, sp_loc)
    ):
        cross_table_match = True
        # Prefer the scraper's own rank-1 locator as the primary when it matched
        if sp_loc and _in_primary_sets(sp_strat, sp_loc):
            primary = {
                "strategy": sp_strat,
                "locator": sp_loc,
                "prepared_locator": scraped_primary.get("prepared_locator") or sp_loc,
                "rank": 1,
                "by": scraped_primary.get("by") or sp_strat,
                "wrapped": bool(scraped_primary.get("wrapped")),
            }
        else:
            # Fall back to searching ordered list for the matching item
            for locator_item in ordered:
                s = str(locator_item.get("strategy") or "").strip().lower()
                v = str(locator_item.get("locator") or "").strip()
                if _in_primary_sets(s, v):
                    primary = locator_item
                    break

    # Cross-table match bypasses the element-type filter — a recorded click
    # proves interactivity regardless of scraper classification.
    if not cross_table_match:
        if et not in _INTERACTIVE_TYPES:
            return None
        if tag_name_raw in _STRUCTURAL_TAGS:
            return None

    # For non-matched elements, pick the best fallback locator to display.
    if not primary:
        _FALLBACK_PREFERENCE = ("xpath", "css", "id", "href", "name", "value", "placeholder")
        by_strategy = {item.get("strategy", ""): item for item in ordered}
        # Prefer scraped_primary if available as the fallback display
        if sp_loc and sp_strat and sp_strat in by_strategy:
            primary = by_strategy[sp_strat]
        else:
            for strat in _FALLBACK_PREFERENCE:
                if strat in by_strategy:
                    primary = by_strategy[strat]
                    break
        if not primary and ordered:
            primary = ordered[0]
    if not primary:
        return None

    # is_stable_primary = True ONLY for genuine recording matches.
    # Heuristics (CSS #id, id attr, name attr) are removed — they produced
    # false positives on pages with no recording history.
    is_stable_primary = cross_table_match

    # Label: use locator_property.text (visible text content) as the primary
    # display, falling back to placeholder / name / id / href.
    raw_text = str(locator_property.get("text") or locators_block.get("text") or "").strip()
    label_text = re.sub(r"\s+", " ", raw_text)
    if not label_text:
        label_text = re.sub(r"\s+", " ", str(
            locators_block.get("placeholder") or
            locators_block.get("name") or
            locators_block.get("id") or
            locators_block.get("href") or
            locators_block.get("ariaLabel") or
            ""
        ).strip())
    if len(label_text) > 60:
        label_text = label_text[:57].rstrip() + "..."

    descriptor_parts = [part for part in [element_type.strip(), tag_name] if part]
    if label_text:
        descriptor_parts.append(label_text)
    label = " | ".join(descriptor_parts) if descriptor_parts else f"element | #{row_id}"

    def _flow_text(value: Any) -> str:
        return re.sub(r"\s+", " ", str(value or "")).strip()

    return {
        "id": row_id,
        "page_name": page_name,
        "page_url": page_url,
        "element_type": element_type,
        "tag_name": tag_name,
        "label": label,
        "is_stable_primary": is_stable_primary,
        "text": _flow_text(locator_property.get("text") or locators_block.get("text")),
        "name": _flow_text(locators_block.get("name")),
        "placeholder": _flow_text(locators_block.get("placeholder")),
        "href": _flow_text(locators_block.get("href")),
        "primary_strategy": primary.get("strategy") if primary else "",
        "primary_locator": primary.get("locator") if primary else "",
        "locator_summary": f"{primary['strategy']}:{primary['locator']}" if primary else "",
    }


def _flow_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _flow_page_name(value: Any) -> str:
    return _flow_text(value) or "Untitled page"


def _build_workflow_sync_element_label(
    row_id: int,
    page_name: str,
    page_url: str,
    element_type: str,
    locator_property: dict[str, Any],
    ordered_locators: list[dict[str, Any]],
) -> str:
    label_source = (
        locator_property.get("text")
        or locator_property.get("aria_label")
        or locator_property.get("ariaLabel")
        or locator_property.get("placeholder")
        or locator_property.get("name")
        or locator_property.get("id")
        or locator_property.get("title")
        or locator_property.get("href")
        or page_name
        or page_url
    )
    label_text = _flow_text(label_source)
    if len(label_text) > 60:
        label_text = label_text[:57].rstrip() + "..."
    tag_name = _flow_text(locator_property.get("tag_name")).lower()
    descriptor_parts = [part for part in [_flow_text(element_type), tag_name, label_text] if part]
    primary = ordered_locators[0] if ordered_locators else None
    if primary:
        descriptor_parts.append(f"{primary['strategy']}={primary['locator']}")
    return " | ".join(descriptor_parts) if descriptor_parts else f"element | #{row_id}"


def _get_ai_workflow_name_options(cursor) -> list[str]:
    cursor.execute(
        """
        SELECT workflow_name
        FROM ai_workflow
        WHERE TRIM(COALESCE(workflow_name, '')) <> ''
        ORDER BY LOWER(workflow_name), workflow_name
        """
    )
    return [str(row[0]).strip() for row in cursor.fetchall() if str(row[0]).strip()]


def _get_ai_workflow_source_sessions(cursor) -> list[dict[str, Any]]:
    cursor.execute(
        """
        SELECT
            m.record_id::text AS record_id,
            COALESCE(NULLIF(TRIM(m.record_name), ''), m.record_id::text) AS record_name,
            COUNT(DISTINCT s.step_no) AS step_count,
            MAX(s.created_at) AS last_seen_at
        FROM session_meta m
        LEFT JOIN steps s ON s.record_id = m.record_id
        WHERE m.is_baseline = TRUE
        GROUP BY m.record_id, m.record_name
        ORDER BY LOWER(COALESCE(NULLIF(TRIM(m.record_name), ''), m.record_id::text)), MAX(s.created_at) DESC
        """
    )
    sessions: list[dict[str, Any]] = []
    for row in cursor.fetchall():
        sessions.append({
            "record_id": str(row[0] or "").strip(),
            "record_name": str(row[1] or "").strip(),
            "step_count": int(row[2] or 0),
            "last_seen_at": row[3].isoformat() if row[3] else None,
        })
    return sessions


def _get_ai_workflow_by_name(cursor, workflow_name: str) -> dict[str, Any] | None:
    cursor.execute(
        """
        SELECT id, workflow_name, page_connections, page_sequence, workflow_payload, created_at, updated_at
        FROM ai_workflow
        WHERE workflow_name = %s
        LIMIT 1
        """,
        [workflow_name],
    )
    row = cursor.fetchone()
    if not row:
        return None
    pc = _safe_json_value(row[2])
    ps = _safe_json_value(row[3])
    wp = _safe_json_value(row[4])
    return {
        "id": row[0],
        "workflow_name": row[1],
        "page_connections": pc if isinstance(pc, list) else [],
        "page_sequence": ps if isinstance(ps, list) else [],
        "workflow_payload": wp if isinstance(wp, dict) else {},
        "created_at": row[5].isoformat() if row[5] else None,
        "updated_at": row[6].isoformat() if row[6] else None,
    }


def _get_ai_workflow_by_source_record_id(cursor, record_id: str) -> dict[str, Any] | None:
    cursor.execute(
        """
        SELECT id, workflow_name, page_connections, page_sequence, workflow_payload, created_at, updated_at
        FROM ai_workflow
        WHERE workflow_payload ->> 'source_record_id' = %s
        ORDER BY updated_at DESC, id DESC
        LIMIT 1
        """,
        [str(record_id or "").strip()],
    )
    row = cursor.fetchone()
    if not row:
        return None
    pc = _safe_json_value(row[2])
    ps = _safe_json_value(row[3])
    wp = _safe_json_value(row[4])
    return {
        "id": row[0],
        "workflow_name": row[1],
        "page_connections": pc if isinstance(pc, list) else [],
        "page_sequence": ps if isinstance(ps, list) else [],
        "workflow_payload": wp if isinstance(wp, dict) else {},
        "created_at": row[5].isoformat() if row[5] else None,
        "updated_at": row[6].isoformat() if row[6] else None,
    }


def _resolve_workflow_name_for_record(cursor, record_id: str, record_name: str) -> str:
    existing = _get_ai_workflow_by_source_record_id(cursor, record_id)
    if existing and _flow_text(existing.get("workflow_name")):
        return _flow_text(existing.get("workflow_name"))

    base_name = _flow_text(record_name) or str(record_id).strip()
    existing_by_name = _get_ai_workflow_by_name(cursor, base_name)
    if not existing_by_name:
        return base_name

    payload = existing_by_name.get("workflow_payload") if isinstance(existing_by_name.get("workflow_payload"), dict) else {}
    if str(payload.get("source_record_id") or "").strip() == str(record_id).strip():
        return base_name
    return f"{base_name} [{str(record_id).strip()[:8]}]"


def _load_record_steps_for_workflow(cursor, record_id: str) -> list[dict[str, Any]]:
    cursor.execute(
        """
        SELECT
            s.step_no,
            COALESCE(s.page_url, ''),
            COALESCE(s.page_title, ''),
            COALESCE(s.action, ''),
            COALESCE(s.element_tag, ''),
            COALESCE(NULLIF(TRIM(s.strategy), ''), NULLIF(TRIM(l.strategy), ''), ''),
            COALESCE(NULLIF(TRIM(s.locator), ''), NULLIF(TRIM(l.locator), ''), ''),
            s.raw_event,
            s.locators_raw,
            COALESCE(NULLIF(TRIM(m.record_name), ''), s.record_id::text)
        FROM steps s
        LEFT JOIN locators l ON l.record_id = s.record_id
                             AND l.step_no = s.step_no
                             AND l.is_primary = TRUE
        LEFT JOIN session_meta m ON m.record_id = s.record_id
        WHERE s.record_id = %s
        ORDER BY s.step_no ASC, s.id ASC
        """,
        [str(record_id).strip()],
    )
    rows: list[dict[str, Any]] = []
    for row in cursor.fetchall():
        rows.append({
            "step_no": int(row[0] or 0),
            "page_url": _flow_text(row[1]),
            "page_title": _flow_text(row[2]),
            "action": _flow_text(row[3]),
            "element_tag": _flow_text(row[4]),
            "strategy": _flow_text(row[5]),
            "locator": _flow_text(row[6]),
            "raw_event": row[7] if isinstance(row[7], dict) else _safe_json_value(row[7]),
            "locators_raw": row[8] if isinstance(row[8], dict) else _safe_json_value(row[8]),
            "record_name": _flow_text(row[9]),
        })
    return rows


def _load_ai_databank_rows_for_workflow_record(cursor, page_urls: list[str], page_titles: list[str]) -> list[dict[str, Any]]:
    filters: list[str] = []
    params: list[Any] = []
    unique_urls = [item for item in dict.fromkeys(_flow_text(value) for value in page_urls) if item]
    unique_titles = [item for item in dict.fromkeys(_flow_page_name(value) for value in page_titles) if item]

    if unique_urls:
        filters.append("COALESCE(NULLIF(TRIM(page_url), ''), '') IN (" + ", ".join(["%s"] * len(unique_urls)) + ")")
        params.extend(unique_urls)
    if unique_titles:
        filters.append("COALESCE(NULLIF(TRIM(page_name), ''), 'Untitled page') IN (" + ", ".join(["%s"] * len(unique_titles)) + ")")
        params.extend(unique_titles)

    where_sql = "WHERE " + " OR ".join(filters) if filters else ""
    cursor.execute(
        f"""
        SELECT id,
               COALESCE(NULLIF(TRIM(page_name), ''), 'Untitled page') AS normalized_page_name,
               COALESCE(NULLIF(TRIM(page_url), ''), '') AS page_url,
               COALESCE(element_type, '') AS element_type,
               locator_property
        FROM ai_databank
        {where_sql}
        ORDER BY normalized_page_name ASC, page_url ASC, updated_at DESC, id DESC
        """,
        params,
    )
    rows: list[dict[str, Any]] = []
    for row in cursor.fetchall():
        locator_property = _normalize_ai_databank_locator_property(row[4])
        ordered = _build_ai_databank_ordered_locators(locator_property)
        locator_keys: set[tuple[str, str]] = set()
        for item in ordered:
            strategy = _flow_text(item.get("strategy"))
            locator = _flow_text(item.get("locator"))
            prepared = _flow_text(item.get("prepared_locator"))
            if strategy and locator:
                locator_keys.add((strategy, locator))
            if strategy and prepared:
                locator_keys.add((strategy, prepared))
        rows.append({
            "id": int(row[0]),
            "page_name": _flow_page_name(row[1]),
            "page_url": _flow_text(row[2]),
            "element_type": _flow_text(row[3]),
            "locator_property": locator_property,
            "ordered_locators": ordered,
            "locator_keys": locator_keys,
            "label": _build_workflow_sync_element_label(
                int(row[0]),
                _flow_page_name(row[1]),
                _flow_text(row[2]),
                _flow_text(row[3]),
                locator_property,
                ordered,
            ),
        })
    return rows


def _step_locator_keys(step: dict[str, Any]) -> set[tuple[str, str]]:
    keys: set[tuple[str, str]] = set()

    def add_pair(strategy: Any, locator: Any) -> None:
        cleaned_strategy = _flow_text(strategy)
        cleaned_locator = _flow_text(locator)
        if not cleaned_strategy or not cleaned_locator:
            return
        keys.add((cleaned_strategy, cleaned_locator))
        normalized = _normalize_ai_databank_strategy_locator(cleaned_strategy, cleaned_locator)
        if normalized:
            keys.add((cleaned_strategy, normalized))
            prepared = _prepare_ai_workflow_locator(cleaned_strategy, normalized)
            if prepared:
                keys.add((cleaned_strategy, prepared))

    add_pair(step.get("strategy"), step.get("locator"))
    locators_raw = step.get("locators_raw") if isinstance(step.get("locators_raw"), dict) else {}
    for strategy, locator in locators_raw.items():
        add_pair(strategy, locator)
    raw_event = step.get("raw_event") if isinstance(step.get("raw_event"), dict) else {}
    raw_event_locators = raw_event.get("locators") if isinstance(raw_event.get("locators"), dict) else {}
    for strategy, locator in raw_event_locators.items():
        add_pair(strategy, locator)
    return keys


def _prepare_ai_workflow_locator(strategy: str, locator: str) -> str:
    strategy_key = _flow_text(strategy)
    cleaned_locator = _flow_text(locator)
    if not strategy_key or not cleaned_locator:
        return ""
    wrap_map = {
        "value": '[value="{value}"]',
        "placeholder": '[placeholder="{value}"]',
        "type": '[type="{value}"]',
        "role": '[role="{value}"]',
        "title": '[title="{value}"]',
        "alt": '[alt="{value}"]',
        "href": '[href="{value}"]',
        "text": '//*[normalize-space(text())="{value}"]',
        "name":       '[name="{value}"]',
        "ariaLabel":  '[aria-label="{value}"]',
        "dataTestId": '[data-testid="{value}"]',
    }
    template = wrap_map.get(strategy_key)
    return template.format(value=cleaned_locator) if template else cleaned_locator


def _match_ai_databank_row_for_step(step: dict[str, Any], databank_rows: list[dict[str, Any]]) -> dict[str, Any] | None:
    step_keys = _step_locator_keys(step)
    step_page_url = _flow_text(step.get("page_url"))
    step_page_title = _flow_page_name(step.get("page_title"))
    best_row = None
    best_score = -1

    for row in databank_rows:
        score = 0
        if step_page_url and row.get("page_url"):
            if row["page_url"] == step_page_url:
                score += 120
            elif row["page_url"].rstrip("/") == step_page_url.rstrip("/"):
                score += 100
        if step_page_title and row.get("page_name") == step_page_title:
            score += 35
        if step_keys and row.get("locator_keys"):
            overlap = step_keys.intersection(row["locator_keys"])
            if overlap:
                score += 400 + (10 * len(overlap))
        if score > best_score:
            best_score = score
            best_row = row

    return best_row if best_score > 0 else None


def _build_ai_workflow_from_record(cursor, record_id: str) -> dict[str, Any]:
    steps = _load_record_steps_for_workflow(cursor, record_id)
    if not steps:
        raise WorkflowGenerationError("No steps were found for the selected record_id.")

    record_name = _flow_text(steps[0].get("record_name")) or str(record_id).strip()
    workflow_name = _resolve_workflow_name_for_record(cursor, record_id, record_name)
    databank_rows = _load_ai_databank_rows_for_workflow_record(
        cursor,
        [step.get("page_url") for step in steps],
        [step.get("page_title") for step in steps],
    )

    step_items: list[dict[str, Any]] = []
    for step in steps:
        matched_row = _match_ai_databank_row_for_step(step, databank_rows)
        page_name = matched_row.get("page_name") if matched_row else (_flow_page_name(step.get("page_title")) or _flow_page_name(step.get("page_url")))
        page_url = matched_row.get("page_url") if matched_row and _flow_text(matched_row.get("page_url")) else _flow_text(step.get("page_url"))
        selected = {
            "selected_element_id": matched_row.get("id") if matched_row else None,
            "selected_element_label": _flow_text(matched_row.get("label")) if matched_row else "",
            "selected_element_strategy": _flow_text(step.get("strategy")),
            "selected_element_locator": _flow_text(step.get("locator")),
        }
        if matched_row:
            ordered = matched_row.get("ordered_locators") or []
            if ordered:
                selected["selected_element_strategy"] = _flow_text(ordered[0].get("strategy")) or selected["selected_element_strategy"]
                selected["selected_element_locator"] = _flow_text(ordered[0].get("locator")) or selected["selected_element_locator"]
        step_items.append({
            "step_no": step.get("step_no"),
            "action": _flow_text(step.get("action")),
            "page_name": page_name,
            "page_url": page_url,
            "selected": selected,
        })

    page_sequence: list[dict[str, Any]] = []
    cards_map: dict[str, dict[str, Any]] = {}
    selected_elements: list[dict[str, Any]] = []
    page_order: list[str] = []
    last_step_for_page: dict[str, dict[str, Any]] = {}
    page_seen: set[str] = set()
    current_page_name = ""
    previous_item: dict[str, Any] | None = None
    page_connections: list[dict[str, Any]] = []
    element_relationships: list[dict[str, Any]] = []
    seen_connections: set[tuple[str, str]] = set()

    for index, item in enumerate(step_items):
        page_name = _flow_page_name(item.get("page_name"))
        page_url = _flow_text(item.get("page_url"))
        if page_name not in page_seen:
            card_id = f"page-card-{len(page_order) + 1}"
            position_index = len(page_order)
            x = 28 + ((position_index % 3) * 300)
            y = 32 + ((position_index // 3) * 170)
            cards_map[page_name] = {
                "id": card_id,
                "page_name": page_name,
                "page_url": page_url,
                "item_count": 0,
                "x": x,
                "y": y,
                "selected_element_id": None,
                "selected_element_label": "",
                "selected_element_strategy": "",
                "selected_element_locator": "",
            }
            page_order.append(page_name)
            page_seen.add(page_name)
            page_sequence.append({
                "order": len(page_sequence) + 1,
                "card_id": card_id,
                "page_name": page_name,
                "page_url": page_url,
                "x": x,
                "y": y,
            })

        cards_map[page_name]["item_count"] += 1
        if item["selected"].get("selected_element_id"):
            cards_map[page_name].update(item["selected"])
            last_step_for_page[page_name] = item

        if current_page_name and current_page_name != page_name:
            connection_key = (current_page_name, page_name)
            if connection_key not in seen_connections:
                from_card = cards_map[current_page_name]
                to_card = cards_map[page_name]
                from_step = last_step_for_page.get(current_page_name) or previous_item or {}
                to_step = item
                connection = {
                    "from": from_card["id"],
                    "to": to_card["id"],
                    "from_page_name": current_page_name,
                    "to_page_name": page_name,
                    "from_element_id": from_step.get("selected", {}).get("selected_element_id"),
                    "from_element_label": from_step.get("selected", {}).get("selected_element_label", ""),
                    "from_element_strategy": from_step.get("selected", {}).get("selected_element_strategy", ""),
                    "from_element_locator": from_step.get("selected", {}).get("selected_element_locator", ""),
                    "to_element_id": to_step.get("selected", {}).get("selected_element_id"),
                    "to_element_label": to_step.get("selected", {}).get("selected_element_label", ""),
                    "to_element_strategy": to_step.get("selected", {}).get("selected_element_strategy", ""),
                    "to_element_locator": to_step.get("selected", {}).get("selected_element_locator", ""),
                }
                page_connections.append(connection)
                element_relationships.append({
                    "from_card_id": from_card["id"],
                    "to_card_id": to_card["id"],
                    **{key: value for key, value in connection.items() if key not in {"from", "to"}},
                })
                seen_connections.add(connection_key)

        current_page_name = page_name
        previous_item = item

    for page_name in page_order:
        card = cards_map[page_name]
        selected_elements.append({
            "card_id": card["id"],
            "page_name": card["page_name"],
            "selected_element_id": card["selected_element_id"],
            "selected_element_label": card["selected_element_label"],
            "selected_element_strategy": card["selected_element_strategy"],
            "selected_element_locator": card["selected_element_locator"],
        })

    workflow_payload = {
        "source_type": "steps",
        "source_record_id": str(record_id).strip(),
        "source_record_name": record_name,
        "cards": [cards_map[page_name] for page_name in page_order],
        "connections": [{"from": item["from"], "to": item["to"]} for item in page_connections],
        "selected_elements": selected_elements,
        "element_relationships": element_relationships,
        "view_state": {
            "scale": 1,
            "pan_x": 0,
            "pan_y": 0,
            "compact": False,
        },
    }

    return {
        "workflow_name": workflow_name,
        "record_id": str(record_id).strip(),
        "record_name": record_name,
        "page_connections": page_connections,
        "page_sequence": page_sequence,
        "workflow_payload": workflow_payload,
        "step_count": len(steps),
    }


def _normalize_ai_databank_strategy_locator(strategy: str, raw_value: Any) -> str:
    locator = str(raw_value or "").strip()
    if not locator:
        return ""
    if strategy == "id" and locator.startswith("#"):
        return locator[1:].strip()
    if strategy == "name":
        match = re.search(r'\[name="([^"]+)"\]', locator)
        if match:
            return match.group(1).strip()
    if strategy == "ariaLabel":
        match = re.search(r'\[aria-label="([^"]+)"\]', locator)
        if match:
            return match.group(1).strip()
    if strategy == "dataTestId":
        match = re.search(r'\[data-testid="([^"]+)"\]', locator)
        if match:
            return match.group(1).strip()
    return locator


def _build_ai_databank_ordered_locators(locator_property: dict[str, Any]) -> list[dict[str, Any]]:
    strategy_order = locator_property.get("fallback_order") or [
        "xpath", "id", "name", "value", "placeholder", "class", "className",
        "tagName", "css", "href", "text", "linkText", "partialLinkText",
        "type", "role", "title", "alt", "ariaLabel", "dataTestId",
    ]
    mapping = locator_property.get("mapping") or {}
    locator_block = locator_property.get("locators") or {}
    wrap_map = {
        "value": '[value="{value}"]',
        "placeholder": '[placeholder="{value}"]',
        "type": '[type="{value}"]',
        "role": '[role="{value}"]',
        "title": '[title="{value}"]',
        "alt": '[alt="{value}"]',
        "href": '[href="{value}"]',
        "text": '//*[normalize-space(text())="{value}"]',
        "name":       '[name="{value}"]',
        "ariaLabel":  '[aria-label="{value}"]',
        "dataTestId": '[data-testid="{value}"]',
    }
    ordered: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    for rank, strategy in enumerate(strategy_order, start=1):
        locator = _normalize_ai_databank_strategy_locator(strategy, locator_block.get(strategy))
        if not locator or (strategy, locator) in seen:
            continue
        seen.add((strategy, locator))
        ordered.append({
            "strategy": strategy,
            "locator": locator,
            "prepared_locator": wrap_map[strategy].format(value=locator) if strategy in wrap_map else locator,
            "rank": rank,
            "by": mapping.get(strategy, "css selector"),
            "wrapped": strategy in wrap_map,
        })

    return ordered


def _build_ai_databank_highlight(locator_property: dict[str, Any]) -> dict[str, float] | None:
    bounds = locator_property.get("bounds") or {}
    viewport = locator_property.get("viewport") or {}
    try:
        viewport_width = float(viewport.get("width") or 0)
        viewport_height = float(viewport.get("height") or 0)
        left = max(0.0, float(bounds.get("left") or 0))
        top = max(0.0, float(bounds.get("top") or 0))
        width = max(0.0, float(bounds.get("width") or 0))
        height = max(0.0, float(bounds.get("height") or 0))
        if viewport_width > 0 and viewport_height > 0 and width > 0 and height > 0:
            clamped_left = min(left, viewport_width)
            clamped_top = min(top, viewport_height)
            clamped_width = min(width, max(0.0, viewport_width - clamped_left))
            clamped_height = min(height, max(0.0, viewport_height - clamped_top))
            if clamped_width > 0 and clamped_height > 0:
                return {
                    "left_pct": (clamped_left / viewport_width) * 100.0,
                    "top_pct": (clamped_top / viewport_height) * 100.0,
                    "width_pct": (clamped_width / viewport_width) * 100.0,
                    "height_pct": (clamped_height / viewport_height) * 100.0,
                }
    except (TypeError, ValueError):
        return None
    return None


def _auto_detect_cdp_port(configured: int = 0) -> int:
    """Return a live CDP port, preferring the configured one when it is reachable."""
    import socket

    def _is_live(port: int) -> bool:
        try:
            with socket.create_connection(("127.0.0.1", port), timeout=0.4):
                return True
        except OSError:
            return False

    if configured > 0 and _is_live(configured):
        return configured
    for probe in (9222, 9223, 9224, 9230):
        try:
            with socket.create_connection(("127.0.0.1", probe), timeout=0.4):
                return probe
        except OSError:
            continue
    return 0


def _scrape_browser_candidates(browser_name: str) -> tuple[list[str], int, str]:
    browser = (browser_name or "chrome").strip().lower()
    if browser == "msedge":
        return ([
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        ], 9223, "Edge")
    return ([
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    ], 9222, "Chrome")


def _open_url_in_debug_browser(port: int, launch_url: str) -> bool:
    if port <= 0 or not launch_url:
        return False
    encoded = requests.utils.quote(launch_url, safe="")
    endpoints = [
        ("PUT", f"http://127.0.0.1:{port}/json/new?{encoded}"),
        ("GET", f"http://127.0.0.1:{port}/json/new?{encoded}"),
    ]
    for method, endpoint in endpoints:
        try:
            resp = requests.request(method, endpoint, timeout=3)
            if resp.ok:
                return True
        except Exception:
            continue
    return False


def _launch_scrape_browser(browser_name: str, launch_url: str, configured_port: int = 0) -> tuple[bool, int, str]:
    if sys.platform != "win32":
        return False, 0, "Launching a scrape browser is only implemented on Windows."

    import ctypes

    candidates, default_port, browser_label = _scrape_browser_candidates(browser_name)
    launch_port = configured_port if configured_port > 0 else default_port
    profile_dir = os.path.join(_RECORDING_LOGS_DIR, f"scrape_browser_{browser_name or 'chrome'}")
    os.makedirs(profile_dir, exist_ok=True)

    for exe in candidates:
        if not os.path.isfile(exe):
            continue
        args = (
            f'--remote-debugging-port={launch_port} '
            f'--user-data-dir="{profile_dir}" '
            f'"{launch_url}"'
        )
        ret = ctypes.windll.shell32.ShellExecuteW(None, "open", exe, args, None, 1)
        if ret > 32:
            for _ in range(40):
                time.sleep(0.25)
                if _auto_detect_cdp_port(launch_port) == launch_port:
                    return True, launch_port, ""
            return False, 0, f"{browser_label} opened, but its remote debugging port {launch_port} did not become reachable."
    return False, 0, f"Could not launch {browser_label}. Check that the browser is installed."


def _ensure_scrape_attach_port(browser_name: str, launch_url: str, configured_port: int = 0) -> tuple[int, str, str]:
    browser = (browser_name or "chrome").strip().lower()
    browser_label = "Edge" if browser == "msedge" else "Chrome"
    attach_port = _auto_detect_cdp_port(configured_port)
    if attach_port:
        return attach_port, "attach", f"Using the existing {browser_label} debug browser."
    if not launch_url:
        return 0, "attach", f"No live {browser_label} debug browser was found, and no launch URL is configured to start one."
    launched, launch_port, err = _launch_scrape_browser(browser, launch_url, configured_port=configured_port)
    if not launched:
        return 0, "attach", err or f"Could not launch {browser_label} in debug mode."
    attach_port = _auto_detect_cdp_port(launch_port)
    if not attach_port:
        return 0, "attach", f"{browser_label} was launched, but its debug port is still not reachable."
    return attach_port, "attach", f"Launched {browser_label} debug browser and opened the requested URL."


# ── Scrape job management ─────────────────────────────────────────────────────
_scrape_jobs: dict = {}
_scrape_jobs_lock = threading.Lock()


def _get_scrape_job(job_id: str):
    with _scrape_jobs_lock:
        return _scrape_jobs.get(job_id)


def _make_scrape_job(job_id: str) -> dict:
    stop_ev = threading.Event()
    pause_ev = threading.Event()
    pause_ev.set()  # not paused initially
    job: dict = {
        "id": job_id,
        "status": "starting",
        "phase": "starting",
        "progress": 0,
        "total": 0,
        "message": "Starting scrape…",
        "saved_rows": 0,
        "page_name": "",
        "page_url": "",
        "attach_port": 0,
        "browser_name": "chrome",
        "item_delay_ms": 0,
        "mode": "attach",
        "launch_url": "",
        "paused": False,
        "_stop_event": stop_ev,
        "_pause_event": pause_ev,
        "_created": time.time(),
    }
    # Evict jobs older than 10 min
    now = time.time()
    with _scrape_jobs_lock:
        old = [k for k, v in _scrape_jobs.items() if now - v.get("_created", now) > 600]
        for k in old:
            del _scrape_jobs[k]
        _scrape_jobs[job_id] = job
    return job


def _run_scrape_job(
    job_id: str,
    attach_port: int,
    db_config: dict,
    webdriver_filename: str,
    browser_name: str = "chrome",
    item_delay_seconds: float = 0.0,
    launch_url: str = "",
    mode: str = "attach",
) -> None:
    job = _get_scrape_job(job_id)
    if not job:
        return
    stop_ev: threading.Event = job["_stop_event"]
    pause_ev: threading.Event = job["_pause_event"]
    job["attach_port"] = attach_port
    job["browser_name"] = browser_name
    job["item_delay_ms"] = int(round(max(0.0, item_delay_seconds) * 1000))
    job["launch_url"] = launch_url or ""
    job["mode"] = mode or "attach"
    if mode == "launch":
        job.update(status="starting", phase="starting", message="Launching browser for scrape…")

    def on_progress(phase: str, count: int, total: int) -> None:
        if phase == "scanning":
            job.update(status="running", phase="scanning", progress=0, total=0,
                       message="Scanning page elements…")
        else:
            pct = round(count / total * 100) if total else 0
            job.update(status="running", phase="saving", progress=count, total=total,
                       pct=pct, done=count,
                       message=f"Saving element {count} of {total}…")

    try:
        from web_scraper import scrape_once
        result = scrape_once(
            url=launch_url,
            attach_port=attach_port,
            db_config=db_config,
            webdriver_filename=webdriver_filename,
            browser_name=browser_name,
            item_delay_seconds=item_delay_seconds,
            on_progress=on_progress,
            stop_event=stop_ev,
            pause_event=pause_ev,
        )
        if stop_ev.is_set():
            job.update(status="stopped", phase="stopped", message="Scrape stopped by user.",
                       pct=job.get("pct", 0), done=job.get("done", 0))
        else:
            saved = result.get("saved_rows", 0)
            page = result.get("page_name") or result.get("page_url") or "current page"
            job.update(
                status="done", phase="done",
                message=f"Scraped {saved} objects from {page}.",
                saved_rows=saved, progress=saved, total=saved,
                pct=100, done=saved,
                page_name=result.get("page_name", ""),
                page_url=result.get("page_url", ""),
            )
    except Exception as exc:
        job.update(status="error", phase="error", message=str(exc))


@login_required
@csrf_exempt
@require_POST
def ai_databank_open_browser(request):
    browser_name = (get_config("recorder.browser", "chrome") or "chrome").strip().lower()
    if browser_name not in ("chrome", "msedge"):
        return JsonResponse({
            "ok": False,
            "error": "AI Databank browser preparation currently supports Chrome or Edge only.",
            "config_url": "/configuration/",
        }, status=400)

    is_edge = browser_name == "msedge"
    browser_label = "Edge" if is_edge else "Chrome"
    rdp_key = "edge.remote_debugging_port" if is_edge else "chrome.remote_debugging_port"
    launch_url = (request.POST.get("launch_url") or get_config("ai_databank.launch_url", "https://demoqa.com/") or "").strip()
    if not launch_url:
        return JsonResponse({
            "ok": False,
            "error": "No launch URL was provided.",
        }, status=400)

    rdp_raw = (get_config(rdp_key) or get_config("chrome.remote_debugging_port") or "").strip()
    configured_port = int(rdp_raw) if rdp_raw.isdigit() and int(rdp_raw) > 0 else 0
    existing_attach_port = _auto_detect_cdp_port(configured_port)
    attach_port, mode, message = _ensure_scrape_attach_port(browser_name, launch_url, configured_port=configured_port)
    if not attach_port:
        return JsonResponse({
            "ok": False,
            "error": message or f"Could not prepare a {browser_label} debug browser.",
            "config_url": "/configuration/",
        }, status=500)

    if existing_attach_port:
        if not _open_url_in_debug_browser(attach_port, launch_url):
            return JsonResponse({
                "ok": False,
                "error": f"Found a live {browser_label} debug browser, but could not open the requested URL in it.",
            }, status=500)
        return JsonResponse({
            "ok": True,
            "browser_name": browser_name,
            "used_attach_port": attach_port,
            "launch_url": launch_url,
            "mode": mode,
            "message": f"Opened URL in the existing {browser_label} debug browser.",
        })

    return JsonResponse({
        "ok": True,
        "browser_name": browser_name,
        "used_attach_port": attach_port,
        "launch_url": launch_url,
        "mode": mode,
        "message": message,
    })


@login_required
@csrf_exempt
@require_POST
def ai_databank_scrape(request):
    _ensure_ai_databank_schema()

    browser_name = (get_config("recorder.browser", "chrome") or "chrome").strip().lower()
    if browser_name not in ("chrome", "msedge"):
        return JsonResponse({
            "ok": False,
            "error": "AI Databank scraping currently supports Chrome or Edge only. Update Recorder -> Browser in Configuration.",
            "config_url": "/configuration/",
        }, status=400)

    is_edge = browser_name == "msedge"
    browser_label = "Edge" if is_edge else "Chrome"
    rdp_key = "edge.remote_debugging_port" if is_edge else "chrome.remote_debugging_port"
    webdriver_key = "edge.webdriver_path" if is_edge else "chrome.webdriver_path"

    rdp_raw = (get_config(rdp_key) or get_config("chrome.remote_debugging_port") or "").strip()
    webdriver_filename = (get_config(webdriver_key) or "").strip()
    delay_raw = (get_config("ai_databank.scrape_item_delay_ms", "100") or "100").strip()
    launch_url = (request.POST.get("launch_url") or get_config("ai_databank.launch_url", "https://demoqa.com/") or "").strip()
    configured_port = int(rdp_raw) if rdp_raw.isdigit() and int(rdp_raw) > 0 else 0
    try:
        item_delay_seconds = max(0.0, float(delay_raw) / 1000.0)
    except (TypeError, ValueError):
        item_delay_seconds = 0.0
    attach_port, mode, attach_message = _ensure_scrape_attach_port(browser_name, launch_url, configured_port=configured_port)

    if not attach_port:
        return JsonResponse({
            "ok": False,
            "error": attach_message or f"Could not prepare a {browser_label} debug browser for attach mode.",
            "config_url": "/configuration/",
        }, status=400)

    db = settings.DATABASES["default"]
    db_config = {
        "dbname": db.get("NAME", "automation_db"),
        "user": db.get("USER", "postgres"),
        "password": db.get("PASSWORD", ""),
        "host": db.get("HOST", "localhost") or "localhost",
        "port": str(db.get("PORT", "5432") or "5432"),
    }

    import secrets as _sec
    job_id = _sec.token_urlsafe(12)
    _make_scrape_job(job_id)
    t = threading.Thread(
        target=_run_scrape_job,
        args=(job_id, attach_port, db_config, webdriver_filename, browser_name, item_delay_seconds, launch_url, mode),
        daemon=True,
    )
    t.start()
    return JsonResponse({
        "ok": True,
        "job_id": job_id,
        "used_attach_port": attach_port,
        "browser_name": browser_name,
        "item_delay_ms": int(round(item_delay_seconds * 1000)),
        "mode": mode,
        "launch_url": launch_url,
        "message": attach_message,
    })


@login_required
def ai_scrape_status(request, job_id):
    """Poll endpoint — returns current state of a scrape job."""
    job = _get_scrape_job(job_id)
    if not job:
        return JsonResponse({"ok": False, "error": "Job not found."}, status=404)
    safe = {k: v for k, v in job.items() if not k.startswith("_")}
    safe["ok"] = True
    # Ensure monitor-expected fields exist
    progress = safe.get("progress", 0)
    total    = safe.get("total", 0)
    safe.setdefault("done", progress)
    safe.setdefault("pct",  round(progress / total * 100) if total else 0)
    safe.setdefault("attach_port", 0)
    safe.setdefault("browser_name", "chrome")
    safe.setdefault("item_delay_ms", 0)
    safe.setdefault("mode", "attach")
    safe.setdefault("launch_url", "")
    return JsonResponse(safe)


@login_required
@csrf_exempt
@require_POST
def ai_scrape_stop(request, job_id):
    """Stop a running scrape job."""
    job = _get_scrape_job(job_id)
    if not job:
        return JsonResponse({"ok": False, "error": "Job not found."}, status=404)
    job["_stop_event"].set()
    job["_pause_event"].set()  # unblock any pause so the thread can exit
    return JsonResponse({"ok": True})


@login_required
@csrf_exempt
@require_POST
def ai_scrape_pause(request, job_id):
    """Toggle pause/resume on a running scrape job."""
    job = _get_scrape_job(job_id)
    if not job:
        return JsonResponse({"ok": False, "error": "Job not found."}, status=404)
    if job.get("paused"):
        job["_pause_event"].set()   # resume
        job["paused"] = False
        job["status"] = "running"
        action = "resumed"
    else:
        job["_pause_event"].clear()  # pause
        job["paused"] = True
        job["status"] = "paused"
        action = "paused"
    return JsonResponse({"ok": True, "action": action, "paused": job["paused"]})


@login_required
@csrf_exempt
@require_POST
def ai_workflow_save(request):
    _ensure_ai_workflow_schema()

    try:
        payload = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid workflow payload."}, status=400)

    workflow_name = str(payload.get("workflow_name") or "").strip()
    if not workflow_name:
        return JsonResponse({"ok": False, "error": "Workflow name is required."}, status=400)

    page_connections = payload.get("page_connections")
    page_sequence = payload.get("page_sequence")
    workflow_payload = payload.get("workflow_payload")

    if not isinstance(page_connections, list):
        page_connections = []
    if not isinstance(page_sequence, list):
        page_sequence = []
    if not isinstance(workflow_payload, dict):
        workflow_payload = {}

    with connection.cursor() as cur:
        cur.execute(
            """
            INSERT INTO ai_workflow (workflow_name, page_connections, page_sequence, workflow_payload, created_at, updated_at)
            VALUES (%s, %s::jsonb, %s::jsonb, %s::jsonb, NOW(), NOW())
            ON CONFLICT (workflow_name)
            DO UPDATE SET
                page_connections = EXCLUDED.page_connections,
                page_sequence = EXCLUDED.page_sequence,
                workflow_payload = EXCLUDED.workflow_payload,
                updated_at = NOW()
            RETURNING id, created_at, updated_at
            """,
            [
                workflow_name,
                json.dumps(page_connections),
                json.dumps(page_sequence),
                json.dumps(workflow_payload),
            ],
        )
        row = cur.fetchone()

    return JsonResponse({
        "ok": True,
        "workflow_id": row[0] if row else None,
        "workflow_name": workflow_name,
        "created_at": row[1].isoformat() if row and row[1] else None,
        "updated_at": row[2].isoformat() if row and row[2] else None,
        "page_connections_count": len(page_connections),
        "page_sequence_count": len(page_sequence),
    })


@login_required
@csrf_exempt
@require_POST
def ai_workflow_update_from_record(request):
    _ensure_ai_databank_schema()

    try:
        payload = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid JSON payload."}, status=400)

    record_id = str(payload.get("record_id") or "").strip()

    # ── Sync ALL sessions when record_id is blank ──────────────
    if not record_id:
        try:
            with connection.cursor() as cur:
                sessions = _get_ai_workflow_source_sessions(cur)
                if not sessions:
                    return JsonResponse({"ok": False, "error": "No recorded sessions found."}, status=404)

                results = []
                for session in sessions:
                    try:
                        workflow = _build_ai_workflow_from_record(cur, session["record_id"])
                        results.append({
                            "workflow_name": workflow["workflow_name"],
                            "record_id": workflow["record_id"],
                            "record_name": workflow["record_name"],
                            "step_count": workflow["step_count"],
                            "page_count": len(workflow["page_sequence"]),
                            "workflow_id": None,
                        })
                    except Exception:
                        # Skip failing sessions but continue with others
                        continue

        except Exception as exc:
            return JsonResponse({"ok": False, "error": str(exc)}, status=500)

        return JsonResponse({
            "ok": True,
            "results": results,
            "synced_count": len(results),
            "total_sessions": len(sessions),
        })

    # ── Single session sync ────────────────────────────────────

    try:
        uuid.UUID(record_id)
    except Exception:
        return JsonResponse({"ok": False, "error": "record_id must be a valid UUID."}, status=400)

    try:
        with connection.cursor() as cur:
            workflow = _build_ai_workflow_from_record(cur, record_id)
    except WorkflowGenerationError as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=400)
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({
        "ok": True,
        "workflow_id": None,
        "workflow_name": workflow["workflow_name"],
        "record_id": workflow["record_id"],
        "record_name": workflow["record_name"],
        "step_count": workflow["step_count"],
        "page_count": len(workflow["page_sequence"]),
        "connection_count": len(workflow["page_connections"]),
        "updated_at": None,
    })


@login_required
def ai_workflow_load(request):
    _ensure_ai_workflow_schema()

    workflow_name = str(request.GET.get("workflow_name") or "").strip()
    if not workflow_name:
        return JsonResponse({"ok": False, "error": "workflow_name is required."}, status=400)

    with connection.cursor() as cur:
        workflow = _get_ai_workflow_by_name(cur, workflow_name)

    if not workflow:
        return JsonResponse({"ok": False, "error": "Workflow not found."}, status=404)

    return JsonResponse({"ok": True, "workflow": workflow})


@login_required
@csrf_exempt
@require_POST
def ai_workflow_delete(request):
    _ensure_ai_workflow_schema()

    try:
        payload = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid workflow payload."}, status=400)

    workflow_name = str(payload.get("workflow_name") or "").strip()
    if not workflow_name:
        return JsonResponse({"ok": False, "error": "Workflow name is required."}, status=400)

    with connection.cursor() as cur:
        cur.execute(
            "DELETE FROM ai_workflow WHERE workflow_name = %s RETURNING id",
            [workflow_name],
        )
        row = cur.fetchone()

    if not row:
        return JsonResponse({"ok": False, "error": "Workflow not found."}, status=404)

    return JsonResponse({
        "ok": True,
        "workflow_id": row[0],
        "workflow_name": workflow_name,
    })


@login_required
def ai_workflow_source_sessions(request):
    _ensure_ai_databank_schema()

    with connection.cursor() as cur:
        sessions = _get_ai_workflow_source_sessions(cur)

    return JsonResponse({
        "ok": True,
        "sessions": sessions,
        "count": len(sessions),
    })


_RESERVED_PROJECT_FOLDER_KEYS = {"", "baseline", "unfiled"}


def _ensure_ai_databank_schema() -> None:
    global _AI_DATABANK_SCHEMA_ENSURED
    if _AI_DATABANK_SCHEMA_ENSURED:
        return

    with connection.cursor() as cur:
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS ai_databank (
                id BIGSERIAL PRIMARY KEY,
                page_url TEXT NOT NULL,
                page_name TEXT NOT NULL DEFAULT '',
                element_type VARCHAR(80) NOT NULL DEFAULT 'element',
                element_fingerprint TEXT NOT NULL DEFAULT '',
                locator_property JSONB NOT NULL DEFAULT '{}'::jsonb,
                screenshot_png BYTEA,
                created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
            )
            """
        )
        cur.execute("ALTER TABLE ai_databank ADD COLUMN IF NOT EXISTS element_fingerprint TEXT")
        cur.execute("ALTER TABLE ai_databank ADD COLUMN IF NOT EXISTS screenshot_png BYTEA")
        cur.execute("ALTER TABLE ai_databank ADD COLUMN IF NOT EXISTS updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()")
        cur.execute(
            "CREATE INDEX IF NOT EXISTS ai_databank_page_url_idx ON ai_databank (page_url)"
        )
        cur.execute(
            "CREATE INDEX IF NOT EXISTS ai_databank_created_at_idx ON ai_databank (created_at DESC)"
        )
        cur.execute(
            """
            CREATE UNIQUE INDEX IF NOT EXISTS ai_databank_page_fingerprint_uniq
            ON ai_databank (page_url, element_fingerprint)
            WHERE element_fingerprint IS NOT NULL AND element_fingerprint <> ''
            """
        )
        # Ensure locators_stat exists and has the tenant_id column the Django model expects.
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS locators_stat (
                id           BIGSERIAL PRIMARY KEY,
                run_id       UUID,
                record_id    UUID       NOT NULL,
                step_no      INTEGER    NOT NULL,
                strategy     TEXT       NOT NULL,
                locator      TEXT       NOT NULL,
                is_primary   BOOLEAN    NOT NULL DEFAULT FALSE,
                locator_rank INTEGER,
                pos_x        FLOAT,
                pos_y        FLOAT,
                action       TEXT,
                page_url     TEXT,
                runner       TEXT,
                author       TEXT,
                folder_name  TEXT,
                created_at   TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                tenant_id    UUID
            )
            """
        )
        cur.execute("ALTER TABLE locators_stat ADD COLUMN IF NOT EXISTS tenant_id UUID")

    _AI_DATABANK_SCHEMA_ENSURED = True


def _ensure_ai_workflow_schema() -> None:
    global _AI_WORKFLOW_SCHEMA_ENSURED
    if _AI_WORKFLOW_SCHEMA_ENSURED:
        return

    with connection.cursor() as cur:
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS ai_workflow (
                id BIGSERIAL PRIMARY KEY,
                workflow_name TEXT NOT NULL,
                page_connections JSONB NOT NULL DEFAULT '[]'::jsonb,
                page_sequence JSONB NOT NULL DEFAULT '[]'::jsonb,
                workflow_payload JSONB NOT NULL DEFAULT '{}'::jsonb,
                created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
            )
            """
        )
        cur.execute("ALTER TABLE ai_workflow ADD COLUMN IF NOT EXISTS page_connections JSONB NOT NULL DEFAULT '[]'::jsonb")
        cur.execute("ALTER TABLE ai_workflow ADD COLUMN IF NOT EXISTS page_sequence JSONB NOT NULL DEFAULT '[]'::jsonb")
        cur.execute("ALTER TABLE ai_workflow ADD COLUMN IF NOT EXISTS workflow_payload JSONB NOT NULL DEFAULT '{}'::jsonb")
        cur.execute("ALTER TABLE ai_workflow ADD COLUMN IF NOT EXISTS updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()")
        cur.execute(
            "CREATE UNIQUE INDEX IF NOT EXISTS ai_workflow_workflow_name_uniq ON ai_workflow (workflow_name)"
        )
        cur.execute(
            "CREATE INDEX IF NOT EXISTS ai_workflow_updated_at_idx ON ai_workflow (updated_at DESC)"
        )

    _AI_WORKFLOW_SCHEMA_ENSURED = True


def _normalize_folder_path(folder_name: str) -> str:
    raw = (folder_name or "").replace("\\", "/")
    parts = [part.strip() for part in raw.split("/") if part.strip()]
    return "/".join(parts)


def _folder_ancestors(folder_name: str) -> list[str]:
    normalized = _normalize_folder_path(folder_name)
    if not normalized:
        return []
    parts = normalized.split("/")
    return ["/".join(parts[:idx]) for idx in range(1, len(parts) + 1)]


def _get_recordings_folder_label() -> str:
    folder_name = _normalize_folder_path(get_config("projects.recordings_folder_label", "Baseline"))
    if not folder_name or folder_name.lower() in  {"", "baseline", "unfiled"}:
        return "Baseline"
    return folder_name


_recordings_aliases_cache: "frozenset[str] | None" = None


def _get_recordings_aliases_from_db() -> frozenset:
    """Return a frozenset of lowercased parent_folder names where is_baseline=True.
    Result is cached for the lifetime of the server process."""
    global _recordings_aliases_cache
    if _recordings_aliases_cache is None:
        try:
            with connection.cursor() as _cur:
                _cur.execute(
                    "SELECT LOWER(parent_folder) FROM parent_folders WHERE is_baseline = TRUE"
                )
                _recordings_aliases_cache = frozenset(row[0] for row in _cur.fetchall())
        except Exception:
            return frozenset()
    return _recordings_aliases_cache


def _clear_recordings_aliases_cache() -> None:
    global _recordings_aliases_cache
    _recordings_aliases_cache = None


def _recordings_sql_aliases() -> frozenset:
    """Return the full set of lowercased folder names that map to the Baseline/Recordings root."""
    return {"baseline", "", _get_recordings_folder_label().lower()}


def _is_recordings_folder_name(folder_name: str) -> bool:
    normalized = _normalize_folder_path(folder_name)
    if not normalized:
        return False
    lower = normalized.lower()
    return lower in {
        "baseline",
        _get_recordings_folder_label().lower(),
    }


def _project_folders_are_public() -> bool:
    return get_config("parent_folders.public", "true") == "true"


def _user_can_view_public_project_folders(user) -> bool:
    if not getattr(user, "is_authenticated", False) or getattr(user, "is_superuser", False):
        return True
    pref_value = get_user_pref(user.id, "projects.view_public_folders", "true")
    return str(pref_value).strip().lower() == "true"


def _user_visible_project_roots(user) -> set[str] | None:
    if not getattr(user, "is_authenticated", False) or getattr(user, "is_superuser", False):
        return None
    if _project_folders_are_public():
        return None

    username = (getattr(user, "username", "") or "").strip().lower()
    if not username:
        return set()

    can_view_public = _user_can_view_public_project_folders(user)

    with connection.cursor() as cur:
        if can_view_public:
            cur.execute(
                """
                SELECT DISTINCT TRIM(COALESCE(parent_folder, ''))
                FROM parent_folders
                WHERE (
                        LOWER(TRIM(COALESCE(author, ''))) = %s
                     OR COALESCE(public, FALSE) = TRUE
                      )
                  AND TRIM(COALESCE(parent_folder, '')) <> ''
                """,
                [username],
            )
        else:
            cur.execute(
                """
                SELECT DISTINCT TRIM(COALESCE(parent_folder, ''))
                FROM parent_folders
                WHERE LOWER(TRIM(COALESCE(author, ''))) = %s
                  AND TRIM(COALESCE(parent_folder, '')) <> ''
                """,
                [username],
            )
        return {str(row[0]).strip() for row in cur.fetchall() if str(row[0]).strip()}


def _is_project_path_visible_to_user(user, folder_name: str) -> bool:
    normalized = _normalize_folder_path(folder_name)
    if not normalized or _is_recordings_folder_name(normalized):
        return True

    visible_roots = _user_visible_project_roots(user)
    if visible_roots is None:
        return True

    root_name = normalized.split("/", 1)[0].strip()
    return root_name in visible_roots


def _ensure_project_folders_table() -> None:
    with connection.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS project_folders (
                folder_name TEXT PRIMARY KEY,
                folder_order INTEGER,
                created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW()
            )
        """)
        cur.execute("ALTER TABLE project_folders ADD COLUMN IF NOT EXISTS folder_order INTEGER")
        cur.execute("ALTER TABLE project_folders ALTER COLUMN folder_order SET DEFAULT 1")
        cur.execute("""
            WITH ordered AS (
                SELECT
                    folder_name,
                    ROW_NUMBER() OVER (ORDER BY created_at, LOWER(folder_name), folder_name) AS new_folder_order
                FROM project_folders
                WHERE TRIM(COALESCE(folder_name, '')) <> ''
                  AND (folder_order IS NULL OR folder_order < 1)
            )
            UPDATE project_folders pf
            SET folder_order = ordered.new_folder_order
            FROM ordered
            WHERE pf.folder_name = ordered.folder_name
        """)


def _list_registered_project_folders() -> list[dict[str, Any]]:
    _ensure_project_folders_table()
    with connection.cursor() as cur:
        cur.execute("""
            SELECT folder_name, COALESCE(folder_order, 2147483647) AS folder_order
            FROM project_folders
            WHERE TRIM(COALESCE(folder_name, '')) <> ''
            ORDER BY folder_order, LOWER(folder_name), folder_name
        """)
        return [
            {"folder_name": row[0], "folder_order": int(row[1]) if row[1] is not None else None}
            for row in cur.fetchall()
        ]


def _list_parent_folder_order_map() -> dict[str, int]:
    with connection.cursor() as cur:
        cur.execute(
            """
            SELECT TRIM(COALESCE(parent_folder, '')),
                   COALESCE(parent_folder_order, 2147483647)
            FROM parent_folders
            WHERE TRIM(COALESCE(parent_folder, '')) <> ''
            ORDER BY parent_folder_order, LOWER(parent_folder), parent_folder
            """
        )
        return {
            str(row[0]).strip(): int(row[1])
            for row in cur.fetchall()
            if str(row[0]).strip()
        }


def _list_sub_folder_order_map() -> dict[str, int]:
    with connection.cursor() as cur:
        cur.execute(
            """
            SELECT TRIM(COALESCE(pf.parent_folder, '')) || '/' || TRIM(COALESCE(sf.sub_folder, '')),
                   COALESCE(sf.sub_folder_order, 2147483647)
            FROM sub_folders sf
            JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
            WHERE TRIM(COALESCE(pf.parent_folder, '')) <> ''
              AND TRIM(COALESCE(sf.sub_folder, '')) <> ''
            ORDER BY sf.sub_folder_order, LOWER(sf.sub_folder), sf.sub_folder
            """
        )
        return {
            _normalize_folder_path(str(row[0])): int(row[1])
            for row in cur.fetchall()
            if _normalize_folder_path(str(row[0]))
        }


def _get_project_folder_metadata_map(folder_names: list[str]) -> dict[str, dict[str, Any]]:
    normalized_names = {
        _normalize_folder_path(folder_name)
        for folder_name in folder_names
        if _normalize_folder_path(folder_name)
    }
    if not normalized_names:
        return {}

    meta_map: dict[str, dict[str, Any]] = {}
    with connection.cursor() as cur:
        cur.execute(
            """
            SELECT TRIM(COALESCE(parent_folder, '')), COALESCE(author, ''), COALESCE(public, FALSE)
            FROM parent_folders
            WHERE TRIM(COALESCE(parent_folder, '')) <> ''
            """
        )
        for folder_name, author, public in cur.fetchall():
            normalized = _normalize_folder_path(folder_name)
            if normalized in normalized_names:
                meta_map[normalized] = {"author": author or "", "public": bool(public)}

        cur.execute(
            """
            SELECT
                TRIM(COALESCE(pf.parent_folder, '')) || '/' || TRIM(COALESCE(sf.sub_folder, '')),
                COALESCE(sf.author, ''),
                COALESCE(sf.public, FALSE)
            FROM sub_folders sf
            JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
            WHERE TRIM(COALESCE(pf.parent_folder, '')) <> ''
              AND TRIM(COALESCE(sf.sub_folder, '')) <> ''
            """
        )
        for folder_name, author, public in cur.fetchall():
            normalized = _normalize_folder_path(folder_name)
            if normalized in normalized_names:
                meta_map[normalized] = {"author": author or "", "public": bool(public)}

        cur.execute(
            """
            SELECT
                TRIM(COALESCE(pf.parent_folder, '')) || '/' ||
                TRIM(COALESCE(sf.sub_folder, '')) || '/' ||
                TRIM(COALESCE(ef.end_folder, '')),
                COALESCE(ef.author, ''),
                COALESCE(ef.public, FALSE)
            FROM end_folders ef
            JOIN sub_folders sf ON sf.sub_folder_id = ef.end_folder_parent
            JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
            WHERE TRIM(COALESCE(pf.parent_folder, '')) <> ''
              AND TRIM(COALESCE(sf.sub_folder, '')) <> ''
              AND TRIM(COALESCE(ef.end_folder, '')) <> ''
            """
        )
        for folder_name, author, public in cur.fetchall():
            normalized = _normalize_folder_path(folder_name)
            if normalized in normalized_names:
                meta_map[normalized] = {"author": author or "", "public": bool(public)}

    return meta_map


def _annotate_project_tree_metadata(nodes: list[dict], user, folder_meta_map: dict[str, dict[str, Any]]) -> None:
    username = (getattr(user, "username", "") or "").strip().lower()

    def _walk(node: dict) -> None:
        if node.get("is_special"):
            node["author"] = ""
            node["public"] = False
            node["can_toggle_public"] = False
        else:
            meta = folder_meta_map.get(node.get("path", ""), {})
            author = (meta.get("author") or "").strip()
            node["author"] = author
            node["public"] = bool(meta.get("public"))
            node["can_toggle_public"] = bool(
                getattr(user, "is_superuser", False)
                or (username and author and author.lower() == username)
            )
        for child in node.get("children", []):
            _walk(child)

    for node in nodes:
        _walk(node)


def _next_project_folder_order(cursor) -> int:
    _ensure_project_folders_table()
    cursor.execute("SELECT COALESCE(MAX(folder_order), 0) + 1 FROM project_folders")
    row = cursor.fetchone()
    return int(row[0] or 1)


def _next_parent_folder_order(cursor) -> int:
    cursor.execute("SELECT COALESCE(MAX(parent_folder_order), 0) + 1 FROM parent_folders")
    row = cursor.fetchone()
    return int(row[0] or 1)


def _next_sub_folder_order(cursor, parent_folder_id: str) -> int:
    cursor.execute(
        "SELECT COALESCE(MAX(sub_folder_order), 0) + 1 FROM sub_folders WHERE sub_folder_parent = %s",
        [parent_folder_id],
    )
    row = cursor.fetchone()
    return int(row[0] or 1)


def _resequence_project_folder_orders(cursor) -> None:
    _ensure_project_folders_table()
    cursor.execute("""
        WITH ordered AS (
            SELECT
                folder_name,
                ROW_NUMBER() OVER (
                    ORDER BY COALESCE(folder_order, 2147483647), created_at, LOWER(folder_name), folder_name
                ) AS new_folder_order
            FROM project_folders
            WHERE TRIM(COALESCE(folder_name, '')) <> ''
        )
        UPDATE project_folders pf
        SET folder_order = ordered.new_folder_order
        FROM ordered
        WHERE pf.folder_name = ordered.folder_name
    """)


def _project_folder_names_in_order(cursor) -> list[str]:
    _ensure_project_folders_table()
    cursor.execute("""
        SELECT folder_name
        FROM project_folders
        WHERE TRIM(COALESCE(folder_name, '')) <> ''
        ORDER BY COALESCE(folder_order, 2147483647), LOWER(folder_name), folder_name
    """)
    return [row[0] for row in cursor.fetchall()]


def _apply_project_folder_global_order(cursor, ordered_names: list[str]) -> int:
    current_order = _project_folder_names_in_order(cursor)
    current_set = set(current_order)
    seen: set[str] = set()
    sanitized: list[str] = []

    for folder_name in ordered_names:
        normalized = _normalize_folder_path(folder_name)
        if not normalized or normalized in seen or normalized not in current_set:
            continue
        seen.add(normalized)
        sanitized.append(normalized)

    if not sanitized:
        return 0

    sanitized.extend(name for name in current_order if name not in seen)

    for order, folder_name in enumerate(sanitized, start=1):
        cursor.execute(
            "UPDATE project_folders SET folder_order = %s WHERE folder_name = %s",
            [order, folder_name],
        )

    return len(sanitized)


def _folder_parent_path(folder_name: str) -> str:
    normalized = _normalize_folder_path(folder_name)
    return normalized.rsplit("/", 1)[0] if "/" in normalized else ""


def _apply_project_folder_sibling_order(cursor, parent_folder: str, ordered_names: list[str]) -> int:
    current_order = _project_folder_names_in_order(cursor)
    sibling_names = [name for name in current_order if _folder_parent_path(name) == parent_folder]
    if not sibling_names:
        return 0

    sibling_set = set(sibling_names)
    seen: set[str] = set()
    sanitized: list[str] = []
    for folder_name in ordered_names:
        normalized = _normalize_folder_path(folder_name)
        if not normalized or normalized in seen or normalized not in sibling_set:
            continue
        seen.add(normalized)
        sanitized.append(normalized)

    if not sanitized:
        return 0

    sanitized.extend(name for name in sibling_names if name not in seen)
    first_index = min(current_order.index(name) for name in sibling_names)
    remaining = [name for name in current_order if name not in sibling_set]
    new_order = remaining[:first_index] + sanitized + remaining[first_index:]

    for order, folder_name in enumerate(new_order, start=1):
        cursor.execute(
            "UPDATE project_folders SET folder_order = %s WHERE folder_name = %s",
            [order, folder_name],
        )

    return len(sanitized)


def _project_folder_sort_key(folder_name: str, folder_order_map: dict[str, int | None]) -> tuple[int, str, str]:
    normalized = _normalize_folder_path(folder_name)
    folder_order = folder_order_map.get(normalized)
    return (
        folder_order if isinstance(folder_order, int) and folder_order > 0 else 2147483647,
        normalized.lower(),
        normalized,
    )


def _project_folder_exists(cursor, folder_name: str) -> bool:
    _ensure_project_folders_table()
    normalized = _normalize_folder_path(folder_name)
    if not normalized:
        return False
    if _is_recordings_folder_name(normalized):
        return True
    cursor.execute("""
        SELECT 1
        FROM (
            SELECT DISTINCT TRIM(COALESCE(folder_name, '')) AS folder_name
            FROM steps
            WHERE TRIM(COALESCE(folder_name, '')) <> ''
            UNION
            SELECT DISTINCT folder_name
            FROM project_folders
            WHERE TRIM(COALESCE(folder_name, '')) <> ''
        ) folders
        WHERE LOWER(folder_name) = %s
        LIMIT 1
    """, [normalized.lower()])
    return cursor.fetchone() is not None


def _register_project_folder(cursor, folder_name: str, insert_after_folder: str = "") -> None:
    _ensure_project_folders_table()
    normalized_folder = _normalize_folder_path(folder_name)
    if not normalized_folder:
        return

    paths = _folder_ancestors(normalized_folder)
    for path in paths[:-1]:
        if path.lower() in _RESERVED_PROJECT_FOLDER_KEYS or path.lower() == _get_recordings_folder_label().lower():
            continue
        cursor.execute("SELECT 1 FROM project_folders WHERE folder_name = %s LIMIT 1", [path])
        if cursor.fetchone() is not None:
            continue
        cursor.execute("""
            INSERT INTO project_folders (folder_name, folder_order)
            VALUES (%s, %s)
        """, [path, _next_project_folder_order(cursor)])

    final_path = paths[-1]
    if final_path.lower() in _RESERVED_PROJECT_FOLDER_KEYS or final_path.lower() == _get_recordings_folder_label().lower():
        return

    cursor.execute("SELECT 1 FROM project_folders WHERE folder_name = %s LIMIT 1", [final_path])
    if cursor.fetchone() is not None:
        return

    cursor.execute("""
        INSERT INTO project_folders (folder_name, folder_order)
        VALUES (%s, %s)
    """, [final_path, _next_project_folder_order(cursor)])

    # Per-container 1-based ordering: only reorder siblings at the same parent level
    parent_folder = _folder_parent_path(final_path)
    all_names = _project_folder_names_in_order(cursor)
    sibling_names = [name for name in all_names if _folder_parent_path(name) == parent_folder]

    insert_after = _normalize_folder_path(insert_after_folder)
    existing = [name for name in sibling_names if name != final_path]
    if insert_after and insert_after in existing:
        idx = existing.index(insert_after) + 1
        ordered_siblings = existing[:idx] + [final_path] + existing[idx:]
    else:
        ordered_siblings = existing + [final_path]

    for i, name in enumerate(ordered_siblings, start=1):
        cursor.execute(
            "UPDATE project_folders SET folder_order = %s WHERE folder_name = %s",
            [i, name],
        )


def _rename_folder_path_in_table(cursor, table_name: str, column_name: str, old_folder: str, new_folder: str) -> None:
    like_pattern = old_folder + "/%"
    cursor.execute(f"""
        UPDATE {table_name}
        SET {column_name} = CASE
            WHEN TRIM(COALESCE({column_name}, '')) = %s THEN %s
            ELSE %s || SUBSTRING(TRIM(COALESCE({column_name}, '')) FROM %s)
        END
        WHERE TRIM(COALESCE({column_name}, '')) = %s
           OR TRIM(COALESCE({column_name}, '')) LIKE %s
    """, [old_folder, new_folder, new_folder, len(old_folder) + 1, old_folder, like_pattern])


def _delete_folder_path_from_registry(cursor, folder_name: str) -> None:
    _ensure_project_folders_table()
    like_pattern = folder_name + "/%"
    cursor.execute("""
        DELETE FROM project_folders
        WHERE TRIM(COALESCE(folder_name, '')) = %s
           OR TRIM(COALESCE(folder_name, '')) LIKE %s
    """, [folder_name, like_pattern])


def _delete_folder_path_from_hierarchy(cursor, folder_name: str) -> None:
        parts = [part.strip() for part in _normalize_folder_path(folder_name).split("/") if part.strip()]
        if not parts:
                return

        if len(parts) == 1:
                cursor.execute(
                        """
                        DELETE FROM end_folders ef
                        USING sub_folders sf, parent_folders pf
                        WHERE ef.end_folder_parent = sf.sub_folder_id
                            AND sf.sub_folder_parent = pf.parent_folder_id
                            AND pf.parent_folder = %s
                        """,
                        [parts[0]],
                )
                cursor.execute(
                        """
                        DELETE FROM sub_folders sf
                        USING parent_folders pf
                        WHERE sf.sub_folder_parent = pf.parent_folder_id
                            AND pf.parent_folder = %s
                        """,
                        [parts[0]],
                )
                cursor.execute("DELETE FROM parent_folders WHERE parent_folder = %s", [parts[0]])
                # Resequence parent_folder_order for the remaining rows so
                # there are no gaps after the deletion.
                cursor.execute(
                        """
                        WITH ordered AS (
                            SELECT parent_folder_id,
                                   ROW_NUMBER() OVER (
                                       ORDER BY COALESCE(parent_folder_order, 2147483647),
                                                LOWER(parent_folder), parent_folder
                                   ) AS new_order
                            FROM parent_folders
                        )
                        UPDATE parent_folders pf
                        SET parent_folder_order = ordered.new_order
                        FROM ordered
                        WHERE pf.parent_folder_id = ordered.parent_folder_id
                        """,
                )
                return

        if len(parts) == 2:
                cursor.execute(
                        """
                        DELETE FROM end_folders ef
                        USING sub_folders sf, parent_folders pf
                        WHERE ef.end_folder_parent = sf.sub_folder_id
                            AND sf.sub_folder_parent = pf.parent_folder_id
                            AND pf.parent_folder = %s
                            AND sf.sub_folder = %s
                        """,
                        [parts[0], parts[1]],
                )
                cursor.execute(
                        """
                        DELETE FROM sub_folders sf
                        USING parent_folders pf
                        WHERE sf.sub_folder_parent = pf.parent_folder_id
                            AND pf.parent_folder = %s
                            AND sf.sub_folder = %s
                        """,
                        [parts[0], parts[1]],
                )
                # Resequence sub_folder_order for the remaining siblings so
                # there are no gaps after the deletion.
                cursor.execute(
                        """
                        WITH ordered AS (
                            SELECT sf.sub_folder_id,
                                   ROW_NUMBER() OVER (
                                       ORDER BY COALESCE(sf.sub_folder_order, 2147483647),
                                                LOWER(sf.sub_folder), sf.sub_folder
                                   ) AS new_order
                            FROM sub_folders sf
                            JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
                            WHERE pf.parent_folder = %s
                        )
                        UPDATE sub_folders sf
                        SET sub_folder_order = ordered.new_order
                        FROM ordered
                        WHERE sf.sub_folder_id = ordered.sub_folder_id
                        """,
                        [parts[0]],
                )
                return

        cursor.execute(
                """
                DELETE FROM end_folders ef
                USING sub_folders sf, parent_folders pf
                WHERE ef.end_folder_parent = sf.sub_folder_id
                    AND sf.sub_folder_parent = pf.parent_folder_id
                    AND pf.parent_folder = %s
                    AND sf.sub_folder = %s
                    AND ef.end_folder = %s
                """,
                [parts[0], parts[1], parts[2]],
        )


def _rename_folder_path_in_hierarchy(cursor, old_folder: str, new_folder: str) -> None:
    old_parts = [part.strip() for part in _normalize_folder_path(old_folder).split("/") if part.strip()]
    new_parts = [part.strip() for part in _normalize_folder_path(new_folder).split("/") if part.strip()]
    if not old_parts or not new_parts or len(old_parts) != len(new_parts):
        return

    if len(old_parts) == 1:
        cursor.execute(
            "UPDATE parent_folders SET parent_folder = %s WHERE parent_folder = %s",
            [new_parts[0], old_parts[0]],
        )
        return

    if len(old_parts) == 2:
        cursor.execute(
            """
            UPDATE sub_folders sf
            SET sub_folder = %s
            FROM parent_folders pf
            WHERE sf.sub_folder_parent = pf.parent_folder_id
                AND pf.parent_folder = %s
                AND sf.sub_folder = %s
            """,
            [new_parts[1], old_parts[0], old_parts[1]],
        )
        return

    if len(old_parts) == 3:
        cursor.execute(
            """
            UPDATE end_folders ef
            SET end_folder = %s
            FROM sub_folders sf
            JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
            WHERE ef.end_folder_parent = sf.sub_folder_id
                AND pf.parent_folder = %s
                AND sf.sub_folder = %s
                AND ef.end_folder = %s
            """,
            [new_parts[2], old_parts[0], old_parts[1], old_parts[2]],
        )


def _folder_dom_id(prefix: str, folder_name: str) -> str:
    normalized = _normalize_folder_path(folder_name)
    token = uuid.uuid5(uuid.NAMESPACE_URL, f"{prefix}:{normalized or 'root'}").hex[:10]
    return f"{prefix}-{token}"


def _build_special_folder_group(folder_name: str, sessions: list[dict]) -> dict:
    label = folder_name or "Unfiled"
    return {
        "path": folder_name,
        "name": label,
        "display_name": label,
        "sessions": sessions,
        "children": [],
        "total_sessions": len(sessions),
        "dom_id": _folder_dom_id("folder-group", folder_name or "unfiled"),
        "collapse_id": _folder_dom_id("folder-body", folder_name or "unfiled"),
        "is_special": True,
        "is_recordings": _is_recordings_folder_name(folder_name),
        "is_unfiled": folder_name == "",
    }


def _build_folder_tree(
    folder_map: dict[str, list[dict]],
    registered_folders: list[str],
    folder_order_map: dict[str, int | None],
) -> list[dict]:
    # Expand every path to include all intermediate ancestor paths so that
    # parent nodes (e.g. "Project001") are always created even when only
    # sub-folders (e.g. "Project001/Sub001") have sessions/runs.
    _raw_paths = set(list(folder_map.keys()) + registered_folders)
    _expanded: set[str] = set()
    for _p in _raw_paths:
        for _ancestor in _folder_ancestors(_p):
            _expanded.add(_ancestor)
    folder_paths = sorted(
        _expanded,
        key=lambda value: _project_folder_sort_key(value, folder_order_map),
    )
    nodes: dict[str, dict] = {}
    roots: list[dict] = []

    for folder_path in folder_paths:
        normalized = _normalize_folder_path(folder_path)
        if not normalized:
            continue
        nodes[normalized] = {
            "path": normalized,
            "name": normalized.rsplit("/", 1)[-1],
            "display_name": normalized.rsplit("/", 1)[-1],
            "sessions": folder_map.get(normalized, []),
            "children": [],
            "total_sessions": 0,
            "folder_order": folder_order_map.get(normalized),
            "dom_id": _folder_dom_id("folder-group", normalized),
            "collapse_id": _folder_dom_id("folder-body", normalized),
            "is_special": False,
            "is_recordings": False,
            "is_unfiled": False,
        }

    for folder_path in folder_paths:
        normalized = _normalize_folder_path(folder_path)
        if not normalized:
            continue
        parent_path = normalized.rsplit("/", 1)[0] if "/" in normalized else ""
        if parent_path and parent_path in nodes:
            nodes[parent_path]["children"].append(nodes[normalized])
        else:
            roots.append(nodes[normalized])

    def _finalize(node: dict) -> int:
        node["children"].sort(key=lambda child: _project_folder_sort_key(child["path"], folder_order_map))
        total = len(node["sessions"])
        for child in node["children"]:
            total += _finalize(child)
        node["total_sessions"] = total
        return total

    roots.sort(key=lambda node: _project_folder_sort_key(node["path"], folder_order_map))
    for root in roots:
        _finalize(root)
    return roots


def _resequence_file_order_for_folder(folder_name: str) -> None:
    """Compact file_order to 1..N for the affected folder/group."""
    folder_key = (folder_name or "").strip().lower()

    with connection.cursor() as cur:
        if not folder_key or _is_recordings_folder_name(folder_name):
            _rff_aliases = list(_recordings_sql_aliases())
            _rff_ph = ",".join(["%s"] * len(_rff_aliases))
            cur.execute(f"""
                WITH ordered AS (
                    SELECT
                        s.record_id,
                        ROW_NUMBER() OVER (
                            ORDER BY MIN(COALESCE(s.file_order, 1)), MIN(s.created_at), s.record_id
                        ) AS new_file_order
                    FROM steps s
                    WHERE LOWER(TRIM(COALESCE(s.folder_name, ''))) IN ({_rff_ph})
                    GROUP BY s.record_id
                )
                UPDATE steps s
                SET file_order = ordered.new_file_order
                FROM ordered
                WHERE s.record_id = ordered.record_id
                  AND LOWER(TRIM(COALESCE(s.folder_name, ''))) IN ({_rff_ph})
            """, _rff_aliases + _rff_aliases)
            return

        if folder_key == "unfiled":
            cur.execute("""
                WITH ordered AS (
                    SELECT
                        s.record_id,
                        ROW_NUMBER() OVER (
                            ORDER BY MIN(COALESCE(s.file_order, 1)), MIN(s.created_at), s.record_id
                        ) AS new_file_order
                    FROM steps s
                    WHERE LOWER(TRIM(COALESCE(s.folder_name, ''))) IN ('', 'unfiled')
                    GROUP BY s.record_id
                )
                UPDATE steps s
                SET file_order = ordered.new_file_order
                FROM ordered
                WHERE s.record_id = ordered.record_id
                  AND LOWER(TRIM(COALESCE(s.folder_name, ''))) IN ('', 'unfiled')
            """)
            return

        cur.execute("""
            WITH ordered AS (
                SELECT
                    s.record_id,
                    ROW_NUMBER() OVER (
                        ORDER BY MIN(COALESCE(s.file_order, 1)), MIN(s.created_at), s.record_id
                    ) AS new_file_order
                FROM steps s
                WHERE TRIM(COALESCE(s.folder_name, '')) = %s
                GROUP BY s.record_id
            )
            UPDATE steps s
            SET file_order = ordered.new_file_order
            FROM ordered
            WHERE s.record_id = ordered.record_id
              AND TRIM(COALESCE(s.folder_name, '')) = %s
        """, [folder_name, folder_name])


def _next_distinct_file_order(cursor, folder_name: str | None = None, end_folder_id: str | None = None) -> int:
    """Return the next file_order based on MAX(file_order) over distinct session-folder pairs.

    When *end_folder_id* is supplied (depth-2 folders) the MAX is scoped by that
    UUID column instead of the folder_name string, giving true hierarchy-aware ordering.
    """

    # Depth-2 (A/B/C) folder: use end_folder_id UUID for precise scoping.
    if end_folder_id:
        cursor.execute("""
            SELECT COALESCE(MAX(file_order), 0) FROM (
                SELECT DISTINCT record_id, COALESCE(file_order, 1) AS file_order
                FROM steps
                WHERE end_folder_id = %s::uuid
            ) AS distinct_files
        """, [end_folder_id])
        return int(cursor.fetchone()[0] or 0) + 1

    folder_key = (folder_name or "").strip().lower()

    if folder_name is not None and folder_key not in {"", "unfiled"} and not _is_recordings_folder_name(folder_name):
        cursor.execute("""
            SELECT COALESCE(MAX(file_order), 0) FROM (
                SELECT DISTINCT
                    record_id,
                    LOWER(TRIM(COALESCE(folder_name, ''))) AS folder_key,
                    COALESCE(file_order, 1) AS file_order
                FROM steps
                WHERE LOWER(TRIM(COALESCE(folder_name, ''))) = %s
            ) AS distinct_files
        """, [folder_key])
        return int(cursor.fetchone()[0] or 0) + 1

    if folder_name is None:
        cursor.execute("""
            SELECT COALESCE(MAX(file_order), 0) FROM (
                SELECT DISTINCT
                    record_id,
                    LOWER(TRIM(COALESCE(folder_name, ''))) AS folder_key,
                    COALESCE(file_order, 1) AS file_order
                FROM steps
            ) AS distinct_files
        """)
        return int(cursor.fetchone()[0] or 0) + 1

    if not folder_key or _is_recordings_folder_name(folder_name):
        _ndfo_aliases = list(_recordings_sql_aliases())
        _ndfo_ph = ",".join(["%s"] * len(_ndfo_aliases))
        cursor.execute(f"""
            SELECT COALESCE(MAX(file_order), 0) FROM (
                SELECT DISTINCT
                    record_id,
                    LOWER(TRIM(COALESCE(folder_name, ''))) AS folder_key,
                    COALESCE(file_order, 1) AS file_order
                FROM steps
                WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_ndfo_ph})
            ) AS distinct_files
        """, _ndfo_aliases)
        return int(cursor.fetchone()[0] or 0) + 1

    if folder_key == "unfiled":
        cursor.execute("""
            SELECT COALESCE(MAX(file_order), 0) FROM (
                SELECT DISTINCT
                    record_id,
                    LOWER(TRIM(COALESCE(folder_name, ''))) AS folder_key,
                    COALESCE(file_order, 1) AS file_order
                FROM steps
                WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ('', 'unfiled')
            ) AS distinct_files
        """)
        return int(cursor.fetchone()[0] or 0) + 1

    return 1


def _record_name_exists_in_folder(cursor, folder_name: str, record_name: str, exclude_record_id: str | None = None) -> bool:
    """Return True when a session with the same display name already exists in the target folder."""
    folder_key = (folder_name or "").strip().lower()
    name_key = (record_name or "").strip().lower()
    if not folder_key or not name_key:
        return False

    _rnf_aliases = list(_recordings_sql_aliases())
    _rnf_ph = ",".join(["%s"] * len(_rnf_aliases))
    _recordings_label = _get_recordings_folder_label().lower()
    if _is_recordings_folder_name(folder_name):
        folder_key = _recordings_label
    params = _rnf_aliases + [_recordings_label, folder_key, name_key]
    exclude_sql = ""
    if exclude_record_id:
        exclude_sql = "AND folders.record_id <> %s::uuid"
        params.append(exclude_record_id)

    cursor.execute(f"""
        SELECT 1
        FROM (
            SELECT DISTINCT
                record_id,
                CASE
                    WHEN LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_rnf_ph}) THEN %s
                    WHEN LOWER(TRIM(COALESCE(folder_name, ''))) IN ('', 'unfiled') THEN 'unfiled'
                    ELSE LOWER(TRIM(COALESCE(folder_name, '')))
                END AS folder_key
            FROM steps
        ) folders
        JOIN session_meta m ON m.record_id = folders.record_id
        WHERE folders.folder_key = %s
          AND LOWER(TRIM(COALESCE(m.record_name, ''))) = %s
          {exclude_sql}
        LIMIT 1
    """, params)
    return cursor.fetchone() is not None


def _resolve_copy_record_name(cursor, folder_name: str, record_name: str, exclude_record_id: str | None = None) -> str:
    """Return an available record name, preferring the original name and then Copy of ... variants."""
    base_name = (record_name or "").strip() or "Untitled"
    if not _record_name_exists_in_folder(cursor, folder_name, base_name, exclude_record_id=exclude_record_id):
        return base_name

    copy_base = f"Copy of {base_name}"
    candidate = copy_base
    suffix = 2
    while _record_name_exists_in_folder(cursor, folder_name, candidate, exclude_record_id=exclude_record_id):
        candidate = f"{copy_base} ({suffix})"
        suffix += 1
    return candidate


# ---------------------------------------------------------------------------
# Copy a session into a folder
# ---------------------------------------------------------------------------

@login_required
def copy_to_folder(request, record_id):
    """POST: copy a recording into a named folder under a brand-new record_id."""
    if request.method != "POST":
        return redirect("sessions_list")
    folder = _normalize_folder_path(request.POST.get("folder", ""))
    if _is_recordings_folder_name(folder):
        folder = _get_recordings_folder_label()
    if not folder:
        messages.error(request, "Folder name is required to copy a recording.")
        return redirect("sessions_list")

    sid = str(record_id)

    # Look up source metadata
    try:
        src_meta = SessionMeta.objects.get(record_id=sid)
        record_name = src_meta.record_name or ""
        recorder    = src_meta.recorder    or ""
    except SessionMeta.DoesNotExist:
        record_name = ""
        recorder    = ""

    with connection.cursor() as _cur:
        _cur.execute("""
            SELECT COALESCE(MIN(NULLIF(TRIM(folder_name), '')), 'Baseline')
            FROM steps
            WHERE record_id = %s
        """, [sid])
        src_folder = _cur.fetchone()[0] or "Baseline"

    with transaction.atomic():
        with connection.cursor() as cursor:
            # ------------------------------------------------------------------
            # Generate a fresh record_id for this copy
            # ------------------------------------------------------------------
            new_sid = str(uuid.uuid4())

            # ------------------------------------------------------------------
            # Read source steps from steps table.
            # ------------------------------------------------------------------
            cursor.execute("""
                SELECT DISTINCT ON (s.step_no)
                       s.step_no, s.action, s.page_url, s.element_tag,
                       COALESCE(s.field_name, d.field_name) AS field_name,
                       COALESCE(s.field_value, d.value)     AS field_value,
                       s.raw_event, s.recorder, s.runner, s.locators_raw,
                       s.pos_x, s.pos_y, s.is_primary, s.locator_rank,
                       s.strategy, s.locator,
                       s.author, s.last_updated_by,
                       s.parent_record_id::text, s.sub_record_id::text,
                       s.end_record::text, s.is_baseline,
                       s.tenant_id, s.validation, s.steps_description,
                       s.page_title, s.engine, s.playwright_code
                FROM steps s
                LEFT JOIN data d ON d.record_id = s.record_id AND d.step_no = s.step_no
                WHERE s.record_id = %s
                ORDER BY s.step_no, s.id
            """, [sid])
            rec_rows = cursor.fetchall()
            src_loc_folder = src_folder

            if not rec_rows:
                messages.error(request, "Could not load source steps for this recording.")
                return redirect("sessions_list")

            # ------------------------------------------------------------------
            # Read locators from the matching source folder.
            # ------------------------------------------------------------------
            cursor.execute("""
                SELECT step_no, strategy, locator, is_primary,
                       locator_rank, pos_x, pos_y
                FROM locators
                WHERE record_id = %s AND folder_name = %s
                ORDER BY step_no, COALESCE(locator_rank, 999), id
            """, [sid, src_loc_folder])
            loc_rows = cursor.fetchall()

            from collections import defaultdict
            locs_by_step: dict = defaultdict(list)
            for (l_step, l_strat, l_loc, l_primary, l_rank, l_px, l_py) in loc_rows:
                locs_by_step[l_step].append((l_strat, l_loc, l_primary, l_rank, l_px, l_py))

            # Destination folder for locators, data and steps
            dest_is_recordings = _is_recordings_folder_name(folder)
            dest_loc_folder    = folder
            meta_folder        = folder
            if not dest_is_recordings:
                _register_project_folder(cursor, meta_folder)
            _dest_pfid, _dest_sfid, _dest_efid = _resolve_folder_ids(
                meta_folder, recorder, dest_is_recordings
            )
            # Always name the copy as "Copy of {name}", with (2)/(3)/… suffix
            # when that name is already taken in the destination folder.
            _ctf_base_name = (record_name or "Untitled").strip()
            _ctf_copy_base = f"Copy of {_ctf_base_name}"
            target_record_name = _ctf_copy_base
            _ctf_suffix = 2
            while _record_name_exists_in_folder(cursor, meta_folder, target_record_name):
                target_record_name = f"{_ctf_copy_base} ({_ctf_suffix})"
                _ctf_suffix += 1

            # Compute file_order scoped by end_folder_id UUID (depth-2) or
            # folder_name string (depth-0/1) for the destination folder.
            _new_file_order = _next_distinct_file_order(cursor, meta_folder, end_folder_id=_dest_efid)
            _src_engine = getattr(src_meta, 'engine', None) or 'selenium'

            # Copy session_meta under the new record_id
            _src_baseline = getattr(src_meta, 'is_baseline', False)
            cursor.execute("""
                INSERT INTO session_meta
                    (record_id, record_name, recorder, folder_name,
                     parent_folder_id, sub_folder_id, end_folder_id, engine, is_baseline, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                ON CONFLICT DO NOTHING
            """, [new_sid, target_record_name, recorder, meta_folder,
                  _dest_pfid, _dest_sfid, _dest_efid, _src_engine, _src_baseline])

            for (step_no, action, page_url, element_tag,
                 field_name, field_value,
                 raw_event, recorder_val, runner, locators_raw,
                 pos_x, pos_y, is_primary, locator_rank,
                 strategy, locator_val,
                 author, last_updated_by,
                 parent_record_id, sub_record_id,
                 end_record, is_baseline,
                 tenant_id, validation, steps_description,
                 page_title, engine, playwright_code) in rec_rows:

                # 1. Copy data entry under new_sid
                new_data_id = None
                if field_name is not None or field_value is not None:
                    cursor.execute("""
                        INSERT INTO data (record_id, step_no, field_name, value, folder_name, created_at)
                        VALUES (%s, %s, %s, %s, %s, NOW())
                        RETURNING id
                    """, [new_sid, step_no, field_name, field_value, dest_loc_folder])
                    new_data_id = cursor.fetchone()[0]

                # 2. Copy locators under new_sid
                new_locator_id = None
                for (l_strat, l_loc, l_primary, l_rank, l_px, l_py) in locs_by_step.get(step_no, []):
                    cursor.execute("""
                        INSERT INTO locators
                            (record_id, step_no, strategy, locator, is_primary,
                             locator_rank, pos_x, pos_y, folder_name, created_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                        RETURNING id
                    """, [new_sid, step_no, l_strat, l_loc, l_primary, l_rank, l_px, l_py, dest_loc_folder])
                    row_id = cursor.fetchone()[0]
                    if l_primary and new_locator_id is None:
                        new_locator_id = row_id

                # 3. Insert into steps table for all destinations
                cursor.execute("""
                    INSERT INTO steps
                        (record_id, step_no, action, page_url, element_tag,
                         locator_id, data_id, raw_event, recorder, runner, folder_name,
                         file_order, parent_folder_id, sub_folder_id, end_folder_id,
                         locators_raw, field_name, field_value,
                         pos_x, pos_y, is_primary, locator_rank,
                         strategy, locator,
                         author, last_updated_by,
                         parent_record_id, sub_record_id, end_record, is_baseline,
                         tenant_id, validation, steps_description,
                         page_title, engine, playwright_code,
                         created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                            %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                            %s, %s, %s, %s, %s, %s, NOW())
                """, [
                    new_sid, step_no, action, page_url, element_tag,
                    new_locator_id, new_data_id,
                    raw_event, recorder_val, runner, folder, _new_file_order,
                    _dest_pfid, _dest_sfid, _dest_efid,
                    locators_raw, field_name, field_value,
                    pos_x, pos_y, is_primary, locator_rank,
                    strategy, locator_val,
                    author, last_updated_by,
                    parent_record_id, sub_record_id, end_record, is_baseline,
                    tenant_id, validation, steps_description,
                    page_title, engine, playwright_code,
                ])

    messages.success(request, f'"{record_name or sid}" copied to folder "{folder}" as "{target_record_name}".')
    return redirect("sessions_list")


# ---------------------------------------------------------------------------
# Copy an entire folder → new folder (each session gets a new record_id)
# ---------------------------------------------------------------------------

@login_required
def copy_folder(request):
    """POST: copy every session in src_folder into dest_folder, each with a fresh record_id.

    Non-special project folders are duplicated as full subtrees so root-level
    copies also preserve nested subfolders/end-folders and their contents.
    """
    if request.method != "POST":
        return redirect("sessions_list")

    src_folder  = _normalize_folder_path(request.POST.get("src_folder",  ""))
    dest_folder = _normalize_folder_path(request.POST.get("dest_folder", ""))
    insert_after_folder = _normalize_folder_path(request.POST.get("insert_after_folder", ""))

    if _is_recordings_folder_name(src_folder):
        src_folder = _get_recordings_folder_label()
    if _is_recordings_folder_name(dest_folder):
        dest_folder = _get_recordings_folder_label()

    if not src_folder:
        messages.error(request, "Source folder is missing.")
        return redirect("sessions_list")

    if not dest_folder:
        src_parts = [part.strip() for part in src_folder.split("/") if part.strip()]
        if len(src_parts) > 1:
            parent_path = "/".join(src_parts[:-1])
            leaf_name = src_parts[-1]
            copy_base = _normalize_folder_path(f"{parent_path}/Copy of {leaf_name}")
        else:
            copy_base = f"Copy of {src_folder}"
        dest_folder = copy_base
        suffix = 2
        with connection.cursor() as cur:
            while _project_folder_exists(cur, dest_folder):
                dest_folder = f"{copy_base} ({suffix})"
                suffix += 1

    if _is_recordings_folder_name(dest_folder):
        messages.error(request, "Please choose a different destination folder name.")
        return redirect("sessions_list")
    if src_folder == dest_folder:
        messages.error(request, "Destination folder must be different from the source folder.")
        return redirect("sessions_list")

    src_parts = [part.strip() for part in src_folder.split("/") if part.strip()]
    duplicate_subtree = bool(src_parts) and not _is_recordings_folder_name(src_folder)
    folder_prefix = src_folder + "/"

    with connection.cursor() as cur:
        if not _is_recordings_folder_name(dest_folder) and _project_folder_exists(cur, dest_folder):
            messages.error(request, f'Folder "{dest_folder}" already exists.')
            return redirect("sessions_list")

    # Collect all sessions that belong to the source folder. Project-folder
    # duplication includes descendant folders and preserves their relative paths
    # under the new destination root, including root-level parent folders.
    with connection.cursor() as cur:
        if _is_recordings_folder_name(src_folder):
            _cf_aliases = list(_recordings_sql_aliases())
            _cf_ph = ",".join(["%s"] * len(_cf_aliases))
            cur.execute(f"""
                SELECT DISTINCT record_id, TRIM(COALESCE(folder_name, '')) AS source_folder
                FROM steps
                WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_cf_ph})
            """, _cf_aliases)
        else:
            if duplicate_subtree:
                cur.execute("""
                    SELECT DISTINCT record_id, TRIM(COALESCE(folder_name, '')) AS source_folder
                    FROM steps
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [src_folder, folder_prefix + "%"])
            else:
                cur.execute("""
                    SELECT DISTINCT record_id, TRIM(COALESCE(folder_name, '')) AS source_folder
                    FROM steps
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                """, [src_folder])

        copy_sources: list[tuple[str, str]] = []
        for row in cur.fetchall():
            source_record_id = str(row[0])
            source_record_folder = _normalize_folder_path(row[1])
            if _is_recordings_folder_name(source_record_folder):
                source_record_folder = _get_recordings_folder_label()
            copy_sources.append((source_record_id, source_record_folder))

    if not copy_sources:
        messages.warning(request, f'No recordings found in folder "{src_folder}".')
        return redirect("sessions_list")

    copied = 0
    for sid, source_record_folder in copy_sources:
        if duplicate_subtree and source_record_folder.startswith(folder_prefix):
            relative_suffix = source_record_folder[len(folder_prefix):]
            target_folder = _normalize_folder_path(f"{dest_folder}/{relative_suffix}")
        else:
            target_folder = dest_folder

        target_insert_after = target_folder
        if duplicate_subtree and source_record_folder == src_folder:
            target_insert_after = insert_after_folder or src_folder

        # Resolve metadata
        try:
            src_meta    = SessionMeta.objects.get(record_id=sid)
            record_name = src_meta.record_name or ""
            recorder    = src_meta.recorder    or ""
        except SessionMeta.DoesNotExist:
            record_name = ""
            recorder    = ""

        with transaction.atomic():
            with connection.cursor() as cursor:
                if not _is_recordings_folder_name(target_folder):
                    _register_project_folder(cursor, target_folder, insert_after_folder=target_insert_after)
                new_sid = str(uuid.uuid4())

                # Resolve destination hierarchy IDs now so that:
                #  - new parent_folders / sub_folders / end_folders rows are
                #    created with fresh UUIDs for the duplicate project, and
                #  - every INSERT below carries the correct FK values directly.
                # No explicit_parent_id: _resolve_folder_ids looks up the new
                # parent by name so sub_folder_parent links to the NEW project row,
                # not the source project's row.
                _p_id, _s_id, _e_id = _resolve_folder_ids(
                    target_folder, recorder or request.user.username, False,
                )
                new_file_order = _next_distinct_file_order(cursor, target_folder, end_folder_id=_e_id)
                _src_engine = getattr(src_meta, 'engine', None) or 'selenium'
                _src_baseline = getattr(src_meta, 'is_baseline', False)

                cursor.execute("""
                    INSERT INTO session_meta
                        (record_id, record_name, recorder, folder_name,
                         parent_folder_id, sub_folder_id, end_folder_id, engine, is_baseline, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                    ON CONFLICT DO NOTHING
                """, [new_sid, record_name, recorder, target_folder, _p_id, _s_id, _e_id, _src_engine, _src_baseline])

                cursor.execute("""
                    SELECT DISTINCT ON (s.step_no)
                           s.step_no, s.action, s.page_url, s.element_tag,
                           d.field_name, d.value AS field_value,
                           s.raw_event, s.recorder, s.runner,
                           s.author, s.last_updated_by, s.headless_state,
                           s.is_baseline, s.file_type
                    FROM steps s
                    LEFT JOIN data d ON d.record_id = s.record_id AND d.step_no = s.step_no
                    WHERE s.record_id = %s
                    ORDER BY s.step_no, s.id
                """, [sid])
                rec_rows = cursor.fetchall()

                cursor.execute("""
                    SELECT step_no, strategy, locator, is_primary,
                           locator_rank, pos_x, pos_y
                    FROM locators
                    WHERE record_id = %s
                    ORDER BY step_no, COALESCE(locator_rank, 999), id
                """, [sid])
                loc_rows = cursor.fetchall()

                from collections import defaultdict
                locs_by_step: dict = defaultdict(list)
                for (l_step, l_strat, l_loc, l_primary, l_rank, l_px, l_py) in loc_rows:
                    locs_by_step[l_step].append((l_strat, l_loc, l_primary, l_rank, l_px, l_py))

                for (step_no, action, page_url, element_tag,
                     field_name, field_value, raw_event, recorder_val, runner,
                     author, last_updated_by, headless_state, is_baseline, file_type) in rec_rows:

                    new_data_id = None
                    if field_name is not None or field_value is not None:
                        cursor.execute("""
                            INSERT INTO data (record_id, step_no, field_name, value, folder_name, created_at)
                            VALUES (%s, %s, %s, %s, %s, NOW())
                            RETURNING id
                        """, [new_sid, step_no, field_name, field_value, target_folder])
                        new_data_id = cursor.fetchone()[0]

                    new_locator_id = None
                    for (l_strat, l_loc, l_primary, l_rank, l_px, l_py) in locs_by_step.get(step_no, []):
                        cursor.execute("""
                            INSERT INTO locators
                                (record_id, step_no, strategy, locator, is_primary,
                                 locator_rank, pos_x, pos_y, folder_name, created_at)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                            RETURNING id
                        """, [new_sid, step_no, l_strat, l_loc, l_primary, l_rank, l_px, l_py, target_folder])
                        row_id = cursor.fetchone()[0]
                        if l_primary and new_locator_id is None:
                            new_locator_id = row_id

                    cursor.execute("""
                        INSERT INTO steps
                            (record_id, step_no, action, page_url, element_tag,
                             locator_id, data_id, raw_event, recorder, runner,
                             author, last_updated_by, headless_state, is_baseline,
                             file_type, folder_name, file_order,
                             parent_folder_id, sub_folder_id, end_folder_id,
                             created_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                    """, [
                        new_sid, step_no, action, page_url, element_tag,
                        new_locator_id, new_data_id,
                        raw_event, recorder_val, runner,
                        author, last_updated_by, headless_state, is_baseline,
                        file_type, target_folder, new_file_order,
                        _p_id, _s_id, _e_id,
                    ])

                # Backfill recordings and run_table which have no direct INSERT above.
                for _tbl in ("recordings", "run_table"):
                    cursor.execute(
                        f"""UPDATE {_tbl}
                               SET parent_folder_id = %s,
                                   sub_folder_id    = %s,
                                   end_folder_id    = %s
                             WHERE record_id = %s""",
                        [_p_id, _s_id, _e_id, new_sid],
                    )

        copied += 1

    messages.success(
        request,
        f'{copied} recording{"s" if copied != 1 else ""} copied from '
        f'"{src_folder}" to "{dest_folder}" — each with a new session ID.'
    )
    return redirect("sessions_list")


@login_required
def create_folder(request):
    """POST: create an empty folder or subfolder entry for the projects page."""
    if request.method != "POST":
        return redirect("sessions_list")

    parent_folder = _normalize_folder_path(request.POST.get("parent_folder", ""))
    insert_after_folder = _normalize_folder_path(request.POST.get("insert_after_folder", ""))
    requested_public = request.POST.get("public") == "on"
    requested_is_baseline = request.POST.get("is_baseline", "") == "true" and request.user.is_superuser
    folder_name = (request.POST.get("folder_name", "") or "").strip()
    folder_parts = [part.strip() for part in folder_name.replace("\\", "/").split("/") if part.strip()]

    if not folder_parts:
        messages.error(request, "Folder name is required.")
        return redirect("sessions_list")
    if len(folder_parts) != 1:
        messages.error(request, "Enter one folder name only. Use Create Subfolder for nested folders.")
        return redirect("sessions_list")
    if folder_parts[0] in {".", ".."}:
        messages.error(request, "Please choose a valid folder name.")
        return redirect("sessions_list")
    if _is_recordings_folder_name(parent_folder):
        parent_folder = _get_recordings_folder_label()
    if parent_folder.lower() in {"unfiled"}:
        messages.error(request, "Subfolders cannot be created under this folder.")
        return redirect("sessions_list")

    full_folder = _normalize_folder_path(f"{parent_folder}/{folder_parts[0]}" if parent_folder else folder_parts[0])
    if full_folder.lower() in _RESERVED_PROJECT_FOLDER_KEYS:
        messages.error(request, "Please choose a different folder name.")
        return redirect("sessions_list")

    with transaction.atomic():
        with connection.cursor() as cur:
            _ensure_project_folders_table()
            if parent_folder and not _project_folder_exists(cur, parent_folder):
                messages.error(request, f'Parent folder "{parent_folder}" was not found.')
                return redirect("sessions_list")
            if _project_folder_exists(cur, full_folder):
                messages.error(request, f'Folder "{full_folder}" already exists.')
                return redirect("sessions_list")
            _register_project_folder(cur, full_folder, insert_after_folder=insert_after_folder)

    # Ensure parent_folders / sub_folders / end_folders rows exist for the
    # newly created path.  _resolve_folder_ids looks up parts[0] in
    # parent_folders internally; that same parent_folder_id is then used
    # consistently as sub_folder_parent — no separate lookup needed here.
    _rf_author, _rf_public, _rf_is_baseline = request.user.username, requested_public, requested_is_baseline
    if parent_folder:
        # Inherit author, public, is_baseline from the direct parent folder row.
        # depth 1 new folder (file_type=sub-folder)  → copy from parent_folders
        # depth 2 new folder (file_type=end-folder)  → copy from sub_folders
        _parent_parts = [p.strip() for p in parent_folder.split("/") if p.strip()]
        try:
            with connection.cursor() as _pf_cur:
                if len(_parent_parts) >= 2:
                    # Parent is a sub_folder row: look up via sub_folder name + parent name
                    _pf_cur.execute(
                        """
                        SELECT sf.author, sf.public, sf.is_baseline
                        FROM sub_folders sf
                        JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
                        WHERE pf.parent_folder = %s AND sf.sub_folder = %s
                        LIMIT 1
                        """,
                        [_parent_parts[0], _parent_parts[1]],
                    )
                else:
                    # Parent is a parent_folder row
                    _pf_cur.execute(
                        "SELECT author, public, is_baseline FROM parent_folders WHERE parent_folder = %s",
                        [_parent_parts[0]],
                    )
                _pf_row = _pf_cur.fetchone()
                if _pf_row:
                    _rf_author      = _pf_row[0] or request.user.username
                    _rf_public      = bool(_pf_row[1])
                    _rf_is_baseline = bool(_pf_row[2])
        except Exception:
            pass
    _resolve_folder_ids(full_folder, _rf_author, _rf_is_baseline, public=_rf_public)
    if not parent_folder and request.user.is_superuser:
        with connection.cursor() as _bf_cur:
            _bf_cur.execute(
                "UPDATE parent_folders SET is_baseline = %s WHERE parent_folder = %s",
                [requested_is_baseline, folder_parts[0]],
            )
        _clear_recordings_aliases_cache()
    # Update last_updated on any ancestor folder
    if parent_folder:
        _touch_folder_hierarchy(parent_folder)

    messages.success(request, f'Folder "{full_folder}" created.')
    return redirect("sessions_list")


@login_required
def set_folder_public(request):
    """POST: toggle public visibility for a folder hierarchy branch."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)

    folder_name = _normalize_folder_path(request.POST.get("folder_name", ""))
    if not folder_name:
        return JsonResponse({"error": "folder_name required"}, status=400)

    target_public = str(request.POST.get("public", "")).strip().lower() in {"1", "true", "on", "yes"}
    meta = _get_project_folder_metadata_map([folder_name]).get(folder_name)
    if not meta:
        return JsonResponse({"error": "Folder not found"}, status=404)

    username = (request.user.username or "").strip().lower()
    author = (meta.get("author") or "").strip().lower()
    if not request.user.is_superuser and (not username or username != author):
        return JsonResponse({"error": "Access denied"}, status=403)

    parts = [part.strip() for part in folder_name.split("/") if part.strip()]
    if not parts:
        return JsonResponse({"error": "Invalid folder"}, status=400)

    with transaction.atomic():
        with connection.cursor() as cur:
            cur.execute(
                "UPDATE parent_folders SET public = %s WHERE parent_folder = %s",
                [target_public, parts[0]],
            )

            if len(parts) == 1:
                cur.execute(
                    """
                    UPDATE sub_folders sf
                    SET public = %s
                    FROM parent_folders pf
                    WHERE pf.parent_folder_id = sf.sub_folder_parent
                      AND pf.parent_folder = %s
                    """,
                    [target_public, parts[0]],
                )
                cur.execute(
                    """
                    UPDATE end_folders ef
                    SET public = %s
                    FROM sub_folders sf
                    JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
                    WHERE ef.end_folder_parent = sf.sub_folder_id
                      AND pf.parent_folder = %s
                    """,
                    [target_public, parts[0]],
                )
            elif len(parts) == 2:
                cur.execute(
                    """
                    UPDATE sub_folders sf
                    SET public = %s
                    FROM parent_folders pf
                    WHERE pf.parent_folder_id = sf.sub_folder_parent
                      AND pf.parent_folder = %s
                      AND sf.sub_folder = %s
                    """,
                    [target_public, parts[0], parts[1]],
                )
                cur.execute(
                    """
                    UPDATE end_folders ef
                    SET public = %s
                    FROM sub_folders sf
                    JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
                    WHERE ef.end_folder_parent = sf.sub_folder_id
                      AND pf.parent_folder = %s
                      AND sf.sub_folder = %s
                    """,
                    [target_public, parts[0], parts[1]],
                )
            else:
                cur.execute(
                    """
                    UPDATE sub_folders sf
                    SET public = %s
                    FROM parent_folders pf
                    WHERE pf.parent_folder_id = sf.sub_folder_parent
                      AND pf.parent_folder = %s
                      AND sf.sub_folder = %s
                    """,
                    [target_public, parts[0], parts[1]],
                )
                cur.execute(
                    """
                    UPDATE end_folders ef
                    SET public = %s
                    FROM sub_folders sf
                    JOIN parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
                    WHERE ef.end_folder_parent = sf.sub_folder_id
                      AND pf.parent_folder = %s
                      AND sf.sub_folder = %s
                      AND ef.end_folder = %s
                    """,
                    [target_public, parts[0], parts[1], parts[2]],
                )

    return JsonResponse({"ok": True, "folder_name": folder_name, "public": target_public})


# ---------------------------------------------------------------------------
# Duplicate a session (full clone under a new record_id, same folder)
# ---------------------------------------------------------------------------

@login_required
def duplicate_session(request, record_id):
    """POST: create a full clone of a session under a new record_id."""
    if request.method != "POST":
        return redirect("sessions_list")

    sid         = str(record_id)
    requested_name = request.POST.get("new_name", "").strip()

    # Resolve source metadata
    try:
        src_meta    = SessionMeta.objects.get(record_id=sid)
        record_name = src_meta.record_name or ""
        recorder    = src_meta.recorder    or ""
    except SessionMeta.DoesNotExist:
        record_name = ""
        recorder    = ""

    # Get folder from steps table (matches how sessions_list groups sessions)
    with connection.cursor() as _cur:
        _cur.execute("""
            SELECT COALESCE(MIN(NULLIF(TRIM(folder_name), '')), 'Baseline')
            FROM steps WHERE record_id = %s
        """, [sid])
        src_folder = _cur.fetchone()[0] or "Baseline"

    # Resolve folder hierarchy IDs from source session_meta
    _pfid = getattr(src_meta, 'parent_folder_id', None)
    _sfid = getattr(src_meta, 'sub_folder_id', None)
    _efid = getattr(src_meta, 'end_folder_id', None)
    if not _pfid:
        _pfid, _sfid, _efid = _resolve_folder_ids(src_folder, recorder, False)

    with transaction.atomic():
        with connection.cursor() as cursor:
            new_sid = str(uuid.uuid4())
            source_name = (record_name or sid).strip() or "Untitled"
            if requested_name:
                copy_name = requested_name
            else:
                copy_base = f"Copy of {source_name}"
                copy_name = copy_base
                suffix = 2
                while _record_name_exists_in_folder(cursor, src_folder, copy_name):
                    copy_name = f"{copy_base} ({suffix})"
                    suffix += 1

            cursor.execute("""
                SELECT DISTINCT ON (s.step_no)
                       s.step_no, s.action, s.page_url, s.element_tag,
                       COALESCE(s.field_name, d.field_name)   AS field_name,
                       COALESCE(s.field_value, d.value)       AS field_value,
                       s.raw_event, s.recorder, s.runner,
                       s.locators_raw,
                       s.pos_x, s.pos_y,
                       s.strategy, s.locator, s.is_primary, s.locator_rank,
                       s.author,
                       s.tenant_id, s.validation, s.steps_description,
                       s.page_title, s.engine, s.playwright_code
                FROM steps s
                LEFT JOIN data d ON d.record_id = s.record_id AND d.step_no = s.step_no
                WHERE s.record_id = %s
                ORDER BY s.step_no, s.id
            """, [sid])
            step_rows = cursor.fetchall()

            dest_folder_name = src_folder
            _src_engine = getattr(src_meta, 'engine', None) or 'selenium'
            _src_baseline = getattr(src_meta, 'is_baseline', False)

            # ── session_meta ────────────────────────────────────────────────
            cursor.execute("""
                INSERT INTO session_meta
                    (record_id, record_name, recorder, folder_name,
                     parent_folder_id, sub_folder_id, end_folder_id, engine, is_baseline, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                ON CONFLICT DO NOTHING
            """, [new_sid, copy_name, recorder, dest_folder_name, _pfid, _sfid, _efid, _src_engine, _src_baseline])

            # ── locators for this session ────────────────────────────────────
            cursor.execute("""
                SELECT step_no, strategy, locator, is_primary,
                       locator_rank, pos_x, pos_y
                FROM locators
                WHERE record_id = %s
                ORDER BY step_no, COALESCE(locator_rank, 999), id
            """, [sid])
            from collections import defaultdict as _dd
            locs_by_step = _dd(list)
            for row in cursor.fetchall():
                locs_by_step[row[0]].append(row[1:])

            # ── compute file_order for the duplicate ────────────────────────
            # Scope by end_folder_id UUID when available (depth-2 folders)
            # so ordering is precise and not subject to folder_name string
            # collisions across the hierarchy.
            _new_file_order = _next_distinct_file_order(cursor, dest_folder_name, end_folder_id=_efid)

            # ── copy each step ───────────────────────────────────────────────
            for (step_no, action, page_url, element_tag,
                 field_name, field_value, raw_event, recorder_val, runner,
                 locators_raw, pos_x, pos_y,
                 strategy, locator_val, is_primary, locator_rank,
                 author,
                 tenant_id, validation, steps_description,
                 page_title, engine, playwright_code) in step_rows:

                new_data_id = None
                if field_name is not None or field_value is not None:
                    cursor.execute("""
                        INSERT INTO data (record_id, step_no, field_name, value, folder_name, created_at)
                        VALUES (%s, %s, %s, %s, %s, NOW()) RETURNING id
                    """, [new_sid, step_no, field_name, field_value, dest_folder_name])
                    new_data_id = cursor.fetchone()[0]

                new_locator_id = None
                for (l_strat, l_loc, l_primary, l_rank, l_px, l_py) in locs_by_step.get(step_no, []):
                    cursor.execute("""
                        INSERT INTO locators
                            (record_id, step_no, strategy, locator, is_primary,
                             locator_rank, pos_x, pos_y, folder_name, created_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW()) RETURNING id
                    """, [new_sid, step_no, l_strat, l_loc, l_primary, l_rank, l_px, l_py, dest_folder_name])
                    row_id = cursor.fetchone()[0]
                    if l_primary and new_locator_id is None:
                        new_locator_id = row_id

                cursor.execute("""
                    INSERT INTO steps
                        (record_id, step_no, action, page_url, element_tag,
                         locator_id, data_id, raw_event, recorder, runner, folder_name,
                         file_order, parent_folder_id, sub_folder_id, end_folder_id,
                         locators_raw, field_name, field_value,
                         pos_x, pos_y, strategy, locator, is_primary, locator_rank,
                         author,
                         tenant_id, validation, steps_description,
                         page_title, engine, playwright_code,
                         created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                """, [
                    new_sid, step_no, action, page_url, element_tag,
                    new_locator_id, new_data_id,
                    raw_event, recorder_val, runner, src_folder, _new_file_order,
                    _pfid, _sfid, _efid,
                    locators_raw, field_name, field_value,
                    pos_x, pos_y, strategy, locator_val, is_primary, locator_rank,
                    author,
                    tenant_id, validation, steps_description,
                    page_title, engine, playwright_code,
                ])

    messages.success(request, f'Duplicated as "{copy_name}".')
    return redirect("sessions_list")


# ---------------------------------------------------------------------------
# Start / stop recording
# ---------------------------------------------------------------------------

def _resolve_folder_ids(
    folder_name: str, author: str, is_baseline: bool,
    explicit_parent_id: str | None = None,
    public: bool = False,
) -> tuple[str | None, str | None, str | None]:
    """Ensure parent_folders / sub_folders / end_folders rows exist for the
    given slash-delimited path and return (parent_folder_id, sub_folder_id, end_folder_id).

    Rules:
      depth 0  →  "A"        →  parent only
      depth 1  →  "A/B"      →  parent + sub
      depth 2  →  "A/B/C"    →  parent + sub + end
    """
    parts = [p.strip() for p in folder_name.split("/") if p.strip()]
    if not parts:
        return None, None, None

    parent_id = sub_id = end_id = None

    # ── parent_folders ───────────────────────────────────────────────────────
    with connection.cursor() as cur:
        cur.execute(
            "SELECT parent_folder_id FROM parent_folders WHERE parent_folder = %s",
            [parts[0]],
        )
        row = cur.fetchone()
        if row:
            parent_id = str(row[0])
            if public:
                cur.execute(
                    "UPDATE parent_folders SET public = TRUE WHERE parent_folder_id = %s::uuid",
                    [parent_id],
                )
        else:
            parent_id = str(uuid.uuid4())
            cur.execute(
                """
                INSERT INTO parent_folders
                    (parent_folder_id, parent_folder, parent_order, parent_folder_order,
                     file_type, author, public, is_baseline)
                VALUES (
                    %s, %s,
                    (SELECT COALESCE(MAX(parent_order), 0) + 1 FROM parent_folders),
                    %s, 'folder', %s, %s, %s
                )
                ON CONFLICT (parent_folder_id) DO NOTHING
                """,
                [parent_id, parts[0], _next_parent_folder_order(cur), author, public, is_baseline],
            )
            if is_baseline:
                _clear_recordings_aliases_cache()

    if len(parts) < 2:
        return parent_id, None, None

    # ── sub_folders ──────────────────────────────────────────────────────────
    # explicit_parent_id comes from the root of insert_after_folder so
    # sub_folder_parent links to the authoritative parent_folders row.
    _sub_parent = explicit_parent_id or parent_id
    with connection.cursor() as cur:
        cur.execute(
            "SELECT sub_folder_id FROM sub_folders WHERE sub_folder = %s AND sub_folder_parent = %s",
            [parts[1], _sub_parent],
        )
        row = cur.fetchone()
        if row:
            sub_id = str(row[0])
        else:
            sub_id = str(uuid.uuid4())
            cur.execute(
                """
                INSERT INTO sub_folders
                    (sub_folder_id, sub_folder, sub_folder_parent, sub_folder_order,
                     file_type, author, public, is_baseline)
                VALUES (
                    %s, %s, %s,
                    %s,
                    'sub-folder', %s, %s, %s
                )
                ON CONFLICT (sub_folder_id) DO NOTHING
                """,
                [sub_id, parts[1], _sub_parent, _next_sub_folder_order(cur, _sub_parent), author, public, is_baseline],
            )
            cur.execute(
                "UPDATE parent_folders SET last_updated = NOW() WHERE parent_folder_id = %s::uuid",
                [parent_id],
            )

    if len(parts) < 3:
        return parent_id, sub_id, None

    # ── end_folders ──────────────────────────────────────────────────────────
    with connection.cursor() as cur:
        cur.execute(
            "SELECT end_folder_id FROM end_folders WHERE end_folder = %s AND end_folder_parent = %s",
            [parts[2], sub_id],
        )
        row = cur.fetchone()
        if row:
            end_id = str(row[0])
        else:
            end_id = str(uuid.uuid4())
            cur.execute(
                """
                INSERT INTO end_folders
                    (end_folder_id, end_folder, end_folder_parent,
                     end_folder_order, end_file_order,
                     file_type, author, public, is_baseline)
                VALUES (
                    %s, %s, %s,
                    (SELECT COALESCE(MAX(end_folder_order), 0) + 1
                     FROM end_folders WHERE end_folder_parent = %s),
                    1, 'end-folder', %s, %s, %s
                )
                ON CONFLICT (end_folder_id) DO NOTHING
                """,
                [end_id, parts[2], sub_id, sub_id, author, public, is_baseline],
            )
            cur.execute(
                "UPDATE sub_folders SET last_updated = NOW() WHERE sub_folder_id = %s::uuid",
                [sub_id],
            )
            cur.execute(
                "UPDATE parent_folders SET last_updated = NOW() WHERE parent_folder_id = %s::uuid",
                [parent_id],
            )

    return parent_id, sub_id, end_id


def _touch_folder_hierarchy(folder_name: str) -> None:
    """Update last_updated on every ancestor folder row when content beneath changes."""
    parts = [p.strip() for p in folder_name.split("/") if p.strip()]
    if not parts:
        return
    try:
        with connection.cursor() as cur:
            cur.execute(
                "UPDATE parent_folders SET last_updated = NOW() WHERE parent_folder = %s",
                [parts[0]],
            )
            if len(parts) >= 2:
                cur.execute("""
                    UPDATE sub_folders SET last_updated = NOW()
                    WHERE sub_folder = %s
                      AND sub_folder_parent IN (
                          SELECT parent_folder_id FROM parent_folders
                          WHERE parent_folder = %s
                      )
                """, [parts[1], parts[0]])
            if len(parts) >= 3:
                cur.execute("""
                    UPDATE end_folders SET last_updated = NOW()
                    WHERE end_folder = %s
                      AND end_folder_parent IN (
                          SELECT sub_folder_id FROM sub_folders
                          WHERE sub_folder = %s
                            AND sub_folder_parent IN (
                                SELECT parent_folder_id FROM parent_folders
                                WHERE parent_folder = %s
                            )
                      )
                """, [parts[2], parts[1], parts[0]])
    except Exception:
        pass


@login_required
def start_recording(request):
    if request.method != "POST":
        return redirect("sessions_list")

    wants_json = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    url = request.POST.get("url", "").strip()
    record_name = request.POST.get("record_name", "").strip()
    folder_name = request.POST.get("folder_name", "").strip()
    engine = request.POST.get("engine", "selenium").strip().lower()  # "selenium" or "playwright"
    open_url_before_recording = (request.POST.get("open_url_before_recording", "true").strip().lower()
                                 in {"1", "true", "on", "yes"})
    is_baseline = request.POST.get("is_baseline", "") == "true" and request.user.is_superuser
    if not record_name:
        if wants_json:
            return JsonResponse({"ok": False, "error": "Please enter a record name before starting a recording."}, status=400)
        messages.error(request, "Please enter a record name before starting a recording.")
        return redirect("sessions_list")
    if open_url_before_recording and not url:
        if wants_json:
            return JsonResponse({"ok": False, "error": "Please enter a URL to record."}, status=400)
        messages.error(request, "Please enter a URL to record.")
        return redirect("sessions_list")
    if not folder_name:
        if wants_json:
            return JsonResponse({"ok": False, "error": "Please select a folder before starting a recording."}, status=400)
        messages.error(request, "Please select a folder before starting a recording.")
        return redirect("sessions_list")

    _rdp = (get_config("chrome.remote_debugging_port") or "").strip()
    attach_to_open_page = False
    if not open_url_before_recording:
        if engine == "playwright":
            if not url:
                if wants_json:
                    return JsonResponse({"ok": False, "error": "Record open page is currently supported only for Selenium recordings."}, status=400)
                messages.error(request, "Record open page is currently supported only for Selenium recordings.")
                return redirect("sessions_list")
        elif _rdp.isdigit() and int(_rdp) > 0:
            try:
                import socket as _socket
                with _socket.create_connection(("127.0.0.1", int(_rdp)), timeout=2):
                    attach_to_open_page = True
            except OSError:
                attach_to_open_page = False
        if not attach_to_open_page and not url:
            if not (_rdp.isdigit() and int(_rdp) > 0):
                if wants_json:
                    return JsonResponse({
                        "ok": False,
                        "error": "chrome.remote_debugging_port must be configured before recording the currently open page.",
                    }, status=400)
                messages.error(
                    request,
                    "chrome.remote_debugging_port must be configured before recording the currently open page.",
                )
            else:
                if wants_json:
                    return JsonResponse({
                        "ok": False,
                        "error": "No Chrome session is listening on the configured remote debugging port. Open the application in that Chrome session or provide a URL.",
                    }, status=400)
                messages.error(
                    request,
                    "No Chrome session is listening on the configured remote debugging port. Open the application in that Chrome session or provide a URL.",
                )
            return redirect("sessions_list")

    should_launch_url = open_url_before_recording or not attach_to_open_page

    # Ensure the URL has a scheme so the browser can navigate to it
    if should_launch_url and url and not url.startswith(("http://", "https://")):
        url = "https://" + url

    # Generate record_id here so the session can track it for post-stop linking
    record_id = str(uuid.uuid4())

    # Upsert folder hierarchy tables and resolve IDs
    if folder_name:
        _resolve_folder_ids(folder_name, request.user.username, is_baseline)

    # Kill any existing recording first (Python process only, leave Chrome alive)
    existing_pid = request.session.get("recording_pid")
    if existing_pid:
        try:
            if sys.platform == "win32":
                subprocess.call(
                    ["taskkill", "/F", "/PID", str(existing_pid)],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            else:
                os.kill(existing_pid, signal.SIGTERM)
        except (OSError, ProcessLookupError):
            pass

    # Forward the exact DB credentials Django is using so main.py can reach app_config.
    _db = settings.DATABASES["default"]
    os.makedirs(_RECORDING_LOGS_DIR, exist_ok=True)
    _log_path = os.path.join(_RECORDING_LOGS_DIR, f"rec_{record_id}.log")
    log_fh = open(_log_path, "w", encoding="utf-8")
    _recorder_script = _PLAYWRIGHT_PY if engine == "playwright" else _MAIN_PY
    _popen_args = [
        sys.executable, "-u", _recorder_script,
        "--record-name", record_name,
        "--recorder", request.user.username,
        "--folder-name", folder_name,
        "--record-id", record_id,
        "--db-host",     _db.get("HOST", "localhost"),
        "--db-port",     str(_db.get("PORT", "5432")),
        "--db-name",     _db.get("NAME", "automation_db"),
        "--db-user",     _db.get("USER", "postgres"),
        "--db-password", _db.get("PASSWORD", ""),
    ]
    if should_launch_url and url:
        _popen_args += ["--url", url]
    # Expose a remote-debugging port so "Add Step" can reattach later (Selenium only)
    if engine != "playwright":
        if should_launch_url:
            if _rdp.isdigit() and int(_rdp) > 0:
                _popen_args += ["--remote-debug-port", _rdp]
        else:
            _popen_args += ["--attach-port", _rdp]
    elif not should_launch_url:
        _popen_args.append("--no-navigate")
    if is_baseline:
        _popen_args.append("--is-baseline")
    proc = subprocess.Popen(
        _popen_args,
        stdout=log_fh,
        stderr=log_fh,
        creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if sys.platform == "win32" else 0,
    )
    log_fh.close()  # child process inherited its own fd; safe to close in parent
    request.session["recording_pid"]          = proc.pid
    request.session["recording_url"]          = url if should_launch_url else ""
    request.session["recording_name"]         = record_name
    request.session["recording_folder"]       = folder_name
    request.session["recording_id"]           = record_id
    request.session["recording_is_baseline"]  = is_baseline
    request.session["recording_engine"]       = engine
    with _ACTIVE_RECORDING_LOCK:
        _ACTIVE_RECORDING[str(record_id)] = {
            "pid": proc.pid, "folder": folder_name,
            "user_id": request.user.pk, "paused": False,
            "engine": engine,
        }

    # Persist PID to a temp file so stop_recording_ajax can always find it
    # even if session and _ACTIVE_RECORDING are lost (e.g. server auto-reload).
    import tempfile as _tf
    _pid_file = os.path.join(_tf.gettempdir(), f"recorder_pid_{record_id}.txt")
    try:
        with open(_pid_file, "w") as _pf:
            _pf.write(f"{proc.pid}\n{folder_name}\n{request.user.pk}")
    except OSError:
        pass

    try:
        _open_recording_monitor(request, record_id, record_name)
    except Exception:
        pass
    _engine_label = "Playwright" if engine == "playwright" else "Selenium"
    if wants_json:
        return JsonResponse({
            "ok": True,
            "record_id": record_id,
            "redirect_url": reverse("session_steps", args=[record_id]) + "?recording=1",
            "message": (
                f"Recording started ({_engine_label}) for: {url}"
                if should_launch_url else
                f"Recording started ({_engine_label}) on the already-open page."
            ),
        })
    if should_launch_url:
        messages.success(request, f"Recording started ({_engine_label}) for: {url}")
    else:
        messages.success(request, f"Recording started ({_engine_label}) on the already-open page.")
    return redirect("sessions_list")


@login_required
def stop_recording(request):
    if request.method != "POST":
        return redirect("sessions_list")

    pid          = request.session.pop("recording_pid",          None)
    folder_name  = request.session.pop("recording_folder",       "")
    record_id    = request.session.pop("recording_id",           "")
    is_baseline  = request.session.pop("recording_is_baseline",  False)
    request.session.pop("recording_url",    None)
    request.session.pop("recording_name",   None)
    request.session.pop("recording_paused", None)

    # Clean up the Python-level pause flag (in case recording was paused when stopped).
    if record_id:
        import tempfile as _tf
        _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{record_id}.flag")
        try:
            os.remove(_flag)
        except OSError:
            pass

    if pid:
        try:
            if sys.platform == "win32":
                # First try a graceful console break so the recorder can flush
                # any last buffered browser events before exiting.
                _stopped_gracefully = False
                try:
                    os.kill(int(pid), signal.CTRL_BREAK_EVENT)
                    for _ in range(20):
                        time.sleep(0.1)
                        try:
                            os.kill(int(pid), 0)
                        except OSError:
                            _stopped_gracefully = True
                            break
                except (AttributeError, OSError, ValueError, ProcessLookupError):
                    _stopped_gracefully = False

                if not _stopped_gracefully:
                    # Fall back to a hard stop, but still avoid /T so Chrome
                    # stays open on its remote-debugging port for Add Step.
                    subprocess.call(["taskkill", "/F", "/PID", str(pid)],
                                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            else:
                os.kill(pid, signal.SIGTERM)
            messages.success(request, "Recording stopped.")
        except (OSError, ProcessLookupError):
            messages.warning(request, "Recording process was already stopped.")
    else:
        messages.warning(request, "No active recording found.")

    # Backfill folder hierarchy IDs into all tables for this recording session
    if record_id and folder_name:
        try:
            parent_id, sub_id, end_id = _resolve_folder_ids(
                folder_name, request.user.username, is_baseline
            )
            _tables = ["recordings", "steps", "run_table", "session_meta"]
            with connection.cursor() as cur:
                # Compute file_order BEFORE updating the FK columns so that
                # the current session (end_folder_id still NULL) is excluded
                # from the MAX calculation.
                _rec_file_order = _next_distinct_file_order(cur, folder_name, end_folder_id=end_id)

                for _tbl in _tables:
                    cur.execute(
                        f"""
                        UPDATE {_tbl}
                           SET parent_folder_id = %s,
                               sub_folder_id    = %s,
                               end_folder_id    = %s
                         WHERE record_id = %s
                        """,
                        [parent_id, sub_id, end_id, record_id],
                    )

                # Set file_order for the new recording (tables that carry it).
                for _tbl_fo in ("recordings", "steps", "run_table"):
                    cur.execute(
                        f"UPDATE {_tbl_fo} SET file_order = %s WHERE record_id = %s",
                        [_rec_file_order, record_id],
                    )
        except Exception as _exc:
            # Non-fatal — recording data is still intact
            messages.warning(request, f"Recording stopped but folder linking failed: {_exc}")

        # Propagate last_updated up the folder hierarchy
        _touch_folder_hierarchy(folder_name)

    if record_id:
        request.session["show_record_more"] = str(record_id)
        return redirect("session_steps", record_id=record_id)
    return redirect("sessions_list")


# ---------------------------------------------------------------------------
# Continue recording — append new steps to an existing session
# ---------------------------------------------------------------------------

@login_required
def continue_recording(request, record_id):
    """POST: reopen the browser and keep recording into an existing session.

    New steps are appended starting after the current MAX(step_no) so they
    never overwrite existing data.  The folder, record_name, and recorder
    are taken from the existing session_meta / steps rows.
    """
    if request.method != "POST":
        return redirect("session_steps", record_id=record_id)

    sid = str(record_id)

    # ── Resolve session metadata ──────────────────────────────────────────
    try:
        meta = SessionMeta.objects.get(record_id=sid)
        record_name = meta.record_name or ""
        folder_name = meta.folder_name or ""
    except SessionMeta.DoesNotExist:
        record_name = ""
        folder_name = ""

    # Derive folder from steps if session_meta is missing it
    if not folder_name:
        _fn_row = Step.objects.filter(record_id=sid).values_list("folder_name", flat=True).first()
        folder_name = (_fn_row or "").strip()

    is_baseline = _is_recordings_folder_name(folder_name)

    # ── Determine starting step_no (append after last existing step) ──────
    with connection.cursor() as _cur:
        _cur.execute(
            "SELECT COALESCE(MAX(step_no), 0) FROM steps WHERE record_id = %s",
            [sid],
        )
        _max_step = int(_cur.fetchone()[0] or 0)

    # ── Find the last page URL — kept for the session message only ───────
    _last_url_row = (
        Step.objects.filter(record_id=sid, page_url__isnull=False)
        .exclude(page_url="")
        .order_by("-step_no")
        .values_list("page_url", flat=True)
        .first()
    )
    _last_url = (_last_url_row or "").strip()

    # ── Kill any already-running recording for this user ─────────────────
    existing_pid = request.session.get("recording_pid")
    if existing_pid:
        try:
            if sys.platform == "win32":
                subprocess.call(
                    ["taskkill", "/F", "/PID", str(existing_pid)],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            else:
                os.kill(existing_pid, signal.SIGTERM)
        except (OSError, ProcessLookupError):
            pass

    # ── Launch main.py with the SAME record_id, offset step_no ───────────
    _rdp = (get_config("chrome.remote_debugging_port") or "").strip()
    if not (_rdp.isdigit() and int(_rdp) > 0):
        # Cannot attach without a remote debugging port — tell the user immediately.
        err = ("chrome.remote_debugging_port is not configured. "
               "Go to /configuration/ and set it (e.g. 9222) so Add Step "
               "can attach to the already-running Chrome window.")
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "error": err})
        messages.error(request, err)
        return redirect("session_steps", record_id=record_id)

    _db = settings.DATABASES["default"]
    os.makedirs(_RECORDING_LOGS_DIR, exist_ok=True)
    _log_path = os.path.join(_RECORDING_LOGS_DIR, f"rec_{sid}_append.log")
    log_fh = open(_log_path, "w", encoding="utf-8")
    _popen_args = [
        sys.executable, "-u", _MAIN_PY,
        "--record-name", record_name,
        "--recorder",    request.user.username,
        "--folder-name", folder_name,
        "--record-id",   sid,
        "--start-step",  str(_max_step),
        "--attach-port", _rdp,
        "--db-host",     _db.get("HOST", "localhost"),
        "--db-port",     str(_db.get("PORT", "5432")),
        "--db-name",     _db.get("NAME", "automation_db"),
        "--db-user",     _db.get("USER", "postgres"),
        "--db-password", _db.get("PASSWORD", ""),
    ]
    if is_baseline:
        _popen_args.append("--is-baseline")
    proc = subprocess.Popen(
        _popen_args,
        stdout=log_fh,
        stderr=log_fh,
        creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if sys.platform == "win32" else 0,
    )
    log_fh.close()

    request.session["recording_pid"]         = proc.pid
    request.session["recording_url"]         = _last_url
    request.session["recording_name"]        = record_name
    request.session["recording_folder"]      = folder_name
    request.session["recording_id"]          = sid
    request.session["recording_is_baseline"] = is_baseline
    request.session.pop("recording_paused",  None)  # clear any stale paused state
    with _ACTIVE_RECORDING_LOCK:
        _ACTIVE_RECORDING[str(sid)] = {
            "pid": proc.pid, "folder": folder_name,
            "user_id": request.user.pk, "paused": False,
        }

    # Persist PID to a temp file so stop_recording_ajax can always find it
    # even if session and _ACTIVE_RECORDING are lost (e.g. server auto-reload).
    import tempfile as _tf
    _pid_file = os.path.join(_tf.gettempdir(), f"recorder_pid_{sid}.txt")
    try:
        with open(_pid_file, "w") as _pf:
            _pf.write(f"{proc.pid}\n{folder_name}\n{request.user.pk}")
    except OSError:
        pass

    try:
        _open_recording_monitor(request, sid, record_name)
    except Exception:
        pass

    # Ensure no stale pause flag exists for this record_id before the new main.py starts.
    _stale_flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{sid}.flag")
    try:
        os.remove(_stale_flag)
    except OSError:
        pass

    # Un-pause RECORDER_JS in the browser so the new recording process can capture events.
    _rdp_resume = (get_config("chrome.remote_debugging_port") or "").strip()
    if _rdp_resume.isdigit() and int(_rdp_resume) > 0:
        _cdp_evaluate(int(_rdp_resume),
            "if(window.__webActionRecorder){"
            "  window.__webActionRecorder.paused=false;"
            "  window.__webActionRecorder.events=[];"
            "}"
            "try{sessionStorage.removeItem('__webActionRecorder_paused');}catch(e){}")

    # Return JSON when called via AJAX (fetch from the steps page), otherwise
    # fall back to a redirect so the old form-based path still works.
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return JsonResponse({
            "ok": True,
            "message": (
                f'Continuing recording for "{record_name or sid[:8]}…" '
                f"— new steps will be appended after step {_max_step}."
            ),
        })

    messages.success(
        request,
        f'Continuing recording for \u201c{record_name or sid[:8]}\u2026\u201d '
        f'— new steps will be appended after step {_max_step}.',
    )
    return redirect("sessions_list")


@login_required
def add_events_recording(request, record_id):
    """POST: start a new recording session attached to the already-open page."""
    if request.method != "POST":
        return redirect("session_steps", record_id=record_id)

    sid = str(record_id)

    try:
        meta = SessionMeta.objects.get(record_id=sid)
        folder_name = (meta.folder_name or "").strip()
        engine = (meta.engine or "selenium").strip().lower()
    except SessionMeta.DoesNotExist:
        folder_name = ""
        engine = "selenium"

    record_name = (request.POST.get("record_name", "") or "").strip()
    if not record_name:
        err = "Record name is required."
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "error": err}, status=400)
        messages.error(request, err)
        return redirect("session_steps", record_id=record_id)

    is_baseline = request.POST.get("is_baseline", "") == "true" and request.user.is_superuser

    if not folder_name:
        _fn_row = Step.objects.filter(record_id=sid).values_list("folder_name", flat=True).first()
        folder_name = (_fn_row or "").strip()

    if not folder_name:
        err = "This session has no folder assigned, so a new Add Events recording cannot be created."
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "error": err}, status=400)
        messages.error(request, err)
        return redirect("session_steps", record_id=record_id)

    _rdp = (get_config("chrome.remote_debugging_port") or "").strip()
    if engine == "playwright":
        err = "Add Events currently supports Selenium sessions only."
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "error": err}, status=400)
        messages.error(request, err)
        return redirect("session_steps", record_id=record_id)
    if not (_rdp.isdigit() and int(_rdp) > 0):
        err = ("chrome.remote_debugging_port is not configured. "
               "Go to /configuration/ and set it (e.g. 9222) so Add Events "
               "can attach to the already-running Chrome window.")
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "error": err}, status=400)
        messages.error(request, err)
        return redirect("session_steps", record_id=record_id)

    try:
        import socket as _socket
        with _socket.create_connection(("127.0.0.1", int(_rdp)), timeout=2):
            pass
    except OSError:
        err = (
            "No Chrome session is listening on the configured remote debugging port. "
            "Open the application in that Chrome session before starting Add Events."
        )
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "error": err}, status=400)
        messages.error(request, err)
        return redirect("session_steps", record_id=record_id)

    new_record_id = str(uuid.uuid4())

    if folder_name:
        _resolve_folder_ids(folder_name, request.user.username, is_baseline)

    existing_pid = request.session.get("recording_pid")
    if existing_pid:
        try:
            if sys.platform == "win32":
                subprocess.call(
                    ["taskkill", "/F", "/PID", str(existing_pid)],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            else:
                os.kill(existing_pid, signal.SIGTERM)
        except (OSError, ProcessLookupError):
            pass

    _db = settings.DATABASES["default"]
    os.makedirs(_RECORDING_LOGS_DIR, exist_ok=True)
    _log_path = os.path.join(_RECORDING_LOGS_DIR, f"rec_{new_record_id}.log")
    log_fh = open(_log_path, "w", encoding="utf-8")
    _popen_args = [
        sys.executable, "-u", _MAIN_PY,
        "--record-name", record_name,
        "--recorder", request.user.username,
        "--folder-name", folder_name,
        "--record-id", new_record_id,
        "--attach-port", _rdp,
        "--db-host", _db.get("HOST", "localhost"),
        "--db-port", str(_db.get("PORT", "5432")),
        "--db-name", _db.get("NAME", "automation_db"),
        "--db-user", _db.get("USER", "postgres"),
        "--db-password", _db.get("PASSWORD", ""),
    ]
    if is_baseline:
        _popen_args.append("--is-baseline")
    proc = subprocess.Popen(
        _popen_args,
        stdout=log_fh,
        stderr=log_fh,
        creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if sys.platform == "win32" else 0,
    )
    log_fh.close()

    request.session["recording_pid"] = proc.pid
    request.session["recording_url"] = ""
    request.session["recording_name"] = record_name
    request.session["recording_folder"] = folder_name
    request.session["recording_id"] = new_record_id
    request.session["recording_is_baseline"] = is_baseline
    request.session["recording_engine"] = engine
    request.session.pop("recording_paused", None)
    with _ACTIVE_RECORDING_LOCK:
        _ACTIVE_RECORDING[str(new_record_id)] = {
            "pid": proc.pid,
            "folder": folder_name,
            "user_id": request.user.pk,
            "paused": False,
            "engine": engine,
        }

    # Persist PID to a temp file so stop_recording_ajax can always find it
    # even if session and _ACTIVE_RECORDING are lost (e.g. server auto-reload).
    import tempfile as _tf_ae
    _pid_file_ae = os.path.join(_tf_ae.gettempdir(), f"recorder_pid_{new_record_id}.txt")
    try:
        with open(_pid_file_ae, "w") as _pf:
            _pf.write(f"{proc.pid}\n{folder_name}\n{request.user.pk}")
    except OSError:
        pass

    try:
        _open_recording_monitor(request, new_record_id, record_name)
    except Exception:
        pass

    try:
        import tempfile as _tf
        _stale_flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{new_record_id}.flag")
        os.remove(_stale_flag)
    except OSError:
        pass

    _cdp_evaluate(int(_rdp),
        "if(window.__webActionRecorder){"
        "  window.__webActionRecorder.paused=false;"
        "  window.__webActionRecorder.events=[];"
        "}"
        "try{sessionStorage.removeItem('__webActionRecorder_paused');}catch(e){}")

    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return JsonResponse({
            "ok": True,
            "record_id": new_record_id,
            "redirect_url": reverse("session_steps", args=[new_record_id]) + "?recording=1",
            "message": f'Add Events recording started for "{record_name or new_record_id[:8]}...".',
        })

    messages.success(request, "Add Events recording started on the already-open page.")
    return redirect("session_steps", record_id=new_record_id)


@login_required
def attach_steps(request, record_id):
    """POST /sessions/<uuid>/attach-steps/
    Copy all steps from a source recording and append them to the end of this session.
    Expects POST param 'source_record_id'.
    """
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)

    source_id = (request.POST.get("source_record_id") or "").strip()
    if not source_id:
        return JsonResponse({"ok": False, "error": "No source recording selected."}, status=400)

    target_id = str(record_id)

    if source_id == target_id:
        return JsonResponse({"ok": False, "error": "Cannot attach a recording to itself."}, status=400)

    with connection.cursor() as cur:
        # Get current max step_no in the target session
        cur.execute("SELECT COALESCE(MAX(step_no), 0) FROM steps WHERE record_id = %s", [target_id])
        max_step = cur.fetchone()[0]

        # Get target folder info from existing steps first, fall back to session_meta
        cur.execute("""
            SELECT COALESCE(folder_name, 'Baseline'), file_order,
                   parent_folder_id, sub_folder_id, end_folder_id
            FROM steps WHERE record_id = %s ORDER BY step_no LIMIT 1
        """, [target_id])
        target_row = cur.fetchone()
        if not target_row:
            # No steps yet — resolve folder info from session_meta
            cur.execute("""
                SELECT COALESCE(folder_name, 'Baseline'),
                       parent_folder_id, sub_folder_id, end_folder_id
                FROM session_meta WHERE record_id = %s
            """, [target_id])
            meta_row = cur.fetchone()
            if meta_row:
                t_folder = meta_row[0]
                t_pfid = meta_row[1]
                t_sfid = meta_row[2]
                t_efid = meta_row[3]
                t_file_order = _next_distinct_file_order(cur, t_folder, end_folder_id=t_efid)
            else:
                t_folder = "Baseline"
                t_file_order = 1
                t_pfid = t_sfid = t_efid = None
        else:
            t_folder, t_file_order, t_pfid, t_sfid, t_efid = target_row

        # Read source steps
        cur.execute("""
            SELECT step_no, action, page_url, element_tag,
                   field_name, field_value, raw_event, recorder, runner,
                   locators_raw, pos_x, pos_y, is_primary, locator_rank,
                   strategy, locator, author, last_updated_by,
                   is_baseline, tenant_id, validation, steps_description,
                   page_title, engine, playwright_code
            FROM steps
            WHERE record_id = %s
            ORDER BY step_no
        """, [source_id])
        src_steps = cur.fetchall()

        if not src_steps:
            return JsonResponse({"ok": False, "error": "Source recording has no steps."}, status=400)

        # Read source locators
        cur.execute("""
            SELECT step_no, strategy, locator, is_primary, locator_rank, pos_x, pos_y
            FROM locators
            WHERE record_id = %s
            ORDER BY step_no, COALESCE(locator_rank, 999), id
        """, [source_id])
        src_locs = cur.fetchall()

        from collections import defaultdict
        locs_by_step: dict = defaultdict(list)
        for (l_step, l_strat, l_loc, l_prim, l_rank, l_px, l_py) in src_locs:
            locs_by_step[l_step].append((l_strat, l_loc, l_prim, l_rank, l_px, l_py))

        # Read source data entries
        cur.execute("""
            SELECT step_no, field_name, value
            FROM data
            WHERE record_id = %s
            ORDER BY step_no
        """, [source_id])
        data_by_step: dict = {}
        for (d_step, d_name, d_val) in cur.fetchall():
            data_by_step[d_step] = (d_name, d_val)

        # Append source steps with offset
        attached_count = 0
        for (step_no, action, page_url, element_tag,
             field_name, field_value, raw_event, recorder_val, runner,
             locators_raw, pos_x, pos_y, is_primary, locator_rank,
             strategy, locator_val, author, last_updated_by,
             is_baseline, tenant_id, validation, steps_description,
             page_title, engine, playwright_code) in src_steps:

            new_step_no = max_step + step_no

            # Insert data entry
            new_data_id = None
            d_entry = data_by_step.get(step_no)
            if d_entry or field_name or field_value:
                d_name = d_entry[0] if d_entry else field_name
                d_val = d_entry[1] if d_entry else field_value
                cur.execute("""
                    INSERT INTO data (record_id, step_no, field_name, value, folder_name, created_at)
                    VALUES (%s, %s, %s, %s, %s, NOW())
                    RETURNING id
                """, [target_id, new_step_no, d_name, d_val, t_folder])
                new_data_id = cur.fetchone()[0]

            # Insert locators
            new_locator_id = None
            for (l_strat, l_loc, l_prim, l_rank, l_px, l_py) in locs_by_step.get(step_no, []):
                cur.execute("""
                    INSERT INTO locators
                        (record_id, step_no, strategy, locator, is_primary,
                         locator_rank, pos_x, pos_y, folder_name, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                    RETURNING id
                """, [target_id, new_step_no, l_strat, l_loc, l_prim, l_rank, l_px, l_py, t_folder])
                row_id = cur.fetchone()[0]
                if l_prim and new_locator_id is None:
                    new_locator_id = row_id

            # Insert step
            cur.execute("""
                INSERT INTO steps
                    (record_id, step_no, action, page_url, element_tag,
                     locator_id, data_id, raw_event, recorder, runner, folder_name,
                     file_order, parent_folder_id, sub_folder_id, end_folder_id,
                     locators_raw, field_name, field_value,
                     pos_x, pos_y, is_primary, locator_rank,
                     strategy, locator, author, last_updated_by,
                     is_baseline, tenant_id, validation, steps_description,
                     page_title, engine, playwright_code, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
            """, [
                target_id, new_step_no, action, page_url, element_tag,
                new_locator_id, new_data_id,
                raw_event, recorder_val, runner, t_folder,
                t_file_order, t_pfid, t_sfid, t_efid,
                locators_raw, field_name, field_value,
                pos_x, pos_y, is_primary, locator_rank,
                strategy, locator_val, author, last_updated_by,
                is_baseline, tenant_id, validation, steps_description,
                page_title, engine, playwright_code,
            ])
            attached_count += 1

    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return JsonResponse({
            "ok": True,
            "attached_count": attached_count,
            "message": f"Attached {attached_count} step(s) from source recording.",
        })
    messages.success(request, f"Attached {attached_count} step(s) from source recording.")
    return redirect("session_steps", record_id=record_id)


@login_required
def rename_session(request, record_id):
    """POST: update record_name in session_meta for a given session."""
    if request.method != "POST":
        return redirect("sessions_list")

    new_name = request.POST.get("new_name", "").strip()
    if not new_name:
        messages.error(request, "Name cannot be blank.")
        return redirect("sessions_list")

    sid = str(record_id)
    with connection.cursor() as cursor:
        cursor.execute("""
            SELECT COALESCE(MIN(NULLIF(TRIM(folder_name), '')), 'Baseline')
            FROM steps WHERE record_id = %s
        """, [sid])
        source_folder = cursor.fetchone()[0] or "Baseline"
        if _record_name_exists_in_folder(cursor, source_folder, new_name, exclude_record_id=sid):
            messages.error(request, f'A file named "{new_name}" already exists in this folder.')
            return redirect("sessions_list")

    obj, _ = SessionMeta.objects.get_or_create(record_id=sid)
    obj.record_name = new_name
    obj.save(update_fields=["record_name"])
    messages.success(request, f'Renamed to “{new_name}”.')
    return redirect("sessions_list")


# ---------------------------------------------------------------------------
# Delete a session
# ---------------------------------------------------------------------------

@login_required
def delete_session(request, record_id):
    if request.method != "POST":
        return redirect("sessions_list")
    sid = str(record_id)
    folder = request.POST.get("folder_name", "").strip()
    # Resolve a human-readable name for the success message before deleting
    try:
        record_name = SessionMeta.objects.get(record_id=record_id).record_name or sid
    except SessionMeta.DoesNotExist:
        record_name = sid
    with transaction.atomic():
        with connection.cursor() as cur:
            if folder in ("Baseline", "") or _is_recordings_folder_name(folder):
                # Null out FK refs first so locators/data can be deleted
                cur.execute("UPDATE run_table   SET locator_id = NULL WHERE record_id = %s AND locator_id IS NOT NULL", [sid])
                cur.execute("UPDATE run_table   SET data_id    = NULL WHERE record_id = %s AND data_id    IS NOT NULL", [sid])
                cur.execute("UPDATE steps       SET locator_id = NULL WHERE record_id = %s AND locator_id IS NOT NULL", [sid])
                cur.execute("UPDATE steps       SET data_id    = NULL WHERE record_id = %s AND data_id    IS NOT NULL", [sid])
                cur.execute("UPDATE recordings  SET locator_id = NULL WHERE record_id = %s AND locator_id IS NOT NULL", [sid])
                cur.execute("UPDATE recordings  SET data_id    = NULL WHERE record_id = %s AND data_id    IS NOT NULL", [sid])
                # Delete all rows for this record across every table
                cur.execute("DELETE FROM locators          WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM data              WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM steps             WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM recordings        WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM run_table         WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM locators_stat     WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM remote_executions WHERE record_id = %s", [sid])
            else:
                # Folder copy — delete only rows belonging to this session + folder
                cur.execute("DELETE FROM run_table WHERE record_id = %s AND folder_name = %s", [sid, folder])
                cur.execute("DELETE FROM steps     WHERE record_id = %s AND folder_name = %s", [sid, folder])
                # Null out FK refs from any remaining rows before removing locators/data
                cur.execute("UPDATE run_table  SET locator_id = NULL WHERE record_id = %s AND locator_id IN (SELECT id FROM locators WHERE record_id = %s AND folder_name = %s)", [sid, sid, folder])
                cur.execute("UPDATE run_table  SET data_id    = NULL WHERE record_id = %s AND data_id    IN (SELECT id FROM data     WHERE record_id = %s AND folder_name = %s)", [sid, sid, folder])
                cur.execute("UPDATE steps      SET locator_id = NULL WHERE record_id = %s AND locator_id IN (SELECT id FROM locators WHERE record_id = %s AND folder_name = %s)", [sid, sid, folder])
                cur.execute("UPDATE steps      SET data_id    = NULL WHERE record_id = %s AND data_id    IN (SELECT id FROM data     WHERE record_id = %s AND folder_name = %s)", [sid, sid, folder])
                cur.execute("UPDATE recordings SET locator_id = NULL WHERE record_id = %s AND locator_id IN (SELECT id FROM locators WHERE record_id = %s AND folder_name = %s)", [sid, sid, folder])
                cur.execute("UPDATE recordings SET data_id    = NULL WHERE record_id = %s AND data_id    IN (SELECT id FROM data     WHERE record_id = %s AND folder_name = %s)", [sid, sid, folder])
                cur.execute("DELETE FROM locators      WHERE record_id = %s AND folder_name = %s", [sid, folder])
                cur.execute("DELETE FROM data          WHERE record_id = %s AND folder_name = %s", [sid, folder])
                cur.execute("DELETE FROM recordings    WHERE record_id = %s AND folder_name = %s", [sid, folder])
                cur.execute("DELETE FROM locators_stat WHERE record_id = %s AND folder_name = %s", [sid, folder])
                # Remove session_meta once this record_id no longer has any surviving
                # session rows after the file delete.
                cur.execute("""
                    DELETE FROM session_meta WHERE record_id = %s
                    AND NOT EXISTS (SELECT 1 FROM steps      WHERE record_id = %s LIMIT 1)
                    AND NOT EXISTS (SELECT 1 FROM recordings WHERE record_id = %s LIMIT 1)
                    AND NOT EXISTS (SELECT 1 FROM run_table  WHERE record_id = %s LIMIT 1)
                """, [sid, sid, sid, sid])
            # For Baseline/full deletes (the if-branch above), always remove session_meta
            if folder in ("Baseline", "") or _is_recordings_folder_name(folder):
                cur.execute("DELETE FROM session_meta WHERE record_id = %s", [sid])
            _resequence_file_order_for_folder(folder or "Baseline")
    messages.success(request, f'"{record_name}" deleted.')
    return redirect("sessions_list")


@login_required
def delete_folder(request):
    """POST: delete every session in a named folder."""
    if request.method != "POST":
        return redirect("sessions_list")
    folder = _normalize_folder_path(request.POST.get("folder_name", ""))
    if not folder:
        messages.warning(request, "No folder specified.")
        return redirect("sessions_list")
    is_recordings = _is_recordings_folder_name(folder)
    if is_recordings and not request.user.is_superuser:
        messages.error(request, "Only superusers can delete the Recordings folder.")
        return redirect("sessions_list")
    db_folder = "Records" if is_recordings else folder
    with transaction.atomic():
        with connection.cursor() as cur:
            if not is_recordings:
                like_pattern = db_folder + "/%"
                cur.execute("""
                    UPDATE run_table SET locator_id = NULL
                    WHERE locator_id IN (
                        SELECT id FROM locators
                        WHERE TRIM(COALESCE(folder_name, '')) = %s
                           OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    )
                """, [db_folder, like_pattern])
                cur.execute("""
                    UPDATE run_table SET data_id = NULL
                    WHERE data_id IN (
                        SELECT id FROM data
                        WHERE TRIM(COALESCE(folder_name, '')) = %s
                           OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    )
                """, [db_folder, like_pattern])
                cur.execute("""
                    UPDATE steps SET locator_id = NULL
                    WHERE locator_id IN (
                        SELECT id FROM locators
                        WHERE TRIM(COALESCE(folder_name, '')) = %s
                           OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    )
                """, [db_folder, like_pattern])
                cur.execute("""
                    UPDATE steps SET data_id = NULL
                    WHERE data_id IN (
                        SELECT id FROM data
                        WHERE TRIM(COALESCE(folder_name, '')) = %s
                           OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    )
                """, [db_folder, like_pattern])
                cur.execute("""
                    UPDATE recordings SET locator_id = NULL
                    WHERE locator_id IN (
                        SELECT id FROM locators
                        WHERE TRIM(COALESCE(folder_name, '')) = %s
                           OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    )
                """, [db_folder, like_pattern])
                cur.execute("""
                    UPDATE recordings SET data_id = NULL
                    WHERE data_id IN (
                        SELECT id FROM data
                        WHERE TRIM(COALESCE(folder_name, '')) = %s
                           OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    )
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM run_table
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM steps
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM recordings
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM locators_stat
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM remote_executions
                    WHERE record_id IN (
                        SELECT record_id FROM session_meta
                        WHERE TRIM(COALESCE(folder_name, '')) = %s
                           OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    )
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM locators
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM data
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                cur.execute("""
                    DELETE FROM session_meta
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                _delete_folder_path_from_hierarchy(cur, db_folder)
                _delete_folder_path_from_registry(cur, db_folder)
                _resequence_project_folder_orders(cur)
                messages.success(request, f'Folder "{folder}" and its subfolders have been deleted.')
                next_url = request.POST.get("next", "").strip()
                return redirect(next_url if next_url.startswith("/") else "sessions_list")

            # NULL out FK refs by actual locator/data IDs being deleted,
            # not by folder_name — other folders' run_table rows may also
            # reference the same locators/data entries.
            cur.execute("UPDATE run_table  SET locator_id = NULL WHERE locator_id IN (SELECT id FROM locators WHERE folder_name = %s)", [db_folder])
            cur.execute("UPDATE run_table  SET data_id    = NULL WHERE data_id    IN (SELECT id FROM data     WHERE folder_name = %s)", [db_folder])
            cur.execute("UPDATE steps      SET locator_id = NULL WHERE locator_id IN (SELECT id FROM locators WHERE folder_name = %s)", [db_folder])
            cur.execute("UPDATE steps      SET data_id    = NULL WHERE data_id    IN (SELECT id FROM data     WHERE folder_name = %s)", [db_folder])
            cur.execute("UPDATE recordings SET locator_id = NULL WHERE locator_id IN (SELECT id FROM locators WHERE folder_name = %s)", [db_folder])
            cur.execute("UPDATE recordings SET data_id    = NULL WHERE data_id    IN (SELECT id FROM data     WHERE folder_name = %s)", [db_folder])
            cur.execute("DELETE FROM remote_executions WHERE record_id IN (SELECT record_id FROM session_meta WHERE folder_name = %s)", [db_folder])
            cur.execute("DELETE FROM run_table      WHERE folder_name = %s", [db_folder])
            cur.execute("DELETE FROM steps          WHERE folder_name = %s", [db_folder])
            cur.execute("DELETE FROM recordings     WHERE folder_name = %s", [db_folder])
            cur.execute("DELETE FROM locators_stat  WHERE folder_name = %s", [db_folder])
            cur.execute("DELETE FROM locators       WHERE folder_name = %s", [db_folder])
            cur.execute("DELETE FROM data           WHERE folder_name = %s", [db_folder])
            if is_recordings:
                cur.execute("DELETE FROM session_meta WHERE folder_name = %s", [db_folder])
    messages.success(request, f'Folder \u201c{folder}\u201d and all its sessions have been deleted.')
    next_url = request.POST.get("next", "").strip()
    return redirect(next_url if next_url.startswith("/") else "sessions_list")


@login_required
def rename_folder(request):
    """POST: rename a named folder across folder-scoped tables."""
    if request.method != "POST":
        return redirect("sessions_list")

    old_folder = _normalize_folder_path(request.POST.get("old_folder_name", ""))
    new_folder = _normalize_folder_path(request.POST.get("new_folder_name", ""))

    old_parts = [part.strip() for part in old_folder.split("/") if part.strip()]
    new_parts = [part.strip() for part in new_folder.split("/") if part.strip()]
    if len(old_parts) > 1 and new_parts:
        parent_path = "/".join(old_parts[:-1])
        new_leaf = new_parts[-1]
        new_folder = _normalize_folder_path(f"{parent_path}/{new_leaf}")

    if not old_folder:
        messages.warning(request, "No folder specified.")
        return redirect("sessions_list")
    if not new_folder:
        messages.warning(request, "New folder name is required.")
        return redirect("sessions_list")
    if old_folder == new_folder:
        messages.info(request, "Folder name is unchanged.")
        return redirect("sessions_list")

    old_key = old_folder.lower()
    new_key = new_folder.lower()
    is_recordings = _is_recordings_folder_name(old_folder)
    if is_recordings and not request.user.is_superuser:
        messages.error(request, "Only superusers can rename the Recordings folder.")
        return redirect("sessions_list")
    if new_key in {"", "baseline", "unfiled"}:
        messages.error(request, "Please choose a different folder name.")
        return redirect("sessions_list")
    if not is_recordings and new_key.startswith(old_key + "/"):
        messages.error(request, "A folder cannot be renamed into one of its own subfolders.")
        return redirect("sessions_list")

    with connection.cursor() as cur:
        if _project_folder_exists(cur, new_folder):
            messages.error(request, f'Folder "{new_folder}" already exists.')
            return redirect("sessions_list")

    if is_recordings:
        _ensure_config_table()
        with connection.cursor() as cur:
            cur.execute("""
                INSERT INTO app_config (key, value, label, description, group_name, input_type)
                VALUES (%s, %s, %s, %s, %s, %s)
                ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value
            """, [
                "projects.recordings_folder_label",
                new_folder,
                "Projects Recordings Folder Label",
                "Display label for the root recordings folder on the Projects page.",
                "UI",
                "text",
            ])
        messages.success(request, f'Folder "{old_folder}" renamed to "{new_folder}".')
        return redirect("sessions_list")

    with transaction.atomic():
        with connection.cursor() as cur:
            _rename_folder_path_in_hierarchy(cur, old_folder, new_folder)
            _rename_folder_path_in_table(cur, "steps", "folder_name", old_folder, new_folder)
            _rename_folder_path_in_table(cur, "locators", "folder_name", old_folder, new_folder)
            _rename_folder_path_in_table(cur, "data", "folder_name", old_folder, new_folder)
            _rename_folder_path_in_table(cur, "run_table", "folder_name", old_folder, new_folder)
            _rename_folder_path_in_table(cur, "session_meta", "folder_name", old_folder, new_folder)
            _rename_folder_path_in_table(cur, "project_folders", "folder_name", old_folder, new_folder)
            _register_project_folder(cur, new_folder)

    messages.success(request, f'Folder "{old_folder}" renamed to "{new_folder}".')
    return redirect("sessions_list")


@login_required
def folder_stats(request, folder_name=None):
    """GET: return JSON stats for all sessions in folder_name."""
    try:
        folder = _normalize_folder_path(folder_name or request.GET.get("folder_name", ""))
        if not folder:
            return JsonResponse({"error": "folder_name required"}, status=400)
        is_recordings = _is_recordings_folder_name(folder)
        db_folder = "Records" if is_recordings else folder

        with connection.cursor() as cur:
            # Distinct record_ids for this folder
            if is_recordings:
                _rids_al = list(_recordings_sql_aliases())
                _rids_ph = ",".join(["%s"] * len(_rids_al))
                cur.execute(
                    f"SELECT DISTINCT record_id FROM steps WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_rids_ph})",
                    _rids_al,
                )
            else:
                like_pattern = db_folder + "/%"
                cur.execute(
                    """
                    SELECT DISTINCT record_id FROM steps
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                    """,
                    [db_folder, like_pattern],
                )
            record_ids = [row[0] for row in cur.fetchall()]

            # Get readable names from session_meta
            if record_ids:
                placeholders = ",".join(["%s"] * len(record_ids))
                cur.execute(
                    f"SELECT record_id, record_name FROM session_meta WHERE record_id IN ({placeholders})",
                    record_ids,
                )
                name_map = {str(row[0]): row[1] for row in cur.fetchall()}
                sessions = [name_map.get(str(sid), str(sid)) for sid in record_ids]
                sessions = sorted(set(sessions))
            else:
                sessions = []

            if is_recordings:
                _fs_aliases = _recordings_sql_aliases()
                _fs_ph = ",".join(["%s"] * len(_fs_aliases))
                _fs_vals = list(_fs_aliases)
                cur.execute(f"SELECT COUNT(*) FROM run_table WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_fs_ph})", _fs_vals)
                runs = cur.fetchone()[0] or 0
                cur.execute(f"SELECT COUNT(*) FROM steps WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_fs_ph})", _fs_vals)
                steps = cur.fetchone()[0] or 0
                cur.execute(f"SELECT COUNT(*) FROM locators WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_fs_ph})", _fs_vals)
                locators = cur.fetchone()[0] or 0
                cur.execute(f"SELECT COUNT(*) FROM data WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_fs_ph})", _fs_vals)
                data_entries = cur.fetchone()[0] or 0
            else:
                cur.execute("""
                    SELECT COUNT(*) FROM run_table
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                runs = cur.fetchone()[0] or 0
                cur.execute("""
                    SELECT COUNT(*) FROM steps
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                steps = cur.fetchone()[0] or 0
                cur.execute("""
                    SELECT COUNT(*) FROM locators
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                locators = cur.fetchone()[0] or 0
                cur.execute("""
                    SELECT COUNT(*) FROM data
                    WHERE TRIM(COALESCE(folder_name, '')) = %s
                       OR TRIM(COALESCE(folder_name, '')) LIKE %s
                """, [db_folder, like_pattern])
                data_entries = cur.fetchone()[0] or 0

        return JsonResponse({
            "folder": folder,
            "sessions": sessions,
            "runs": runs,
            "steps": steps,
            "locators": locators,
            "data_entries": data_entries,
        })
    except Exception as exc:
        import traceback
        return JsonResponse({"error": str(exc), "trace": traceback.format_exc()}, status=500)


# ---------------------------------------------------------------------------
# Folder data entries — all test data for every test case in a project
# ---------------------------------------------------------------------------

@login_required
def folder_data_page(request):
    """Standalone page showing all test data for a project folder."""
    folder = _normalize_folder_path(request.GET.get("folder_name", ""))
    if not folder:
        return redirect("sessions_list")

    is_recordings = _is_recordings_folder_name(folder)
    db_folder = "Records" if is_recordings else folder

    entries = []
    try:
        with connection.cursor() as cur:
            if is_recordings:
                _al = list(_recordings_sql_aliases())
                _ph = ",".join(["%s"] * len(_al))
                cur.execute(f"""
                    SELECT d.id, d.record_id, d.step_no, d.field_name, d.value,
                           COALESCE(m.record_name, '') AS record_name,
                           COALESCE(d.is_global, FALSE) AS is_global,
                           COALESCE(d.formula, '') AS formula,
                           d.category, d.sub_category,
                           d.increment_value, d.increment_frequency,
                           d.decrement_value, d.decrement_frequency,
                           d.calculate_on, d.calculate_mode
                      FROM data d
                      LEFT JOIN session_meta m ON m.record_id = d.record_id
                     WHERE LOWER(TRIM(COALESCE(d.folder_name, ''))) IN ({_ph})
                     ORDER BY d.record_id, d.step_no
                """, _al)
            else:
                like_pattern = db_folder + "/%"
                cur.execute("""
                    SELECT d.id, d.record_id, d.step_no, d.field_name, d.value,
                           COALESCE(m.record_name, '') AS record_name,
                           COALESCE(d.is_global, FALSE) AS is_global,
                           COALESCE(d.formula, '') AS formula,
                           d.category, d.sub_category,
                           d.increment_value, d.increment_frequency,
                           d.decrement_value, d.decrement_frequency,
                           d.calculate_on, d.calculate_mode
                      FROM data d
                      LEFT JOIN session_meta m ON m.record_id = d.record_id
                     WHERE TRIM(COALESCE(d.folder_name, '')) = %s
                        OR TRIM(COALESCE(d.folder_name, '')) LIKE %s
                     ORDER BY d.record_id, d.step_no
                """, [db_folder, like_pattern])

            for row in cur.fetchall():
                entries.append({
                    "data_id": row[0],
                    "record_id": str(row[1]),
                    "step_no": row[2],
                    "field_name": row[3] or "",
                    "value": row[4] or "",
                    "record_name": row[5] or "",
                    "is_global": bool(row[6]),
                    "formula": row[7] or "",
                    "category": row[8] or "",
                    "sub_category": row[9] or "",
                    "increment_value": row[10],
                    "increment_frequency": row[11] or "",
                    "decrement_value": row[12],
                    "decrement_frequency": row[13] or "",
                    "calculate_on": row[14] or "",
                    "calculate_mode": row[15] or "",
                })
    except Exception:
        pass

    return render(request, "recorder/folder_data.html", {
        "folder_name": folder,
        "entries": entries,
        "total": len(entries),
    })


@login_required
def folder_data_entries(request):
    """GET ?folder_name=... — return all data rows for every session in the folder (recursive)."""
    folder = _normalize_folder_path(request.GET.get("folder_name", ""))
    if not folder:
        return JsonResponse({"error": "folder_name required"}, status=400)

    is_recordings = _is_recordings_folder_name(folder)
    db_folder = "Records" if is_recordings else folder

    try:
        with connection.cursor() as cur:
            if is_recordings:
                _al = list(_recordings_sql_aliases())
                _ph = ",".join(["%s"] * len(_al))
                cur.execute(f"""
                    SELECT d.id, d.record_id, d.step_no, d.field_name, d.value,
                           COALESCE(m.record_name, '') AS record_name,
                           COALESCE(d.folder_name, '') AS folder_name
                      FROM data d
                      LEFT JOIN session_meta m ON m.record_id = d.record_id
                     WHERE LOWER(TRIM(COALESCE(d.folder_name, ''))) IN ({_ph})
                     ORDER BY d.record_id, d.step_no
                """, _al)
            else:
                like_pattern = db_folder + "/%"
                cur.execute("""
                    SELECT d.id, d.record_id, d.step_no, d.field_name, d.value,
                           COALESCE(m.record_name, '') AS record_name,
                           COALESCE(d.folder_name, '') AS folder_name
                      FROM data d
                      LEFT JOIN session_meta m ON m.record_id = d.record_id
                     WHERE TRIM(COALESCE(d.folder_name, '')) = %s
                        OR TRIM(COALESCE(d.folder_name, '')) LIKE %s
                     ORDER BY d.record_id, d.step_no
                """, [db_folder, like_pattern])

            rows = cur.fetchall()

        entries = []
        for row in rows:
            entries.append({
                "data_id": row[0],
                "record_id": str(row[1]),
                "step_no": row[2],
                "field_name": row[3] or "",
                "value": row[4] or "",
                "record_name": row[5] or "",
                "folder_name": row[6] or "",
            })

        return JsonResponse({"ok": True, "folder": folder, "entries": entries, "total": len(entries)})
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)


# ---------------------------------------------------------------------------
# Reorder steps
# ---------------------------------------------------------------------------

@login_required
def reorder_steps(request, record_id):
    """POST JSON {"order": [step_no, ...], "folder_name": "..."} — reassigns step_no 1,2,3..."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload   = json.loads(request.body)
        order     = [int(x) for x in payload["order"]]  # original step_nos in new sequence
        folder    = (payload.get("folder_name") or "Baseline").strip()
    except (KeyError, ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    sid = str(record_id)
    OFFSET = 100_000

    is_recordings = _is_recordings_folder_name(folder)
    if is_recordings:
        _ro_aliases = list(_recordings_sql_aliases())
        _ro_ph = ",".join(["%s"] * len(_ro_aliases))
        fn_cond = f"LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_ro_ph})"
        fn_args = _ro_aliases
    else:
        fn_cond = "folder_name = %s"
        fn_args = [folder]

    all_tables = ["steps", "locators", "data", "run_table"]

    with transaction.atomic():
        with connection.cursor() as cur:
            # Phase 1 — shift step_nos out of range to avoid conflicts
            for tbl in all_tables:
                cur.execute(
                    f"UPDATE {tbl} SET step_no = step_no + %s WHERE record_id = %s AND {fn_cond}",
                    [OFFSET, sid] + fn_args,
                )
            # Phase 2 — reassign sequential step_nos
            for new_no, old_no in enumerate(order, start=1):
                original_in_db = old_no + OFFSET
                for tbl in all_tables:
                    cur.execute(
                        f"UPDATE {tbl} SET step_no = %s WHERE record_id = %s AND {fn_cond} AND step_no = %s",
                        [new_no, sid] + fn_args + [original_in_db],
                    )
    return JsonResponse({"ok": True, "total": len(order)})


# ---------------------------------------------------------------------------
# Update a data entry value
# ---------------------------------------------------------------------------

def _resolve_unique_data_field_name(cursor, field_name, exclude_data_id=None):
    candidate = (field_name or "").strip()
    if not candidate:
        return ""

    if exclude_data_id is None:
        cursor.execute(
            "SELECT field_name FROM data WHERE field_name IS NOT NULL AND field_name <> '' AND field_name LIKE %s",
            [f"{candidate}%"],
        )
    else:
        cursor.execute(
            "SELECT field_name FROM data WHERE field_name IS NOT NULL AND field_name <> '' AND id <> %s AND field_name LIKE %s",
            [exclude_data_id, f"{candidate}%"],
        )

    pattern = re.compile(rf"^{re.escape(candidate)}(?:-(\d+))?$")
    used_names = {
        str(row[0]).strip()
        for row in cursor.fetchall()
        if row and row[0] and pattern.fullmatch(str(row[0]).strip())
    }
    if candidate not in used_names:
        return candidate

    suffix = 2
    while True:
        numbered = f"{candidate}-{suffix}"
        if numbered not in used_names:
            return numbered
        suffix += 1

@login_required
def update_data_entry(request, record_id, step_no):
    """POST JSON updates a data row's field_name and/or value for this step."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload = json.loads(request.body)
        has_value = "value" in payload
        has_field_name = "field_name" in payload
        new_value = payload.get("value", "") if has_value else None
        new_field_name = payload.get("field_name", "") if has_field_name else None
        data_id = payload.get("data_id")
    except (ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    if not has_value and not has_field_name:
        return JsonResponse({"error": "No data changes provided"}, status=400)

    sid = str(record_id)
    resolved_name = ""
    resolved_value = ""

    with transaction.atomic():
        with connection.cursor() as cur:
            if data_id:
                cur.execute(
                    "SELECT id, field_name, value FROM data WHERE id = %s AND record_id = %s",
                    [data_id, sid],
                )
            else:
                cur.execute(
                    "SELECT id, field_name, value FROM data WHERE record_id = %s AND step_no = %s ORDER BY id LIMIT 1",
                    [sid, step_no],
                )

            row = cur.fetchone()
            if not row:
                return JsonResponse({"error": "No data row found"}, status=404)

            data_row_id, current_field_name, current_value = row
            resolved_value = new_value if has_value else (current_value or "")
            if has_field_name:
                requested_name = (new_field_name or "").strip()
                if requested_name:
                    # Check for duplicate name within the same record (session)
                    cur.execute(
                        "SELECT id FROM data WHERE field_name = %s AND record_id = %s AND id <> %s LIMIT 1",
                        [requested_name, sid, data_row_id],
                    )
                    if cur.fetchone():
                        return JsonResponse({
                            "ok": False,
                            "duplicate": True,
                            "message": "Name already taken",
                        })
                resolved_name = requested_name
            else:
                resolved_name = (current_field_name or "").strip()

            # When field_name changes, look up the matching data row to link via data_id
            _linked_data_id = data_row_id
            if has_field_name and resolved_name:
                # Find data entry with this field_name (prefer global, then any)
                cur.execute("""
                    SELECT id, value FROM data
                     WHERE field_name = %s AND id <> %s
                     ORDER BY is_global DESC, id
                     LIMIT 1
                """, [resolved_name, data_row_id])
                _match = cur.fetchone()
                if _match:
                    _linked_data_id = _match[0]
                    if _match[1] is not None:
                        resolved_value = _match[1]

            if not (has_field_name and _linked_data_id != data_row_id):
                cur.execute(
                    "UPDATE data SET field_name = %s, value = %s WHERE id = %s",
                    [resolved_name or None, resolved_value, data_row_id],
                )

            # Any name picked in Steps should be available as project-global test data.
            if has_field_name and resolved_name:
                cur.execute(
                    "SELECT folder_name FROM steps WHERE record_id = %s AND step_no = %s LIMIT 1",
                    [sid, step_no],
                )
                _step_folder_row = cur.fetchone()
                _step_folder = ((_step_folder_row[0] if _step_folder_row else None) or "").strip()
                if _step_folder:
                    _like = _step_folder + "/%"
                    cur.execute(
                        """
                        UPDATE data
                           SET is_global = TRUE
                         WHERE field_name = %s
                           AND (
                               TRIM(COALESCE(folder_name, '')) = %s
                               OR TRIM(COALESCE(folder_name, '')) LIKE %s
                           )
                        """,
                        [resolved_name, _step_folder, _like],
                    )
                else:
                    cur.execute(
                        "UPDATE data SET is_global = TRUE WHERE field_name = %s",
                        [resolved_name],
                    )

            cur.execute(
                """UPDATE steps
                      SET field_name      = %s,
                          field_value     = %s,
                          data_id         = %s,
                          last_updated_by = %s,
                          updated_at      = NOW()
                    WHERE record_id = %s AND step_no = %s""",
                [resolved_name or None, resolved_value, _linked_data_id,
                 request.user.username, sid, step_no],
            )

    # Regenerate playwright_code for this step so PW Code column stays in sync
    new_pw_code = ""
    _step = None
    try:
        _step = Step.objects.filter(record_id=sid, step_no=step_no).first()
        if _step:
            _loc = Locator.objects.filter(id=_step.locator_id).first() if _step.locator_id else None
            _step_dict = {
                "step_no": _step.step_no,
                "action": _step.action,
                "page_url": _step.page_url,
                "element_tag": _step.element_tag,
                "raw_event": _step.raw_event if isinstance(_step.raw_event, dict) else {},
                "locator_strategy": _loc.strategy if _loc else None,
                "locator_value": _loc.locator if _loc else None,
                "data_field_name": resolved_name,
                "field_name": resolved_name,
                "data_value": resolved_value,
                "field_value": resolved_value,
            }
            _pw_line = _pw_step_line(_step_dict)
            new_pw_code = _pw_line.strip() if _pw_line else ""
            with connection.cursor() as _cur2:
                _cur2.execute(
                    "UPDATE steps SET playwright_code = %s WHERE record_id = %s AND step_no = %s",
                    [new_pw_code, sid, step_no],
                )
    except Exception:
        pass

    # Regenerate steps_description when the data value changes (input/change actions)
    new_steps_description = ""
    try:
        if _step and _step.action in ("input", "change"):
            new_steps_description = f"User input recorded: '{resolved_value}'"
            with connection.cursor() as _cur3:
                _cur3.execute(
                    "UPDATE steps SET steps_description = %s WHERE record_id = %s AND step_no = %s",
                    [new_steps_description, sid, step_no],
                )
        else:
            new_steps_description = getattr(_step, "steps_description", None) or ""
    except Exception:
        pass

    return JsonResponse({
        "ok": True,
        "field_name": resolved_name,
        "field_value": resolved_value,
        "playwright_code": new_pw_code,
        "steps_description": new_steps_description,
    })


# ---------------------------------------------------------------------------
# Toggle is_global flag on a data entry
# ---------------------------------------------------------------------------
@login_required
def toggle_data_global(request, data_id):
    """POST — toggle the is_global flag for a data row."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        with connection.cursor() as cur:
            cur.execute(
                "UPDATE data SET is_global = NOT COALESCE(is_global, FALSE) WHERE id = %s RETURNING is_global",
                [data_id],
            )
            row = cur.fetchone()
            if not row:
                return JsonResponse({"error": "Not found"}, status=404)
            return JsonResponse({"ok": True, "is_global": bool(row[0])})
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)


# ---------------------------------------------------------------------------
# Add a new global data entry
# ---------------------------------------------------------------------------
@login_required
def add_data_entry(request):
    """POST JSON {field_name, value, folder_name, category, sub_category, ...} — create a new global data entry."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload = json.loads(request.body)
        field_name = (payload.get("field_name", "") or "").strip()
        value = (payload.get("value", "") or "").strip()
        folder_name = (payload.get("folder_name", "") or "").strip()
        increment_value = payload.get("increment_value") or None
        increment_frequency = (payload.get("increment_frequency") or "").strip() or None
        decrement_value = payload.get("decrement_value") or None
        decrement_frequency = (payload.get("decrement_frequency") or "").strip() or None
        calculate_on = (payload.get("calculate_on") or "").strip() or None
        calculate_mode = (payload.get("calculate_mode") or "").strip() or None
    except (ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    if not field_name:
        return JsonResponse({"ok": False, "error": "Name is required"})

    # Check for duplicate name in the same folder
    try:
        with connection.cursor() as cur:
            if folder_name:
                cur.execute("""
                    SELECT id FROM data
                     WHERE field_name = %s
                       AND (TRIM(COALESCE(folder_name, '')) = %s
                            OR TRIM(COALESCE(folder_name, '')) LIKE %s)
                     LIMIT 1
                """, [field_name, folder_name, folder_name + "/%"])
            else:
                cur.execute(
                    "SELECT id FROM data WHERE field_name = %s LIMIT 1",
                    [field_name],
                )
            if cur.fetchone():
                return JsonResponse({"ok": False, "error": "Name already taken"})

            # Insert new entry with is_global = TRUE
            cur.execute("""
                INSERT INTO data (record_id, step_no, field_name, value, folder_name, is_global,
                                  increment_value, increment_frequency,
                                  decrement_value, decrement_frequency,
                                  calculate_on, calculate_mode, created_at)
                VALUES (%s, 0, %s, %s, %s, TRUE, %s, %s, %s, %s, %s, %s, NOW())
                RETURNING id
            """, [
                "00000000-0000-0000-0000-000000000000",
                field_name,
                value or None,
                folder_name or None,
                increment_value,
                increment_frequency,
                decrement_value,
                decrement_frequency,
                calculate_on,
                calculate_mode,
            ])
            new_id = cur.fetchone()[0]
        return JsonResponse({"ok": True, "data_id": new_id})
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)


# ---------------------------------------------------------------------------
# Delete data entries (bulk)
# ---------------------------------------------------------------------------
@login_required
def delete_data_entries(request):
    """POST JSON {ids: [1,2,...]} — delete data entries by id."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload = json.loads(request.body)
        ids = payload.get("ids", [])
    except (ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    if not ids or not isinstance(ids, list):
        return JsonResponse({"ok": False, "error": "No entries specified"})

    # Sanitize: only allow ints
    safe_ids = [int(i) for i in ids if str(i).isdigit() or isinstance(i, int)]
    if not safe_ids:
        return JsonResponse({"ok": False, "error": "Invalid IDs"})

    try:
        with connection.cursor() as cur:
            placeholders = ",".join(["%s"] * len(safe_ids))
            cur.execute(
                f"DELETE FROM data WHERE id IN ({placeholders})",
                safe_ids,
            )
        return JsonResponse({"ok": True, "deleted": len(safe_ids)})
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)


# ---------------------------------------------------------------------------
# Update formula for a data entry
# ---------------------------------------------------------------------------
@login_required
def update_data_formula(request, data_id):
    """POST JSON {formula, increment_value, increment_frequency, decrement_value, decrement_frequency, calculate_on, calculate_mode}."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload = json.loads(request.body)
        formula = (payload.get("formula", "") or "").strip()
        increment_value = payload.get("increment_value") or None
        increment_frequency = (payload.get("increment_frequency") or "").strip() or None
        decrement_value = payload.get("decrement_value") or None
        decrement_frequency = (payload.get("decrement_frequency") or "").strip() or None
        calculate_on = (payload.get("calculate_on") or "").strip() or None
        calculate_mode = (payload.get("calculate_mode") or "").strip() or None
    except (ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    # If formula is empty/cleared, just save metadata and return
    if not formula:
        try:
            with connection.cursor() as cur:
                cur.execute(
                    """UPDATE data SET formula = NULL,
                       increment_value = %s, increment_frequency = %s,
                       decrement_value = %s, decrement_frequency = %s,
                       calculate_on = %s, calculate_mode = %s
                     WHERE id = %s RETURNING id""",
                    [increment_value, increment_frequency,
                     decrement_value, decrement_frequency,
                     calculate_on, calculate_mode, data_id],
                )
                if not cur.fetchone():
                    return JsonResponse({"error": "Not found"}, status=404)
            return JsonResponse({"ok": True, "formula": ""})
        except Exception as exc:
            return JsonResponse({"error": str(exc)}, status=500)

    # Formula must start with =
    if not formula.startswith("="):
        return JsonResponse({"ok": False, "error": "Formula must start with ="})

    expression = formula[1:].strip()  # Remove leading =
    if not expression:
        return JsonResponse({"ok": False, "error": "Formula expression is empty"})

    # Find the folder_name and current value for this data entry
    try:
        with connection.cursor() as cur:
            cur.execute(
                "SELECT folder_name, field_name, value FROM data WHERE id = %s", [data_id]
            )
            row = cur.fetchone()
            if not row:
                return JsonResponse({"error": "Not found"}, status=404)
            entry_folder = (row[0] or "").strip()
            entry_own_name = (row[1] or "").strip()
            entry_own_value = row[2]
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)

    # Extract referenced names from expression (identifiers not part of operators/numbers)
    import re as _re
    # Tokenize: find all word sequences that could be variable names
    tokens = _re.findall(r'[A-Za-z_]\w*', expression)
    referenced_names = list(set(tokens))

    # Look up values for referenced names in the same folder
    # Pre-seed with self value if referenced (one-shot compute, not circular)
    name_values = {}
    if entry_own_name and entry_own_name in referenced_names and entry_own_value is not None:
        name_values[entry_own_name] = entry_own_value
    if referenced_names:
        try:
            with connection.cursor() as cur:
                _ph = ",".join(["%s"] * len(referenced_names))
                # Get values for referenced field_names in the same folder
                if entry_folder:
                    cur.execute(f"""
                        SELECT field_name, value FROM data
                         WHERE field_name IN ({_ph})
                           AND (TRIM(COALESCE(folder_name, '')) = %s
                                OR TRIM(COALESCE(folder_name, '')) LIKE %s)
                         ORDER BY id
                    """, referenced_names + [entry_folder, entry_folder + "/%"])
                else:
                    cur.execute(f"""
                        SELECT field_name, value FROM data
                         WHERE field_name IN ({_ph})
                         ORDER BY id
                    """, referenced_names)
                for fname, fval in cur.fetchall():
                    if fname and fname not in name_values:
                        name_values[fname] = fval
        except Exception as exc:
            return JsonResponse({"error": str(exc)}, status=500)

    # Validate that all referenced names have numerical values
    errors = []
    resolved_values = {}
    for name in referenced_names:
        if name not in name_values:
            errors.append(f"'{name}' not found in data")
        else:
            raw = name_values[name]
            try:
                resolved_values[name] = float(raw)
            except (TypeError, ValueError):
                errors.append(f"'{name}' value '{raw}' is not numerical")

    if errors:
        return JsonResponse({"ok": False, "error": "; ".join(errors)})

    # Safely evaluate the expression by substituting names with their values
    # Build a safe expression with only numbers and operators
    safe_expr = expression
    # Sort names by length descending to avoid partial replacements
    for name in sorted(resolved_values.keys(), key=len, reverse=True):
        safe_expr = safe_expr.replace(name, str(resolved_values[name]))

    # Convert ^ to ** (power operator)
    safe_expr = safe_expr.replace("^", "**")

    # Convert percentage notation: e.g. 5% → (5/100)
    safe_expr = _re.sub(r'([\d\.]+)\s*%', r'(\1/100)', safe_expr)

    # Validate the expression only contains safe characters (numbers, operators, parens, dots, spaces)
    if not _re.match(r'^[\d\s\+\-\*\/\(\)\.]+$', safe_expr):
        return JsonResponse({"ok": False, "error": f"Invalid characters in expression: {safe_expr}"})

    # Compute
    try:
        computed = eval(safe_expr)  # noqa: S307 - expression is sanitized above
        computed_value = str(round(computed, 10) if isinstance(computed, float) else computed)
    except ZeroDivisionError:
        return JsonResponse({"ok": False, "error": "Division by zero"})
    except Exception as exc:
        return JsonResponse({"ok": False, "error": f"Computation error: {exc}"})

    # Save formula, computed value, and metadata
    try:
        with connection.cursor() as cur:
            cur.execute(
                """UPDATE data SET formula = %s, value = %s,
                   increment_value = %s, increment_frequency = %s,
                   decrement_value = %s, decrement_frequency = %s,
                   calculate_on = %s, calculate_mode = %s
                 WHERE id = %s RETURNING id""",
                [formula, computed_value,
                 increment_value, increment_frequency, decrement_value,
                 decrement_frequency, calculate_on, calculate_mode, data_id],
            )
            if not cur.fetchone():
                return JsonResponse({"error": "Not found"}, status=404)
        return JsonResponse({"ok": True, "formula": formula, "computed_value": computed_value})
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)


# ---------------------------------------------------------------------------
# Pre-run formula computation: compute all formulas in a folder before replay
# ---------------------------------------------------------------------------
def compute_folder_formulas(folder_name):
    """
    Recompute all data entries that have a formula, increment, or decrement
    in the given folder.  Called before each replay run so that formula-driven
    values (increment_value, decrement_value, calculate_on) are up-to-date.

    Returns a list of {id, field_name, new_value, error} dicts.
    """
    import re as _re
    from datetime import date, timedelta

    if not folder_name:
        return []

    is_recordings = _is_recordings_folder_name(folder_name)
    db_folder = "Records" if is_recordings else folder_name

    # Load all data entries in this folder
    entries = []
    try:
        with connection.cursor() as cur:
            if is_recordings:
                _al = list(_recordings_sql_aliases())
                _ph = ",".join(["%s"] * len(_al))
                cur.execute(f"""
                    SELECT id, field_name, value, formula,
                           increment_value, increment_frequency,
                           decrement_value, decrement_frequency,
                           calculate_on, calculate_mode
                      FROM data
                     WHERE LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_ph})
                     ORDER BY id
                """, _al)
            else:
                like_pattern = db_folder + "/%"
                cur.execute("""
                    SELECT id, field_name, value, formula,
                           increment_value, increment_frequency,
                           decrement_value, decrement_frequency,
                           calculate_on, calculate_mode
                      FROM data
                     WHERE TRIM(COALESCE(folder_name, '')) = %s
                        OR TRIM(COALESCE(folder_name, '')) LIKE %s
                     ORDER BY id
                """, [db_folder, like_pattern])
            for row in cur.fetchall():
                entries.append({
                    "id": row[0],
                    "field_name": row[1] or "",
                    "value": row[2] or "",
                    "formula": row[3] or "",
                    "increment_value": row[4],
                    "increment_frequency": row[5] or "",
                    "decrement_value": row[6],
                    "decrement_frequency": row[7] or "",
                    "calculate_on": row[8] or "",
                    "calculate_mode": row[9] or "",
                })
    except Exception:
        return []

    if not entries:
        return []

    # Build name→value map for formula resolution
    name_values = {}
    for e in entries:
        if e["field_name"]:
            name_values[e["field_name"]] = e["value"]

    results = []
    today = date.today()

    for e in entries:
        new_value = e["value"]
        error = None
        changed = False

        # --- Apply increment ---
        if e["increment_value"] is not None and e["increment_frequency"]:
            try:
                current = float(new_value) if new_value else 0
                inc = float(e["increment_value"])
                new_value = str(round(current + inc, 10)).rstrip('0').rstrip('.')
                changed = True
            except (TypeError, ValueError):
                error = f"Cannot increment non-numeric value '{new_value}'"

        # --- Apply decrement ---
        if e["decrement_value"] is not None and e["decrement_frequency"] and not error:
            try:
                current = float(new_value) if new_value else 0
                dec = float(e["decrement_value"])
                new_value = str(round(current - dec, 10)).rstrip('0').rstrip('.')
                changed = True
            except (TypeError, ValueError):
                error = f"Cannot decrement non-numeric value '{new_value}'"

        # --- Apply calculate_on with calculate_mode ---
        if e["calculate_on"] and e["calculate_mode"] and not error:
            try:
                target_date = date.fromisoformat(e["calculate_on"])
                mode = e["calculate_mode"]
                if mode == "On":
                    new_value = e["calculate_on"] if today == target_date else new_value
                elif mode == "Not on":
                    new_value = e["calculate_on"] if today != target_date else new_value
                elif mode == "Before":
                    new_value = e["calculate_on"] if today < target_date else new_value
                elif mode == "After":
                    new_value = e["calculate_on"] if today > target_date else new_value
                elif mode == "On or before":
                    new_value = e["calculate_on"] if today <= target_date else new_value
                elif mode == "On or after":
                    new_value = e["calculate_on"] if today >= target_date else new_value
                elif mode == "Today":
                    new_value = str(today)
                    changed = True
                elif mode == "Yesterday":
                    new_value = str(today - timedelta(days=1))
                    changed = True
                elif mode == "Last 7 days":
                    new_value = str(today - timedelta(days=7))
                    changed = True
                elif mode == "Last 30 days":
                    new_value = str(today - timedelta(days=30))
                    changed = True
                elif mode == "This month":
                    new_value = str(today.replace(day=1))
                    changed = True
                elif mode == "This year":
                    new_value = str(today.replace(month=1, day=1))
                    changed = True
                if new_value != e["value"]:
                    changed = True
            except (ValueError, TypeError):
                pass  # invalid date format — skip

        # --- Compute formula ---
        if e["formula"] and e["formula"].startswith("=") and not error:
            expression = e["formula"][1:].strip()
            if expression:
                # Update name_values with latest computed value from this pass
                if e["field_name"] and changed:
                    name_values[e["field_name"]] = new_value

                tokens = _re.findall(r'[A-Za-z_]\w*', expression)
                referenced_names = list(set(tokens))

                # Pre-seed with self value
                local_values = {}
                if e["field_name"] and e["field_name"] in referenced_names:
                    local_values[e["field_name"]] = new_value

                for name in referenced_names:
                    if name not in local_values and name in name_values:
                        local_values[name] = name_values[name]

                # Check all references resolved
                missing = [n for n in referenced_names if n not in local_values]
                if missing:
                    error = f"Unresolved: {', '.join(missing)}"
                else:
                    # Substitute and evaluate
                    resolved = {}
                    bad = False
                    for name in referenced_names:
                        try:
                            resolved[name] = float(local_values[name])
                        except (TypeError, ValueError):
                            error = f"'{name}' value '{local_values[name]}' is not numeric"
                            bad = True
                            break
                    if not bad:
                        safe_expr = expression
                        for name in sorted(resolved.keys(), key=len, reverse=True):
                            safe_expr = safe_expr.replace(name, str(resolved[name]))
                        safe_expr = safe_expr.replace("^", "**")
                        safe_expr = _re.sub(r'([\d\.]+)\s*%', r'(\1/100)', safe_expr)
                        if _re.match(r'^[\d\s\+\-\*\/\(\)\.]+$', safe_expr):
                            try:
                                computed = eval(safe_expr)  # noqa: S307
                                new_value = str(round(computed, 10) if isinstance(computed, float) else computed).rstrip('0').rstrip('.')
                                changed = True
                            except ZeroDivisionError:
                                error = "Division by zero"
                            except Exception as exc:
                                error = f"Eval error: {exc}"
                        else:
                            error = f"Invalid expression: {safe_expr}"

        # Update DB if value changed
        if changed and not error and new_value != e["value"]:
            try:
                with connection.cursor() as cur:
                    cur.execute(
                        "UPDATE data SET value = %s WHERE id = %s",
                        [new_value, e["id"]],
                    )
                # Update local map for subsequent formula lookups
                if e["field_name"]:
                    name_values[e["field_name"]] = new_value
            except Exception as exc:
                error = str(exc)

        results.append({
            "id": e["id"],
            "field_name": e["field_name"],
            "new_value": new_value,
            "error": error,
        })

    return results


# ---------------------------------------------------------------------------
# Get all global data names for a folder (for autocomplete)
# ---------------------------------------------------------------------------
@login_required
def global_data_names(request):
    """GET ?folder_name=X — returns distinct global field_names in the folder."""
    folder = _normalize_folder_path(request.GET.get("folder_name", ""))
    if not folder:
        return JsonResponse({"names": []})

    is_recordings = _is_recordings_folder_name(folder)
    db_folder = "Records" if is_recordings else folder

    names = []
    try:
        with connection.cursor() as cur:
            if is_recordings:
                _al = list(_recordings_sql_aliases())
                _ph = ",".join(["%s"] * len(_al))
                cur.execute(f"""
                    SELECT DISTINCT field_name FROM data
                     WHERE is_global = TRUE
                       AND field_name IS NOT NULL AND field_name <> ''
                       AND LOWER(TRIM(COALESCE(folder_name, ''))) IN ({_ph})
                     ORDER BY field_name
                """, _al)
            else:
                like_pattern = db_folder + "/%"
                cur.execute("""
                    SELECT DISTINCT field_name FROM data
                     WHERE is_global = TRUE
                       AND field_name IS NOT NULL AND field_name <> ''
                       AND (TRIM(COALESCE(folder_name, '')) = %s
                            OR TRIM(COALESCE(folder_name, '')) LIKE %s)
                     ORDER BY field_name
                """, [db_folder, like_pattern])
            names = [r[0] for r in cur.fetchall()]
    except Exception:
        pass
    return JsonResponse({"names": names})


# ---------------------------------------------------------------------------
# Update validation text for a step
# ---------------------------------------------------------------------------

@login_required
def update_validation(request, record_id, step_no):
    """POST JSON {"validation": "..."} — saves the validation text for this step.
    Empty string clears the validation (no check will run during replay).
    """
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload = json.loads(request.body)
        validation = payload.get("validation", "") or ""
    except (ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    sid = str(record_id)
    validation_value = validation.strip() or None

    with connection.cursor() as cur:
        cur.execute(
            """UPDATE steps
                  SET validation      = %s,
                      last_updated_by = %s,
                      updated_at      = NOW()
                WHERE record_id = %s AND step_no = %s""",
            [validation_value, request.user.username, sid, step_no],
        )
        cur.execute(
            """UPDATE run_table
                  SET validation      = %s,
                      updated_at      = NOW()
                WHERE record_id = %s AND step_no = %s""",
            [validation_value, sid, step_no],
        )

    return JsonResponse({"ok": True})


# ---------------------------------------------------------------------------
# Update a locator value
# ---------------------------------------------------------------------------

@login_required
def set_primary_locator(request, record_id, step_no):
    """POST JSON {"strategy": "xpath"} — sets is_primary=TRUE for that strategy on (record_id, step_no)."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload  = json.loads(request.body)
        strategy = payload.get("strategy", "").strip()
    except (json.JSONDecodeError, KeyError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    if not strategy:
        return JsonResponse({"error": "strategy required"}, status=400)
    _sid = str(record_id)
    _step = Step.objects.filter(record_id=_sid, step_no=step_no).first()
    if not _step:
        return JsonResponse({"error": "Step not found"}, status=404)
    with connection.cursor() as cur:
        cur.execute(
            "UPDATE locators SET is_primary = FALSE WHERE record_id = %s AND step_no = %s",
            [_sid, step_no],
        )
        cur.execute(
            "UPDATE locators SET is_primary = TRUE  WHERE record_id = %s AND step_no = %s AND strategy = %s",
            [_sid, step_no, strategy],
        )
        cur.execute(
            "SELECT id, locator, locator_rank FROM locators WHERE record_id = %s AND step_no = %s AND strategy = %s ORDER BY id LIMIT 1",
            [_sid, step_no, strategy],
        )
        row = cur.fetchone()
        if not row:
            synthetic_locator, synthetic_rank = _synthetic_step_locator(_step, strategy)
            if not synthetic_locator:
                return JsonResponse({"error": "Strategy not found"}, status=404)
            cur.execute(
                """
                INSERT INTO locators (record_id, step_no, strategy, locator, is_primary, locator_rank, pos_x, pos_y, folder_name, engine)
                VALUES (%s, %s, %s, %s, TRUE, %s, %s, %s, %s, %s)
                RETURNING id, locator, locator_rank
                """,
                [
                    _sid,
                    step_no,
                    strategy,
                    synthetic_locator,
                    synthetic_rank,
                    getattr(_step, "pos_x", None),
                    getattr(_step, "pos_y", None),
                    getattr(_step, "folder_name", None),
                    getattr(_step, "engine", None) or "selenium",
                ],
            )
            row = cur.fetchone()
        new_locator_id, new_locator_val, _rank = row
        # Update steps (denormalised columns)
        cur.execute(
            """
            UPDATE steps
               SET strategy     = %s,
                   locator      = %s,
                   is_primary   = TRUE,
                   locator_rank = %s,
                   locator_id   = %s
             WHERE record_id = %s AND step_no = %s
            """,
            [strategy, new_locator_val, _rank, new_locator_id, _sid, step_no],
        )
        # Keep recordings table in sync
        cur.execute(
            """
            UPDATE recordings
               SET strategy     = %s,
                   locator      = %s,
                   is_primary   = TRUE,
                   locator_rank = %s,
                   locator_id   = %s
             WHERE record_id = %s AND step_no = %s
            """,
            [strategy, new_locator_val, _rank, new_locator_id, _sid, step_no],
        )

    # Regenerate playwright_code for this step now that the primary locator changed.
    # Strip raw_event.locators so _pw_locator falls through to the explicit
    # locator_strategy/locator_value fallback — otherwise it would re-derive
    # the locator from auto-priority (id > ariaLabel > …) ignoring the user's choice.
    new_pw_code = ""
    try:
        _step = Step.objects.filter(record_id=_sid, step_no=step_no).first()
        if _step:
            _raw_ev = dict(_step.raw_event) if isinstance(_step.raw_event, dict) else {}
            _raw_ev.pop("locators", None)   # remove auto-detected locs
            _data = DataEntry.objects.filter(id=_step.data_id).first() if _step.data_id else None
            # field_value is a DB column not declared on the Django model; use getattr to avoid AttributeError
            _field_value = getattr(_step, "field_value", None) or (_data.value if _data else None)
            _step_dict = {
                "step_no":          _step.step_no,
                "action":           _step.action,
                "page_url":         _step.page_url,
                "element_tag":      _step.element_tag,
                "raw_event":        _raw_ev,
                "locators_raw":     {},          # suppress any stored raw locs
                "locator_strategy": strategy,
                "locator_value":    new_locator_val,
                "data_value":       _data.value if _data else None,
                "field_value":      _field_value,
            }
            _pw_line = _pw_step_line(_step_dict)
            new_pw_code = _pw_line.strip() if _pw_line else ""
            with connection.cursor() as _cur2:
                _cur2.execute(
                    "UPDATE steps SET playwright_code = %s WHERE record_id = %s AND step_no = %s",
                    [new_pw_code, _sid, step_no],
                )
    except Exception as _exc:
        import logging as _logging
        _logging.getLogger(__name__).exception("set_primary_locator: pw_code regen failed: %s", _exc)

    return JsonResponse({"ok": True, "locator_id": new_locator_id, "locator": new_locator_val, "playwright_code": new_pw_code})


def update_locator(request, record_id, step_no):
    """POST JSON {"locator_id": 123, "value": "..."} — updates locators.locator by id."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        payload     = json.loads(request.body)
        loc_id      = int(payload["locator_id"])
        new_value   = payload.get("value", "")
    except (KeyError, ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    with connection.cursor() as cur:
        cur.execute(
            "UPDATE locators SET locator = %s WHERE id = %s AND record_id = %s",
            [new_value, loc_id, str(record_id)],
        )
        if cur.rowcount == 0:
            return JsonResponse({"error": "No locator row found"}, status=404)

        # Sync steps.locator for this step if this locator is the primary one
        cur.execute(
            """
            UPDATE steps s
               SET locator = %s
             WHERE s.record_id = %s
               AND s.step_no   = %s
               AND s.locator_id = %s
            """,
            [new_value, str(record_id), step_no, loc_id],
        )

    return JsonResponse({"ok": True})


# ---------------------------------------------------------------------------
# Delete a single step
# ---------------------------------------------------------------------------

@login_required
def delete_step(request, record_id, step_no):
    if request.method != "POST":
        return redirect("session_steps", record_id=record_id)
    with transaction.atomic():
        with connection.cursor() as cur:
            _sid = str(record_id)

            # run_table rows reference locators/data — remove them first
            cur.execute(
                "DELETE FROM run_table WHERE record_id = %s AND step_no = %s",
                [_sid, step_no],
            )
            cur.execute(
                "DELETE FROM steps WHERE record_id = %s AND step_no = %s",
                [_sid, step_no],
            )
            if cur.rowcount == 0:
                # Step not found (stale step_no) — bail without deleting anything
                if request.headers.get("Accept") == "application/json":
                    return JsonResponse({"ok": False, "error": "Step not found"}, status=404)
                messages.error(request, f"Step {step_no} not found.")
                return redirect("session_steps", record_id=record_id)

            # Delete all locators and data for this (record_id, step_no)
            cur.execute(
                "DELETE FROM locators WHERE record_id = %s AND step_no = %s",
                [_sid, step_no],
            )
            cur.execute(
                "DELETE FROM data WHERE record_id = %s AND step_no = %s",
                [_sid, step_no],
            )

            # Resequence remaining step_nos (1, 2, 3, ...) after the gap left by deletion
            cur.execute(
                "SELECT step_no FROM steps WHERE record_id = %s ORDER BY step_no",
                [_sid],
            )
            remaining = [r[0] for r in cur.fetchall()]
            _OFFSET = 100_000
            _all_tbls = ["steps", "locators", "data", "run_table"]
            for _tbl in _all_tbls:
                cur.execute(
                    f"UPDATE {_tbl} SET step_no = step_no + %s WHERE record_id = %s",
                    [_OFFSET, _sid],
                )
            for new_no, old_no in enumerate(remaining, start=1):
                for _tbl in _all_tbls:
                    cur.execute(
                        f"UPDATE {_tbl} SET step_no = %s WHERE record_id = %s AND step_no = %s",
                        [new_no, _sid, old_no + _OFFSET],
                    )

    if request.headers.get("Accept") == "application/json":
        return JsonResponse({"ok": True})
    messages.success(request, f"Step {step_no} deleted.")
    return redirect("session_steps", record_id=record_id)


@login_required
def bulk_delete_steps(request, record_id):
    """Delete multiple steps in one transaction and resequence."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)
    try:
        body = json.loads(request.body)
        step_nos = [int(n) for n in body.get("step_nos", [])]
    except (json.JSONDecodeError, ValueError, TypeError):
        return JsonResponse({"ok": False, "error": "Invalid request"}, status=400)
    if not step_nos:
        return JsonResponse({"ok": False, "error": "No steps specified"}, status=400)

    with transaction.atomic():
        with connection.cursor() as cur:
            _sid = str(record_id)
            placeholders = ",".join(["%s"] * len(step_nos))
            args_list = [_sid] + step_nos

            cur.execute(
                f"DELETE FROM run_table WHERE record_id = %s AND step_no IN ({placeholders})",
                args_list,
            )
            cur.execute(
                f"DELETE FROM steps WHERE record_id = %s AND step_no IN ({placeholders})",
                args_list,
            )
            cur.execute(
                f"DELETE FROM locators WHERE record_id = %s AND step_no IN ({placeholders})",
                args_list,
            )
            cur.execute(
                f"DELETE FROM data WHERE record_id = %s AND step_no IN ({placeholders})",
                args_list,
            )

            # Resequence remaining steps
            cur.execute(
                "SELECT step_no FROM steps WHERE record_id = %s ORDER BY step_no",
                [_sid],
            )
            remaining = [r[0] for r in cur.fetchall()]
            _OFFSET = 100_000
            _all_tbls = ["steps", "locators", "data", "run_table"]
            for _tbl in _all_tbls:
                cur.execute(
                    f"UPDATE {_tbl} SET step_no = step_no + %s WHERE record_id = %s",
                    [_OFFSET, _sid],
                )
            for new_no, old_no in enumerate(remaining, start=1):
                for _tbl in _all_tbls:
                    cur.execute(
                        f"UPDATE {_tbl} SET step_no = %s WHERE record_id = %s AND step_no = %s",
                        [new_no, _sid, old_no + _OFFSET],
                    )

    return JsonResponse({"ok": True, "deleted": len(step_nos)})


# ---------------------------------------------------------------------------
# Live-recording helpers (steps since N, stop-ajax)
# ---------------------------------------------------------------------------

def steps_since_api(request, record_id, since):
    """GET /api/sessions/<uuid>/steps/since/<int>/
    Returns steps with step_no > since as JSON for live-polling the table
    while a recording is in progress (Add Step feature).
    Joins the locators table (primary locator) and the data table so the
    client can render fully-functional editable rows.
    Accepts Django session auth OR X-Monitor-Token header for the detached monitor.
    """
    if not request.user.is_authenticated:
        user = _resolve_monitor_user(request)
        if user is None:
            return JsonResponse({"error": "Unauthorized"}, status=401)
    sid = str(record_id)
    with connection.cursor() as cur:
        # Main row: steps + primary locator + data entry
        cur.execute("""
            SELECT DISTINCT ON (s.step_no)
                   s.step_no, s.action, s.page_url, s.element_tag,
                   COALESCE(s.field_name,  d.field_name) AS field_name,
                   COALESCE(s.field_value, d.value)      AS field_value,
                   s.recorder, s.runner,
                   COALESCE(s.strategy, l.strategy)      AS strategy,
                   COALESCE(s.locator,  l.locator)       AS locator,
                   l.id                                  AS locator_id,
                   d.id                                  AS data_id,
                 s.created_at, s.raw_event,
                 COALESCE(s.validation, '')            AS validation,
                 COALESCE(s.steps_description, '')     AS steps_description
            FROM steps s
            LEFT JOIN locators l ON  l.record_id = s.record_id
                                 AND l.step_no   = s.step_no
                                 AND l.is_primary = TRUE
            LEFT JOIN data d     ON  d.record_id = s.record_id
                                 AND d.step_no   = s.step_no
            WHERE s.record_id = %s AND s.step_no > %s
            ORDER BY s.step_no, s.id
        """, [sid, since])
        cols = [c.name for c in cur.description]
        rows_raw = list(cur.fetchall())

        # All strategies per newly found step_no (for the strategy dropdown)
        step_nos = [row[0] for row in rows_raw]
        strats_by_step: dict = {}
        if step_nos:
            ph = ",".join(["%s"] * len(step_nos))
            cur.execute(f"""
                SELECT DISTINCT ON (step_no, strategy)
                       step_no, strategy, is_primary, id, locator
                FROM locators
                WHERE record_id = %s AND step_no IN ({ph})
                ORDER BY step_no, strategy, is_primary DESC
            """, [sid] + step_nos)
            for sno, strat, isprim, lid, loc_val in cur.fetchall():
                strats_by_step.setdefault(sno, []).append({
                    "strategy":   strat,
                    "is_primary": isprim,
                    "locator_id": lid,
                    "locator":    loc_val,
                })

        rows = []
        for row in rows_raw:
            d = dict(zip(cols, row))
            if d["created_at"]:
                d["created_at"] = d["created_at"].isoformat()
            d["strategies"] = strats_by_step.get(d["step_no"], [])
            rows.append(d)

    # ── Check if the recording process has died and surface any error ─────
    recording_error = None
    recording_stopped = False
    _ses_id  = request.session.get("recording_id", "")
    _ses_pid = request.session.get("recording_pid")
    if _ses_id == sid and _ses_pid and not _is_pid_alive(_ses_pid):
        recording_stopped = True
        request.session.pop("recording_pid",  None)
        request.session.pop("recording_id",   None)
        # Only report an error if the recording was NOT intentionally stopped
        # (i.e. it's still in _ACTIVE_RECORDING — meaning it crashed, not user-stopped).
        with _ACTIVE_RECORDING_LOCK:
            _was_still_active = sid in _ACTIVE_RECORDING
            _ACTIVE_RECORDING.pop(sid, None)
        if _was_still_active:
            # Process died on its own (crash) — read last log line as error
            _log_path = os.path.join(_RECORDING_LOGS_DIR, f"rec_{sid}_append.log")
            try:
                with open(_log_path, "r", encoding="utf-8", errors="replace") as _lf:
                    _tail = _lf.read()[-4096:]   # last 4 KB is plenty
                for _line in reversed(_tail.splitlines()):
                    _line = _line.strip()
                    if _line and not _line.startswith("[pool]"):
                        recording_error = _line
                        break
            except OSError:
                pass

    # Fallback for monitor windows (which have no Django session):
    # If the recording is NOT in _ACTIVE_RECORDING and the PID temp file is
    # gone or the PID is dead, consider it stopped.
    if not recording_stopped and _ses_id != sid:
        with _ACTIVE_RECORDING_LOCK:
            _still_active = sid in _ACTIVE_RECORDING
        if not _still_active:
            import tempfile as _tf_poll
            _pid_file_poll = os.path.join(_tf_poll.gettempdir(), f"recorder_pid_{sid}.txt")
            _poll_pid = None
            try:
                with open(_pid_file_poll, "r", encoding="utf-8") as _pf:
                    _poll_pid = int(_pf.readline().strip())
            except (OSError, ValueError):
                pass
            if _poll_pid is None or not _is_pid_alive(_poll_pid):
                recording_stopped = True

    response = JsonResponse({
        "steps": rows,
        "recording_stopped": recording_stopped,
        "recording_error":   recording_error,
    })
    response["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response["Pragma"] = "no-cache"
    response["Expires"] = "0"
    return response


def _cdp_evaluate(rdp_port: int, js: str) -> bool:
    """Broadcast a Runtime.evaluate call to ALL Chrome page tabs via CDP.

    Targeting every tab is essential because Selenium's current window (used
    by main.py's pop_events) may be a different tab than the one the user is
    actively browsing — injecting into only the first tab silently misses the
    tab that actually has RECORDER_JS buffering events.
    """
    import urllib.request, json as _json
    try:
        with urllib.request.urlopen(f"http://127.0.0.1:{rdp_port}/json", timeout=3) as r:
            tabs = _json.loads(r.read())
        page_tabs = [t for t in tabs if t.get("type") == "page" and t.get("webSocketDebuggerUrl")]
        if not page_tabs:
            return False
        import websocket as _ws
        ok = False
        for tab in page_tabs:
            try:
                conn = _ws.create_connection(tab["webSocketDebuggerUrl"], timeout=5)
                conn.send(_json.dumps({"id": 1, "method": "Runtime.evaluate", "params": {"expression": js}}))
                conn.recv()
                conn.close()
                ok = True
            except Exception:
                pass
        return ok
    except Exception:
        return False


@login_required
def pause_recording_ajax(request):
    """POST /record/pause-ajax/ — disable the in-browser recorder so no new events are buffered."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)
    _rec_pid = request.session.get("recording_pid")
    if not _rec_pid:
        return JsonResponse({"ok": False, "error": "No active recording"})
    # Python-level flag: keyed on record_id (UUID) so both Django and main.py
    # agree on exactly the same path regardless of tempdir quirks.
    _rec_id = request.session.get("recording_id", "")
    import tempfile as _tf
    _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{_rec_id}.flag")
    try:
        open(_flag, 'w').close()
    except OSError:
        pass
    rdp = (get_config("chrome.remote_debugging_port") or "").strip()
    if rdp and rdp.isdigit():
        _cdp_evaluate(int(rdp),
            "if(window.__webActionRecorder){"
            "  window.__webActionRecorder.paused=true;"
            "  window.__webActionRecorder.events=[];"
            "}"
            "try{sessionStorage.setItem('__webActionRecorder_paused','true');}catch(e){}")
    request.session["recording_paused"] = True
    return JsonResponse({"ok": True})


@login_required
def resume_recording_ajax(request):
    """POST /record/resume-ajax/ — re-enable the in-browser recorder."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)
    _rec_pid = request.session.get("recording_pid")
    if not _rec_pid:
        return JsonResponse({"ok": False, "error": "No active recording"})
    _rec_id = request.session.get("recording_id", "")
    import tempfile as _tf
    _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{_rec_id}.flag")
    try:
        os.remove(_flag)
    except OSError:
        pass
    rdp = (get_config("chrome.remote_debugging_port") or "").strip()
    if rdp and rdp.isdigit():
        _cdp_evaluate(int(rdp),
            "if(window.__webActionRecorder){"
            "  window.__webActionRecorder.paused=false;"
            "}"
            "try{sessionStorage.removeItem('__webActionRecorder_paused');}catch(e){}")
    request.session["recording_paused"] = False
    return JsonResponse({"ok": True})


@csrf_exempt
def stop_recording_ajax(request):
    """POST /record/stop-ajax/
    Stop the current recording for the logged-in user; return JSON.
    Mirrors stop_recording but does not redirect so it can be called
    via fetch() from the steps-page live-recording UI.
    """
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)

    # Authenticate: accept Django session OR monitor token
    if not request.user.is_authenticated:
        user = _resolve_monitor_user(request)
        if user is None:
            return JsonResponse({"ok": False, "error": "Authentication required"}, status=401)
    else:
        user = request.user

    requested_record_id = ""
    try:
        if request.body:
            body = json.loads(request.body.decode("utf-8"))
            requested_record_id = str(body.get("record_id") or "").strip()
    except Exception:
        requested_record_id = ""
    if not requested_record_id:
        requested_record_id = str(request.POST.get("record_id") or "").strip()

    pid         = request.session.pop("recording_pid",         None)
    folder_name = request.session.pop("recording_folder",      "")
    record_id   = request.session.pop("recording_id",          "")
    is_baseline = request.session.pop("recording_is_baseline", False)
    request.session.pop("recording_url",    None)
    request.session.pop("recording_name",   None)
    was_paused  = request.session.pop("recording_paused", None)

    # Fallback for steps pages that still know the active record_id even when
    # the Django session no longer carries recording_pid/recording_id.
    if requested_record_id and (not record_id or str(record_id) != requested_record_id):
        with _ACTIVE_RECORDING_LOCK:
            rec = _ACTIVE_RECORDING.get(requested_record_id)
        if rec and (user.is_superuser or rec.get("user_id") == user.pk):
            record_id = requested_record_id
            pid = pid or rec.get("pid")
            folder_name = folder_name or rec.get("folder", "")
            if was_paused is None:
                was_paused = bool(rec.get("paused"))

    if not pid and record_id:
        with _ACTIVE_RECORDING_LOCK:
            rec = _ACTIVE_RECORDING.get(str(record_id))
        if rec and (user.is_superuser or rec.get("user_id") == user.pk):
            pid = rec.get("pid")
            folder_name = folder_name or rec.get("folder", "")
            if was_paused is None:
                was_paused = bool(rec.get("paused"))

    # Final fallback for dev-server reloads: read the recording PID persisted
    # at launch time so Stop still works even if session and in-memory state reset.
    if requested_record_id and not pid:
        import tempfile as _tf
        _pid_file = os.path.join(_tf.gettempdir(), f"recorder_pid_{requested_record_id}.txt")
        try:
            with open(_pid_file, "r", encoding="utf-8") as _pf:
                _lines = [line.strip() for line in _pf.readlines()]
            _file_pid = int(_lines[0]) if _lines and (_lines[0] or "").isdigit() else None
            _file_folder = _lines[1] if len(_lines) > 1 else ""
            _file_user_id = int(_lines[2]) if len(_lines) > 2 and (_lines[2] or "").isdigit() else None
            if _file_pid and (user.is_superuser or _file_user_id == user.pk):
                record_id = requested_record_id
                pid = _file_pid
                folder_name = folder_name or _file_folder
        except (OSError, ValueError):
            pass

    if requested_record_id and not record_id:
        record_id = requested_record_id

    # If the recording process is already gone by the time the user clicks Stop,
    # treat the request as a successful cleanup so the UI can recover and show
    # Record More again instead of staying stuck in live-recording mode.
    if not pid and record_id:
        with _ACTIVE_RECORDING_LOCK:
            _ACTIVE_RECORDING.pop(str(record_id), None)
        return JsonResponse({"ok": True, "already_stopped": True})

    if not pid and not record_id:
        return JsonResponse({"ok": False, "error": "No active recording"}, status=400)

    # Clean up the Python-level pause flag file (in case recording was paused when stopped).
    if record_id:
        import tempfile as _tf
        _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{record_id}.flag")
        _pid_file = os.path.join(_tf.gettempdir(), f"recorder_pid_{record_id}.txt")
        try:
            os.remove(_flag)
        except OSError:
            pass
        try:
            os.remove(_pid_file)
        except OSError:
            pass

    # If recording was paused when stopped, un-pause Chrome so it's clean
    # for any subsequent Replay or Record More action.
    if was_paused:
        _rdp_stop = (get_config("chrome.remote_debugging_port") or "").strip()
        if _rdp_stop.isdigit() and int(_rdp_stop) > 0:
            _cdp_evaluate(int(_rdp_stop),
                "if(window.__webActionRecorder){"
                "  window.__webActionRecorder.paused=false;"
                "  window.__webActionRecorder.events=[];"
                "}"
                "try{sessionStorage.removeItem('__webActionRecorder_paused');}catch(e){}")

    if pid:
        try:
            if sys.platform == "win32":
                # Kill only the Python process (/F), NOT its children (/T omitted)
                # so Chrome stays open on its remote-debugging port for Add Step.
                subprocess.call(
                    ["taskkill", "/F", "/PID", str(pid)],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            else:
                os.kill(pid, signal.SIGTERM)
        except (OSError, ProcessLookupError):
            pass

    if record_id and folder_name:
        try:
            parent_id, sub_id, end_id = _resolve_folder_ids(
                folder_name, user.username, is_baseline
            )
            _tables = ["recordings", "steps", "run_table", "session_meta"]
            with connection.cursor() as cur:
                _rec_file_order = _next_distinct_file_order(cur, folder_name, end_folder_id=end_id)
                for _tbl in _tables:
                    cur.execute(
                        f"UPDATE {_tbl} SET parent_folder_id=%s, sub_folder_id=%s,"
                        f" end_folder_id=%s WHERE record_id=%s",
                        [parent_id, sub_id, end_id, record_id],
                    )
                for _tbl_fo in ("recordings", "steps", "run_table"):
                    cur.execute(
                        f"UPDATE {_tbl_fo} SET file_order=%s WHERE record_id=%s",
                        [_rec_file_order, record_id],
                    )
        except Exception:
            pass
        _touch_folder_hierarchy(folder_name)

    # Keep the recorded Chrome window alive so Record More / continue_recording
    # can reattach to the same remote-debugging session.

    with _ACTIVE_RECORDING_LOCK:
        _ACTIVE_RECORDING.pop(str(record_id), None)

    return JsonResponse({"ok": True})


# ---------------------------------------------------------------------------
# Monitor-compatible recording control (no Django session required)
# ---------------------------------------------------------------------------

@csrf_exempt
def monitor_pause_recording(request):
    """POST /api/recording/pause/ — pause via monitor_token + record_id."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"ok": False, "error": "Unauthorized"}, status=401)

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        body = {}
    rec_id = str(body.get("record_id") or "").strip()
    if not rec_id:
        return JsonResponse({"ok": False, "error": "record_id required"})

    with _ACTIVE_RECORDING_LOCK:
        rec = _ACTIVE_RECORDING.get(rec_id)
    if not rec:
        import tempfile as _tf_fb
        _pid_file = os.path.join(_tf_fb.gettempdir(), f"recorder_pid_{rec_id}.txt")
        try:
            with open(_pid_file, "r", encoding="utf-8") as _pf:
                _lines = [line.strip() for line in _pf.readlines()]
            _file_pid = int(_lines[0]) if _lines and (_lines[0] or "").isdigit() else None
            _file_folder = _lines[1] if len(_lines) > 1 else ""
            _file_user_id = int(_lines[2]) if len(_lines) > 2 and (_lines[2] or "").isdigit() else None
            if _file_pid and (user.is_superuser or _file_user_id == user.pk):
                rec = {"pid": _file_pid, "folder": _file_folder, "user_id": _file_user_id, "paused": False}
                with _ACTIVE_RECORDING_LOCK:
                    _ACTIVE_RECORDING[rec_id] = rec
        except (OSError, ValueError):
            pass
    if not rec:
        return JsonResponse({"ok": False, "error": "No active recording for this id"})

    import tempfile as _tf
    _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{rec_id}.flag")
    try:
        open(_flag, "w").close()
    except OSError:
        pass
    rdp = (get_config("chrome.remote_debugging_port") or "").strip()
    if rdp and rdp.isdigit():
        _cdp_evaluate(int(rdp),
            "if(window.__webActionRecorder){"
            "  window.__webActionRecorder.paused=true;"
            "  window.__webActionRecorder.events=[];"
            "}"
            "try{sessionStorage.setItem('__webActionRecorder_paused','true');}catch(e){}")
    with _ACTIVE_RECORDING_LOCK:
        if rec_id in _ACTIVE_RECORDING:
            _ACTIVE_RECORDING[rec_id]["paused"] = True
    return JsonResponse({"ok": True})


@csrf_exempt
def monitor_resume_recording(request):
    """POST /api/recording/resume/ — resume via monitor_token + record_id."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"ok": False, "error": "Unauthorized"}, status=401)

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        body = {}
    rec_id = str(body.get("record_id") or "").strip()
    if not rec_id:
        return JsonResponse({"ok": False, "error": "record_id required"})

    with _ACTIVE_RECORDING_LOCK:
        rec = _ACTIVE_RECORDING.get(rec_id)
    if not rec:
        import tempfile as _tf_fb
        _pid_file = os.path.join(_tf_fb.gettempdir(), f"recorder_pid_{rec_id}.txt")
        try:
            with open(_pid_file, "r", encoding="utf-8") as _pf:
                _lines = [line.strip() for line in _pf.readlines()]
            _file_pid = int(_lines[0]) if _lines and (_lines[0] or "").isdigit() else None
            _file_folder = _lines[1] if len(_lines) > 1 else ""
            _file_user_id = int(_lines[2]) if len(_lines) > 2 and (_lines[2] or "").isdigit() else None
            if _file_pid and (user.is_superuser or _file_user_id == user.pk):
                rec = {"pid": _file_pid, "folder": _file_folder, "user_id": _file_user_id, "paused": False}
                with _ACTIVE_RECORDING_LOCK:
                    _ACTIVE_RECORDING[rec_id] = rec
        except (OSError, ValueError):
            pass
    if not rec:
        return JsonResponse({"ok": False, "error": "No active recording for this id"})

    import tempfile as _tf
    _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{rec_id}.flag")
    try:
        os.remove(_flag)
    except OSError:
        pass
    rdp = (get_config("chrome.remote_debugging_port") or "").strip()
    if rdp and rdp.isdigit():
        _cdp_evaluate(int(rdp),
            "if(window.__webActionRecorder){"
            "  window.__webActionRecorder.paused=false;"
            "}"
            "try{sessionStorage.removeItem('__webActionRecorder_paused');}catch(e){}")
    with _ACTIVE_RECORDING_LOCK:
        if rec_id in _ACTIVE_RECORDING:
            _ACTIVE_RECORDING[rec_id]["paused"] = False
    return JsonResponse({"ok": True})


@csrf_exempt
def monitor_stop_recording(request):
    """POST /api/recording/stop/ — stop via monitor_token + record_id."""
    if request.method not in {"POST", "GET"}:
        return JsonResponse({"ok": False, "error": "GET or POST required"}, status=405)

    try:
        body = json.loads(request.body.decode("utf-8")) if request.body else {}
    except Exception:
        body = {}

    user = _resolve_monitor_user(request)
    if user is None:
        _body_token = str(body.get("monitor_token") or "").strip()
        if _body_token:
            try:
                _payload = signing.loads(_body_token, salt="recording-monitor", max_age=86400)
                _user_id = int(_payload.get("uid"))
                user = User.objects.filter(pk=_user_id, is_active=True).first()
            except Exception:
                user = None
    if user is None:
        return JsonResponse({"ok": False, "error": "Unauthorized"}, status=401)

    rec_id = str(body.get("record_id") or request.GET.get("record_id") or "").strip()
    if not rec_id:
        return JsonResponse({"ok": False, "error": "record_id required"})

    with _ACTIVE_RECORDING_LOCK:
        rec = _ACTIVE_RECORDING.get(rec_id)

    # Fallback: if _ACTIVE_RECORDING was lost (e.g. server auto-reload),
    # try to recover the PID from the persisted temp file.
    if not rec:
        import tempfile as _tf_fb
        _pid_file = os.path.join(_tf_fb.gettempdir(), f"recorder_pid_{rec_id}.txt")
        try:
            with open(_pid_file, "r", encoding="utf-8") as _pf:
                _lines = [line.strip() for line in _pf.readlines()]
            _file_pid = int(_lines[0]) if _lines and (_lines[0] or "").isdigit() else None
            _file_folder = _lines[1] if len(_lines) > 1 else ""
            _file_user_id = int(_lines[2]) if len(_lines) > 2 and (_lines[2] or "").isdigit() else None
            if _file_pid and (user.is_superuser or _file_user_id == user.pk):
                rec = {"pid": _file_pid, "folder": _file_folder, "user_id": _file_user_id, "paused": False}
        except (OSError, ValueError):
            pass

    if not rec:
        return JsonResponse({"ok": False, "error": "No active recording for this id"})

    pid = rec.get("pid")
    folder_name = rec.get("folder", "")

    # Clean up pause flag
    import tempfile as _tf
    _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{rec_id}.flag")
    try:
        os.remove(_flag)
    except OSError:
        pass

    # Un-pause Chrome JS if it was paused
    if rec.get("paused"):
        _rdp_stop = (get_config("chrome.remote_debugging_port") or "").strip()
        if _rdp_stop.isdigit() and int(_rdp_stop) > 0:
            _cdp_evaluate(int(_rdp_stop),
                "if(window.__webActionRecorder){"
                "  window.__webActionRecorder.paused=false;"
                "  window.__webActionRecorder.events=[];"
                "}"
                "try{sessionStorage.removeItem('__webActionRecorder_paused');}catch(e){}")

    # Kill the recorder process
    if pid:
        try:
            if sys.platform == "win32":
                subprocess.call(
                    ["taskkill", "/F", "/PID", str(pid)],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            else:
                os.kill(pid, signal.SIGTERM)
        except (OSError, ProcessLookupError):
            pass

    # Backfill folder hierarchy
    if folder_name:
        try:
            parent_id, sub_id, end_id = _resolve_folder_ids(
                folder_name, user.username, False
            )
            _tables = ["recordings", "steps", "run_table", "session_meta"]
            with connection.cursor() as cur:
                _rec_file_order = _next_distinct_file_order(cur, folder_name, end_folder_id=end_id)
                for _tbl in _tables:
                    cur.execute(
                        f"UPDATE {_tbl} SET parent_folder_id=%s, sub_folder_id=%s,"
                        f" end_folder_id=%s WHERE record_id=%s",
                        [parent_id, sub_id, end_id, rec_id],
                    )
                for _tbl_fo in ("recordings", "steps", "run_table"):
                    cur.execute(
                        f"UPDATE {_tbl_fo} SET file_order=%s WHERE record_id=%s",
                        [_rec_file_order, rec_id],
                    )
        except Exception:
            pass
        _touch_folder_hierarchy(folder_name)

    # Clean up PID temp file
    import tempfile as _tf_cleanup
    _pid_file_cleanup = os.path.join(_tf_cleanup.gettempdir(), f"recorder_pid_{rec_id}.txt")
    try:
        os.remove(_pid_file_cleanup)
    except OSError:
        pass

    with _ACTIVE_RECORDING_LOCK:
        _ACTIVE_RECORDING.pop(rec_id, None)

    return JsonResponse({"ok": True})


def _auto_step_description(step) -> str:
    """Generate a human-readable step description from raw_event data.
    Used as a fallback when steps_description is blank/null."""
    action = step.action or ""
    re = step.raw_event or {}
    tag = (re.get("tag") or getattr(step, "element_tag", None) or "").lower()
    name = re.get("name") or re.get("id") or ""
    text = re.get("text") or ""
    value = re.get("value") or ""
    key = re.get("key") or ""

    if action == "click":
        if text:
            return f"Left-mouse-click on '{text}'"
        elif tag and name:
            return f"Left-mouse-click on {tag} '{name}'"
        elif tag:
            return f"Left-mouse-click on <{tag}>"
        return "Left-mouse-click"
    elif action == "dblclick":
        if text:
            return f"Double-click on '{text}'"
        elif tag and name:
            return f"Double-click on {tag} '{name}'"
        return "Double-click"
    elif action == "contextmenu":
        if text:
            return f"Right-click on '{text}'"
        elif tag and name:
            return f"Right-click on {tag} '{name}'"
        return "Right-click"
    elif action in ("input", "change"):
        val = text or value
        return f"User input recorded: '{val}'"
    elif action == "keydown":
        return f"Key pressed: '{key}'"
    elif action == "submit":
        return f"Form submitted on <{tag}>" if tag else "Form submitted"
    elif action in ("navigate", "open", "goto"):
        url = getattr(step, "page_url", "") or ""
        return f"Navigate to {url}" if url else "Navigate"
    elif action == "navigate_back":
        return "Browser back"
    elif action == "navigate_forward":
        return "Browser forward"
    elif action == "scroll":
        dy = int(re.get("delta_y") or 0)
        return "Scroll up" if dy < 0 else "Scroll down"
    return action.replace("_", " ").capitalize()


@login_required
def session_steps(request, record_id):
    """Show all steps for a specific session."""
    # Handle PW code inline save
    if request.method == "POST" and request.POST.get("action") == "save_pw_code":
        step_no = request.POST.get("step_no")
        pw_code = request.POST.get("playwright_code", "")
        Step.objects.filter(record_id=record_id, step_no=step_no).update(playwright_code=pw_code)
        return JsonResponse({"ok": True})

    steps = list(Step.objects.filter(record_id=record_id).order_by("step_no"))
    if not steps:
        steps = list(Recording.objects.filter(record_id=record_id).order_by("step_no"))

    # Auto-populate playwright_code if missing
    if steps and not getattr(steps[0], "playwright_code", None):
        try:
            _populate_playwright_code(str(record_id))
            # Refresh steps to get the populated values
            steps = list(Step.objects.filter(record_id=record_id).order_by("step_no"))
            if not steps:
                steps = list(Recording.objects.filter(record_id=record_id).order_by("step_no"))
        except Exception:
            pass

    # Eager-load related locators and data entries
    locator_ids = [s.locator_id for s in steps if s.locator_id]
    data_ids = [s.data_id for s in steps if s.data_id]

    locators_map = {
        loc.id: loc
        for loc in Locator.objects.filter(id__in=locator_ids)
    }
    data_map = {
        de.id: de
        for de in DataEntry.objects.filter(id__in=data_ids)
    }
    dataset_rows = []
    try:
        _sid_str_ds = str(record_id)
        with connection.cursor() as _ds_cur:
            _ds_cur.execute("""
                    SELECT s.step_no,
                           sd.id AS session_data_id,
                           s.data_id AS linked_data_id,
                           COALESCE(d.field_name, sd.field_name, s.field_name, '') AS display_field_name,
                           COALESCE(d.value, sd.value, s.field_value, '') AS display_value
                  FROM steps s
                  LEFT JOIN data d ON d.id = s.data_id
                  LEFT JOIN data sd ON sd.record_id = s.record_id AND sd.step_no = s.step_no
                 WHERE s.record_id = %s
                   AND (s.data_id IS NOT NULL OR sd.id IS NOT NULL)
                 ORDER BY s.step_no, s.id
            """, [_sid_str_ds])
            _seen_steps = set()
            for _sno, _session_data_id, _linked_data_id, _fname, _dval in _ds_cur.fetchall():
                if _sno in _seen_steps:
                    continue
                _seen_steps.add(_sno)
                dataset_rows.append({
                    "data_id": _session_data_id or _linked_data_id,
                    "linked_data_id": _linked_data_id,
                    "step_no": _sno,
                    "field_name": (_fname or "").strip(),
                    "value": _dval or "",
                })
    except Exception:
        dataset_rows = []

    # Load all locators per step with values and rank (for dropdown + Element Identity table)
    _sid_str = str(record_id)
    with connection.cursor() as _cur:
        _cur.execute("""
            SELECT id, step_no, strategy, locator, is_primary, COALESCE(locator_rank, 99)
            FROM locators
            WHERE record_id = %s
            ORDER BY step_no, COALESCE(locator_rank, 99), id
        """, [_sid_str])
        _loc_rows = _cur.fetchall()
    _strategies_by_step: dict = {}
    _all_locators_by_step: dict = {}
    for _loc_id, _sno, _strat, _loc_val, _isprim, _rank in _loc_rows:
        # For strategy dropdown (existing — deduplicated by strategy name)
        if not any(s["strategy"] == _strat for s in _strategies_by_step.get(_sno, [])):
            _strategies_by_step.setdefault(_sno, []).append({
                "strategy": _strat,
                "is_primary": _isprim,
                "locator": _loc_val,
                "locator_id": _loc_id,
                "rank": _rank,
            })
        # For Element Identity table (all rows in rank order)
        _all_locators_by_step.setdefault(_sno, []).append({
            "locator_id": _loc_id,
            "strategy": _strat,
            "locator":  _loc_val,
            "is_primary": _isprim,
            "rank": _rank,
        })

    enriched = []
    for step in steps:
        _rendered_locators = _step_identity_locators(step, _all_locators_by_step.get(step.step_no, []))
        _rendered_strategies = _step_strategy_options(step, _rendered_locators)
        enriched.append({
            "step": step,
            "locator": locators_map.get(step.locator_id),
            "data": data_map.get(step.data_id),
            "strategies": _rendered_strategies,
            "all_locators": _rendered_locators,
            "display_description": step.steps_description or _auto_step_description(step),
        })

    try:
        _meta = SessionMeta.objects.get(record_id=record_id)
        record_name = _meta.record_name
        session_engine = _meta.engine or 'selenium'
    except SessionMeta.DoesNotExist:
        record_name = ""
        session_engine = 'selenium'

    # Derive folder_name for scoped reorder queries
    _raw_fn = (getattr(steps[0], "folder_name", None) or "") if steps else ""
    _raw_fn_s = _raw_fn.strip().lower()
    folder_name = "Baseline" if _raw_fn_s in {"baseline", ""} else _raw_fn.strip()

    paginator = Paginator(enriched, 100)
    page_obj  = paginator.get_page(request.GET.get("page"))

    # Max step_no for the live-recording poll baseline
    max_step_no = max((s.step_no for s in steps), default=0)

    # Is this user currently recording into this exact session?
    _rec_id  = request.session.get("recording_id", "")
    _rec_pid = request.session.get("recording_pid")
    is_recording_this = (
        str(_rec_id) == str(record_id)
        and bool(_rec_pid)
        and _is_pid_alive(int(_rec_pid))
    )

    # Fallback: check _ACTIVE_RECORDING in case PID check failed or session is stale
    if not is_recording_this:
        with _ACTIVE_RECORDING_LOCK:
            _ar = _ACTIVE_RECORDING.get(str(record_id))
        if _ar and (request.user.is_superuser or _ar.get("user_id") == request.user.pk):
            _ar_pid = _ar.get("pid")
            if _ar_pid and _is_pid_alive(int(_ar_pid)):
                is_recording_this = True

    # Final fallback: PID temp file (survives server restarts)
    if not is_recording_this:
        import tempfile as _tf_check
        _pid_file_check = os.path.join(_tf_check.gettempdir(), f"recorder_pid_{record_id}.txt")
        try:
            with open(_pid_file_check, "r", encoding="utf-8") as _pfc:
                _lines_check = [l.strip() for l in _pfc.readlines()]
            _check_pid = int(_lines_check[0]) if _lines_check and _lines_check[0].isdigit() else None
            _check_uid = int(_lines_check[2]) if len(_lines_check) > 2 and _lines_check[2].isdigit() else None
            if _check_pid and (request.user.is_superuser or _check_uid == request.user.pk):
                if _is_pid_alive(_check_pid):
                    is_recording_this = True
        except (OSError, ValueError, IndexError):
            pass

    # Is the active recording currently paused?
    is_recording_paused = is_recording_this and bool(request.session.get("recording_paused"))

    # All baseline sessions for the Attach Event dropdown
    # Only include sessions that actually have steps (avoids "no steps" error).
    _attach_sessions = []
    try:
        with connection.cursor() as _aec:
            _aec.execute("""
                SELECT m.record_id, COALESCE(m.record_name, ''), COALESCE(m.folder_name, ''), COALESCE(m.engine, 'selenium')
                FROM session_meta m
                WHERE m.is_baseline = TRUE
                  AND EXISTS (SELECT 1 FROM steps s WHERE s.record_id = m.record_id)
                ORDER BY m.record_name ASC
            """)
            _seen_names = {}
            _rows = _aec.fetchall()
            # Detect duplicate display names so we can disambiguate
            for _rid, _rname, _rfolder, _rengine in _rows:
                _display = _rname or _rid
                _seen_names[_display] = _seen_names.get(_display, 0) + 1
            for _rid, _rname, _rfolder, _rengine in _rows:
                _display = _rname or _rid
                _label = _rname
                if _seen_names.get(_display, 0) > 1:
                    _label = f"{_rname} ({_rid[:8]})" if _rname else _rid[:8]
                _attach_sessions.append({
                    "record_id": _rid,
                    "record_name": _label,
                    "folder": _rfolder,
                    "engine": _rengine,
                })
    except Exception:
        pass

    return render(request, "recorder/steps.html", {
        "record_id": record_id,
        "record_name": record_name,
        "folder_name": folder_name,
        "session_engine": session_engine,
        "enriched": page_obj,
        "page_obj": page_obj,
        "max_step_no": max_step_no,
        "dataset_rows": dataset_rows,
        "is_recording_this": is_recording_this,
        "is_recording_paused": is_recording_paused,
        "sessions": _attach_sessions,
    })


def download_session(request, record_id, fmt):
    """Download session steps as CSV, PDF, or DOC."""
    sid = str(record_id)

    # Load steps — prefer folder copies (steps), fall back to raw recordings
    steps = list(Step.objects.filter(record_id=sid).order_by("step_no"))
    if not steps:
        steps = list(Recording.objects.filter(record_id=sid).order_by("step_no"))

    locator_ids = [s.locator_id for s in steps if s.locator_id]
    data_ids    = [s.data_id    for s in steps if s.data_id]
    locators_map = {loc.id: loc for loc in Locator.objects.filter(id__in=locator_ids)}
    data_map     = {de.id:  de  for de  in DataEntry.objects.filter(id__in=data_ids)}

    # ── Meta fields ────────────────────────────────────────────────
    try:
        meta = SessionMeta.objects.get(record_id=sid)
        record_name = meta.record_name or sid
        recorder    = meta.recorder or ""
    except SessionMeta.DoesNotExist:
        record_name = sid
        recorder    = ""

    # folder_name: from first step (normalised), fall back to 'Recordings'
    _raw_fn = (getattr(steps[0], "folder_name", None) or "") if steps else ""
    _raw_fn_s = _raw_fn.strip().lower()
    folder_name = "Baseline" if _raw_fn_s in {"baseline", ""} else _raw_fn.strip()

    # runner: from the most recent RunResult for this session
    last_run = RunResult.objects.filter(record_id=sid).order_by("-created_at").first()
    runner = (last_run.runner or "") if last_run else ""

    # Screenshots from the last run (keyed by step_no)
    if last_run:
        _last_run_id_str  = str(last_run.run_id)
        _last_rec_id_str  = str(last_run.record_id)
        _last_run_results = RunResult.objects.filter(run_id=last_run.run_id, record_id=sid)
        _ss_map = {rr.step_no: bytes(rr.screenshot)
                   for rr in _last_run_results if rr.screenshot}
    else:
        _last_run_id_str = _last_rec_id_str = ""
        _ss_map = {}

    meta_fields = [
        ("Project",  folder_name),
        ("Filename", record_name),
        ("Designer", recorder),
        ("Tester",   runner),
    ]

    safe_name = "".join(c for c in record_name if c.isalnum() or c in " _-").strip() or sid

    col_headers = ["#", "Page URL", "Action", "Element", "Identity", "Label", "Fieldname", "Data", "Recorder", "Script Runner", "Timestamp", "Screenshot URL"]
    _trunc = lambda v, n=30: (str(v)[:n] + "…") if len(str(v or "")) > n else str(v or "")
    rows = []
    for s in steps:
        loc = locators_map.get(s.locator_id)
        dat = data_map.get(s.data_id)
        _ss_url = (
            f"/run/{_last_run_id_str}/screenshot/{_last_rec_id_str}/{s.step_no}/"
            if s.step_no in _ss_map else ""
        )
        rows.append([
            s.step_no,
            _trunc(s.page_url         or ""),
            s.action           or "",
            s.element_tag      or "",
            loc.strategy if loc else "",
            _trunc(loc.locator  if loc else ""),
            dat.field_name if dat else "",
            dat.value      if dat else "",
            s.recorder         or "",
            s.runner           or "",
            s.created_at.strftime("%Y-%m-%d %H:%M:%S") if s.created_at else "",
            _ss_url,
        ])

    if fmt == "csv":
        from django.http import HttpResponse as _HR
        response = _HR(content_type="text/csv; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{safe_name}.csv"'
        writer = csv.writer(response)
        # Header block
        for label, value in meta_fields:
            writer.writerow([f"{label}:", value])
        writer.writerow([])  # blank separator
        # Stats block
        _csv_total  = len(rows)
        _csv_pages  = len({r[1] for r in rows if r[1]})
        _csv_acts   = len({r[2] for r in rows if r[2]})
        _csv_elems  = len({r[3] for r in rows if r[3]})
        writer.writerow(["Total Steps", "Unique Pages", "Unique Actions", "Unique Elements"])
        writer.writerow([_csv_total, _csv_pages, _csv_acts, _csv_elems])
        writer.writerow([])  # blank separator
        writer.writerow(col_headers)
        writer.writerows(rows)
        return response

    if fmt == "pdf":
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                        Paragraph, Spacer, HRFlowable,
                                        Image as RLImage)
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from django.http import HttpResponse as _HR

        # ── Brand / palette ────────────────────────────────────────────
        _NAVY       = colors.HexColor("#1e3c5e")
        _BLUE       = colors.HexColor("#2563a8")
        _BLUE_LIGHT = colors.HexColor("#dbeafe")
        _GREEN      = colors.HexColor("#198754")
        _GREEN_BG   = colors.HexColor("#d1fae5")
        _GREY       = colors.HexColor("#6c757d")
        _GREY_BG    = colors.HexColor("#f3f4f6")
        _AMBER      = colors.HexColor("#f59e0b")
        _AMBER_BG   = colors.HexColor("#fef3c7")
        _ROW_ALT    = colors.HexColor("#f0f6fc")
        _BORDER     = colors.HexColor("#cbd5e1")
        _WHITE      = colors.white
        _DARK       = colors.HexColor("#0f172a")

        # 11 cols: #, Page URL, Action, Element, Identity, Label, Fieldname, Data, Recorder, Script Runner, Timestamp
        # A4 landscape frame ≈ 801 pt; keep sum ≤ 785 for breathing room
        _COL_WIDTHS = [22, 138, 52, 52, 58, 110, 68, 68, 56, 64, 77]   # sum = 765

        buffer = io.BytesIO()

        def _add_page_number(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(_GREY)
            canvas.drawRightString(
                landscape(A4)[0] - 20, 14,
                f"Page {canvas.getPageNumber()}  ·  {record_name}"
            )
            canvas.restoreState()

        doc = SimpleDocTemplate(
            buffer, pagesize=landscape(A4),
            leftMargin=20, rightMargin=20, topMargin=28, bottomMargin=28,
            onFirstPage=_add_page_number, onLaterPages=_add_page_number,
        )

        _tbl_hdr = ParagraphStyle(
            "tbl_hdr", fontName="Helvetica-Bold", fontSize=7.5, leading=9,
            textColor=_WHITE, wordWrap="LTR",
        )
        _cell = ParagraphStyle(
            "cell", fontName="Helvetica", fontSize=7.5, leading=9,
            textColor=_DARK, wordWrap="LTR", splitLongWords=True,
        )
        _meta_label = ParagraphStyle(
            "meta_lbl", fontName="Helvetica-Bold", fontSize=8.5, leading=11,
            textColor=_NAVY,
        )
        _meta_value = ParagraphStyle(
            "meta_val", fontName="Helvetica", fontSize=8.5, leading=11,
            textColor=_DARK,
        )

        def _action_para(text):
            _ACTION_COLOURS = {
                "click":    (_BLUE,  _BLUE_LIGHT),
                "dblclick": (_BLUE,  _BLUE_LIGHT),
                "input":    (_AMBER, _AMBER_BG),
                "change":   (_AMBER, _AMBER_BG),
                "keydown":  (_GREY,  _GREY_BG),
                "submit":   (_GREEN, _GREEN_BG),
                "navigate": (colors.HexColor("#7c3aed"), colors.HexColor("#ede9fe")),
                "navigation": (colors.HexColor("#7c3aed"), colors.HexColor("#ede9fe")),
            }
            clr, _ = _ACTION_COLOURS.get(text.lower(), (_DARK, _WHITE))
            st = ParagraphStyle(
                f"act_{text}", fontName="Helvetica-Bold", fontSize=7.5, leading=9,
                textColor=clr, wordWrap="LTR",
            )
            return Paragraph(text.capitalize(), st)

        # ── Stats ──────────────────────────────────────────────────────
        _total        = len(rows)
        _unique_pages = len({r[1] for r in rows if r[1]})
        _unique_acts  = len({r[2] for r in rows if r[2]})
        _unique_elems = len({r[3] for r in rows if r[3]})

        # ── Banner ─────────────────────────────────────────────────────
        banner_data = [[
            Paragraph(f"<b>{record_name}</b>", ParagraphStyle(
                "bn", fontName="Helvetica-Bold", fontSize=15, leading=18,
                textColor=_WHITE,
            )),
            Paragraph("SESSION STEPS REPORT", ParagraphStyle(
                "bn2", fontName="Helvetica-Bold", fontSize=9, leading=11,
                textColor=colors.HexColor("#93c5fd"), alignment=TA_RIGHT,
            )),
        ]]
        banner = Table(banner_data, colWidths=[560, 208])
        banner.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, -1), _NAVY),
            ("LINEBELOW",     (0, 0), (-1, -1), 3, _BLUE),
            ("TOPPADDING",    (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ("LEFTPADDING",   (0, 0), (0, -1),  14),
            ("RIGHTPADDING",  (-1, 0), (-1, -1), 14),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ]))

        # ── Meta block ─────────────────────────────────────────────────
        meta_rows_tbl = [[Paragraph(f"{lbl}:", _meta_label), Paragraph(str(val), _meta_value)]
                         for lbl, val in meta_fields]
        meta_tbl = Table(meta_rows_tbl, colWidths=[75, 240], hAlign="LEFT")
        meta_tbl.setStyle(TableStyle([
            ("TOPPADDING",    (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("ALIGN",         (0, 0), (-1, -1), "LEFT"),
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ]))

        # ── Stats tiles (4 tiles) ──────────────────────────────────────
        _STAT_CONFIGS = [
            ("TOTAL STEPS",    str(_total),        _NAVY,  colors.HexColor("#e0eaf8")),
            ("UNIQUE PAGES",   str(_unique_pages),  _BLUE,  _BLUE_LIGHT),
            ("UNIQUE ACTIONS", str(_unique_acts),   _AMBER, _AMBER_BG),
            ("UNIQUE ELEMENTS",str(_unique_elems),  _GREEN, _GREEN_BG),
        ]
        cw4 = 768 / 4
        stats_hdr = [Paragraph(lbl, ParagraphStyle(
            f"sh_{i}", fontName="Helvetica-Bold", fontSize=7, leading=8,
            textColor=clr, alignment=TA_CENTER,
        )) for i, (lbl, _, clr, _bg) in enumerate(_STAT_CONFIGS)]
        stats_val = [Paragraph(val, ParagraphStyle(
            f"sv_{i}", fontName="Helvetica-Bold", fontSize=20, leading=24,
            textColor=clr, alignment=TA_CENTER,
        )) for i, (_lbl, val, clr, _bg) in enumerate(_STAT_CONFIGS)]

        stats_tbl = Table([stats_hdr, stats_val], colWidths=[cw4] * 4)
        _stat_styles = [
            ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, 0),  5),
            ("BOTTOMPADDING", (0, 0), (-1, 0),  3),
            ("TOPPADDING",    (0, 1), (-1, 1),  4),
            ("BOTTOMPADDING", (0, 1), (-1, 1),  8),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
            ("BOX",           (0, 0), (-1, -1), 1.5, _BORDER),
            ("LINEAFTER",     (0, 0), (2, -1),  0.5, _BORDER),
        ]
        for i, (_lbl, _val, _clr, bg) in enumerate(_STAT_CONFIGS):
            _stat_styles.append(("BACKGROUND", (i, 0), (i, -1), bg))
        stats_tbl.setStyle(TableStyle(_stat_styles))

        # ── Step table ─────────────────────────────────────────────────
        _MUTED = ParagraphStyle(
            "cell_muted", fontName="Helvetica", fontSize=7, leading=9,
            textColor=_GREY, wordWrap="LTR", splitLongWords=True,
        )
        _N_COLS = 11
        from reportlab.platypus import PageBreak

        def _make_hdr_row_session():
            return [Paragraph(h, _tbl_hdr) for h in col_headers[:_N_COLS]]

        step_flowables = []
        for step_idx, r in enumerate(rows):
            step_no = r[0]
            bg = colors.white if step_idx % 2 == 0 else _ROW_ALT
            row_cells = [
                Paragraph(str(r[0]), _cell),
                Paragraph(str(r[1] or ""), _cell),
                _action_para(str(r[2])),
                Paragraph(str(r[3] or ""), _cell),
                Paragraph(str(r[4] or ""), _MUTED),
                Paragraph(str(r[5] or ""), _cell),
                Paragraph(str(r[6] or ""), _MUTED),
                Paragraph(str(r[7] or ""), _cell),
                Paragraph(str(r[8] or ""), _MUTED),
                Paragraph(str(r[9] or ""), _cell),
                Paragraph(str(r[10] or ""), _MUTED),
            ]
            _ss_bytes = _ss_map.get(step_no)
            if _ss_bytes:
                try:
                    _ss_img = RLImage(io.BytesIO(_ss_bytes), width=720, height=430, kind="bound")
                except Exception:
                    _ss_img = Paragraph("", _cell)
            else:
                _ss_img = Paragraph("", _cell)

            step_tbl = Table(
                [_make_hdr_row_session(), row_cells, [_ss_img] + [Paragraph("", _cell)] * (_N_COLS - 1)],
                colWidths=_COL_WIDTHS, hAlign="CENTER",
            )
            step_tbl.setStyle(TableStyle([
                # Header row (0)
                ("BACKGROUND",    (0, 0), (-1, 0), _NAVY),
                ("LINEBELOW",     (0, 0), (-1, 0), 2, _BLUE),
                ("ALIGN",         (0, 0), (-1, 0), "CENTER"),
                ("VALIGN",        (0, 0), (-1, 0), "MIDDLE"),
                ("GRID",          (0, 0), (-1, 0), 0.4, _BORDER),
                ("LEFTPADDING",   (0, 0), (-1, 0), 6),
                ("RIGHTPADDING",  (0, 0), (-1, 0), 6),
                ("TOPPADDING",    (0, 0), (-1, 0), 6),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                # Data row (1)
                ("BACKGROUND",    (0, 1), (-1, 1), bg),
                ("GRID",          (0, 1), (-1, 1), 0.4, _BORDER),
                ("LINEBELOW",     (0, 1), (-1, 1), 1.5, _BORDER),
                ("ALIGN",         (0, 1), (-1, 1), "LEFT"),
                ("VALIGN",        (0, 1), (-1, 1), "MIDDLE"),
                ("LEFTPADDING",   (0, 1), (-1, 1), 6),
                ("RIGHTPADDING",  (0, 1), (-1, 1), 6),
                ("TOPPADDING",    (0, 1), (-1, 1), 5),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 5),
                # Screenshot row (2) — spans full width
                ("BACKGROUND",    (0, 2), (-1, 2), bg),
                ("SPAN",          (0, 2), (_N_COLS - 1, 2)),
                ("ALIGN",         (0, 2), (-1, 2), "CENTER"),
                ("VALIGN",        (0, 2), (-1, 2), "MIDDLE"),
                ("TOPPADDING",    (0, 2), (-1, 2), 16),
                ("BOTTOMPADDING", (0, 2), (-1, 2), 16),
                ("LEFTPADDING",   (0, 2), (-1, 2), 0),
                ("RIGHTPADDING",  (0, 2), (-1, 2), 0),
                ("LINEBELOW",     (0, 2), (-1, 2), 0.4, _BORDER),
            ]))
            step_flowables.append(PageBreak())
            step_flowables.append(step_tbl)

        elements = [
            banner,
            Spacer(1, 8),
            meta_tbl,
            Spacer(1, 10),
            HRFlowable(width="100%", thickness=1, color=_BORDER),
            Spacer(1, 8),
            stats_tbl,
        ] + step_flowables

        doc.build(elements, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
        buffer.seek(0)
        response = _HR(buffer.read(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{safe_name}.pdf"'
        return response

    if fmt == "doc":
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Mm, Inches
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from django.http import HttpResponse as _HR

        _ACTION_RGB = {
            "click":      RGBColor(0x25, 0x63, 0xa8),
            "dblclick":   RGBColor(0x25, 0x63, 0xa8),
            "input":      RGBColor(0xd9, 0x77, 0x06),
            "change":     RGBColor(0xd9, 0x77, 0x06),
            "keydown":    RGBColor(0x6c, 0x75, 0x7d),
            "submit":     RGBColor(0x19, 0x87, 0x54),
            "navigate":   RGBColor(0x7c, 0x3a, 0xed),
            "navigation": RGBColor(0x7c, 0x3a, 0xed),
        }
        _HDR_BG  = "1e3c5e"   # navy
        _HDR_FG  = "ffffff"

        def _set_cell_bg(cell, hex_color):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color)
            tcPr.append(shd)

        document = DocxDocument()

        # Landscape A4
        section = document.sections[0]
        section.orientation  = WD_ORIENT.LANDSCAPE
        section.page_width   = Mm(297)
        section.page_height  = Mm(210)
        section.left_margin  = Mm(12)
        section.right_margin = Mm(12)
        section.top_margin   = Mm(14)
        section.bottom_margin = Mm(14)

        # Title
        hdg = document.add_heading(f"{record_name} — Steps", 0)
        hdg.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Meta block
        meta_tbl = document.add_table(rows=len(meta_fields), cols=2)
        meta_tbl.autofit = False
        meta_tbl.allow_autofit = False
        for i, (label, value) in enumerate(meta_fields):
            # Label column — fixed 30mm wide
            lc = meta_tbl.rows[i].cells[0]
            lc.width = Mm(30)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lr = lp.add_run(f"{label}:")
            lr.bold = True
            lr.font.size = Pt(9)
            # Value column
            vc = meta_tbl.rows[i].cells[1]
            vc.width = Mm(100)
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            vp.add_run(value).font.size = Pt(9)
        document.add_paragraph()

        # ── Stats tiles ────────────────────────────────────────────────
        _stat_total  = len(rows)
        _stat_pages  = len({r[1] for r in rows if r[1]})
        _stat_acts   = len({r[2] for r in rows if r[2]})
        _stat_elems  = len({r[3] for r in rows if r[3]})

        _STAT_DATA = [
            ("TOTAL STEPS",     str(_stat_total), "1e3c5e", "e0eaf8"),
            ("UNIQUE PAGES",    str(_stat_pages), "2563a8", "dbeafe"),
            ("UNIQUE ACTIONS",  str(_stat_acts),  "d97706", "fef3c7"),
            ("UNIQUE ELEMENTS", str(_stat_elems), "198754", "d1fae5"),
        ]

        stats_tbl = document.add_table(rows=2, cols=4)
        stats_tbl.autofit = False
        stats_tbl.allow_autofit = False
        _tile_w = Mm(65)
        for ci in range(4):
            for ri in range(2):
                stats_tbl.rows[ri].cells[ci].width = _tile_w

        for ci, (label, value, fg, bg) in enumerate(_STAT_DATA):
            # Label row
            lc = stats_tbl.rows[0].cells[ci]
            _set_cell_bg(lc, bg)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(7)
            lr.font.color.rgb = RGBColor(
                int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
            )
            # Value row
            vc = stats_tbl.rows[1].cells[ci]
            _set_cell_bg(vc, bg)
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr = vp.add_run(value)
            vr.bold = True
            vr.font.size = Pt(20)
            vr.font.color.rgb = RGBColor(
                int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
            )

        document.add_paragraph()

        # One page per step — each step gets a page break + its own 3-row table
        _N_COLS   = 11
        _COL_W_MM = [8, 38, 16, 16, 18, 32, 20, 20, 18, 22, 27]  # total ≈ 235 mm

        def _add_step_table_session(step_idx, row_data):
            document.add_page_break()
            tbl = document.add_table(rows=3, cols=_N_COLS)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            tbl.allow_autofit = False
            for ci, mm in enumerate(_COL_W_MM):
                for r in tbl.rows:
                    r.cells[ci].width = Mm(mm)

            # Header row (row 0)
            for ci, h in enumerate(col_headers[:_N_COLS]):
                cell = tbl.rows[0].cells[ci]
                _set_cell_bg(cell, _HDR_BG)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

            # Data row (row 1)
            bg_hex = "f0f6fc" if step_idx % 2 == 0 else "ffffff"
            for ci, val in enumerate(row_data[:_N_COLS]):
                cell = tbl.rows[1].cells[ci]
                _set_cell_bg(cell, bg_hex)
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(7)
                if ci == 2:
                    rgb = _ACTION_RGB.get(str(val).lower())
                    if rgb:
                        run.font.color.rgb = rgb
                        run.bold = True

            # Screenshot row (row 2) — merged full width
            ss_cells = tbl.rows[2].cells
            merged = ss_cells[0].merge(ss_cells[_N_COLS - 1])
            _set_cell_bg(merged, bg_hex)
            return merged

        for step_idx, row_data in enumerate(rows):
            step_no = row_data[0]
            merged  = _add_step_table_session(step_idx, row_data)
            _ss_bytes = _ss_map.get(step_no)
            if _ss_bytes:
                try:
                    ss_para = merged.paragraphs[0]
                    ss_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ss_run  = ss_para.add_run()
                    ss_run.add_picture(io.BytesIO(_ss_bytes), width=Inches(8.0))
                except Exception:
                    merged.paragraphs[0].add_run("").font.size = Pt(1)
            else:
                merged.paragraphs[0].add_run("").font.size = Pt(1)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = _HR(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{safe_name}.docx"'
        return response

    if fmt == "playwright":
        bundle = _load_record_bundle(str(record_id))
        script = _build_playwright_script(bundle)

        from django.http import HttpResponse as _HR2
        response = _HR2(script, content_type="text/x-python; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{safe_name}_playwright.py"'
        return response

    if fmt == "selenium":
        bundle = _load_record_bundle(str(record_id))
        script = _build_selenium_script(bundle)
        from django.http import HttpResponse as _HR3
        response = _HR3(script, content_type="text/x-python; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{safe_name}_selenium.py"'
        return response

    raise Http404(f"Unknown format: {fmt}")
# ---------------------------------------------------------------------------

@login_required
def download_run(request, run_id, fmt):
    """Download a replay run's step results as CSV, PDF, or DOC."""
    rid = str(run_id)

    with connection.cursor() as cur:
        cur.execute("""
            SELECT
                r.step_no,
                r.action,
                r.page_url,
                r.element_tag,
                r.status,
                r.message,
                r.runner,
                r.run_date,
                r.folder_name,
                r.screenshot,
                r.run_id::text                AS run_id_str,
                r.record_id::text             AS record_id_str,
                COALESCE(m.record_name, '')    AS record_name,
                COALESCE(m.recorder, '')       AS recorder,
                COALESCE(m.engine, r.engine, 'selenium') AS engine
            FROM run_table r
            LEFT JOIN session_meta m ON m.record_id = r.record_id
            WHERE r.run_id = %s
            ORDER BY r.step_no
        """, [rid])
        cols = [c.name for c in cur.description]
        rows_raw = [dict(zip(cols, row)) for row in cur.fetchall()]

    if not rows_raw:
        raise Http404("Run not found")

    first       = rows_raw[0]
    record_name = first["record_name"] or rid
    engine      = (first.get("engine") or "selenium").strip().lower()
    runner      = first["runner"] or ""
    run_date    = first["run_date"]
    folder_name = first["folder_name"] or "Recordings"
    recorder    = first["recorder"] or ""
    run_date_str = run_date.strftime("%Y-%m-%d %H:%M") if run_date else ""

    meta_fields = [
        ("Project",  folder_name),
        ("Filename", record_name),
        ("Designer", recorder),
        ("Tester",   runner),
        ("Run Date", run_date_str),
    ]

    safe_name = "".join(c for c in record_name if c.isalnum() or c in " _-").strip() or rid
    safe_name_run = f"{safe_name}_run"
    engine_suffix = "_playwright" if engine == "playwright" else "_selenium"

    col_headers = ["#", "Page URL", "Action", "Element", "Steps", "Script Runner", "Status", "Screenshot URL"]
    _trunc = lambda v, n=30: (str(v)[:n] + "…") if len(str(v or "")) > n else str(v or "")
    rows = []
    for r in rows_raw:
        _ss_url = (
            f"/run/{r['run_id_str']}/screenshot/{r['record_id_str']}/{r['step_no']}/"
            if r.get("screenshot") else ""
        )
        rows.append([
            r["step_no"],
            _trunc(r["page_url"]     or ""),
            r["action"]       or "",
            r["element_tag"]  or "",
            _trunc(r["message"]  or ""),
            r["runner"]       or "",
            r["status"]       or "",
            _ss_url,
        ])

    if fmt == "csv":
        from django.http import HttpResponse as _HR
        response = _HR(content_type="text/csv; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{safe_name_run}.csv"'
        writer = csv.writer(response)
        for label, value in meta_fields:
            writer.writerow([f"{label}:", value])
        writer.writerow([])
        # Stats block
        _rcsv_total    = len(rows_raw)
        _rcsv_passed   = sum(1 for r in rows_raw if r["status"] == "pass")
        _rcsv_failed   = sum(1 for r in rows_raw if r["status"] == "fail")
        _rcsv_ne       = _rcsv_total - _rcsv_passed - _rcsv_failed
        _rcsv_pct      = round(_rcsv_passed / _rcsv_total * 100) if _rcsv_total else 0
        writer.writerow(["Total Steps", "Passed", "Failed", "Not Executed", "Pass Rate"])
        writer.writerow([_rcsv_total, _rcsv_passed, _rcsv_failed, _rcsv_ne, f"{_rcsv_pct}%"])
        writer.writerow([])
        writer.writerow(col_headers)
        writer.writerows(rows)
        return response

    if fmt == "pdf":
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                        Paragraph, Spacer, HRFlowable,
                                        Image as RLImage)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from django.http import HttpResponse as _HR

        # ── Brand / palette ────────────────────────────────────────────
        _NAVY       = colors.HexColor("#1e3c5e")
        _BLUE       = colors.HexColor("#2563a8")
        _BLUE_LIGHT = colors.HexColor("#dbeafe")
        _GREEN      = colors.HexColor("#198754")
        _GREEN_BG   = colors.HexColor("#d1fae5")
        _RED        = colors.HexColor("#dc3545")
        _RED_BG     = colors.HexColor("#fee2e2")
        _GREY       = colors.HexColor("#6c757d")
        _GREY_BG    = colors.HexColor("#f3f4f6")
        _AMBER      = colors.HexColor("#f59e0b")
        _AMBER_BG   = colors.HexColor("#fef3c7")
        _ROW_ALT    = colors.HexColor("#f0f6fc")
        _BORDER     = colors.HexColor("#cbd5e1")
        _WHITE      = colors.white
        _DARK       = colors.HexColor("#0f172a")

        _STATUS_COLOURS = {
            "pass":         _GREEN,
            "fail":         _RED,
            "not_executed": _GREY,
        }
        _STATUS_BG = {
            "pass":         _GREEN_BG,
            "fail":         _RED_BG,
            "not_executed": _GREY_BG,
        }

        _COL_WIDTHS = [30, 185, 62, 62, 215, 95, 70]   # 7 cols — sum ≈ 719

        buffer = io.BytesIO()

        # Page numbers via canvas callback
        _page_info = {"total": 0}
        def _add_page_number(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 7)
            canvas.setFillColor(_GREY)
            canvas.drawRightString(
                landscape(A4)[0] - 20, 14,
                f"Page {canvas.getPageNumber()}  ·  {record_name}  ·  {run_date_str}"
            )
            canvas.restoreState()

        doc = SimpleDocTemplate(
            buffer, pagesize=landscape(A4),
            leftMargin=20, rightMargin=20, topMargin=28, bottomMargin=28,
            onFirstPage=_add_page_number, onLaterPages=_add_page_number,
        )

        # ── Paragraph styles ───────────────────────────────────────────
        _title_style = ParagraphStyle(
            "rpt_title", fontName="Helvetica-Bold", fontSize=16, leading=20,
            textColor=_NAVY, alignment=TA_LEFT,
        )
        _subtitle_style = ParagraphStyle(
            "rpt_sub", fontName="Helvetica", fontSize=9, leading=12,
            textColor=_GREY, alignment=TA_LEFT,
        )
        _meta_label = ParagraphStyle(
            "meta_lbl", fontName="Helvetica-Bold", fontSize=8.5, leading=11,
            textColor=_NAVY,
        )
        _meta_value = ParagraphStyle(
            "meta_val", fontName="Helvetica", fontSize=8.5, leading=11,
            textColor=_DARK,
        )
        _stat_label = ParagraphStyle(
            "stat_lbl", fontName="Helvetica-Bold", fontSize=7.5, leading=9,
            textColor=_WHITE, alignment=TA_CENTER, spaceAfter=2,
        )
        _tbl_hdr = ParagraphStyle(
            "tbl_hdr", fontName="Helvetica-Bold", fontSize=7.5, leading=9,
            textColor=_WHITE, wordWrap="LTR",
        )
        _cell = ParagraphStyle(
            "cell", fontName="Helvetica", fontSize=7.5, leading=9,
            textColor=_DARK, wordWrap="LTR", splitLongWords=True,
        )

        def _stat_num(text, clr):
            return ParagraphStyle(
                f"sn_{text}", fontName="Helvetica-Bold", fontSize=18, leading=21,
                textColor=clr, alignment=TA_CENTER,
            )

        def _status_para(text):
            clr = _STATUS_COLOURS.get(text.lower(), _GREY)
            st = ParagraphStyle(
                f"sp_{text}", fontName="Helvetica-Bold", fontSize=7.5, leading=9,
                textColor=clr, wordWrap="LTR",
            )
            icons = {"pass": "✔", "fail": "✘", "not_executed": "—"}
            icon = icons.get(text.lower(), "")
            label = {"pass": "Pass", "fail": "Fail", "not_executed": "Not Executed"}.get(text.lower(), text.capitalize())
            return Paragraph(f"{icon} {label}", st)

        def _action_para(text):
            _ACTION_COLOURS = {
                "click":    (_BLUE,  _BLUE_LIGHT),
                "dblclick": (_BLUE,  _BLUE_LIGHT),
                "input":    (_AMBER, _AMBER_BG),
                "change":   (_AMBER, _AMBER_BG),
                "keydown":  (_GREY,  _GREY_BG),
                "submit":   (_GREEN, _GREEN_BG),
                "navigate": (colors.HexColor("#7c3aed"), colors.HexColor("#ede9fe")),
                "navigation": (colors.HexColor("#7c3aed"), colors.HexColor("#ede9fe")),
            }
            clr, _ = _ACTION_COLOURS.get(text.lower(), (_DARK, _WHITE))
            st = ParagraphStyle(
                f"act_{text}", fontName="Helvetica-Bold", fontSize=7.5, leading=9,
                textColor=clr, wordWrap="LTR",
            )
            return Paragraph(text.capitalize(), st)

        # ── Stats ──────────────────────────────────────────────────────
        _total        = len(rows_raw)
        _passed       = sum(1 for r in rows_raw if r["status"] == "pass")
        _failed       = sum(1 for r in rows_raw if r["status"] == "fail")
        _not_executed = _total - _passed - _failed
        _pct_pass     = round(_passed / _total * 100) if _total else 0
        _pass_clr     = _GREEN if _pct_pass >= 80 else (_AMBER if _pct_pass >= 50 else _RED)

        # ── Banner title ───────────────────────────────────────────────
        banner_data = [[
            Paragraph(f"<b>{record_name}</b>", ParagraphStyle(
                "bn", fontName="Helvetica-Bold", fontSize=15, leading=18,
                textColor=_WHITE,
            )),
            Paragraph("RUN RESULTS REPORT", ParagraphStyle(
                "bn2", fontName="Helvetica-Bold", fontSize=9, leading=11,
                textColor=colors.HexColor("#93c5fd"), alignment=TA_RIGHT,
            )),
        ]]
        banner = Table(banner_data, colWidths=[560, 208])
        banner.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, -1), _NAVY),
            ("LINEBELOW",     (0, 0), (-1, -1), 3, _BLUE),
            ("TOPPADDING",    (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ("LEFTPADDING",   (0, 0), (0, -1),  14),
            ("RIGHTPADDING",  (-1, 0), (-1, -1), 14),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ]))

        # ── Meta info block ────────────────────────────────────────────
        meta_rows = [[Paragraph(f"{lbl}:", _meta_label), Paragraph(str(val), _meta_value)]
                     for lbl, val in meta_fields]
        meta_tbl = Table(meta_rows, colWidths=[75, 240], hAlign="LEFT")
        meta_tbl.setStyle(TableStyle([
            ("TOPPADDING",    (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("ALIGN",         (0, 0), (-1, -1), "LEFT"),
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ]))

        # ── Stats cards (5 tiles) ──────────────────────────────────────
        _STAT_CONFIGS = [
            ("TOTAL STEPS",   str(_total),        _NAVY,    colors.HexColor("#e0eaf8")),
            ("PASSED",        str(_passed),        _GREEN,   _GREEN_BG),
            ("FAILED",        str(_failed),        _RED,     _RED_BG),
            ("NOT EXECUTED",  str(_not_executed),  _GREY,    _GREY_BG),
            ("PASS RATE",     f"{_pct_pass}%",     _pass_clr, _AMBER_BG if 50 <= _pct_pass < 80 else (_GREEN_BG if _pct_pass >= 80 else _RED_BG)),
        ]
        cw = 768 / 5
        stats_hdr = [Paragraph(lbl, ParagraphStyle(
            f"sh_{i}", fontName="Helvetica-Bold", fontSize=7, leading=8,
            textColor=clr, alignment=TA_CENTER,
        )) for i, (lbl, _, clr, _bg) in enumerate(_STAT_CONFIGS)]
        stats_val = [Paragraph(val, ParagraphStyle(
            f"sv_{i}", fontName="Helvetica-Bold", fontSize=20, leading=24,
            textColor=clr, alignment=TA_CENTER,
        )) for i, (_lbl, val, clr, _bg) in enumerate(_STAT_CONFIGS)]

        stats_tbl = Table([stats_hdr, stats_val], colWidths=[cw] * 5)
        _stat_styles = [
            ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, 0),  5),
            ("BOTTOMPADDING", (0, 0), (-1, 0),  3),
            ("TOPPADDING",    (0, 1), (-1, 1),  4),
            ("BOTTOMPADDING", (0, 1), (-1, 1),  8),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
            ("BOX",           (0, 0), (-1, -1), 1.5, _BORDER),
            ("LINEAFTER",     (0, 0), (3, -1),  0.5, _BORDER),
        ]
        for i, (_lbl, _val, _clr, bg) in enumerate(_STAT_CONFIGS):
            _stat_styles.append(("BACKGROUND", (i, 0), (i, -1), bg))
        stats_tbl.setStyle(TableStyle(_stat_styles))

        # ── Main step results table ────────────────────────────────────
        _N_COLS = 7
        col_headers_display = ["#", "Page URL", "Action", "Element", "Steps", "Script Runner", "Status"]
        from reportlab.platypus import PageBreak

        def _make_hdr_row_run():
            return [Paragraph(h, _tbl_hdr) for h in col_headers_display]

        step_flowables = []
        for step_idx, (r, rr) in enumerate(zip(rows, rows_raw)):
            status_text = str(r[6])
            action_text = str(r[2])
            if status_text.lower() == "pass":
                bg = colors.HexColor("#f0fdf4")
            elif status_text.lower() == "fail":
                bg = colors.HexColor("#fff1f2")
            else:
                bg = colors.white if step_idx % 2 == 0 else _ROW_ALT
            row_cells = [
                Paragraph(str(r[0]), _cell),
                Paragraph(str(r[1] or ""), _cell),
                _action_para(action_text),
                Paragraph(str(r[3] or ""), _cell),
                Paragraph(str(r[4] or ""), _cell),
                Paragraph(str(r[5] or ""), _cell),
                _status_para(status_text),
            ]
            _ss_raw = rr.get("screenshot")
            if _ss_raw:
                try:
                    _ss_img = RLImage(io.BytesIO(bytes(_ss_raw)), width=720, height=430, kind="bound")
                except Exception:
                    _ss_img = Paragraph("", _cell)
            else:
                _ss_img = Paragraph("", _cell)

            step_tbl = Table(
                [_make_hdr_row_run(), row_cells, [_ss_img] + [Paragraph("", _cell)] * (_N_COLS - 1)],
                colWidths=_COL_WIDTHS, hAlign="CENTER",
            )
            step_tbl.setStyle(TableStyle([
                # Header row (0)
                ("BACKGROUND",    (0, 0), (-1, 0), _NAVY),
                ("LINEBELOW",     (0, 0), (-1, 0), 2, _BLUE),
                ("ALIGN",         (0, 0), (-1, 0), "CENTER"),
                ("VALIGN",        (0, 0), (-1, 0), "MIDDLE"),
                ("GRID",          (0, 0), (-1, 0), 0.4, _BORDER),
                ("LEFTPADDING",   (0, 0), (-1, 0), 6),
                ("RIGHTPADDING",  (0, 0), (-1, 0), 6),
                ("TOPPADDING",    (0, 0), (-1, 0), 6),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                # Data row (1)
                ("BACKGROUND",    (0, 1), (-1, 1), bg),
                ("GRID",          (0, 1), (-1, 1), 0.4, _BORDER),
                ("LINEBELOW",     (0, 1), (-1, 1), 1.5, _BORDER),
                ("ALIGN",         (0, 1), (-1, 1), "LEFT"),
                ("VALIGN",        (0, 1), (-1, 1), "MIDDLE"),
                ("LEFTPADDING",   (0, 1), (-1, 1), 6),
                ("RIGHTPADDING",  (0, 1), (-1, 1), 6),
                ("TOPPADDING",    (0, 1), (-1, 1), 5),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 5),
                # Screenshot row (2) — spans full width
                ("BACKGROUND",    (0, 2), (-1, 2), bg),
                ("SPAN",          (0, 2), (_N_COLS - 1, 2)),
                ("ALIGN",         (0, 2), (-1, 2), "CENTER"),
                ("VALIGN",        (0, 2), (-1, 2), "MIDDLE"),
                ("TOPPADDING",    (0, 2), (-1, 2), 16),
                ("BOTTOMPADDING", (0, 2), (-1, 2), 16),
                ("LEFTPADDING",   (0, 2), (-1, 2), 0),
                ("RIGHTPADDING",  (0, 2), (-1, 2), 0),
                ("LINEBELOW",     (0, 2), (-1, 2), 0.4, _BORDER),
            ]))
            step_flowables.append(PageBreak())
            step_flowables.append(step_tbl)

        # ── Assemble document ──────────────────────────────────────────
        elements = [
            banner,
            Spacer(1, 8),
            meta_tbl,
            Spacer(1, 10),
            HRFlowable(width="100%", thickness=1, color=_BORDER),
            Spacer(1, 8),
            stats_tbl,
        ] + step_flowables

        doc.build(elements, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
        buffer.seek(0)
        response = _HR(buffer.read(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{safe_name_run}{engine_suffix}.pdf"'
        return response

    if fmt == "doc":
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Mm, Inches
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from django.http import HttpResponse as _HR

        def _set_cell_bg(cell, hex_color):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color)
            tcPr.append(shd)

        _STATUS_RGB = {
            "pass":         RGBColor(0x19, 0x87, 0x54),
            "fail":         RGBColor(0xDC, 0x35, 0x45),
            "not_executed": RGBColor(0x6C, 0x75, 0x7D),
        }
        _ACTION_RGB = {
            "click":      RGBColor(0x25, 0x63, 0xa8),
            "dblclick":   RGBColor(0x25, 0x63, 0xa8),
            "input":      RGBColor(0xd9, 0x77, 0x06),
            "change":     RGBColor(0xd9, 0x77, 0x06),
            "keydown":    RGBColor(0x6c, 0x75, 0x7d),
            "submit":     RGBColor(0x19, 0x87, 0x54),
            "navigate":   RGBColor(0x7c, 0x3a, 0xed),
            "navigation": RGBColor(0x7c, 0x3a, 0xed),
        }

        document = DocxDocument()
        section = document.sections[0]
        section.orientation   = WD_ORIENT.LANDSCAPE
        section.page_width    = Mm(297)
        section.page_height   = Mm(210)
        section.left_margin   = Mm(12)
        section.right_margin  = Mm(12)
        section.top_margin    = Mm(14)
        section.bottom_margin = Mm(14)

        hdg = document.add_heading(f"{record_name} \u2014 Run Results", 0)
        hdg.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Meta block
        meta_tbl = document.add_table(rows=len(meta_fields), cols=2)
        meta_tbl.autofit = False
        meta_tbl.allow_autofit = False
        for i, (label, value) in enumerate(meta_fields):
            lc = meta_tbl.rows[i].cells[0]
            lc.width = Mm(30)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lr = lp.add_run(f"{label}:")
            lr.bold = True; lr.font.size = Pt(9)
            vc = meta_tbl.rows[i].cells[1]
            vc.width = Mm(100)
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            vp.add_run(value).font.size = Pt(9)
        document.add_paragraph()

        # Stats tiles — Total / Passed / Failed / Not Executed / Pass Rate
        _doc_total  = len(rows_raw)
        _doc_passed = sum(1 for r in rows_raw if r["status"] == "pass")
        _doc_failed = sum(1 for r in rows_raw if r["status"] == "fail")
        _doc_ne     = _doc_total - _doc_passed - _doc_failed
        _doc_pct    = round(_doc_passed / _doc_total * 100) if _doc_total else 0
        _doc_pct_fg = "198754" if _doc_pct >= 80 else ("d97706" if _doc_pct >= 50 else "dc3545")
        _doc_pct_bg = "d1fae5" if _doc_pct >= 80 else ("fef3c7" if _doc_pct >= 50 else "fee2e2")
        _STAT_DATA = [
            ("TOTAL STEPS",  str(_doc_total),        "1e3c5e", "e0eaf8"),
            ("PASSED",       str(_doc_passed),        "198754", "d1fae5"),
            ("FAILED",       str(_doc_failed),        "dc3545", "fee2e2"),
            ("NOT EXECUTED", str(_doc_ne),            "6c757d", "f3f4f6"),
            ("PASS RATE",    f"{_doc_pct}%",         _doc_pct_fg, _doc_pct_bg),
        ]
        stats_tbl = document.add_table(rows=2, cols=5)
        stats_tbl.autofit = False
        stats_tbl.allow_autofit = False
        _tile_w = Mm(51)
        for ci in range(5):
            for ri in range(2):
                stats_tbl.rows[ri].cells[ci].width = _tile_w
        for ci, (label, value, fg, bg) in enumerate(_STAT_DATA):
            lc = stats_tbl.rows[0].cells[ci]
            _set_cell_bg(lc, bg)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = lp.add_run(label)
            lr.bold = True; lr.font.size = Pt(7)
            lr.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
            vc = stats_tbl.rows[1].cells[ci]
            _set_cell_bg(vc, bg)
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr = vp.add_run(value)
            vr.bold = True; vr.font.size = Pt(18)
            vr.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
        document.add_paragraph()

        # One page per step — each step gets a page break + its own 3-row table
        _N_COLS   = 7
        _COL_W_MM = [9, 58, 22, 22, 75, 34, 22]  # total ≈ 242 mm
        _doc_col_headers = ["#", "Page URL", "Action", "Element", "Steps", "Script Runner", "Status"]

        def _add_step_table_run(step_idx, row_data):
            document.add_page_break()
            tbl = document.add_table(rows=3, cols=_N_COLS)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            tbl.allow_autofit = False
            for ci, mm in enumerate(_COL_W_MM):
                for r in tbl.rows:
                    r.cells[ci].width = Mm(mm)

            # Header row (row 0)
            for ci, h in enumerate(_doc_col_headers):
                cell = tbl.rows[0].cells[ci]
                _set_cell_bg(cell, "1e3c5e")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

            # Data row (row 1)
            status_val = str(row_data[6]).lower()
            bg_hex = "f0fdf4" if status_val == "pass" else ("fff1f2" if status_val == "fail" else ("f0f6fc" if step_idx % 2 == 0 else "ffffff"))
            for ci, val in enumerate(row_data[:7]):
                cell = tbl.rows[1].cells[ci]
                _set_cell_bg(cell, bg_hex)
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(7.5)
                if ci == 2:
                    rgb = _ACTION_RGB.get(str(val).lower())
                    if rgb:
                        run.font.color.rgb = rgb
                        run.bold = True
                elif ci == 6:
                    rgb = _STATUS_RGB.get(str(val).lower())
                    if rgb:
                        run.font.color.rgb = rgb
                        run.bold = True

            # Screenshot row (row 2) — merged full width
            ss_cells = tbl.rows[2].cells
            merged = ss_cells[0].merge(ss_cells[_N_COLS - 1])
            _set_cell_bg(merged, bg_hex)
            return merged

        for step_idx, (row_data, rr) in enumerate(zip(rows, rows_raw)):
            merged  = _add_step_table_run(step_idx, row_data)
            _ss_raw = rr.get("screenshot")
            if _ss_raw:
                try:
                    ss_para = merged.paragraphs[0]
                    ss_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ss_run  = ss_para.add_run()
                    ss_run.add_picture(io.BytesIO(bytes(_ss_raw)), width=Inches(8.0))
                except Exception:
                    merged.paragraphs[0].add_run("").font.size = Pt(1)
            else:
                merged.paragraphs[0].add_run("").font.size = Pt(1)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = _HR(
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{safe_name_run}.docx"'
        return response

    if fmt == "playwright":
        bundle = _load_record_bundle(str(first["record_id_str"]))
        script = _build_playwright_script(bundle)
        from django.http import HttpResponse as _HR2
        response = _HR2(script, content_type="text/x-python; charset=utf-8")
        _pw_name = safe_name_run if 'safe_name_run' in dir() else str(first["record_id_str"])[:8]
        response["Content-Disposition"] = f'attachment; filename="{_pw_name}_playwright.py"'
        return response

    raise Http404(f"Unknown format: {fmt}")


# ---------------------------------------------------------------------------
# Active runs JSON API  (used by the /sessions/ floating panel)
# ---------------------------------------------------------------------------

def active_runs_api(request):
    """GET: return replay jobs as JSON, scoped by tenant.

    Superusers see every job.
    Tenant-aware users see only jobs that share their tenant_id.
    Users without a tenant assignment see their own jobs only (by runner name).
    """
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"detail": "authentication required"}, status=403)

    is_super = user.is_superuser
    _req_tenant = getattr(request, "tenant_id", None) or _get_user_tenant_id(user)
    _req_user   = user.username
    rows = []
    for _run_id, _job in list(_REPLAY_JOBS.items()):
        runner     = _job.get("runner", "")
        _job_tenant = _job.get("tenant_id")

        # Visibility rules (non-superuser path):
        #  1. Same tenant  → visible to all members
        #  2. No tenant assigned → fall back to own jobs only
        if not is_super:
            if _req_tenant and _job_tenant:
                if _job_tenant != _req_tenant:
                    continue  # different tenant – hidden
            elif runner != _req_user:
                continue  # no tenant, show only own jobs

        with _job["lock"]:
            status = _job.get("status", "")
            done = len(_job["results"])
        rows.append({
            "run_id":      _run_id,
            "record_id":  _job.get("record_id", ""),
            "record_name": _job.get("record_name", ""),
            "folder_name": _job.get("folder_name", ""),
            "session_name": _job.get("session_name", ""),
            "status":      status,
            "total":       _job.get("total", 0),
            "done":        done,
            "runner":      runner,
            "started_at":  _job.get("started_at"),
            "tenant_id":   str(_job_tenant) if _job_tenant else None,
        })
    # Prune stale edit locks before returning
    _now = time.time()
    with _EDITING_LOCK:
        stale = [rid for rid, info in _EDITING_SESSIONS.items()
                 if _now - info["acquired_at"] > _EDIT_LOCK_TTL]
        for rid in stale:
            _EDITING_SESSIONS.pop(rid, None)
        editing = {rid: info["user"] for rid, info in _EDITING_SESSIONS.items()}
    return JsonResponse({"runs": rows, "is_superuser": is_super, "editing": editing})


# ---------------------------------------------------------------------------
# Edit-lock API  (acquire / heartbeat / release)
# ---------------------------------------------------------------------------

@login_required
def acquire_edit_lock(request, record_id):
    """POST: claim the edit lock for a session. Returns {"ok": true} if acquired,
    {"ok": false, "editor": "<user>"} if already held by someone else."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    sid  = str(record_id)
    user = request.user.username
    now  = time.time()
    with _EDITING_LOCK:
        existing = _EDITING_SESSIONS.get(sid)
        if existing:
            # Expired lock → reassign
            if now - existing["acquired_at"] > _EDIT_LOCK_TTL:
                _EDITING_SESSIONS[sid] = {"user": user, "acquired_at": now}
                return JsonResponse({"ok": True})
            # Same user refreshes their own lock
            if existing["user"] == user:
                existing["acquired_at"] = now
                return JsonResponse({"ok": True})
            # Held by another user
            return JsonResponse({"ok": False, "editor": existing["user"]})
        _EDITING_SESSIONS[sid] = {"user": user, "acquired_at": now}
    return JsonResponse({"ok": True})


@login_required
def edit_lock_heartbeat(request, record_id):
    """POST: refresh the edit lock TTL so the tab is still considered open."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    sid  = str(record_id)
    user = request.user.username
    now  = time.time()
    with _EDITING_LOCK:
        existing = _EDITING_SESSIONS.get(sid)
        if existing and existing["user"] == user:
            existing["acquired_at"] = now
            return JsonResponse({"ok": True})
    return JsonResponse({"ok": False})


@login_required
def release_edit_lock(request, record_id):
    """POST: release the edit lock when the user navigates away."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    sid  = str(record_id)
    user = request.user.username
    with _EDITING_LOCK:
        existing = _EDITING_SESSIONS.get(sid)
        if existing and existing["user"] == user:
            _EDITING_SESSIONS.pop(sid, None)
    return JsonResponse({"ok": True})


@login_required
def scrape_monitor(request):
    """Standalone scrape progress monitor — floating window like Active Runs Monitor."""
    return render(request, "recorder/scrape_monitor.html")


def recording_monitor(request):
    """Standalone Recording Monitor popup — shows pause/stop controls for the active recording."""
    user = _resolve_monitor_user(request)
    if user is None:
        return redirect("login")
    record_id = request.GET.get("record_id", "")
    record_name = request.GET.get("record_name", "")
    if not record_name and record_id:
        try:
            record_name = SessionMeta.objects.get(record_id=record_id).record_name or ""
        except SessionMeta.DoesNotExist:
            pass
    return render(request, "recorder/recording_monitor.html", {
        "record_id": record_id,
        "record_name": record_name,
        "monitor_token": request.GET.get("monitor_token", ""),
    })


def active_runs_monitor(request):
    """Standalone popup page — polls /api/active-runs/ independently."""
    user = _resolve_monitor_user(request)
    if user is None:
        return redirect("login")
    return render(request, "recorder/active_runs_monitor.html", {
        "is_superuser":    user.is_superuser,
        "current_username": user.username,
        "monitor_token":   request.GET.get("monitor_token", ""),
    })


@csrf_exempt
def clear_run(request, run_id):
    """POST: remove a single finished job from _REPLAY_JOBS."""
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"ok": False, "error": "Authentication required."}, status=403)
    rid = str(run_id)
    with _JOBS_LOCK:
        job = _REPLAY_JOBS.get(rid)
        if job:
            # Non-superusers can only clear their own jobs
            if not user.is_superuser and job.get("runner") != user.username:
                return JsonResponse({"ok": False, "error": "Access denied."}, status=403)
            if job.get("status") in ("done", "error", "stopped"):
                _REPLAY_JOBS.pop(rid, None)
                _SESSION_TO_RUN.pop((job.get("record_id", ""), job.get("runner", "")), None)
    return JsonResponse({"ok": True})


@csrf_exempt
def clear_all_runs(request):
    """POST: remove all finished jobs from _REPLAY_JOBS.
    Superusers clear every finished job; regular users clear only their own.
    """
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"ok": False, "error": "Authentication required."}, status=403)
    with _JOBS_LOCK:
        finished = [
            rid for rid, j in list(_REPLAY_JOBS.items())
            if j.get("status") in ("done", "error", "stopped")
            and (user.is_superuser or j.get("runner") == user.username)
        ]
        for rid in finished:
            job = _REPLAY_JOBS.pop(rid, None)
            if job:
                _SESSION_TO_RUN.pop((job.get("record_id", ""), job.get("runner", "")), None)
    return JsonResponse({"ok": True, "cleared": len(finished)})


def stop_all_runs(request):
    """POST: stop all running/paused jobs (superuser only), then clear every job."""
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"ok": False, "error": "Authentication required."}, status=403)
    if request.method != "POST":
        return JsonResponse({"ok": False}, status=405)
    if not user.is_superuser:
        return JsonResponse({"ok": False, "error": "Superuser required."}, status=403)
    with _JOBS_LOCK:
        stopped = 0
        for _job in list(_REPLAY_JOBS.values()):
            if _job.get("status") in ("running", "paused"):
                _job["pause_event"].clear()  # unblock if paused
                _job["stop_event"].set()
                with _job["lock"]:
                    _job["status"] = "stopped"
                stopped += 1
        # Clear every job (stopped + finished) so the list is empty
        for rid in list(_REPLAY_JOBS.keys()):
            _REPLAY_JOBS.pop(rid, None)
        _SESSION_TO_RUN.clear()
    return JsonResponse({"ok": True, "stopped": stopped})


# ---------------------------------------------------------------------------
# Set headless_state for a session
# ---------------------------------------------------------------------------

@login_required
def set_headless_state(request, record_id):
    """POST (JSON): persist headless_state on steps/recordings for this session."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        body = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": "invalid JSON"}, status=400)
    headless    = bool(body.get("headless", False))
    folder_name = body.get("folder_name", "").strip()
    sid = str(record_id)
    with connection.cursor() as cur:
        if _is_recordings_folder_name(folder_name) or not folder_name:
            cur.execute(
                "UPDATE steps SET headless_state = %s WHERE record_id = %s",
                [headless, sid],
            )
        else:
            cur.execute(
                "UPDATE steps SET headless_state = %s WHERE record_id = %s AND folder_name = %s",
                [headless, sid, folder_name],
            )
    return JsonResponse({"ok": True, "headless": headless})


@login_required
def set_engine(request, record_id):
    """POST: update the engine field on session_meta for this session."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    engine = (request.POST.get("engine") or "selenium").strip().lower()
    if engine not in ("selenium", "playwright"):
        engine = "selenium"
    sid = str(record_id)
    with connection.cursor() as cur:
        cur.execute(
            "UPDATE session_meta SET engine = %s WHERE record_id = %s",
            [engine, sid],
        )
        cur.execute(
            "UPDATE steps SET engine = %s WHERE record_id = %s",
            [engine, sid],
        )
    return JsonResponse({"ok": True, "engine": engine})


@login_required
def save_playwright_script(request, record_id):
    """POST: save the edited Playwright script to disk and populate step codes."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    script = request.POST.get("script", "")
    if not script.strip():
        return JsonResponse({"error": "empty script"}, status=400)
    import os
    scripts_dir = os.path.join(settings.BASE_DIR, "generated_assets", "playwright_scripts")
    os.makedirs(scripts_dir, exist_ok=True)
    fname = str(record_id) + ".py"
    fpath = os.path.join(scripts_dir, fname)
    with open(fpath, "w", encoding="utf-8") as f:
        f.write(script)
    # Also populate per-step playwright_code in the DB
    _populate_playwright_code(str(record_id))
    return JsonResponse({"ok": True, "path": fpath})


@login_required
def save_selenium_script(request, record_id):
    """POST: save the edited Selenium script to disk."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    script = request.POST.get("script", "")
    if not script.strip():
        return JsonResponse({"error": "empty script"}, status=400)
    import os
    scripts_dir = os.path.join(settings.BASE_DIR, "generated_assets", "selenium_scripts")
    os.makedirs(scripts_dir, exist_ok=True)
    fname = str(record_id) + ".py"
    fpath = os.path.join(scripts_dir, fname)
    with open(fpath, "w", encoding="utf-8") as f:
        f.write(script)
    return JsonResponse({"ok": True, "path": fpath})


# ---------------------------------------------------------------------------
# Bulk copy selected sessions into a destination folder
# ---------------------------------------------------------------------------

@login_required
def bulk_copy(request):
    """POST: copy each submitted record_id into dest_folder under new record_ids."""
    if request.method != "POST":
        return redirect("sessions_list")

    raw_ids     = request.POST.getlist("record_ids")
    dest_folder = _normalize_folder_path(request.POST.get("dest_folder", ""))
    if _is_recordings_folder_name(dest_folder):
        dest_folder = _get_recordings_folder_label()

    valid_ids = []
    for rid in raw_ids:
        try:
            valid_ids.append(str(uuid.UUID(rid)))
        except (ValueError, AttributeError):
            pass

    if not valid_ids:
        messages.warning(request, "No sessions selected.")
        return redirect("sessions_list")
    if not dest_folder:
        messages.error(request, "Destination folder name is required.")
        return redirect("sessions_list")

    from collections import defaultdict

    copied = 0
    for sid in valid_ids:
        try:
            src_meta    = SessionMeta.objects.get(record_id=sid)
            record_name = src_meta.record_name or ""
            recorder    = src_meta.recorder    or ""
        except SessionMeta.DoesNotExist:
            src_meta    = None
            record_name = ""
            recorder    = ""

        with transaction.atomic():
            with connection.cursor() as cursor:
                if not _is_recordings_folder_name(dest_folder):
                    _register_project_folder(cursor, dest_folder)
                new_sid        = str(uuid.uuid4())
                new_file_order = _next_distinct_file_order(cursor, dest_folder)
                target_record_name = _resolve_copy_record_name(
                    cursor,
                    dest_folder,
                    record_name or sid,
                    exclude_record_id=sid,
                )

                _src_baseline = getattr(src_meta, 'is_baseline', False) if src_meta else False
                cursor.execute("""
                    INSERT INTO session_meta (record_id, record_name, recorder, folder_name, is_baseline, created_at)
                    VALUES (%s, %s, %s, %s, %s, NOW())
                    ON CONFLICT DO NOTHING
                """, [new_sid, target_record_name, recorder, dest_folder, _src_baseline])

                cursor.execute("""
                    SELECT DISTINCT ON (s.step_no)
                           s.step_no, s.action, s.page_url, s.element_tag,
                           d.field_name, d.value AS field_value,
                           s.raw_event, s.recorder, s.runner,
                           s.author, s.last_updated_by, s.headless_state,
                           s.is_baseline, s.file_type
                    FROM steps s
                    LEFT JOIN data d ON d.record_id = s.record_id AND d.step_no = s.step_no
                    WHERE s.record_id = %s
                    ORDER BY s.step_no, s.id
                """, [sid])
                rec_rows = cursor.fetchall()

                cursor.execute("""
                    SELECT step_no, strategy, locator, is_primary,
                           locator_rank, pos_x, pos_y
                    FROM locators
                    WHERE record_id = %s
                    ORDER BY step_no, COALESCE(locator_rank, 999), id
                """, [sid])
                locs_by_step: dict = defaultdict(list)
                for (l_step, l_strat, l_loc, l_primary, l_rank, l_px, l_py) in cursor.fetchall():
                    locs_by_step[l_step].append((l_strat, l_loc, l_primary, l_rank, l_px, l_py))

                for (step_no, action, page_url, element_tag,
                     field_name, field_value, raw_event, recorder_val, runner,
                     author, last_updated_by, headless_state, is_baseline, file_type) in rec_rows:

                    new_data_id = None
                    if field_name is not None or field_value is not None:
                        cursor.execute("""
                            INSERT INTO data (record_id, step_no, field_name, value, folder_name, created_at)
                            VALUES (%s, %s, %s, %s, %s, NOW())
                            RETURNING id
                        """, [new_sid, step_no, field_name, field_value, dest_folder])
                        new_data_id = cursor.fetchone()[0]

                    new_locator_id = None
                    for (l_strat, l_loc, l_primary, l_rank, l_px, l_py) in locs_by_step.get(step_no, []):
                        cursor.execute("""
                            INSERT INTO locators
                                (record_id, step_no, strategy, locator, is_primary,
                                 locator_rank, pos_x, pos_y, folder_name, created_at)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                            RETURNING id
                        """, [new_sid, step_no, l_strat, l_loc, l_primary, l_rank, l_px, l_py, dest_folder])
                        row_id = cursor.fetchone()[0]
                        if l_primary and new_locator_id is None:
                            new_locator_id = row_id

                    cursor.execute("""
                        INSERT INTO steps
                            (record_id, step_no, action, page_url, element_tag,
                             locator_id, data_id, raw_event, recorder, runner,
                             author, last_updated_by, headless_state, is_baseline,
                             file_type, folder_name, file_order, created_at)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                    """, [
                        new_sid, step_no, action, page_url, element_tag,
                        new_locator_id, new_data_id,
                        raw_event, recorder_val, runner,
                        author, last_updated_by, headless_state, is_baseline,
                        file_type, dest_folder, new_file_order,
                    ])

                _p_id, _s_id, _e_id = _resolve_folder_ids(
                    dest_folder, recorder or request.user.username, False
                )
                for _tbl in ("recordings", "steps", "run_table", "session_meta"):
                    cursor.execute(
                        f"""UPDATE {_tbl}
                               SET parent_folder_id = %s,
                                   sub_folder_id    = %s,
                                   end_folder_id    = %s
                             WHERE record_id = %s""",
                        [_p_id, _s_id, _e_id, new_sid],
                    )
        copied += 1

    messages.success(
        request,
        f'{copied} recording{"s" if copied != 1 else ""} copied to "{dest_folder}".'
    )
    return redirect("sessions_list")


# ---------------------------------------------------------------------------
# Bulk delete stats (AJAX preview)
# ---------------------------------------------------------------------------

@login_required
def bulk_delete_stats(request):
    """POST: return aggregate stats for selected record_ids as JSON."""
    if request.method != "POST":
        return JsonResponse({"error": "Method not allowed"}, status=405)

    raw_ids = request.POST.getlist("record_ids")
    valid_ids = []
    for rid in raw_ids:
        try:
            valid_ids.append(str(uuid.UUID(rid)))
        except (ValueError, AttributeError):
            pass

    if not valid_ids:
        return JsonResponse({"file_count": 0, "runs": 0, "steps": 0, "locators": 0, "data_entries": 0})

    ph = ", ".join(["%s"] * len(valid_ids))
    with connection.cursor() as cur:
        cur.execute(f"SELECT COUNT(DISTINCT run_id) FROM run_table WHERE record_id IN ({ph})", valid_ids)
        runs = cur.fetchone()[0] or 0

        cur.execute(f"SELECT COUNT(*) FROM steps WHERE record_id IN ({ph})", valid_ids)
        steps = cur.fetchone()[0] or 0

        cur.execute(f"SELECT COUNT(*) FROM locators WHERE record_id IN ({ph})", valid_ids)
        locators = cur.fetchone()[0] or 0

        cur.execute(f"SELECT COUNT(*) FROM data WHERE record_id IN ({ph})", valid_ids)
        data_entries = cur.fetchone()[0] or 0

    return JsonResponse({
        "file_count":   len(valid_ids),
        "runs":         runs,
        "steps":        steps,
        "locators":     locators,
        "data_entries": data_entries,
    })


# Bulk delete sessions
# ---------------------------------------------------------------------------

@login_required
def bulk_delete(request):
    """POST: delete every submitted record_id across all related tables."""
    if request.method != "POST":
        return redirect("sessions_list")

    raw_ids = request.POST.getlist("record_ids")

    valid_ids = []
    for rid in raw_ids:
        try:
            sid = str(uuid.UUID(rid))
            if sid not in valid_ids:
                valid_ids.append(sid)
        except (ValueError, AttributeError):
            pass

    if not valid_ids:
        messages.warning(request, "No sessions selected.")
        return redirect("sessions_list")

    deleted = 0
    affected_folders = set()
    with transaction.atomic():
        with connection.cursor() as cur:
            for sid in valid_ids:
                cur.execute(
                    """
                    SELECT DISTINCT TRIM(COALESCE(folder_name, ''))
                    FROM steps
                    WHERE record_id = %s
                    UNION
                    SELECT DISTINCT TRIM(COALESCE(folder_name, ''))
                    FROM session_meta
                    WHERE record_id = %s
                    """,
                    [sid, sid],
                )
                for row in cur.fetchall():
                    folder_name = (row[0] or "").strip() or "Baseline"
                    affected_folders.add(folder_name)

                cur.execute("UPDATE run_table SET locator_id = NULL WHERE record_id = %s AND locator_id IS NOT NULL", [sid])
                cur.execute("UPDATE run_table SET data_id    = NULL WHERE record_id = %s AND data_id    IS NOT NULL", [sid])
                cur.execute("UPDATE steps     SET locator_id = NULL WHERE record_id = %s AND locator_id IS NOT NULL", [sid])
                cur.execute("UPDATE steps     SET data_id    = NULL WHERE record_id = %s AND data_id    IS NOT NULL", [sid])
                cur.execute("UPDATE recordings SET locator_id = NULL WHERE record_id = %s AND locator_id IS NOT NULL", [sid])
                cur.execute("UPDATE recordings SET data_id    = NULL WHERE record_id = %s AND data_id    IS NOT NULL", [sid])
                cur.execute("DELETE FROM locators          WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM data              WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM steps             WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM recordings        WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM run_table         WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM locators_stat     WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM remote_executions WHERE record_id = %s", [sid])
                cur.execute("DELETE FROM session_meta      WHERE record_id = %s", [sid])
                deleted += 1

            for folder_name in affected_folders:
                _resequence_file_order_for_folder(folder_name)

    messages.success(request, f'{deleted} recording{"s" if deleted != 1 else ""} deleted.')
    return redirect("sessions_list")


# Bulk headless replay
# ---------------------------------------------------------------------------

@login_required
def bulk_replay(request):
    """POST: start a headless replay thread for every submitted record_id."""
    if request.method != "POST":
        return redirect("sessions_list")

    raw_ids = request.POST.getlist("record_ids")
    raw_api_ids = request.POST.getlist("api_testcase_ids")
    if not raw_ids and not raw_api_ids:
        messages.warning(request, "No sessions or API test cases selected.")
        return redirect("sessions_list")

    started = 0
    api_started = 0
    headless = request.POST.get("headless") == "on"
    _runner_name = request.user.username
    _tenant_id = getattr(request, "tenant_id", None)

    # Serial mode: constrain the shared executor to 1 concurrent job by
    # wrapping each submission in a tenant-scoped semaphore.
    _exec_mode = (get_config("replay.execution_mode") or "parallel").strip().lower()
    if _exec_mode == "serial":
        _serial_sem: threading.Semaphore | None = threading.Semaphore(1)
    else:
        _serial_sem = None

    def _dedup_steps(qs):
        seen: set[int] = set()
        result = []
        for s in qs:
            if s.step_no not in seen:
                seen.add(s.step_no)
                result.append(s)
        return result

    # --- Sort submitted IDs by folder hierarchy before launching threads --------
    _valid_ids: list[str] = []
    for _raw in raw_ids:
        try:
            _valid_ids.append(str(uuid.UUID(_raw.strip())))
        except ValueError:
            pass

    if _valid_ids:
        with connection.cursor() as _ord_cur:
            _ord_cur.execute("""
                SELECT record_id FROM (
                    SELECT DISTINCT ON (s.record_id)
                           s.record_id::text AS record_id,
                           COALESCE(pf.parent_folder_order, 2147483647) AS pfo,
                           CASE
                               WHEN s.sub_folder_id IS NULL AND s.end_folder_id IS NULL THEN 0
                               ELSE COALESCE(sf.sub_folder_order, 2147483647)
                           END AS sfo,
                           CASE WHEN s.end_folder_id IS NOT NULL THEN 1 ELSE 0 END AS is_end,
                           COALESCE(ef.end_folder_order, 2147483647) AS efo,
                           COALESCE(s.file_order,        2147483647) AS fo
                    FROM steps s
                    LEFT JOIN end_folders    ef ON ef.end_folder_id    = s.end_folder_id
                    LEFT JOIN sub_folders    sf ON sf.sub_folder_id    = COALESCE(s.sub_folder_id, ef.end_folder_parent)
                    LEFT JOIN parent_folders pf ON pf.parent_folder_id = COALESCE(s.parent_folder_id, sf.sub_folder_parent)
                    WHERE s.record_id = ANY(%s::uuid[])
                    ORDER BY s.record_id,
                             COALESCE(pf.parent_folder_order, 2147483647),
                             CASE
                                 WHEN s.sub_folder_id IS NULL AND s.end_folder_id IS NULL THEN 0
                                 ELSE COALESCE(sf.sub_folder_order, 2147483647)
                             END,
                             CASE WHEN s.end_folder_id IS NOT NULL THEN 1 ELSE 0 END,
                             COALESCE(ef.end_folder_order, 2147483647),
                             COALESCE(s.file_order,        2147483647)
                ) sub
                ORDER BY pfo, sfo, is_end, efo, fo
            """, [_valid_ids])
            _sorted_ids = [row[0] for row in _ord_cur.fetchall()]
        # Append any IDs not found in steps (edge case) at the end
        _seen_ids = set(_sorted_ids)
        _sorted_ids.extend(i for i in _valid_ids if i not in _seen_ids)
    else:
        _sorted_ids = []

    for sid in _sorted_ids:

        # Cancel the current user's existing job for this session (if any).
        # Do NOT touch jobs from other users -- they run independently.
        with _JOBS_LOCK:
            existing_run_id = _SESSION_TO_RUN.get((sid, _runner_name))
            if existing_run_id:
                existing = _REPLAY_JOBS.get(existing_run_id)
                if existing:
                    existing["stop_event"].set()

        steps = _dedup_steps(Step.objects.filter(record_id=sid).order_by("step_no", "id"))
        steps = sorted(steps, key=lambda s: s.step_no)
        if not steps:
            continue

        _raw_fn = getattr(steps[0], "folder_name", None)
        _replay_folder = _get_recordings_folder_label() if _is_recordings_folder_name(_raw_fn or "") else (_raw_fn or _get_recordings_folder_label())

        # Pre-run: compute all formulas (increment/decrement/calculate) in this folder
        try:
            compute_folder_formulas(_replay_folder)
        except Exception:
            pass

        try:
            _session_name = SessionMeta.objects.get(record_id=sid).record_name or ""
        except SessionMeta.DoesNotExist:
            _session_name = ""

        try:
            _session_engine = SessionMeta.objects.get(record_id=sid).engine or "selenium"
        except SessionMeta.DoesNotExist:
            _session_engine = "selenium"

        pause_event = threading.Event()
        stop_event  = threading.Event()
        run_id = str(uuid.uuid4())

        job = {
            "run_id":      run_id,
            "record_id":  sid,
            "results":     [],
            "status":      "running",
            "pause_event": pause_event,
            "stop_event":  stop_event,
            "total":       len(steps),
            "lock":        threading.Lock(),
            "started_at":  time.time(),
            "finished_at": None,
            "runner":      _runner_name,
            "record_name": _session_name,
            "folder_name": _replay_folder,
            "session_name": _session_name,
            "tenant_id":   _tenant_id,
            "engine":      _session_engine,
        }

        def _make_run(steps, _job, _sid, _run_id, _folder, _headless, _sem, _engine):
            def _run():
                import traceback
                close_old_connections()
                def _on_step(r):
                    with _job["lock"]:
                        _job["results"].append(r)
                        if r and r.get("ok") is False and _job.get("status") == "running":
                            _job["status"] = "failed"
                            _job["finished_at"] = time.time()
                if _sem is not None:
                    _sem.acquire()
                try:
                    if _engine == "playwright" and playwright_replay_session:
                        _replay_fn = playwright_replay_session
                    else:
                        _replay_fn = replay_session
                    _replay_fn(
                        _sid,
                        headless=_headless,
                        pause_event=_job["pause_event"],
                        stop_event=_job["stop_event"],
                        on_step=_on_step,
                        steps=steps,
                        run_id=_run_id,
                        runner=_job["runner"],
                        folder_name=_folder,
                    )
                    with _job["lock"]:
                        if _job["status"] in ("failed", "error", "stopped"):
                            pass  # already terminal — set by _on_step or stop
                        elif _job["stop_event"].is_set():
                            _job["status"] = "stopped"
                        elif any(not r.get("ok") for r in _job["results"]):
                            _job["status"] = "failed"
                        else:
                            _job["status"] = "done"
                        if not _job.get("finished_at"):
                            _job["finished_at"] = time.time()
                except Exception as exc:
                    tb = traceback.format_exc()
                    print(f"[BULK REPLAY ERROR] {exc}\n{tb}", flush=True)
                    with _job["lock"]:
                        _job["results"].append({
                            "step_no": "-", "action": "ERROR",
                            "page_url": "-", "status": f"{exc}",
                            "ok": False,
                        })
                        _job["status"] = "error"
                        _job["finished_at"] = time.time()
                finally:
                    if _sem is not None:
                        _sem.release()
                    close_old_connections()
            return _run

        # Register job BEFORE submitting to the pool so the run is immediately
        # visible in the active-runs monitor.
        with _JOBS_LOCK:
            _REPLAY_JOBS[run_id] = job
            _SESSION_TO_RUN[(sid, _runner_name)] = run_id
        _submit_replay_job(
            _make_run(steps, job, sid, run_id, _replay_folder, headless, _serial_sem, _session_engine)
        )
        started += 1

    # Execute selected API testcases and register them as monitor jobs.
    _valid_api_ids: list[int] = []
    for _raw_api in raw_api_ids:
        try:
            _valid_api_ids.append(int(str(_raw_api).strip()))
        except (TypeError, ValueError):
            pass

    if _valid_api_ids:
        try:
            from api_testcases.models import TestCase as ApiTestCase, Environment as ApiEnvironment
            from api_testcases.views import _execute_tc as _execute_api_tc

            _api_env = ApiEnvironment.objects.filter(is_active=True).order_by('id').first()
            if not _api_env:
                _api_env = ApiEnvironment.objects.order_by('id').first()

            if _api_env:
                _api_qs = ApiTestCase.objects.filter(pk__in=_valid_api_ids)
                _api_map = {tc.id: tc for tc in _api_qs}
                _ordered_api_tcs = [_api_map[_id] for _id in _valid_api_ids if _id in _api_map]

                def _make_api_run(_tc, _job, _sem):
                    def _run_api():
                        close_old_connections()
                        if _sem is not None:
                            _sem.acquire()
                        try:
                            if _job["stop_event"].is_set():
                                with _job["lock"]:
                                    _job["status"] = "stopped"
                                    _job["finished_at"] = time.time()
                                return

                            result = _execute_api_tc(_tc, _api_env, request.user, request)
                            ok = (str(result.get("result_status", "")).lower() == "passed")
                            with _job["lock"]:
                                _job["results"].append({
                                    "step_no": 1,
                                    "action": "API_EXECUTE",
                                    "page_url": result.get("request_url", "-"),
                                    "status": result.get("result_status", "done"),
                                    "ok": ok,
                                })
                                _job["status"] = "done" if ok else "failed"
                                _job["finished_at"] = time.time()
                        except Exception as _api_exc:
                            print(f"[BULK API EXECUTION ERROR] testcase={getattr(_tc, 'id', '?')} error={_api_exc}", flush=True)
                            with _job["lock"]:
                                _job["results"].append({
                                    "step_no": 1,
                                    "action": "API_EXECUTE",
                                    "page_url": "-",
                                    "status": str(_api_exc),
                                    "ok": False,
                                })
                                _job["status"] = "error"
                                _job["finished_at"] = time.time()
                        finally:
                            if _sem is not None:
                                _sem.release()
                            close_old_connections()
                    return _run_api

                for _tc in _ordered_api_tcs:
                    _api_run_id = str(uuid.uuid4())
                    _api_job = {
                        "run_id": _api_run_id,
                        "record_id": f"api:{_tc.id}",
                        "results": [],
                        "status": "running",
                        "pause_event": threading.Event(),
                        "stop_event": threading.Event(),
                        "total": 1,
                        "lock": threading.Lock(),
                        "started_at": time.time(),
                        "finished_at": None,
                        "runner": _runner_name,
                        "record_name": getattr(_tc, "name", f"API Test Case {_tc.id}"),
                        "folder_name": "API",
                        "session_name": getattr(_tc, "name", f"API Test Case {_tc.id}"),
                        "tenant_id": _tenant_id,
                        "engine": "api",
                    }
                    with _JOBS_LOCK:
                        _REPLAY_JOBS[_api_run_id] = _api_job
                    _submit_replay_job(_make_api_run(_tc, _api_job, _serial_sem if _exec_mode == "serial" else None))
                    api_started += 1
            else:
                messages.warning(request, "No active API environment found. API test cases were skipped.")
        except Exception as _bulk_api_exc:
            print(f"[BULK API EXECUTION SETUP ERROR] {_bulk_api_exc}", flush=True)

    if started or api_started:
        _parts = []
        if started:
            _parts.append(f"{started} replay{'s' if started != 1 else ''} ({'headless' if headless else 'with browser'})")
        if api_started:
            _parts.append(f"{api_started} API test case{'s' if api_started != 1 else ''}")
        messages.success(request, f"Started {' and '.join(_parts)}. Check the dashboard for progress.")
    else:
        messages.warning(request, "No valid sessions or API test cases were found to execute.")
    return redirect("sessions_list")


# ---------------------------------------------------------------------------
# Replay a session  (start / pause / stop / status / results page)
# ---------------------------------------------------------------------------

@csrf_exempt
def replay(request, record_id):
    """Confirmation page (GET) or start replay (POST)."""
    # GET requires authentication (viewing page)
    # POST can be from authenticated user OR remote backend
    if request.method == "GET" and not request.user.is_authenticated:
        return redirect(f"{settings.LOGIN_URL}?next={request.path}")
    
    sid = str(record_id)
    if request.method == "POST":
        headless  = request.POST.get("headless") == "on"
        base_url  = request.POST.get("base_url", "").strip()
        replay_engine = request.POST.get("replay_engine", "selenium").strip().lower()
        # Use authenticated user's name, or "remote" for backend requests
        _runner_name = request.user.username if request.user.is_authenticated else "remote"

        # ----------------------------------------------------------------
        # Kill THIS USER's recording PID, stop THIS USER's replay jobs,
        # and kill the browser — before starting a new replay.
        # ----------------------------------------------------------------
        import tempfile as _tf

        # 1. Kill the current user's active recording process.
        _rec_pid = request.session.pop("recording_pid", None)
        _popped_rec_id = request.session.pop("recording_id", None)
        if _rec_pid:
            try:
                if sys.platform == "win32":
                    subprocess.call(
                        ["taskkill", "/F", "/T", "/PID", str(_rec_pid)],
                        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                    )
                else:
                    os.kill(int(_rec_pid), signal.SIGTERM)
                print(f"[REPLAY] Killed recording PID {_rec_pid} before replay.", flush=True)
            except Exception:
                pass
            # Remove the pause flag for this recording session.
            _flag_rec_id = _popped_rec_id or sid
            _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{_flag_rec_id}.flag")
            try:
                os.remove(_flag)
            except OSError:
                pass

        # Clean up remaining recording session keys.
        request.session.pop("recording_url",          None)
        request.session.pop("recording_name",         None)
        request.session.pop("recording_folder",       None)
        request.session.pop("recording_is_baseline",  None)
        request.session.pop("recording_paused",       None)

        # 2. Stop this user's existing replay job for this session (if any).
        with _JOBS_LOCK:
            existing_run_id = _SESSION_TO_RUN.get((sid, _runner_name))
            if existing_run_id:
                existing = _REPLAY_JOBS.get(existing_run_id)
                if existing:
                    existing["stop_event"].set()

        # 3. Kill the Chrome browser process on the RDP port so a fresh browser
        #    is used for the new replay run (no stale tabs/state left behind).
        _rdp_str_now = (get_config("chrome.remote_debugging_port") or "").strip()
        if _rdp_str_now.isdigit() and int(_rdp_str_now) > 0:
            _kill_browser_on_port(int(_rdp_str_now))
            print(f"[REPLAY] Killed browser on RDP port {_rdp_str_now} before replay.", flush=True)

        pause_event = threading.Event()   # set = paused
        stop_event  = threading.Event()   # set = stop requested
        run_id = str(uuid.uuid4())

        # Load steps NOW in the request thread while DB connection is live
        # Try folder copies first (steps table), fall back to raw recordings.
        # The steps/recordings tables may have multiple rows per step_no when a session
        # has been copied into folders — deduplicate by step_no, keeping the first row
        # (lowest pk / natural DB order) to avoid executing each step N times.
        def _dedup_steps(qs):
            seen: set[int] = set()
            result = []
            for s in qs:
                if s.step_no not in seen:
                    seen.add(s.step_no)
                    result.append(s)
            return result

        steps = _dedup_steps(Step.objects.filter(record_id=record_id).order_by("step_no", "id"))
        steps = sorted(steps, key=lambda s: s.step_no)

        if not steps:
            messages.error(request, f"No recorded steps were found for session {sid}.")
            return redirect("sessions_list")

        _raw_fn = getattr(steps[0], "folder_name", None) if steps else None
        _replay_folder = _get_recordings_folder_label() if _is_recordings_folder_name(_raw_fn or "") else (_raw_fn or _get_recordings_folder_label())

        # Pre-run: compute all formulas (increment/decrement/calculate) in this folder
        try:
            compute_folder_formulas(_replay_folder)
        except Exception:
            pass

        try:
            _session_name = SessionMeta.objects.get(record_id=record_id).record_name or ""
        except SessionMeta.DoesNotExist:
            _session_name = ""

        _add_step_after = request.POST.get("add_step_after") == "1"
        _rdp_for_replay: int | None = None
        if _add_step_after:
            _rdp_str = (get_config("chrome.remote_debugging_port") or "").strip()
            if _rdp_str.isdigit() and int(_rdp_str) > 0:
                _rdp_for_replay = int(_rdp_str)

        _tenant_id = getattr(request, "tenant_id", None)

        job = {
            "run_id":      run_id,
            "record_id":  sid,
            "results":     [],
            "status":      "running",   # running | paused | stopped | done | error
            "pause_event": pause_event,
            "stop_event":  stop_event,
            "total":       len(steps),
            "lock":        threading.Lock(),
            "started_at":  time.time(),
            "finished_at": None,
            "runner":      _runner_name,
            "record_name": _session_name,
            "folder_name": _replay_folder,
            "session_name": _session_name,
            "tenant_id":   _tenant_id,
        }

        def _run(steps=steps, _job=job):
            import traceback
            # Close any inherited connection — thread gets its own fresh one if needed
            close_old_connections()
            # Capture _job directly so this thread never appends to a *later* job
            # that may have replaced _REPLAY_JOBS[sid] while this thread was still running.
            def _on_step(r):
                with _job["lock"]:
                    _job["results"].append(r)
                    if r and r.get("ok") is False and _job.get("status") == "running":
                        _job["status"] = "failed"
                        _job["finished_at"] = time.time()
            try:
                if replay_engine == "playwright" and playwright_replay_session:
                    _replay_fn = playwright_replay_session
                else:
                    _replay_fn = replay_session
                _replay_fn(
                    sid,
                    headless=headless,
                    pause_event=pause_event,
                    stop_event=stop_event,
                    on_step=_on_step,
                    steps=steps,
                    run_id=run_id,
                    runner=_runner_name,
                    folder_name=_replay_folder,
                    keep_open=bool(_rdp_for_replay),
                    rdp_port=_rdp_for_replay,
                )
                with job["lock"]:
                    if job["status"] in ("failed", "error", "stopped"):
                        pass  # already terminal — set by _on_step or stop
                    elif stop_event.is_set():
                        job["status"] = "stopped"
                    elif any(not r.get("ok") for r in job["results"]):
                        job["status"] = "failed"
                    else:
                        job["status"] = "done"
                    if not job.get("finished_at"):
                        job["finished_at"] = time.time()
            except Exception as exc:
                tb = traceback.format_exc()
                print(f"[REPLAY ERROR] {exc}\n{tb}", flush=True)
                with job["lock"]:
                    job["results"].append({
                        "step_no": "-", "action": "ERROR",
                        "page_url": "-", "status": f"{exc}",
                        "ok": False,
                    })
                    job["status"] = "error"
                    job["finished_at"] = time.time()
            finally:
                close_old_connections()

        # Register job BEFORE submitting to pool so it is immediately visible
        with _JOBS_LOCK:
            _REPLAY_JOBS[run_id] = job
            _SESSION_TO_RUN[(sid, _runner_name)] = run_id
        _submit_replay_job(_run)
        
        # Notify all pages to refresh (for remote execution tracking)
        global _LAST_REMOTE_ACTION
        _LAST_REMOTE_ACTION = time.time()

        if _add_step_after:
            request.session["replay_add_step_after"] = run_id

        return redirect("replay_run", run_id=run_id)

    # GET — always show the confirmation/start form
    sid = str(record_id)
    step_count = Step.objects.filter(record_id=record_id).values("step_no").distinct().count()
    try:
        _meta = SessionMeta.objects.get(record_id=record_id)
        record_name = _meta.record_name
        session_engine = _meta.engine or 'selenium'
    except SessionMeta.DoesNotExist:
        record_name = ""
        session_engine = 'selenium'
    return render(request, "recorder/replay.html", {
        "record_id": record_id,
        "record_name": record_name,
        "step_count": step_count,
        "has_job": False,
        "job_status": None,
        "job_run_id": "",
        "add_step_after": False,
        "session_engine": session_engine,
    })


def replay_run(request, run_id):
    """Live progress page for a specific replay run.

    Accepts either a Django session login OR a monitor_token query param so the
    detached monitor window can navigate here via its 'Watch run' button.
    """
    # Resolve user — session login takes priority, then monitor_token fallback.
    if not request.user.is_authenticated:
        user = _resolve_monitor_user(request)
        if user is None:
            from django.contrib.auth.views import redirect_to_login
            return redirect_to_login(request.get_full_path())
        request.user = user
    rid = str(run_id)
    job = _REPLAY_JOBS.get(rid)
    # Helper to preserve monitor_token across redirects
    def _redirect_to_run(r_id):
        from django.urls import reverse
        url = reverse("view_run", kwargs={"run_id": r_id})
        _mt = request.GET.get("monitor_token", "")
        if _mt:
            url += "?monitor_token=" + _mt
        return redirect(url)

    if not job:
        # Run not in memory (server restarted or stale URL).
        # Delegate to the completed-run view, which now handles stale IDs
        # gracefully (active job fallback or sessions redirect with warning).
        return _redirect_to_run(run_id)

    # Job already finished — show the completed run results.
    with job["lock"]:
        _current_status = job.get("status")
        _finished_at = job.get("finished_at") or 0
    if _current_status in ("done", "stopped", "failed", "error"):
        # Do not keep users on the live page once a run is terminal.
        # Route through the completed-run view, which handles stale/missing
        # run rows gracefully and redirects to Projects when needed.
        return _redirect_to_run(rid)

    record_id = job.get("record_id", "")
    step_count = job["total"]
    record_name = ""
    if record_id:
        try:
            record_name = SessionMeta.objects.get(record_id=record_id).record_name
        except SessionMeta.DoesNotExist:
            pass
    _asa = request.session.pop("replay_add_step_after", None)
    add_step_after = (_asa == rid)
    try:
        _meta_rr = SessionMeta.objects.get(record_id=record_id)
        session_engine = _meta_rr.engine or 'selenium'
    except SessionMeta.DoesNotExist:
        session_engine = 'selenium'
    return render(request, "recorder/replay.html", {
        "record_id": record_id,
        "record_name": record_name,
        "step_count": step_count,
        "has_job": True,
        "job_status": _current_status,
        "job_run_id": rid,
        "add_step_after": add_step_after,
        "session_engine": session_engine,
    })


@login_required
def remote_execute(request, record_id):
    """Handle remote execution requests."""
    if request.method != "POST":
        return redirect("sessions_list")
    
    remote_ip = request.POST.get("remote_ip", "").strip()
    remote_port = request.POST.get("remote_port", "8888").strip()
    headless = request.POST.get("headless") == "on"
    
    if not remote_ip:
        messages.error(request, "Remote IP address is required.")
        return redirect("last_run", record_id=record_id)
    
    # Validate IP format
    import re
    ip_pattern = r'^(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}|[a-zA-Z0-9.-]+)$'
    if not re.match(ip_pattern, remote_ip):
        messages.error(request, "Invalid IP address or hostname format.")
        return redirect("last_run", record_id=record_id)
    
    try:
        remote_port = int(remote_port)
        if remote_port < 1 or remote_port > 65535:
            raise ValueError("Port must be between 1 and 65535")
    except ValueError as e:
        messages.error(request, f"Invalid port: {str(e)}")
        return redirect("last_run", record_id=record_id)
    
    sessions_count = Recording.objects.filter(record_id=record_id).values("step_no").distinct().count()
    if not sessions_count:
        sessions_count = Step.objects.filter(record_id=record_id).values("step_no").distinct().count()
    
    try:
        session_meta = SessionMeta.objects.get(record_id=record_id)
        session_name = session_meta.record_name or ""
    except SessionMeta.DoesNotExist:
        session_name = ""
    
    # Save remote execution to database for future quick access
    global _LAST_REMOTE_ACTION
    try:
        RemoteExecution.objects.create(
            user=request.user.username if request.user else None,
            remote_ip=remote_ip,
            remote_port=int(remote_port),
            record_id=record_id,
            headless=headless
        )
        _LAST_REMOTE_ACTION = time.time()  # Track for page auto-refresh
    except Exception as e:
        print(f"[REMOTE] Warning: Could not save remote execution history: {e}", flush=True)

    # Upsert into RemoteTarget (distinct targets list, updates last_used each time)
    try:
        target, _ = RemoteTarget.objects.get_or_create(
            remote_ip=remote_ip,
            remote_port=int(remote_port),
        )
        target.save()  # touches auto_now last_used
    except Exception as e:
        print(f"[REMOTE] Warning: Could not upsert remote target: {e}", flush=True)
    
    # Construct the remote execution URL
    remote_url = f"http://{remote_ip}:{remote_port}/sessions/{record_id}/replay/"
    
    # Prepare POST data for remote execution
    remote_data = {
        "headless": "on" if headless else "off",
    }
    
    print(f"[REMOTE] Initiating remote execution on {remote_ip}:{remote_port} for session {record_id}", flush=True)
    print(f"[REMOTE] Session: {session_name}, Steps: {sessions_count}", flush=True)
    
    # Actually execute on the remote machine
    import requests
    try:
        # Do NOT follow redirects — the 302 Location header contains the run_id
        response = requests.post(remote_url, data=remote_data, timeout=10, allow_redirects=False)
        print(f"[REMOTE] Response: status={response.status_code}, Location={response.headers.get('Location', '')}", flush=True)
        if response.status_code in [302, 301]:
            location = response.headers.get("Location", "")
            print(f"[REMOTE] ✓ Remote execution started, redirecting to: {location}", flush=True)
            messages.success(request,
                f"Remote execution started on {remote_ip}:{remote_port} for '{session_name}' ({sessions_count} steps).")
            # Build absolute URL to the remote replay progress page
            remote_replay_url = f"http://{remote_ip}:{remote_port}{location}"
            return redirect(remote_replay_url)
        elif response.status_code == 200:
            print(f"[REMOTE] ✓ Remote returned 200 (no redirect)", flush=True)
            messages.success(request,
                f"Remote execution initiated on {remote_ip}:{remote_port} for session '{session_name}' ({sessions_count} steps). "
                f"Check the remote instance for progress.")
        elif 400 <= response.status_code < 500:
            print(f"[REMOTE] ✗ Remote returned client error {response.status_code}", flush=True)
            messages.error(request,
                f"Remote rejected request with status {response.status_code}. "
                f"Ensure the session exists on {remote_ip}:{remote_port} and the app is properly configured.")
        else:
            print(f"[REMOTE] ✗ Remote returned status {response.status_code}", flush=True)
            messages.warning(request,
                f"Remote returned status {response.status_code}. "
                f"Check {remote_ip}:{remote_port} is reachable and running.")
    except requests.exceptions.Timeout:
        print(f"[REMOTE] ✗ Remote execution timeout", flush=True)
        messages.error(request,
            f"Connection to {remote_ip}:{remote_port} timed out. "
            f"Ensure the remote machine is running and reachable.")
    except requests.exceptions.ConnectionError as e:
        print(f"[REMOTE] ✗ Connection error: {e}", flush=True)
        messages.error(request,
            f"Cannot connect to {remote_ip}:{remote_port}. "
            f"Ensure the remote machine is running and reachable.")
    except Exception as e:
        print(f"[REMOTE] ✗ Unexpected error: {e}", flush=True)
        messages.error(request,
            f"Error executing on remote: {str(e)}")

    # Fallback redirect (only reached if no 302 from remote)
    return redirect("last_run", record_id=record_id)


@login_required
def get_last_remote_ip(request):
    """Return all distinct remote targets, most-recently-used first."""
    try:
        targets = list(
            RemoteTarget.objects.values("remote_ip", "remote_port").order_by("-last_used")[:20]
        )
        if targets:
            return JsonResponse({
                'status': 'success',
                'remote_ip':   targets[0]['remote_ip'],
                'remote_port': targets[0]['remote_port'],
                'targets':     targets,
            })
        return JsonResponse({
            'status':      'not_found',
            'remote_ip':   '',
            'remote_port': 8888,
            'targets':     [],
        })
    except Exception as e:
        print(f"[API] Error fetching remote targets: {e}", flush=True)
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@csrf_exempt
def get_remote_action_status(request):
    """API endpoint to check if remote action was triggered (for page refresh)."""
    global _LAST_REMOTE_ACTION
    since = float(request.GET.get('since', 0))
    
    return JsonResponse({
        'last_action': _LAST_REMOTE_ACTION,
        'should_refresh': _LAST_REMOTE_ACTION > since
    })


@login_required
def update_session_order(request):
    """Update file_order for sessions based on drag-and-drop.

    Each item in record_ids must be {record_id, folder_name}.
    file_order is applied to ALL steps rows that share the same
    (record_id, folder_name) — i.e. every step of that session in
    that folder receives the same value (the session's position).
    """
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=400)

    try:
        data = json.loads(request.body)
        ordered_ids = data.get("record_ids", [])

        with connection.cursor() as cur:
            for order, item in enumerate(ordered_ids, start=1):
                try:
                    if isinstance(item, dict):
                        record_id   = str(item.get("record_id", "") or "").strip()
                        folder_name = item.get("folder_name", None)
                    else:
                        record_id   = str(item or "").strip()
                        folder_name = None

                    if not record_id:
                        continue

                    if folder_name is None:
                        # Legacy: update everything for this record_id
                        cur.execute(
                            "UPDATE steps      SET file_order = %s WHERE record_id = %s",
                            [order, record_id],
                        )
                        continue

                    folder_name = (folder_name or "").strip()
                    folder_key  = folder_name.lower()

                    if folder_key in {"", "baseline"} or _is_recordings_folder_name(folder_name):
                        # Baseline sessions live in steps (and possibly legacy recordings)
                        cur.execute("""
                            UPDATE steps SET file_order = %s
                            WHERE record_id = %s
                              AND LOWER(TRIM(COALESCE(folder_name, ''))) = ANY(%s)
                        """, [order, record_id, list(_recordings_sql_aliases())])

                    elif folder_key == "unfiled":
                        cur.execute("""
                            UPDATE steps SET file_order = %s
                            WHERE record_id = %s
                              AND LOWER(TRIM(COALESCE(folder_name, ''))) IN ('', 'unfiled')
                        """, [order, record_id])

                    else:
                        cur.execute("""
                            UPDATE steps SET file_order = %s
                            WHERE record_id = %s
                              AND TRIM(COALESCE(folder_name, '')) = %s
                        """, [order, record_id, folder_name])

                except Exception as e:
                    print(f"[ORDER] Error updating {record_id}: {e}", flush=True)

        return JsonResponse({"status": "success", "updated": len(ordered_ids)})
    except Exception as e:
        print(f"[ORDER] Error: {e}", flush=True)
        return JsonResponse({"error": str(e)}, status=500)


@login_required
def update_project_folder_order(request):
    """Update persisted folder order based on drag-and-drop."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=400)

    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON payload"}, status=400)

    parent_folder = _normalize_folder_path(data.get("parent_folder", ""))
    raw_folder_names = data.get("folder_names", [])
    if not isinstance(raw_folder_names, list):
        return JsonResponse({"error": "folder_names must be a list"}, status=400)

    ordered_names: list[str] = []
    seen: set[str] = set()
    for folder_name in raw_folder_names:
        normalized = _normalize_folder_path(folder_name)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        ordered_names.append(normalized)

    if not ordered_names:
        return JsonResponse({"error": "No folder names provided"}, status=400)

    try:
        with transaction.atomic():
            with connection.cursor() as cur:
                _ensure_project_folders_table()

                # Auto-register any folder that is visible in the tree but has
                # not yet been inserted into project_folders (e.g. folders that
                # were created by the recorder before the registry existed).
                for folder_name in ordered_names:
                    cur.execute("""
                        INSERT INTO project_folders (folder_name, folder_order)
                        VALUES (%s, 2147483647)
                        ON CONFLICT (folder_name) DO NOTHING
                    """, [folder_name])

                # Apply 1-based order exactly matching the dragged row positions:
                # first item in the list → order 1, second → order 2, etc.
                # Also propagate to the recorder tables so folder_order is
                # consistent across project_folders, steps, recordings, and run_table.
                for i, folder_name in enumerate(ordered_names, start=1):
                    if not parent_folder and "/" not in folder_name:
                        cur.execute(
                            "UPDATE parent_folders SET parent_folder_order = %s WHERE TRIM(COALESCE(parent_folder,'')) = %s",
                            [i, folder_name],
                        )
                    elif parent_folder and "/" not in parent_folder:
                        _parts = [p.strip() for p in folder_name.split("/") if p.strip()]
                        if len(_parts) == 2 and _parts[0] == parent_folder:
                            cur.execute(
                                """
                                UPDATE sub_folders sf
                                SET    sub_folder_order = %s
                                FROM   parent_folders pf
                                WHERE  sf.sub_folder_parent = pf.parent_folder_id
                                  AND  pf.parent_folder = %s
                                  AND  sf.sub_folder = %s
                                """,
                                [i, parent_folder, _parts[1]],
                            )
                    cur.execute(
                        "UPDATE project_folders SET folder_order = %s WHERE folder_name = %s",
                        [i, folder_name],
                    )
                    cur.execute(
                        "UPDATE steps     SET folder_order = %s WHERE TRIM(COALESCE(folder_name,'')) = %s",
                        [i, folder_name],
                    )
                    cur.execute(
                        "UPDATE run_table  SET folder_order = %s WHERE TRIM(COALESCE(folder_name,'')) = %s",
                        [i, folder_name],
                    )

                    # depth=2 (end-folder): update end_folders.end_folder_order
                    _parts = [p.strip() for p in folder_name.split("/") if p.strip()]
                    if len(_parts) == 3:
                        cur.execute(
                            """
                            UPDATE end_folders ef
                            SET    end_folder_order = %s
                            FROM   sub_folders sf
                            JOIN   parent_folders pf ON pf.parent_folder_id = sf.sub_folder_parent
                            WHERE  ef.end_folder_parent = sf.sub_folder_id
                              AND  ef.end_folder  = %s
                              AND  sf.sub_folder  = %s
                              AND  pf.parent_folder = %s
                            """,
                            [i, _parts[2], _parts[1], _parts[0]],
                        )

        return JsonResponse({"status": "success", "updated": len(ordered_names)})
    except Exception as e:
        print(f"[FOLDER-ORDER] Error: {e}", flush=True)
        return JsonResponse({"error": str(e)}, status=500)


def _append_result(run_id: str, result: dict):
    job = _REPLAY_JOBS.get(run_id)
    if job:
        with job["lock"]:
            job["results"].append(result)


def replay_status(request, run_id):
    """JSON polling endpoint — returns current progress.

    Accepts a Django session login OR a monitor_token query/header param so
    the detached monitor window and Watch-run tabs can poll without a session.
    """
    if not request.user.is_authenticated:
        _u = _resolve_monitor_user(request)
        if _u is None:
            return JsonResponse({"detail": "authentication required"}, status=403)
        request.user = _u
    rid = str(run_id)
    job = _REPLAY_JOBS.get(rid)
    if not job:
        return JsonResponse({"status": "none", "results": [], "total": 0, "run_id": rid})
    with job["lock"]:
        if job["status"] in ("running", "failing"):
            db_rows = list(
                RunResult.objects.filter(run_id=rid)
                .order_by("step_no")
                .values(
                    "step_no",
                    "action",
                    "page_url",
                    "element_tag",
                    "steps_description",
                    "validation",
                    "status",
                    "message",
                )
            )
            if db_rows:
                if len(db_rows) >= len(job["results"]):
                    job["results"] = [
                        {
                            "step_no": row["step_no"],
                            "action": row["action"],
                            "page_url": row["page_url"],
                            "element_tag": row.get("element_tag") or "",
                            "steps_description": row.get("steps_description") or "",
                            "validation": row.get("validation") or "",
                            "status": row.get("message") or row["status"],
                            "ok": row["status"] == RunResult.STATUS_PASS,
                        }
                        for row in db_rows
                    ]

                any_fail = any(row["status"] == RunResult.STATUS_FAIL for row in db_rows)
                any_not_executed = any(
                    row["status"] == RunResult.STATUS_NOT_EXECUTED for row in db_rows
                )
                if any_fail:
                    job["status"] = "failed"
                elif job["stop_event"].is_set():
                    job["status"] = "stopped"
                elif len(db_rows) >= job["total"]:
                    job["status"] = "failed" if any_not_executed else "done"

        return JsonResponse({
            "status":  job["status"],
            "run_id":  job["run_id"],
            "results": job["results"],
            "total":   job["total"],
        })


@csrf_exempt
def pause_replay(request, run_id):
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"action": "noop", "status": "none", "error": "Authentication required."}, status=403)
    rid = str(run_id)
    if request.method != "POST":
        return JsonResponse({"action": "noop", "status": "none"}, status=405)
    job = _REPLAY_JOBS.get(rid)
    if job:
        with job["lock"]:
            if job["status"] == "paused":
                job["pause_event"].clear()
                job["status"] = "running"
                action = "resumed"
            else:
                job["pause_event"].set()
                job["status"] = "paused"
                action = "paused"
        return JsonResponse({"action": action, "status": job["status"]})
    return JsonResponse({"action": "noop", "status": "none"})


@csrf_exempt
def stop_replay(request, run_id):
    user = _resolve_monitor_user(request)
    if user is None:
        return JsonResponse({"status": "none", "error": "Authentication required."}, status=403)
    rid = str(run_id)
    if request.method != "POST":
        return JsonResponse({"status": "none"}, status=405)
    job = _REPLAY_JOBS.get(rid)
    if job:
        job["pause_event"].clear()  # unblock if paused
        job["stop_event"].set()
        with job["lock"]:
            job["status"] = "stopped"
        return JsonResponse({"status": "stopped"})
    return JsonResponse({"status": "none"})


@csrf_exempt
def hotkey_play(request):
    """Start replay for the last replayed session, or newest session if none ran yet."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "detail": "POST required"}, status=405)

    remote_addr = (request.META.get("REMOTE_ADDR") or "").strip()
    if remote_addr not in {"127.0.0.1", "::1", "localhost"}:
        return JsonResponse({"ok": False, "detail": "Local requests only"}, status=403)

    try:
        for _run_id, _job in list(_REPLAY_JOBS.items()):
            with _job["lock"]:
                if _job.get("status") in ("running", "paused"):
                    return JsonResponse({"ok": False, "detail": "already running"})

        record_id = request.POST.get("record_id") or request.GET.get("record_id")

        user = _resolve_monitor_user(request)
        if user is None:
            user = User.objects.filter(is_active=True, is_superuser=True).order_by("id").first()
        if user is None:
            user = User.objects.filter(is_active=True).order_by("id").first()
        if user is None:
            return JsonResponse({"ok": False, "detail": "no active users found"})

        tenant_id = getattr(request, "tenant_id", None) or _get_user_tenant_id(user)

        if record_id:
            sessions = SessionMeta.objects.filter(record_id=record_id)
            if tenant_id:
                sessions = sessions.filter(tenant_id=tenant_id)
            if not sessions.exists():
                return JsonResponse({"ok": False, "detail": "session not found"}, status=404)
            record_id = str(record_id)
        else:
            last_runs = RunResult.objects.order_by("-created_at")
            sessions = SessionMeta.objects.order_by("-created_at")
            if tenant_id:
                last_runs = last_runs.filter(tenant_id=tenant_id)
                sessions = sessions.filter(tenant_id=tenant_id)

            last_run = last_runs.first()
            if last_run:
                record_id = str(last_run.record_id)
            else:
                last_session = sessions.first()
                if not last_session:
                    return JsonResponse({"ok": False, "detail": "no sessions found"})
                record_id = str(last_session.record_id)

        from django.test import RequestFactory

        factory = RequestFactory()
        fake = factory.post(f"/sessions/{record_id}/replay/")
        fake.user = user
        fake.session = {}
        fake.tenant_id = tenant_id

        replay(fake, record_id=record_id)
        # Grab the run_id that was just registered so the client can navigate to it.
        _new_run_id = None
        with _JOBS_LOCK:
            _new_run_id = _SESSION_TO_RUN.get((str(record_id), user.username))
        try:
            minimize_browser(request)
        except Exception:
            pass
        monitor_url = f"http://127.0.0.1:{request.get_port()}/active_runs/?minimized=1"
        replay_url  = f"/replay/{_new_run_id}/" if _new_run_id else None
        return JsonResponse({"ok": True, "detail": f"replay started for {record_id}",
                             "monitor_url": monitor_url, "replay_url": replay_url})
    except Exception as exc:
        return JsonResponse({"ok": False, "detail": str(exc)})


@csrf_exempt
def minimize_browser(request):
    """Minimize visible browser windows and keep the monitor popup on top."""
    try:
        import time as _time

        import psutil
        import win32con
        import win32gui
        import win32process

        browser_exes = {"chrome.exe", "msedge.exe", "firefox.exe"}
        monitor_hwnd = [None]

        def _minimize_others(hwnd, _):
            if not win32gui.IsWindowVisible(hwnd):
                return
            title = win32gui.GetWindowText(hwnd)
            if not title:
                return
            if "Active Runs Monitor" in title:
                monitor_hwnd[0] = hwnd
                return
            try:
                _, pid = win32process.GetWindowThreadProcessId(hwnd)
                proc_name = psutil.Process(pid).name().lower()
            except Exception:
                return
            if proc_name in browser_exes:
                win32gui.ShowWindow(hwnd, win32con.SW_MINIMIZE)

        win32gui.EnumWindows(_minimize_others, None)

        if not monitor_hwnd[0]:
            _time.sleep(0.3)

            def _find_monitor(hwnd, _):
                if win32gui.IsWindowVisible(hwnd):
                    title = win32gui.GetWindowText(hwnd)
                    if title and "Active Runs Monitor" in title:
                        monitor_hwnd[0] = hwnd

            win32gui.EnumWindows(_find_monitor, None)

        if monitor_hwnd[0]:
            try:
                import ctypes

                hwnd = monitor_hwnd[0]
                ws_minimize = 0x20000000
                style = win32gui.GetWindowLong(hwnd, win32con.GWL_STYLE)
                if style & ws_minimize:
                    win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
                    _time.sleep(0.15)

                class _RECT(ctypes.Structure):
                    _fields_ = [
                        ("left", ctypes.c_long),
                        ("top", ctypes.c_long),
                        ("right", ctypes.c_long),
                        ("bottom", ctypes.c_long),
                    ]

                wa = _RECT()
                ctypes.windll.user32.SystemParametersInfoW(48, 0, ctypes.byref(wa), 0)

                win_w = 540
                margin = 16
                try:
                    rect = win32gui.GetWindowRect(hwnd)
                    win_h = max(rect[3] - rect[1], 40)
                except Exception:
                    win_h = 80

                x = wa.right - win_w - margin
                y = wa.bottom - win_h - margin

                win32gui.SetWindowPos(
                    hwnd,
                    win32con.HWND_TOPMOST,
                    x,
                    y,
                    0,
                    0,
                    0x0001 | 0x0040,
                )
                win32gui.SetForegroundWindow(hwnd)
                win32gui.BringWindowToTop(hwnd)
            except Exception:
                pass

        return JsonResponse({"ok": True})
    except Exception:
        return JsonResponse({"ok": True})


def _open_detached_monitor_window(monitor_url: str) -> None:
    import ctypes
    import tempfile
    import time as _time

    profile_dir = os.path.join(tempfile.gettempdir(), "webconx_monitor_profile")

    try:
        import psutil

        for proc in psutil.process_iter(["pid", "name", "cmdline"]):
            try:
                cmdline = proc.info.get("cmdline") or []
                if any(profile_dir in (arg or "") for arg in cmdline):
                    proc.kill()
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass
        _time.sleep(0.8)
    except Exception:
        pass

    prefs_file = os.path.join(profile_dir, "Default", "Preferences")
    try:
        if os.path.isfile(prefs_file):
            os.remove(prefs_file)
    except Exception:
        pass

    class _RECT(ctypes.Structure):
        _fields_ = [
            ("left", ctypes.c_long),
            ("top", ctypes.c_long),
            ("right", ctypes.c_long),
            ("bottom", ctypes.c_long),
        ]

    wa = _RECT()
    ctypes.windll.user32.SystemParametersInfoW(48, 0, ctypes.byref(wa), 0)
    win_w, win_h, margin = 540, 250, 16
    x = wa.right - win_w - margin
    y = wa.bottom - win_h - margin

    chromium_paths = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    launched = False
    for exe in chromium_paths:
        if not os.path.isfile(exe):
            continue
        args = (
            f'--app="{monitor_url}" '
            f'--user-data-dir="{profile_dir}" '
            f'--window-size={win_w},{win_h} '
            f'--window-position={x},{y} '
            f'--no-first-run '
            f'--no-default-browser-check '
            f'--disable-session-crashed-bubble '
            f'--disable-infobars'
        )
        ret = ctypes.windll.shell32.ShellExecuteW(None, "open", exe, args, None, 1)
        if ret > 32:
            launched = True
            break

    if not launched:
        ret = ctypes.windll.shell32.ShellExecuteW(None, "open", monitor_url, None, None, 1)
        if ret <= 32:
            return

    _time.sleep(2.0)
    try:
        import win32con
        import win32gui

        def _size_monitor(hwnd, _):
            if win32gui.IsWindowVisible(hwnd):
                title = win32gui.GetWindowText(hwnd)
                if title and ("Active Runs Monitor" in title or "Recording Monitor" in title):
                    win32gui.SetWindowPos(
                        hwnd,
                        win32con.HWND_TOPMOST,
                        x,
                        y,
                        win_w,
                        win_h,
                        0x0040,
                    )
                    win32gui.SetForegroundWindow(hwnd)
                    win32gui.BringWindowToTop(hwnd)

        win32gui.EnumWindows(_size_monitor, None)
    except Exception:
        pass


def _ensure_detached_monitor_open(request, *, minimize_windows: bool = True) -> str:
    user = _resolve_monitor_user(request)
    if user is None:
        user = getattr(request, "user", None) if getattr(request, "user", None) and request.user.is_authenticated else None
    if user is None:
        user = User.objects.filter(is_active=True, is_superuser=True).order_by("id").first()
    if user is None:
        user = User.objects.filter(is_active=True).order_by("id").first()
    if user is None:
        raise RuntimeError("no active users found")

    if minimize_windows:
        minimize_browser(request)

    try:
        import win32gui

        found = {"open": False}

        def _find(hwnd, _):
            if found["open"]:
                return
            if not win32gui.IsWindowVisible(hwnd):
                return
            title = win32gui.GetWindowText(hwnd)
            if title and "Active Runs Monitor" in title:
                found["open"] = True

        win32gui.EnumWindows(_find, None)
        if found["open"]:
            return "monitor already open"
    except Exception:
        pass

    token = _build_monitor_token(user)
    monitor_url = f"http://127.0.0.1:{request.get_port()}/api/active-runs/local-login/?minimized=1&monitor_token={token}"
    _open_detached_monitor_window(monitor_url)
    return "opened"


def _open_recording_monitor(request, record_id: str, record_name: str = "") -> str:
    """Open the dedicated Recording Monitor popup for the given session."""
    user = _resolve_monitor_user(request)
    if user is None:
        user = getattr(request, "user", None) if getattr(request, "user", None) and request.user.is_authenticated else None
    if user is None:
        user = User.objects.filter(is_active=True, is_superuser=True).order_by("id").first()
    if user is None:
        user = User.objects.filter(is_active=True).order_by("id").first()
    if user is None:
        raise RuntimeError("no active users found")

    # Check if already open
    try:
        import win32gui

        found = {"open": False}

        def _find(hwnd, _):
            if found["open"]:
                return
            if not win32gui.IsWindowVisible(hwnd):
                return
            title = win32gui.GetWindowText(hwnd)
            if title and "Recording Monitor" in title:
                found["open"] = True

        win32gui.EnumWindows(_find, None)
        if found["open"]:
            return "recording monitor already open"
    except Exception:
        pass

    token = _build_monitor_token(user)
    import urllib.parse
    params = urllib.parse.urlencode({
        "record_id": record_id,
        "record_name": record_name,
        "monitor_token": token,
    })
    monitor_url = f"http://127.0.0.1:{request.get_port()}/recording-monitor/?{params}"
    _open_detached_monitor_window(monitor_url)
    return "opened"


@csrf_exempt
def open_detached_monitor(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "detail": "POST required"}, status=405)

    remote_addr = (request.META.get("REMOTE_ADDR") or "").strip()
    if remote_addr not in {"127.0.0.1", "::1", "localhost"}:
        return JsonResponse({"ok": False, "detail": "Local requests only"}, status=403)

    try:
        detail = _ensure_detached_monitor_open(request, minimize_windows=True)
        return JsonResponse({"ok": True, "detail": detail})
    except Exception as exc:
        return JsonResponse({"ok": False, "detail": str(exc)})


@csrf_exempt
def hotkey_pause(request):
    """Pause the first running replay job."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "action": "noop"}, status=405)

    remote_addr = (request.META.get("REMOTE_ADDR") or "").strip()
    if remote_addr not in {"127.0.0.1", "::1", "localhost"}:
        return JsonResponse({"ok": False, "action": "noop"}, status=403)

    for _run_id, _job in list(_REPLAY_JOBS.items()):
        with _job["lock"]:
            status = _job.get("status")
        if status == "running":
            _job["pause_event"].set()
            with _job["lock"]:
                _job["status"] = "paused"
            return JsonResponse({"ok": True, "action": "paused", "run_id": _run_id})
    return JsonResponse({"ok": False, "action": "noop"})


@csrf_exempt
def hotkey_resume(request):
    """Resume the first paused replay job."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "action": "noop"}, status=405)

    remote_addr = (request.META.get("REMOTE_ADDR") or "").strip()
    if remote_addr not in {"127.0.0.1", "::1", "localhost"}:
        return JsonResponse({"ok": False, "action": "noop"}, status=403)

    for _run_id, _job in list(_REPLAY_JOBS.items()):
        with _job["lock"]:
            status = _job.get("status")
        if status == "paused":
            _job["pause_event"].clear()
            with _job["lock"]:
                _job["status"] = "running"
            return JsonResponse({"ok": True, "action": "resumed", "run_id": _run_id})
    return JsonResponse({"ok": False, "action": "noop"})


@csrf_exempt
def hotkey_scrape(request):
    """Scrape the current page in the connected Chrome browser (F6 global hotkey)."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required"}, status=405)

    remote_addr = (request.META.get("REMOTE_ADDR") or "").strip()
    if remote_addr not in {"127.0.0.1", "::1", "localhost"}:
        return JsonResponse({"ok": False, "error": "Local requests only"}, status=403)

    _ensure_ai_databank_schema()

    rdp_raw = (get_config("chrome.remote_debugging_port") or "").strip()
    webdriver_filename = (get_config("chrome.webdriver_path") or "").strip()
    configured_port = int(rdp_raw) if rdp_raw.isdigit() and int(rdp_raw) > 0 else 0
    attach_port = _auto_detect_cdp_port(configured_port)

    if not attach_port:
        return JsonResponse({
            "ok": False,
            "error": "No Chrome remote debugging port found. Start Chrome with --remote-debugging-port=9222.",
        }, status=400)

    db = settings.DATABASES["default"]
    db_config = {
        "dbname": db.get("NAME", "automation_db"),
        "user": db.get("USER", "postgres"),
        "password": db.get("PASSWORD", ""),
        "host": db.get("HOST", "localhost") or "localhost",
        "port": str(db.get("PORT", "5432") or "5432"),
    }

    try:
        from web_scraper import scrape_once
        result = scrape_once(
            attach_port=attach_port,
            db_config=db_config,
            webdriver_filename=webdriver_filename,
        )
    except Exception as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=500)

    return JsonResponse({
        "ok": True,
        "saved_rows": result.get("saved_rows", 0),
        "page_url": result.get("page_url", ""),
        "page_name": result.get("page_name", ""),
    })


@csrf_exempt
def hotkey_stop(request):
    """Stop all active replay jobs AND active recordings immediately."""
    if request.method != "POST":
        return JsonResponse({"ok": False, "stopped": [], "recordings_stopped": []}, status=405)

    remote_addr = (request.META.get("REMOTE_ADDR") or "").strip()
    if remote_addr not in {"127.0.0.1", "::1", "localhost"}:
        return JsonResponse({"ok": False, "stopped": [], "recordings_stopped": []}, status=403)

    stopped = []
    for _run_id, _job in list(_REPLAY_JOBS.items()):
        with _job["lock"]:
            status = _job.get("status")
        if status in ("running", "paused"):
            _job["pause_event"].clear()
            _job["stop_event"].set()
            with _job["lock"]:
                _job["status"] = "stopped"
            stopped.append(_run_id)

    # Also stop any active recordings
    recordings_stopped = []
    with _ACTIVE_RECORDING_LOCK:
        for _rec_id, _rec in list(_ACTIVE_RECORDING.items()):
            _pid = _rec.get("pid")
            if _pid:
                try:
                    if sys.platform == "win32":
                        subprocess.call(
                            ["taskkill", "/F", "/PID", str(_pid)],
                            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                        )
                    else:
                        os.kill(_pid, signal.SIGTERM)
                except (OSError, ProcessLookupError):
                    pass
                recordings_stopped.append(_rec_id)
            # Clean up pause flag
            import tempfile as _tf
            _flag = os.path.join(_tf.gettempdir(), f"recorder_paused_{_rec_id}.flag")
            try:
                os.remove(_flag)
            except OSError:
                pass
        # Clear all active recordings
        for _rid in recordings_stopped:
            _ACTIVE_RECORDING.pop(_rid, None)

    return JsonResponse({
        "ok": bool(stopped) or bool(recordings_stopped),
        "stopped": stopped,
        "recordings_stopped": recordings_stopped,
    })


@login_required
def last_run(request, record_id):
    """Show all steps for the most recent replay run of this session."""
    try:
        record_name = SessionMeta.objects.get(record_id=record_id).record_name
    except SessionMeta.DoesNotExist:
        record_name = ""

    # Find the latest run_id by picking the run with the most recent created_at
    latest = (
        RunResult.objects
        .filter(record_id=record_id)
        .order_by("-created_at")
        .values("run_id")
        .first()
    )

    if not latest:
        step_count = Step.objects.filter(record_id=record_id).count()
        return render(request, "recorder/last_run.html", {
            "record_id": record_id,
            "record_name": record_name,
            "run_id": None,
            "rows": [],
            "step_count": step_count,
            "total": 0, "passed": 0, "failed": 0, "not_executed": 0,
        })

    run_id = latest["run_id"]
    rows, _first_row = _build_run_rows_with_not_executed(run_id, record_id)

    total        = len(rows)
    passed       = sum(1 for r in rows if r.status == RunResult.STATUS_PASS)
    failed       = sum(1 for r in rows if r.status == RunResult.STATUS_FAIL)
    not_executed = sum(1 for r in rows if r.status == RunResult.STATUS_NOT_EXECUTED)
    latest_run_day = None
    if rows:
        _latest_run_dt = rows[0].run_date or rows[0].created_at
        if _latest_run_dt is not None:
            latest_run_day = _latest_run_dt.date().isoformat()

    paginator = Paginator(rows, 100)
    page_obj  = paginator.get_page(request.GET.get("page"))

    return render(request, "recorder/last_run.html", {
        "record_id": record_id,
        "record_name": record_name,
        "run_id": run_id,
        "rows": page_obj,
        "page_obj": page_obj,
        "total": total,
        "passed": passed,
        "failed": failed,
        "not_executed": not_executed,
        "latest_run_day": latest_run_day,
    })

# ---------------------------------------------------------------------------
# User management  (staff-only)
# ---------------------------------------------------------------------------

@login_required
def user_list(request):
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")
    users = User.objects.all().order_by("date_joined")
    user_actions_map = {}
    user_public_folder_map = {}
    for u in users:
        pref_raw = get_user_pref(u.id, "projects.visible_actions", ",".join(sorted(_ALL_PROJ_ACTIONS)))
        visible  = _resolve_project_visible_actions(pref_raw)
        user_actions_map[str(u.pk)] = sorted(visible)
        user_public_folder_map[str(u.pk)] = _user_can_view_public_project_folders(u)
    return render(request, "recorder/user_list.html", {
        "users": users,
        "user_actions_map_json": json.dumps(user_actions_map),
        "user_public_folder_map_json": json.dumps(user_public_folder_map),
    })


def create_user(request):
    if request.user.is_authenticated and not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")

    if request.method == "POST":
        username  = request.POST.get("username", "").strip()
        email     = request.POST.get("email", "").strip()
        password1 = request.POST.get("password1", "")
        password2 = request.POST.get("password2", "")
        is_staff  = str(request.POST.get("is_staff", "")).strip().lower() in {"1", "true", "on", "yes"}
        is_active = str(request.POST.get("is_active", "")).strip().lower() in {"1", "true", "on", "yes"}
        is_superuser = bool(
            request.user.is_authenticated
            and request.user.is_superuser
            and str(request.POST.get("is_superuser", "")).strip().lower() in {"1", "true", "on", "yes"}
        )

        errors = {}
        if not username:
            errors["username"] = "Username is required."
        elif User.objects.filter(username=username).exists():
            errors["username"] = f"Username '{username}' is already taken."
        if not is_staff and not is_superuser:
            errors["role"] = "Select at least one role: Staff or Superuser."
        if not password1:
            errors["password1"] = "Password is required."
        elif len(password1) < 6:
            errors["password1"] = "Password must be at least 6 characters."
        elif password1 != password2:
            errors["password2"] = "Passwords do not match."

        if errors:
            return render(request, "recorder/create_user.html", {
                "errors": errors,
                "form_data": {
                    "username": username,
                    "email": email,
                    "is_staff": is_staff,
                    "is_active": is_active,
                    "is_superuser": is_superuser,
                },
            })

        user = User.objects.create_user(username=username, email=email, password=password1)
        user.is_staff = is_staff
        user.is_active = is_active
        user.is_superuser = is_superuser
        user.save()
        # Save visible_actions pref if provided
        raw_va = request.POST.get("visible_actions", "").strip()
        va_value = ",".join(sorted(a for a in raw_va.split(",") if a.strip() in _ALL_PROJ_ACTIONS)) \
                   if raw_va else ",".join(sorted(_ALL_PROJ_ACTIONS))
        _ensure_user_prefs_table()
        with connection.cursor() as cur:
            cur.execute("""
                INSERT INTO user_prefs (user_id, pref_key, pref_value)
                VALUES (%s, %s, %s)
                ON CONFLICT (user_id, pref_key) DO UPDATE SET pref_value = EXCLUDED.pref_value
            """, [user.id, "projects.visible_actions", va_value])
        messages.success(request, f"User '{username}' created successfully.")
        return redirect("user_list" if request.user.is_authenticated else "login")

    return render(request, "recorder/create_user.html", {
        "errors": {},
        "form_data": {"is_staff": True, "is_active": True, "is_superuser": False},
    })


@login_required
def delete_user(request, pk):
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")
    if request.method != "POST":
        return redirect("user_list")
    try:
        user = User.objects.get(pk=pk)
        if user.is_superuser:
            messages.error(request, "Cannot delete a superuser.")
        elif user == request.user:
            messages.error(request, "Cannot delete your own account.")
        else:
            username = user.username
            user.delete()
            messages.success(request, f"User '{username}' deleted.")
    except User.DoesNotExist:
        messages.error(request, "User not found.")
    return redirect("user_list")


@login_required
def edit_user(request, pk):
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")
    if request.method != "POST":
        return redirect("user_list")

    is_ajax = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    def _err(msg, field=None):
        if is_ajax:
            return JsonResponse({"ok": False, "error": msg, "field": field})
        messages.error(request, msg)
        return redirect("user_list")

    try:
        target = User.objects.get(pk=pk)
    except User.DoesNotExist:
        return _err("User not found.")

    if target.is_superuser and not request.user.is_superuser:
        return _err("Only a superuser can edit another superuser.")

    new_username = request.POST.get("username", "").strip()
    new_email    = request.POST.get("email", "").strip()
    password1    = request.POST.get("password1", "")
    password2    = request.POST.get("password2", "")
    is_staff     = request.POST.get("is_staff")     == "on"
    is_superuser = request.POST.get("is_superuser") == "on"
    is_active    = request.POST.get("is_active")    == "on"

    if not new_username:
        return _err("Username is required.", "username")
    if new_username != target.username and User.objects.filter(username=new_username).exists():
        return _err(f"Username \u2018{new_username}\u2019 is already taken.", "username")
    if not is_staff and not is_superuser:
        return _err("Select at least one role: Staff or Superuser.")
    if password1:
        if len(password1) < 6:
            return _err("Password must be at least 6 characters.", "password1")
        if password1 != password2:
            return _err("Passwords do not match.", "password2")

    target.username = new_username
    target.email    = new_email
    if password1:
        target.set_password(password1)
    if target != request.user:
        target.is_staff     = is_staff
        target.is_superuser = is_superuser
        target.is_active    = is_active
    target.save()

    # If deactivated, stop recordings/replays and invalidate all active sessions
    if not target.is_active and target != request.user:
        _force_stop_user_activity(target.pk, target.username)
        import datetime as _dt
        from django.contrib.sessions.models import Session as _Session
        _now = _dt.datetime.now(_dt.timezone.utc)
        for _s in _Session.objects.filter(expire_date__gte=_now):
            try:
                _d = _s.get_decoded()
            except Exception:
                continue
            if _d.get("_auth_user_id") == str(target.pk):
                _s.delete()

    if is_ajax:
        return JsonResponse({"ok": True, "username": new_username})
    messages.success(request, f"User \u2018{new_username}\u2019 updated successfully.")
    return redirect("user_list")


def session_check(request):
    """Lightweight endpoint polled by the client heartbeat.
    Returns {"ok": true} only when authenticated AND active.
    No @login_required so unauthenticated/inactive users get JSON 200 {"ok": false}
    rather than a redirect loop.
    """
    user = _resolve_monitor_user(request)
    ok = bool(user and getattr(user, "is_active", False))
    return JsonResponse({"ok": ok})


@login_required
def force_logout_user(request, pk):
    """POST: stop recordings/replays and flush all active sessions for the target user."""
    if not request.user.is_superuser:
        return JsonResponse({"ok": False, "error": "Access denied."})
    if request.method != "POST":
        return JsonResponse({"ok": False, "error": "POST required."}, status=405)
    try:
        target = User.objects.get(pk=pk)
    except User.DoesNotExist:
        return JsonResponse({"ok": False, "error": "User not found."})
    if target == request.user:
        return JsonResponse({"ok": False, "error": "Cannot logout your own account."})

    # Stop any active recording / replay jobs first
    stopped_msgs = _force_stop_user_activity(target.pk, target.username)

    # Flush all active sessions
    import datetime as _dt
    from django.contrib.sessions.models import Session as _Session
    _now = _dt.datetime.now(_dt.timezone.utc)
    count = 0
    for _s in _Session.objects.filter(expire_date__gte=_now):
        try:
            _d = _s.get_decoded()
        except Exception:
            continue
        if _d.get("_auth_user_id") == str(target.pk):
            _s.delete()
            count += 1
    return JsonResponse({"ok": True, "sessions": count, "stopped": stopped_msgs})


def _force_stop_user_activity(target_pk, target_username):
    """Kill any active recording and stop all running replays for target_pk/username.
    Returns a list of human-readable message strings describing what was stopped.
    """
    import datetime as _dt
    from django.contrib.sessions.models import Session as _Session

    stopped = []

    # --- Stop recordings (look for recording_pid in all active sessions for this user) ---
    _now = _dt.datetime.now(_dt.timezone.utc)
    for _s in _Session.objects.filter(expire_date__gte=_now):
        try:
            _d = _s.get_decoded()
        except Exception:
            continue
        if _d.get("_auth_user_id") != str(target_pk):
            continue
        _pid = _d.get("recording_pid")
        if not _pid:
            continue
        try:
            if sys.platform == "win32":
                subprocess.call(
                    ["taskkill", "/F", "/T", "/PID", str(_pid)],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            else:
                os.kill(int(_pid), signal.SIGTERM)
            stopped.append(f"Recording stopped and saved (pid {_pid}).")
        except (OSError, ProcessLookupError):
            pass

    # --- Stop replays belonging to this user ---
    with _JOBS_LOCK:
        for _run_id, _job in list(_REPLAY_JOBS.items()):
            if _job.get("runner") != target_username:
                continue
            if _job.get("status") not in ("running", "paused"):
                continue
            _job["pause_event"].clear()   # unblock if paused
            _job["stop_event"].set()
            with _job["lock"]:
                _job["status"] = "stopped"
                _job["finished_at"] = time.time()
            stopped.append(f"Replay run {_run_id} stopped and saved.")
            # Leave the job in _REPLAY_JOBS as "stopped" so the monitor shows it
            # (the thread will exit cleanly at the next step-boundary check).
            # Only remove the dedup key so the user can re-run the same session.
            _SESSION_TO_RUN.pop((_job.get("record_id", ""), target_username), None)

    return stopped


@login_required
def toggle_user_active(request, pk):
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")
    if request.method != "POST":
        return redirect("user_list")

    is_ajax = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    try:
        target = User.objects.get(pk=pk)
        if target == request.user:
            if is_ajax:
                return JsonResponse({"ok": False, "error": "Cannot deactivate your own account."})
            messages.error(request, "Cannot deactivate your own account.")
        elif target.is_superuser and not request.user.is_superuser:
            if is_ajax:
                return JsonResponse({"ok": False, "error": "Only a superuser can toggle another superuser."})
            messages.error(request, "Only a superuser can toggle another superuser.")
        else:
            will_deactivate = target.is_active  # True means we're flipping to inactive
            if will_deactivate:
                stopped_msgs = _force_stop_user_activity(target.pk, target.username)
            else:
                stopped_msgs = []
            target.is_active = not target.is_active
            target.save()
            state = "activated" if target.is_active else "deactivated"
            # Flush all active sessions when deactivating
            if not target.is_active:
                import datetime as _dt
                from django.contrib.sessions.models import Session as _Session
                _now = _dt.datetime.now(_dt.timezone.utc)
                for _s in _Session.objects.filter(expire_date__gte=_now):
                    try:
                        _d = _s.get_decoded()
                    except Exception:
                        continue
                    if _d.get("_auth_user_id") == str(target.pk):
                        _s.delete()
            if is_ajax:
                return JsonResponse({"ok": True, "state": state, "stopped": stopped_msgs})
            messages.success(request, f"User '{target.username}' {state}.")
    except User.DoesNotExist:
        if is_ajax:
            return JsonResponse({"ok": False, "error": "User not found."})
        messages.error(request, "User not found.")
    return redirect("user_list")


@login_required
def run_detail_partial(request, run_id):
    """Return an HTML fragment with step details for a run_id (used by inline expand in history)."""
    rows, first_row = _build_run_rows_with_not_executed(run_id)
    passed       = sum(1 for r in rows if r.status == RunResult.STATUS_PASS)
    failed       = sum(1 for r in rows if r.status == RunResult.STATUS_FAIL)
    not_executed = sum(1 for r in rows if r.status == RunResult.STATUS_NOT_EXECUTED)
    return render(request, "recorder/run_detail_partial.html", {
        "rows": rows,
        "first_row": first_row,
        "total": len(rows),
        "passed": passed,
        "failed": failed,
        "not_executed": not_executed,
    })


def session_last_run_partial(request, record_id):
    """Return run_detail_partial.html for the latest run of a session (inline expand on Projects page)."""
    latest = (
        RunResult.objects
        .filter(record_id=record_id)
        .order_by("-created_at")
        .values("run_id")
        .first()
    )
    if not latest:
        raise Http404("No runs found for this session")
    run_id = latest["run_id"]
    rows, first_row = _build_run_rows_with_not_executed(run_id, record_id)
    passed       = sum(1 for r in rows if r.status == RunResult.STATUS_PASS)
    failed       = sum(1 for r in rows if r.status == RunResult.STATUS_FAIL)
    not_executed = sum(1 for r in rows if r.status == RunResult.STATUS_NOT_EXECUTED)
    return render(request, "recorder/run_detail_partial.html", {
        "rows": rows,
        "first_row": first_row,
        "total": len(rows),
        "passed": passed,
        "failed": failed,
        "not_executed": not_executed,
    })


def _build_run_rows_with_not_executed(run_id, record_id=None):
    rows = list(
        RunResult.objects
        .filter(run_id=run_id)
        .order_by("step_no")
    )
    if not rows:
        raise Http404("Run not found")

    first_row = rows[0]
    record_id = record_id or first_row.record_id

    # If a run was stopped early, trailing steps may never have been inserted
    # into run_table. Synthesize them here so run detail views show them
    # explicitly as Not Executed instead of silently dropping them.
    existing_step_nos = {r.step_no for r in rows}
    source_steps = list(Step.objects.filter(record_id=record_id).order_by("step_no", "id"))
    if not source_steps:
        source_steps = list(Recording.objects.filter(record_id=record_id).order_by("step_no", "id"))

    if source_steps:
        seen_step_nos: set[int] = set()
        missing_rows: list[RunResult] = []
        for step in source_steps:
            if step.step_no in seen_step_nos:
                continue
            seen_step_nos.add(step.step_no)
            if step.step_no in existing_step_nos:
                continue
            missing_rows.append(RunResult(
                run_id=run_id,
                record_id=record_id,
                step_no=step.step_no,
                action=step.action,
                page_url=step.page_url,
                element_tag=step.element_tag,
                locator_id=getattr(step, "locator_id", None),
                data_id=getattr(step, "data_id", None),
                raw_event=step.raw_event,
                status=RunResult.STATUS_NOT_EXECUTED,
                message="Not Executed",
                runner=first_row.runner,
                author=getattr(step, "recorder", None) or getattr(step, "author", None),
                folder_name=getattr(step, "folder_name", None),
                parent_folder_id=getattr(step, "parent_folder_id", None),
                sub_folder_id=getattr(step, "sub_folder_id", None),
                end_folder_id=getattr(step, "end_folder_id", None),
                validation=getattr(step, "validation", None),
                steps_description=getattr(step, "steps_description", None),
                page_title=getattr(step, "page_title", None),
                engine=first_row.engine,
            ))
        if missing_rows:
            rows.extend(missing_rows)
            rows.sort(key=lambda row: row.step_no)

    return rows, first_row


@login_required
def run_step_screenshot(request, run_id, record_id, step_no):
    """Serve the PNG screenshot stored as BYTEA in run_table, looked up by run_id+record_id+step_no."""
    rr = (
        RunResult.objects
        .filter(run_id=run_id, record_id=record_id, step_no=step_no)
        .exclude(screenshot__isnull=True)
        .first()
    )
    if not rr:
        raise Http404("No screenshot for this step")
    return HttpResponse(bytes(rr.screenshot), content_type="image/png")


def view_run(request, run_id):
    """Show all steps for a specific replay run by run_id."""
    if not request.user.is_authenticated:
        user = _resolve_monitor_user(request)
        if user is None:
            from django.contrib.auth.views import redirect_to_login
            return redirect_to_login(request.get_full_path())
        request.user = user
    first = RunResult.objects.filter(run_id=run_id).first()
    if not first:
        # If the run is currently active/in-memory, use the live replay view.
        # Do not redirect for terminal jobs, or /run -> /replay -> /run loops can occur.
        with _JOBS_LOCK:
            _job = _REPLAY_JOBS.get(str(run_id))
        if _job:
            with _job["lock"]:
                _status = _job.get("status")
            if _status in ("running", "paused", "failing"):
                return redirect("replay_run", run_id=run_id)

        # Graceful fallback for stale/deleted run IDs.
        if getattr(request, "user", None) and request.user.is_authenticated:
            messages.warning(request, "Run not found or no longer available.")
        return redirect("sessions_list")

    record_id = first.record_id
    try:
        record_name = SessionMeta.objects.get(record_id=record_id).record_name
    except SessionMeta.DoesNotExist:
        record_name = ""

    rows = list(
        RunResult.objects
        .filter(run_id=run_id)
        .order_by("step_no")
    )

    # If a run was stopped early, trailing steps may never have been inserted
    # into run_table. Synthesize them here so the completed run view shows
    # them explicitly as Not Executed instead of silently dropping them.
    _existing_step_nos = {r.step_no for r in rows}
    _source_steps = list(Step.objects.filter(record_id=record_id).order_by("step_no", "id"))
    if not _source_steps:
        _source_steps = list(Recording.objects.filter(record_id=record_id).order_by("step_no", "id"))

    if _source_steps:
        _seen_step_nos: set[int] = set()
        _missing_rows: list[RunResult] = []
        for _step in _source_steps:
            if _step.step_no in _seen_step_nos:
                continue
            _seen_step_nos.add(_step.step_no)
            if _step.step_no in _existing_step_nos:
                continue
            _missing_rows.append(RunResult(
                run_id=run_id,
                record_id=record_id,
                step_no=_step.step_no,
                action=_step.action,
                page_url=_step.page_url,
                element_tag=_step.element_tag,
                locator_id=getattr(_step, "locator_id", None),
                data_id=getattr(_step, "data_id", None),
                raw_event=_step.raw_event,
                status=RunResult.STATUS_NOT_EXECUTED,
                message="Not Executed",
                runner=first.runner,
                author=getattr(_step, "recorder", None) or getattr(_step, "author", None),
                folder_name=getattr(_step, "folder_name", None),
                parent_folder_id=getattr(_step, "parent_folder_id", None),
                sub_folder_id=getattr(_step, "sub_folder_id", None),
                end_folder_id=getattr(_step, "end_folder_id", None),
                validation=getattr(_step, "validation", None),
                steps_description=getattr(_step, "steps_description", None),
                page_title=getattr(_step, "page_title", None),
                engine=first.engine,
            ))
        if _missing_rows:
            rows.extend(_missing_rows)
            rows.sort(key=lambda r: r.step_no)

    total        = len(rows)
    passed       = sum(1 for r in rows if r.status == RunResult.STATUS_PASS)
    failed       = sum(1 for r in rows if r.status == RunResult.STATUS_FAIL)
    not_executed = sum(1 for r in rows if r.status == RunResult.STATUS_NOT_EXECUTED)

    paginator = Paginator(rows, 100)
    page_obj  = paginator.get_page(request.GET.get("page"))

    return render(request, "recorder/last_run.html", {
        "record_id": record_id,
        "record_name": record_name,
        "run_id": run_id,
        "rows": page_obj,
        "page_obj": page_obj,
        "total": total,
        "passed": passed,
        "failed": failed,
        "not_executed": not_executed,
    })


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

_CONFIG_DEFAULTS = [
    # (key, value, label, description, group_name, input_type)
    ("features.bulk_replay_enabled",     "true",  "Enable Bulk Replay",
     "Allow selecting and replaying multiple sessions at once.",                "Features",  "checkbox"),
    ("features.history_enabled",         "true",  "Enable History Page",
     "Show the History link in the navigation bar.",                            "Features",  "checkbox"),
    ("features.user_management_enabled", "true",  "Enable User Management",
     "Show the Users link in the navigation bar (superusers only).",            "Features",  "checkbox"),
    ("features.licensing_enabled",       "true",  "Enable Licensing Page",
     "Show the Licensing link in the navigation bar.",                          "Features",  "checkbox"),
    ("parent_folders.public",            "true",  "Project Folders Visible To All Users",
     "If enabled, all users can see every parent/sub/end folder. If disabled, only the parent folder author and superadmins can view that hierarchy.", "Features",  "checkbox"),

    ("dashboard.refresh_interval",       "5",     "Refresh Interval (seconds)",
     "How often the dashboard polls for active run updates.",                   "Dashboard", "number"),
    ("dashboard.max_jobs_display",       "20",    "Max Active Jobs Shown",
     "Maximum active run cards visible in the monitor panel.",                  "Dashboard", "number"),

    ("ui.sessions_per_page",             "25",    "Sessions per Page",
     "Rows per page in the Projects view.",                                     "UI",        "number"),
    ("ui.steps_per_page",                "50",    "Steps per Page",
     "Rows per page in the step detail view.",                                  "UI",        "number"),
    ("ui.history_per_page",              "25",    "History Entries per Page",
     "Rows per page in the History view.",                                      "UI",        "number"),

    ("replay.headless_default",          "false", "Default Headless Mode",
     "New replays default to headless mode.",                                   "Replay",    "checkbox"),
    ("replay.implicit_wait",             "10",    "Implicit Wait (seconds, Selenium only)",
     "How long Selenium waits for elements before timing out. Playwright does not use Selenium-style implicit waits.",                 "Replay",    "number"),
    ("replay.page_timeout",              "30",    "Page Load Timeout (seconds)",
     "Maximum time allowed for a page to fully load.",                          "Replay",    "number"),
    ("replay.overlay_timeout",           "60",    "Overlay Wait Timeout (seconds)",
     "Maximum time to wait for blocking overlays (e.g. JSF #inProgressPage) to disappear after a click or form submit. Increase for slow application servers.",
     "Replay",    "number"),
    ("replay.step_delay",                "0",     "Step Delay (milliseconds)",
     "Pause between each replayed step — 0 means no delay.",                   "Replay",    "number"),
    ("replay.execution_mode",            "parallel", "Execution Mode",
     "Run bulk replays in parallel (simultaneous) or serial (one at a time).",  "Replay",    "select"),
    ("replay.max_parallel_replays",      "4",     "Max Parallel Replays",
     "Maximum number of headless replay threads running at the same time during bulk replay.", "Replay", "number"),
    ("replay.step_timeout",              "10",    "Step Timeout (seconds)",
     "How long to wait for an element per locator strategy before trying the next.",  "Replay", "number"),
    ("replay.poll_interval",             "0.5",   "Poll Interval (seconds)",
     "Frequency of element-presence checks within the timeout loop.",                 "Replay", "number"),
    ("replay.step_retries",              "2",     "Step Retries",
     "Outer retry count per step on failure (1 + N = total attempts).",               "Replay", "number"),
    ("replay.retry_delay",               "5",     "Retry Delay (seconds)",
     "Wait time between outer step retries.",                                         "Replay", "number"),
    ("replay.step_settle",               "0.3",   "Step Settle (seconds)",
     "Pause after each successful step before moving to the next.",                   "Replay", "number"),
    ("replay.window_timeout",            "15",    "Page / Window Timeout (seconds)",
     "Maximum wait for a page to load or a window to appear.",                        "Replay", "number"),
    ("replay.nav_retries",               "2",     "Navigation Retries",
     "Extra navigation attempts when the initial page load fails.",                   "Replay", "number"),
    ("replay.nav_retry_wait",            "3",     "Navigation Retry Wait (seconds)",
     "Wait time between navigation retry attempts.",                                  "Replay", "number"),
    ("replay.max_step_delay",             "2",     "Max Recorded Step Delay (seconds)",
     "Cap the recorded inter-step delay replayed between actions. Steps recorded with longer pauses will be clamped to this value. Set 0 to disable (replay exact recorded timing).", "Replay", "number"),

    ("chrome.implicit_wait",                "10",    "Implicit Wait (seconds)",
     "Seconds Selenium waits for elements to appear before timing out.",         "Chrome",    "number"),
    ("chrome.extra_arguments",              "--start-maximized\n--disable-blink-features=AutomationControlled\n--no-sandbox\n--disable-gpu\n--ignore-certificate-errors\n--allow-running-insecure-content",
     "Extra Chrome Arguments",
     "Chrome flags applied on every launch, one per line.",                     "Chrome",    "textarea"),
    ("chrome.experimental_options",          'excludeSwitches=["enable-automation"]',
     "Experimental Chrome Options",
     "Chrome experimental options, one per line as key=json_value.\nExample: excludeSwitches=[\"enable-automation\"]\nuseAutomationExtension=false", "Chrome", "textarea"),
    ("chrome.webdriver_path",               "",      "ChromeDriver Executable",
     "Pick a driver from the webdrivers/chrome/ folder, or leave Auto to use Selenium Manager.", "Chrome", "select"),
    ("chrome.remote_debugging_port",        "",      "Remote Debugging Port",
     "CDP port Chrome was started with (e.g. 9222). Required for Add Step to attach to the already-running browser window. Leave blank to disable Add Step.", "Chrome", "number"),
    ("ai_databank.scrape_item_delay_ms",    "100",   "Scrape Item Delay (ms)",
     "Delay after each saved element during AI Databank scraping. Increase this to make pause/stop easier to observe; 0 disables the delay.", "Chrome", "number"),
    ("ai_databank.launch_url",              "https://demoqa.com/", "Launch Mode URL",
     "Fallback URL for AI Databank scrape launch mode. Used when no live debug browser is available; Selenium will open a fresh browser and navigate here before scraping.", "Chrome", "text"),

    # -- Playwright (browser launch) ----------------------------------------
    ("playwright.slow_mo",                  "0",     "Slow Motion (ms)",
     "Slows down every Playwright operation by this many milliseconds. Useful for debugging.", "Playwright", "number"),
    ("playwright.devtools",                 "false", "Open DevTools",
     "Automatically open Chromium DevTools on every launched page.",           "Playwright", "checkbox"),
    ("playwright.extra_args",               "--start-maximized",  "Extra Chromium Arguments",
     "Additional Chromium flags, one per line. Applied to every Playwright launch.", "Playwright", "textarea"),

    # -- Playwright (browser context) ---------------------------------------
    ("playwright.viewport_width",           "1280",  "Viewport Width (px)",
     "Browser viewport width. Leave blank for no-viewport (maximised) mode.",  "Playwright", "number"),
    ("playwright.viewport_height",          "720",   "Viewport Height (px)",
     "Browser viewport height. Leave blank for no-viewport (maximised) mode.", "Playwright", "number"),
    ("playwright.user_agent",               "",      "User Agent",
     "Custom User-Agent string. Leave blank for the Chromium default.",        "Playwright", "text"),
    ("playwright.locale",                   "en-US", "Locale",
     "Browser locale (e.g. en-US, fr-FR). Leave blank for default.",           "Playwright", "text"),
    ("playwright.timezone_id",              "",      "Timezone ID",
     "IANA timezone (e.g. Asia/Manila, America/New_York). Leave blank for system default.", "Playwright", "text"),
    ("playwright.geolocation_latitude",     "",      "Geolocation Latitude",
     "Latitude for geolocation override (e.g. 14.6). Leave blank to disable.", "Playwright", "text"),
    ("playwright.geolocation_longitude",    "",      "Geolocation Longitude",
     "Longitude for geolocation override (e.g. 121.0). Leave blank to disable.", "Playwright", "text"),
    ("playwright.permissions",              "",      "Browser Permissions",
     "Comma-separated list of permissions to grant (e.g. geolocation, notifications).", "Playwright", "text"),
    ("playwright.record_video",             "false", "Record Video",
     "Record a video of every Playwright replay session.",                     "Playwright", "checkbox"),
    ("playwright.record_video_dir",         "logs/videos", "Video Directory",
     "Folder where recorded videos are saved (relative to project root).",    "Playwright", "text"),
    ("playwright.accept_downloads",         "true",  "Accept Downloads",
     "Automatically accept all file downloads during replay.",                "Playwright", "checkbox"),

    # -- Playwright (page defaults) -----------------------------------------
    ("playwright.default_timeout",          "30000", "Default Timeout (ms)",
     "Default timeout for all Playwright actions (page.set_default_timeout).", "Playwright", "number"),
    ("playwright.default_navigation_timeout", "60000", "Navigation Timeout (ms)",
     "Default navigation timeout (page.set_default_navigation_timeout).",     "Playwright", "number"),

    # -- Playwright (proxy) -------------------------------------------------
    ("playwright.proxy_server",             "",      "Proxy Server",
     "HTTP proxy server URL (e.g. http://proxy.com:8080). Leave blank to disable.", "Playwright", "text"),
    ("playwright.proxy_username",           "",      "Proxy Username",
     "Username for proxy authentication.",                                    "Playwright", "text"),
    ("playwright.proxy_password",           "",      "Proxy Password",
     "Password for proxy authentication.",                                    "Playwright", "text"),

    # -- Playwright (persistent context) ------------------------------------
    ("playwright.user_data_dir",            "",      "User Data Directory",
     "Path to a Chromium user data directory for persistent context (keeps cookies/sessions). Leave blank for a fresh profile.", "Playwright", "text"),

    ("license.product_name",                "WebConX Automation", "Product Name",
     "Display name shown on the Licensing page.",                    "Licensing", "text"),
    ("license.license_key",                 "", "License Key",
     "The current license key shown to administrators.",             "Licensing", "text"),
    ("license.plan",                        "Enterprise", "Plan",
     "The active subscription or support plan.",                     "Licensing", "text"),
    ("license.status",                      "Active", "License Status",
     "Human-readable status label for the current license.",         "Licensing", "text"),
    ("license.seats_total",                 "25", "Seats Purchased",
     "Total number of licensed seats.",                              "Licensing", "number"),
    ("license.seats_used",                  "0", "Seats In Use",
     "Currently assigned seats.",                                     "Licensing", "number"),
    ("license.valid_from",                  "", "Valid From",
     "License effective date in YYYY-MM-DD format.",                 "Licensing", "text"),
    ("license.valid_until",                 "", "Valid Until",
     "License expiry date in YYYY-MM-DD format.",                    "Licensing", "text"),
    ("license.owner",                       "", "Licensed To",
     "Company or person the license is issued to.",                  "Licensing", "text"),
    ("license.contact_email",               "", "Support Email",
     "Support or account contact email address.",                    "Licensing", "text"),
    ("license.notes",                       "", "Notes",
     "Optional freeform notes for licensing and renewals.",          "Licensing", "textarea"),

    ("recorder.browser",                    "chrome", "Browser",
     "Browser used when recording: chrome, firefox, or msedge.",                 "Recorder",  "select"),

    ("firefox.implicit_wait",               "10",    "Implicit Wait (seconds)",
     "Seconds Firefox waits for elements to appear before timing out.",           "Firefox",   "number"),
    ("firefox.extra_arguments",             "--width=1920\n--heigh=1080",      "Firefox Arguments",
     "Additional Firefox flags, one per line (e.g. --width=1920).",              "Firefox",   "textarea"),
    ("firefox.webdriver_path",              "",      "GeckoDriver Executable",
     "Pick a driver from the webdrivers/firefox/ folder, or leave Auto.",        "Firefox",   "select"),

    ("edge.implicit_wait",                  "10",    "Implicit Wait (seconds)",
     "Seconds Edge waits for elements to appear before timing out.",              "Edge",      "number"),
    ("edge.extra_arguments",                "",      "Extra Edge Arguments",
     "Additional Edge flags, one per line (e.g. --window-size=1920,1080).",      "Edge",      "textarea"),
    ("edge.webdriver_path",                 "",      "EdgeDriver Executable",
     "Pick a driver from the webdrivers/edge/ folder, or leave Auto.",           "Edge",      "select"),
    ("edge.remote_debugging_port",          "",      "Remote Debugging Port",
     "CDP port Edge was started with (e.g. 9222). Used by AI Databank scrape when Recorder -> Browser is set to Edge.", "Edge", "number"),
]

_GROUP_ICONS = {
    "Replay":    "bi-play-circle",
    "Dashboard": "bi-speedometer2",
    "UI":        "bi-layout-text-sidebar",
    "Features":  "bi-toggles",
    "Licensing": "bi-key",
    "Recorder":  "bi-record-circle",
    "Chrome":    "bi-browser-chrome",
    "Firefox":   "bi-browser-firefox",
    "Edge":      "bi-browser-edge",
    "Playwright": "bi-lightning-charge",
}

# ---------------------------------------------------------------------------
# WebDriver scanner
# ---------------------------------------------------------------------------

_WEBDRIVERS_DIR        = os.path.join(os.path.dirname(os.path.dirname(__file__)), "webdrivers")
_WEBDRIVERS_CHROME_DIR = os.path.join(_WEBDRIVERS_DIR, "chrome")
_WEBDRIVERS_FIREFOX_DIR = os.path.join(_WEBDRIVERS_DIR, "firefox")
_WEBDRIVERS_EDGE_DIR   = os.path.join(_WEBDRIVERS_DIR, "edge")


def _scan_webdrivers(browser_dir: str) -> list[dict]:
    """Return [{label, value}] for every file found in *browser_dir*.

    The first entry is always "Auto" (empty value = let Selenium Manager decide).
    label  = filename without extension  (e.g. chromedriver_124)
    value  = filename with extension     (e.g. chromedriver_124.exe) – path resolved at runtime
    """
    opts = [{"label": "Auto (Selenium Manager)", "value": ""}]
    try:
        if os.path.isdir(browser_dir):
            for name in sorted(os.listdir(browser_dir)):
                full_path = os.path.join(browser_dir, name)
                if os.path.isfile(full_path):
                    label = os.path.splitext(name)[0]  # strip .exe / .zip etc.
                    opts.append({"label": label, "value": name})
    except Exception:
        pass
    return opts


_SELECT_OPTIONS: dict[str, list[dict]] = {
    "chrome.webdriver_path":  None,  # populated lazily per request
    "firefox.webdriver_path": None,
    "edge.webdriver_path":    None,
    "recorder.browser":       None,
}


def _ensure_schema_columns():
    """Add any optional columns that may be missing from legacy DB installs."""
    with connection.cursor() as cur:
        cur.execute("ALTER TABLE locators ADD COLUMN IF NOT EXISTS pos_x FLOAT;")
        cur.execute("ALTER TABLE locators ADD COLUMN IF NOT EXISTS pos_y FLOAT;")
        cur.execute("ALTER TABLE steps ADD COLUMN IF NOT EXISTS recorder TEXT;")
        cur.execute("ALTER TABLE steps ADD COLUMN IF NOT EXISTS runner  TEXT;")
        cur.execute("ALTER TABLE session_meta ADD COLUMN IF NOT EXISTS recorder TEXT;")
        cur.execute("ALTER TABLE run_table    ADD COLUMN IF NOT EXISTS runner      TEXT;")
        cur.execute("ALTER TABLE run_table    ADD COLUMN IF NOT EXISTS folder_name TEXT;")
        cur.execute("ALTER TABLE run_table    ADD COLUMN IF NOT EXISTS run_date    TIMESTAMPTZ;")
        cur.execute("ALTER TABLE run_table    ALTER COLUMN run_date TYPE TIMESTAMPTZ USING run_date::TIMESTAMPTZ;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS locators_raw JSONB;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS field_name   TEXT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS field_value  TEXT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS pos_x        FLOAT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS pos_y        FLOAT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS strategy     TEXT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS locator      TEXT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS is_primary   BOOLEAN;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS locator_rank INTEGER;")
        cur.execute("ALTER TABLE locators     ADD COLUMN IF NOT EXISTS locator_rank INTEGER;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS headless_state BOOLEAN NOT NULL DEFAULT FALSE;")
        cur.execute("ALTER TABLE steps        ADD COLUMN IF NOT EXISTS headless_state BOOLEAN NOT NULL DEFAULT FALSE;")
        cur.execute("ALTER TABLE recordings ADD COLUMN IF NOT EXISTS headless_state BOOLEAN NOT NULL DEFAULT FALSE")
        cur.execute("ALTER TABLE steps      ADD COLUMN IF NOT EXISTS headless_state BOOLEAN NOT NULL DEFAULT FALSE")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS file_order INTEGER NOT NULL DEFAULT 1;")
        cur.execute("ALTER TABLE steps        ADD COLUMN IF NOT EXISTS file_order INTEGER NOT NULL DEFAULT 1;")
        cur.execute("ALTER TABLE run_table    ADD COLUMN IF NOT EXISTS file_order INTEGER NOT NULL DEFAULT 1;")
        cur.execute("ALTER TABLE run_table    ADD COLUMN IF NOT EXISTS parent_folder_id UUID;")
        cur.execute("ALTER TABLE run_table    ADD COLUMN IF NOT EXISTS sub_folder_id    UUID;")
        cur.execute("ALTER TABLE run_table    ADD COLUMN IF NOT EXISTS end_folder_id    UUID;")
        cur.execute("ALTER TABLE locators     ADD COLUMN IF NOT EXISTS folder_name TEXT;")
        cur.execute("ALTER TABLE data         ADD COLUMN IF NOT EXISTS folder_name TEXT;")
        cur.execute("ALTER TABLE session_meta ADD COLUMN IF NOT EXISTS folder_name TEXT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS runner      TEXT;")
        cur.execute("ALTER TABLE recordings   ADD COLUMN IF NOT EXISTS folder_name TEXT;")
    _ensure_project_folders_table()


def _ensure_config_table():
    """Create app_config table if missing, then upsert default rows.

    The user's saved *value* is preserved on conflict; only metadata
    (label, description, group_name, input_type) is refreshed so that
    changes to _CONFIG_DEFAULTS (e.g. text → select) take effect immediately.
    """
    with connection.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS app_config (
                key         VARCHAR(100) PRIMARY KEY,
                value       TEXT NOT NULL DEFAULT '',
                label       VARCHAR(200) NOT NULL DEFAULT '',
                description TEXT NOT NULL DEFAULT '',
                group_name  VARCHAR(100) NOT NULL DEFAULT 'General',
                input_type  VARCHAR(20)  NOT NULL DEFAULT 'text'
            )
        """)
        for key, value, label, description, group_name, input_type in _CONFIG_DEFAULTS:
            cur.execute("""
                INSERT INTO app_config (key, value, label, description, group_name, input_type)
                VALUES (%s, %s, %s, %s, %s, %s)
                ON CONFLICT (key) DO UPDATE SET
                    label       = EXCLUDED.label,
                    description = EXCLUDED.description,
                    group_name  = EXCLUDED.group_name,
                    input_type  = EXCLUDED.input_type,
                    value       = CASE
                                    WHEN app_config.value = '' AND EXCLUDED.value != ''
                                    THEN EXCLUDED.value
                                    ELSE app_config.value
                                  END
            """, [key, value, label, description, group_name, input_type])
        # Remove keys that no longer exist in _CONFIG_DEFAULTS
        _valid_keys = [row[0] for row in _CONFIG_DEFAULTS]
        cur.execute(
            "DELETE FROM app_config WHERE key != ALL(%s)",
            [_valid_keys],
        )
        # Migration: ensure SSL/cert bypass flags are present in chrome.extra_arguments
        # for existing installations where the value was set before these defaults existed.
        for _ssl_flag in ("--ignore-certificate-errors", "--allow-running-insecure-content"):
            cur.execute(
                "UPDATE app_config SET value = value || %s "
                "WHERE key = 'chrome.extra_arguments' AND position(%s IN value) = 0",
                [f"\n{_ssl_flag}", _ssl_flag],
            )

        # Migration: strip any headless=False / headless=True lines that were
        # baked into the old default for playwright.extra_args.  Headless mode is
        # now controlled exclusively by the per-session toggle on the Projects
        # page (steps.headless_state → replay POST parameter), so the global
        # extra_args should not carry a conflicting headless override.
        cur.execute("""
            UPDATE app_config
            SET value = trim(
                regexp_replace(
                    regexp_replace(value,
                        '(?im)^\\s*headless\\s*=\\s*(true|false)\\s*\n?', '', 'g'),
                    '^\n+|\n+$', '', 'g')
            )
            WHERE key = 'playwright.extra_args'
              AND value ~* 'headless\\s*=\\s*(true|false)'
        """)

        # Cleanup: reset placeholder Playwright values that should remain opt-in.
        # These were example values, not safe runtime defaults.
        for _pw_key, _pw_value in {
            "playwright.user_agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/136.0.0.0 Safari/537.36",
            "playwright.timezone_id": "Asia/Manila",
            "playwright.geolocation_latitude": "14.6",
            "playwright.geolocation_longitude": "121.0",
            "playwright.permissions": "geolocation",
            "playwright.proxy_server": "http://proxy.com:8080",
            "playwright.proxy_username": "user",
            "playwright.proxy_password": "pass",
            "playwright.user_data_dir": "user-data",
        }.items():
            cur.execute(
                "UPDATE app_config SET value = '' WHERE key = %s AND value = %s",
                [_pw_key, _pw_value],
            )


def get_config(key: str, default: str = "") -> str:
    """Return a single config value; safe to call even before the table exists."""
    try:
        with connection.cursor() as cur:
            cur.execute("SELECT value FROM app_config WHERE key = %s", [key])
            row = cur.fetchone()
            return row[0] if row else default
    except Exception:
        return default


@login_required
def configuration(request):
    """Show / save the application configuration table."""
    _ensure_config_table()

    with connection.cursor() as cur:
        cur.execute("SELECT key, value FROM app_config")
        value_map = {row[0]: row[1] for row in cur.fetchall()}

    all_settings = [
        {
            "key": key,
            "value": value_map.get(key, default_value),
            "label": label,
            "description": description,
            "group_name": group_name,
            "input_type": input_type,
        }
        for key, default_value, label, description, group_name, input_type in _CONFIG_DEFAULTS
    ]

    groups: dict[str, list[dict[str, Any]]] = {}
    group_order: list[str] = []
    for setting in all_settings:
        group_name = setting["group_name"]
        if group_name not in groups:
            groups[group_name] = []
            group_order.append(group_name)
        groups[group_name].append(setting)

    _wd_chrome = _scan_webdrivers(_WEBDRIVERS_CHROME_DIR)
    _wd_firefox = _scan_webdrivers(_WEBDRIVERS_FIREFOX_DIR)
    _wd_edge = _scan_webdrivers(_WEBDRIVERS_EDGE_DIR)
    _dynamic_options = {
        "chrome.webdriver_path": _wd_chrome,
        "firefox.webdriver_path": _wd_firefox,
        "edge.webdriver_path": _wd_edge,
        "recorder.browser": [
            {"label": "Chrome", "value": "chrome"},
            {"label": "Firefox", "value": "firefox"},
            {"label": "Edge", "value": "edge"},
        ],
        "replay.execution_mode": [
            {"label": "Parallel (simultaneous)", "value": "parallel"},
            {"label": "Serial (one at a time)", "value": "serial"},
        ],
    }
    def _driver_hint(opts, dirpath):
        n = len(opts) - 1
        return (f'{n} driver{"s" if n != 1 else ""} found &nbsp;·&nbsp; '
                f'Folder: <code>{os.path.relpath(dirpath)}</code>')
    _select_hints = {
        "chrome.webdriver_path":  _driver_hint(_wd_chrome,  _WEBDRIVERS_CHROME_DIR),
        "firefox.webdriver_path": _driver_hint(_wd_firefox, _WEBDRIVERS_FIREFOX_DIR),
        "edge.webdriver_path":    _driver_hint(_wd_edge,    _WEBDRIVERS_EDGE_DIR),
        "recorder.browser":       "Browser used by the standalone recorder (main.py).",
    }
    for s in all_settings:
        if s["input_type"] == "select":
            s["options"] = _dynamic_options.get(s["key"], [])
            s["hint"]    = _select_hints.get(s["key"], "")
            # Normalize: if stored value is an old absolute path, reduce to filename only
            if os.sep in s["value"] or (os.altsep and os.altsep in s["value"]):
                s["value"] = os.path.basename(s["value"])
                # Persist the normalised value so it's clean in the DB going forward
                with connection.cursor() as _nc:
                    _nc.execute("UPDATE app_config SET value = %s WHERE key = %s",
                                [s["value"], s["key"]])

    # Build list of (group_name, icon, settings_list) for the template
    group_list = [
        (g, _GROUP_ICONS.get(g, "bi-sliders"), groups[g])
        for g in group_order
    ]

    return render(request, "recorder/configuration.html", {
        "group_list": group_list,
    })


@login_required
def licensing(request):
    _ensure_config_table()

    def _int_value(key: str, default: int = 0) -> int:
        try:
            return max(0, int((get_config(key, str(default)) or str(default)).strip() or default))
        except (TypeError, ValueError):
            return default

    seats_total = _int_value("license.seats_total", 25)
    seats_used = _int_value("license.seats_used", 0)
    if seats_total > 0:
        seats_used = min(seats_used, seats_total)
    seats_available = max(seats_total - seats_used, 0)
    utilization_pct = int(round((seats_used / seats_total) * 100)) if seats_total > 0 else 0

    return render(request, "recorder/licensing.html", {
        "license_product_name": get_config("license.product_name", "WebConX Automation"),
        "license_key": get_config("license.license_key", ""),
        "license_plan": get_config("license.plan", "Enterprise"),
        "license_status": get_config("license.status", "Active"),
        "license_valid_from": get_config("license.valid_from", ""),
        "license_valid_until": get_config("license.valid_until", ""),
        "license_owner": get_config("license.owner", ""),
        "license_contact_email": get_config("license.contact_email", ""),
        "license_notes": get_config("license.notes", ""),
        "license_seats_total": seats_total,
        "license_seats_used": seats_used,
        "license_seats_available": seats_available,
        "license_utilization_pct": utilization_pct,
    })


@login_required
def configuration_api(request):
    """GET: return all config key/value pairs as JSON.
       POST {key, value}: save a single config entry."""
    if not request.user.is_superuser:
        return JsonResponse({"error": "Access denied."}, status=403)
    try:
        _ensure_config_table()
        if request.method == "POST":
            data = json.loads(request.body)
            key   = data.get("key", "").strip()
            value = data.get("value", "")
            if not key:
                return JsonResponse({"error": "key required"}, status=400)
            with connection.cursor() as cur:
                cur.execute(
                    "UPDATE app_config SET value = %s WHERE key = %s",
                    [value, key],
                )
            return JsonResponse({"ok": True})
        with connection.cursor() as cur:
            cur.execute("SELECT key, value FROM app_config")
            data = {row[0]: row[1] for row in cur.fetchall()}
        return JsonResponse(data)
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)


# ---------------------------------------------------------------------------
# Per-user preferences
# ---------------------------------------------------------------------------

_ALL_PROJ_ACTIONS = {
    "copy",
    "copy_folder_button",
    "copy_folder_modal",
    "create_folder_modal",
    "create_subfolder_button",
    "create_subfolder_modal",
    "delete",
    "delete_folder",
    "delete_folder_modal",
    "download",
    "duplicate",
    "headless",
    "rename",
    "rename_folder_button",
    "rename_folder_modal",
    "replay",
    "view",
}


def _resolve_project_visible_actions(pref_raw: str) -> set[str]:
    visible = {a.strip() for a in (pref_raw or "").split(",") if a.strip()}
    if not visible:
        return set(_ALL_PROJ_ACTIONS)

    # Backward compatibility for older saved preferences.
    if "create_subfolder" in visible:
        visible.add("create_subfolder_button")
        visible.add("create_subfolder_modal")
    if "copy_folder" in visible:
        visible.add("copy_folder_button")
    if "rename_folder" in visible:
        visible.add("rename_folder_button")
        visible.add("rename_folder_modal")

    return {action for action in visible if action in _ALL_PROJ_ACTIONS}


def _ensure_user_prefs_table():
    with connection.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS user_prefs (
                user_id   INTEGER NOT NULL,
                pref_key  VARCHAR(100) NOT NULL,
                pref_value TEXT NOT NULL DEFAULT '',
                PRIMARY KEY (user_id, pref_key)
            )
        """)


def get_user_pref(user_id: int, key: str, default: str = "") -> str:
    try:
        _ensure_user_prefs_table()
        with connection.cursor() as cur:
            cur.execute(
                "SELECT pref_value FROM user_prefs WHERE user_id = %s AND pref_key = %s",
                [user_id, key],
            )
            row = cur.fetchone()
            return row[0] if row else default
    except Exception:
        return default


@login_required
def save_user_pref(request):
    """POST JSON {key, value}: save a single user preference."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        body = json.loads(request.body)
        key   = str(body["key"]).strip()
        value = str(body["value"]).strip()
    except (KeyError, ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    _ensure_user_prefs_table()
    with connection.cursor() as cur:
        cur.execute("""
            INSERT INTO user_prefs (user_id, pref_key, pref_value)
            VALUES (%s, %s, %s)
            ON CONFLICT (user_id, pref_key) DO UPDATE SET pref_value = EXCLUDED.pref_value
        """, [request.user.id, key, value])
    return JsonResponse({"ok": True})


@login_required
def admin_save_user_pref(request, pk):
    """POST JSON {key, value}: superuser saves a preference for any user by pk."""
    if not request.user.is_superuser:
        return JsonResponse({"error": "Forbidden"}, status=403)
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        target = User.objects.get(pk=pk)
    except User.DoesNotExist:
        return JsonResponse({"error": "User not found"}, status=404)
    try:
        body  = json.loads(request.body)
        key   = str(body["key"]).strip()
        value = str(body["value"]).strip()
    except (KeyError, ValueError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    _ensure_user_prefs_table()
    with connection.cursor() as cur:
        cur.execute("""
            INSERT INTO user_prefs (user_id, pref_key, pref_value)
            VALUES (%s, %s, %s)
            ON CONFLICT (user_id, pref_key) DO UPDATE SET pref_value = EXCLUDED.pref_value
        """, [target.id, key, value])
    return JsonResponse({"ok": True})


# ---------------------------------------------------------------------------
# Tenant management  (superuser only)
# ---------------------------------------------------------------------------

@login_required
def tenant_list(request):
    """List all tenants with member counts.  Superuser only."""
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")
    from .models import Tenant, UserProfile
    tenants = (
        Tenant.objects
        .prefetch_related("members__user")
        .order_by("name")
    )
    tenant_rows = []
    for t in tenants:
        member_count = t.members.count()
        tenant_rows.append({
            "id":           str(t.id),
            "name":         t.name,
            "slug":         t.slug,
            "is_active":    t.is_active,
            "created_at":   t.created_at,
            "member_count": member_count,
        })
    all_users = User.objects.select_related("profile__tenant").order_by("username")
    # Ensure every user has a profile row so the template can safely access u.profile
    from .models import UserProfile
    existing_pks = set(UserProfile.objects.values_list("user_id", flat=True))
    new_profiles = [
        UserProfile(user=u) for u in all_users if u.pk not in existing_pks
    ]
    if new_profiles:
        UserProfile.objects.bulk_create(new_profiles, ignore_conflicts=True)
        all_users = User.objects.select_related("profile__tenant").order_by("username")
    return render(request, "recorder/tenant_list.html", {
        "tenant_rows": tenant_rows,
        "all_users":   all_users,
    })


@login_required
def tenant_create(request):
    """POST: create a new tenant.  Superuser only."""
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")
    if request.method != "POST":
        return redirect("tenant_list")
    from django.utils.text import slugify
    from .models import Tenant
    name = request.POST.get("name", "").strip()
    if not name:
        messages.error(request, "Tenant name is required.")
        return redirect("tenant_list")
    slug = slugify(name)[:100]
    if Tenant.objects.filter(slug=slug).exists():
        messages.error(request, f"A tenant with slug '{slug}' already exists.")
        return redirect("tenant_list")
    Tenant.objects.create(name=name, slug=slug, is_active=True)
    messages.success(request, f"Tenant '{name}' created.")
    return redirect("tenant_list")


@login_required
def tenant_delete(request, tenant_id):
    """POST: delete a tenant.  Superuser only."""
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect("dashboard")
    if request.method != "POST":
        return redirect("tenant_list")
    from .models import Tenant
    try:
        t = Tenant.objects.get(id=tenant_id)
        name = t.name
        # Detach members before deleting (SET NULL on FK, just ensure profiles exist)
        t.delete()
        messages.success(request, f"Tenant '{name}' deleted.")
    except Tenant.DoesNotExist:
        messages.error(request, "Tenant not found.")
    return redirect("tenant_list")


@login_required
def tenant_assign_user(request):
    """POST JSON {user_id, tenant_id}: assign a user to a tenant.  Superuser only."""
    if not request.user.is_superuser:
        return JsonResponse({"error": "Access denied."}, status=403)
    if request.method != "POST":
        return JsonResponse({"error": "POST required."}, status=405)
    try:
        body      = json.loads(request.body)
        user_id   = int(body["user_id"])
        tenant_id = body.get("tenant_id")   # None means "unassign"
    except (KeyError, ValueError, TypeError, json.JSONDecodeError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    from .models import Tenant, UserProfile
    try:
        target_user = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return JsonResponse({"error": "User not found."}, status=404)
    tenant_obj = None
    if tenant_id:
        try:
            tenant_obj = Tenant.objects.get(id=tenant_id)
        except Tenant.DoesNotExist:
            return JsonResponse({"error": "Tenant not found."}, status=404)
    profile, _ = UserProfile.objects.get_or_create(user=target_user)
    profile.tenant = tenant_obj
    profile.save()
    return JsonResponse({
        "ok":        True,
        "user_id":   user_id,
        "username":  target_user.username,
        "tenant_id": str(tenant_obj.id) if tenant_obj else None,
        "tenant":    tenant_obj.name if tenant_obj else None,
    })


# ---------------------------------------------------------------------------
# Chatbot API
# ---------------------------------------------------------------------------

@csrf_exempt
@login_required
def chatbot_api(request):
    """POST JSON {"message": "...", "history": [...]} — chatbot agent endpoint."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    try:
        body = json.loads(request.body)
        user_message = (body.get("message") or "").strip()
        history = body.get("history") or []
    except (json.JSONDecodeError, ValueError) as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    if not user_message:
        return JsonResponse({"error": "Empty message"}, status=400)

    # Save user message
    from recorder.models import ChatMessage
    ChatMessage.objects.create(user=request.user, role="user", content=user_message)

    try:
        normalized_message = re.sub(r"\s+", " ", user_message or "").strip()
        _is_generate_workflow = (
            re.search(r"\bgenerate\b", normalized_message, re.IGNORECASE)
            and re.search(r"\b(?:workflow|wrokflow|workfow|worflow|worklow)\b", normalized_message, re.IGNORECASE)
        )
        if _is_generate_workflow:
            from ai_agent import _parse_generate_workflow_request, tool_generate_workflow_test_case

            workflow_request = _parse_generate_workflow_request(normalized_message) or {}
            tool_result = tool_generate_workflow_test_case(
                workflow_query=str(workflow_request.get("workflow_query") or "").strip(),
                folder_name="AI Gen",
                author=(request.user.username or "admin").strip() or "admin",
            )
            result = {
                "reply": str(tool_result.get("text") or ""),
                "download": tool_result.get("download"),
            }
        else:
            from ai_agent import handle_chat_message
            result = handle_chat_message(
                user_message,
                username=request.user.username,
                conversation_history=history,
            )
    except Exception as exc:
        error_reply = f"Agent error: {exc}"
        ChatMessage.objects.create(user=request.user, role="bot", content=error_reply)
        return JsonResponse({"reply": error_reply, "download": None})

    reply = result.get("reply", "")
    download = result.get("download")

    if reply:
        reply = reply.replace("C:\\desktop__automation", str(settings.BASE_DIR))
        reply = reply.replace("C:/desktop__automation", str(settings.BASE_DIR).replace("\\", "/"))

    if isinstance(download, dict):
        for key in ("url", "filename", "content"):
            value = download.get(key)
            if isinstance(value, str):
                download[key] = value.replace("C:\\desktop__automation", str(settings.BASE_DIR))
                download[key] = download[key].replace("C:/desktop__automation", str(settings.BASE_DIR).replace("\\", "/"))

    # Save bot reply
    ChatMessage.objects.create(
        user=request.user, role="bot", content=reply,
        download=download,
    )

    return JsonResponse({"reply": reply, "download": download})


@login_required
def chatbot_history(request):
    """GET — return the last 50 chat messages for the current user."""
    from recorder.models import ChatMessage
    messages = ChatMessage.objects.filter(user=request.user).order_by("-created_at")[:50]
    data = list(reversed([
        {
            "role": m.role,
            "content": m.content,
            "download": m.download,
            "created_at": m.created_at.isoformat(),
        }
        for m in messages
    ]))
    return JsonResponse({"messages": data})


@csrf_exempt
@login_required
def chatbot_clear(request):
    """POST — clear all chat history for the current user."""
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    from recorder.models import ChatMessage
    deleted, _ = ChatMessage.objects.filter(user=request.user).delete()
    return JsonResponse({"cleared": deleted})


