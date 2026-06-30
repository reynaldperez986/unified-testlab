import json
import time
import requests
from requests.auth import HTTPDigestAuth
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.db import connection
from django.db.models import Count, Q, OuterRef, Subquery
from django.urls import reverse
from django.utils import timezone
from datetime import timedelta
from urllib.parse import urlparse
import csv
import os
from pathlib import Path
from pprint import pformat
import re
import subprocess
import sys
import tempfile

try:
    from requests_oauthlib import OAuth1
except ImportError:
    OAuth1 = None

try:
    from requests_aws4auth import AWS4Auth
except ImportError:
    AWS4Auth = None

try:
    from requests_ntlm import HttpNtlmAuth
except ImportError:
    HttpNtlmAuth = None

from .models import Environment, Project, TestCase, TestExecution, AuditLog, UserProfile, ApiModule, ModuleEndpoint, ThemeSettings
from .forms import (
    LoginForm, EnvironmentForm, TestCaseForm, ProjectCreateForm, UserCreateForm, UserEditForm, ThemeSettingsForm
)

def get_client_ip(request):
    """Extract client IP from request."""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        return x_forwarded_for.split(',')[0].strip()
    return request.META.get('REMOTE_ADDR')


def log_action(user, action, target_type='', target_id=None, details='', request=None):
    """Create an audit log entry."""
    AuditLog.objects.create(
        user=user,
        action=action,
        target_type=target_type,
        target_id=target_id,
        details=details,
        ip_address=get_client_ip(request) if request else None,
    )


def role_required(roles):
    """Decorator to restrict access by user role."""
    def decorator(view_func):
        def wrapper(request, *args, **kwargs):
            if not hasattr(request.user, 'api_profile'):
                messages.error(request, 'User profile not found.')
                return redirect('api:dashboard')
            if request.user.api_profile.role not in roles:
                messages.error(request, 'You do not have permission to access this page.')
                return redirect('api:dashboard')
            return view_func(request, *args, **kwargs)
        wrapper.__name__ = view_func.__name__
        return wrapper
    return decorator


def save_response_files(test_case_name, response_body):
    """Save response JSON and CSV files to the Response folder."""
    try:
        # Ensure Response folder exists
        response_dir = Path(__file__).resolve().parent.parent / 'Response'
        response_dir.mkdir(exist_ok=True)
        
        # Sanitize test case name for filename
        safe_name = test_case_name.replace(' ', '_').replace('/', '_').replace('\\', '_')
        
        # Parse response body
        try:
            parsed_body = json.loads(response_body) if response_body else {}
        except (json.JSONDecodeError, TypeError):
            parsed_body = {}
        
        # Save JSON file
        json_path = response_dir / f'{safe_name}_response.json'
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(parsed_body, f, indent=2, ensure_ascii=False)
        
        # Save CSV file
        csv_path = response_dir / f'{safe_name}_response.csv'
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            
            if isinstance(parsed_body, list):
                # If response is a list of objects, write as rows
                if parsed_body and isinstance(parsed_body[0], dict):
                    # Get all keys from all objects
                    fieldnames = set()
                    for item in parsed_body:
                        if isinstance(item, dict):
                            fieldnames.update(item.keys())
                    fieldnames = sorted(list(fieldnames))
                    writer.writerow(fieldnames)
                    for item in parsed_body:
                        row = [str(item.get(field, '')) for field in fieldnames]
                        writer.writerow(row)
                else:
                    # List of primitives
                    writer.writerow(['Value'])
                    for item in parsed_body:
                        writer.writerow([str(item)])
            elif isinstance(parsed_body, dict):
                # If response is a single object, write as key-value pairs
                writer.writerow(['Key', 'Value'])
                for key, value in parsed_body.items():
                    if isinstance(value, (dict, list)):
                        writer.writerow([key, json.dumps(value)])
                    else:
                        writer.writerow([key, str(value)])
            else:
                # Fallback for other types
                writer.writerow(['Response'])
                writer.writerow([str(parsed_body)])
    except Exception as e:
        # Silently fail - don't interrupt test execution if file saving fails
        pass


def fetch_oauth2_access_token(token_url, client_id, client_secret):
    """Fetch OAuth 2.0 access token using client credentials grant."""
    if not token_url or not client_id or not client_secret:
        return None, 'Token URL, Client ID, and Client Secret are required.'

    try:
        response = requests.post(
            token_url,
            data={
                'grant_type': 'client_credentials',
                'client_id': client_id,
                'client_secret': client_secret,
            },
            timeout=20,
            verify=True,
        )
    except requests.exceptions.RequestException as e:
        return None, f'Unable to fetch token: {str(e)}'

    if not response.ok:
        return None, f'Token endpoint returned {response.status_code}: {response.text[:300]}'

    try:
        payload = response.json()
    except ValueError:
        return None, 'Token endpoint returned non-JSON response.'

    access_token = payload.get('access_token')
    if not access_token:
        return None, 'access_token was not found in token response.'

    return access_token, None


def _json_field_match(expected, actual):
    """Recursively validate that actual contains all expected JSON fields and values."""
    if isinstance(expected, dict):
        if not isinstance(actual, dict):
            return False
        for key, value in expected.items():
            if key not in actual:
                return False
            if not _json_field_match(value, actual[key]):
                return False
        return True

    if isinstance(expected, list):
        if not isinstance(actual, list):
            return False
        if len(expected) != len(actual):
            return False
        for idx, value in enumerate(expected):
            if not _json_field_match(value, actual[idx]):
                return False
        return True

    return expected == actual


def _json_dict_match_anywhere(expected_dict, actual):
    """Return True when expected_dict matches at root or any nested object in actual."""
    if _json_field_match(expected_dict, actual):
        return True

    if isinstance(actual, dict):
        for value in actual.values():
            if _json_dict_match_anywhere(expected_dict, value):
                return True
        return False

    if isinstance(actual, list):
        for value in actual:
            if _json_dict_match_anywhere(expected_dict, value):
                return True
        return False

    return False


def _match_expected_content(expected_content, response_body):
    """Match expected content against response body using strict JSON field matching when possible."""
    expected_text = (expected_content or '').strip()
    body_text = (response_body or '').strip()
    if not expected_text:
        return None

    # Accept common fragment inputs like '"chargeDetails": {...}' by coercing to a JSON object.
    expected_json_text = expected_text
    if not expected_json_text.startswith('{') and not expected_json_text.startswith('[') and ':' in expected_json_text:
        fragment_text = expected_json_text
        # Some saved fragments already include a dangling closing brace at the end
        # (for example: '"foo": 1 }'). Normalize before wrapping into an object.
        if fragment_text.endswith('}') and fragment_text.count('{') < fragment_text.count('}'):
            fragment_text = fragment_text[:-1].rstrip()
        expected_json_text = '{' + fragment_text + '}'

    expected_json = None
    try:
        expected_json = json.loads(expected_json_text)
    except (json.JSONDecodeError, TypeError, ValueError):
        # Accept multiple JSON object fragments pasted as: {...}, {...}
        # by coercing them into a JSON array.
        if expected_text.startswith('{') and '},' in expected_text:
            try:
                expected_json = json.loads(f'[{expected_text}]')
            except (json.JSONDecodeError, TypeError, ValueError):
                expected_json = None

    if expected_json is None:
        return expected_text in (response_body or '')

    try:
        body_json = json.loads(body_text)
    except (json.JSONDecodeError, TypeError, ValueError):
        return expected_text in (response_body or '')

    # Backward-compatibility: older expected format wrapped payload under response_body.
    if isinstance(expected_json, dict) and 'response_body' in expected_json:
        wrapped = expected_json.get('response_body')
        if isinstance(wrapped, (dict, list)):
            expected_json = wrapped

    if _json_field_match(expected_json, body_json):
        return True

    # For list fragments, require each expected object to match somewhere in payload.
    if isinstance(expected_json, list):
        all_matched = True
        for expected_item in expected_json:
            if isinstance(expected_item, dict):
                if not _json_dict_match_anywhere(expected_item, body_json):
                    all_matched = False
                    break
            else:
                if not _json_dict_match_anywhere({'value': expected_item}, {'value': body_json}):
                    all_matched = False
                    break
        if all_matched:
            return True

    # Allow flat expected JSON objects to match nested objects in response payloads.
    if isinstance(expected_json, dict):
        return _json_dict_match_anywhere(expected_json, body_json)

    return False


def _json_contains_key_value_pair(payload, key, expected_value):
    """Return True when payload contains key and matching value anywhere in nested JSON."""
    if isinstance(payload, dict):
        if key in payload:
            actual = payload.get(key)
            if expected_value == '':
                return True
            if actual is None:
                return expected_value == ''
            if isinstance(actual, bool):
                actual_text = 'true' if actual else 'false'
            else:
                actual_text = str(actual)
            if actual_text == expected_value or expected_value in actual_text:
                return True
        return any(_json_contains_key_value_pair(v, key, expected_value) for v in payload.values())

    if isinstance(payload, list):
        return any(_json_contains_key_value_pair(item, key, expected_value) for item in payload)

    return False


def _json_path_lookup(payload, path_parts):
    """Resolve dotted JSON path parts against payload and return (found, value)."""
    current = payload
    for part in path_parts:
        if isinstance(current, dict) and part in current:
            current = current.get(part)
            continue
        return False, None
    return True, current


def _json_value_matches(actual, expected_value):
    """Compare a JSON value against expected text value with light coercion."""
    if expected_value == '':
        return True
    if actual is None:
        return expected_value == ''
    if isinstance(actual, bool):
        actual_text = 'true' if actual else 'false'
    else:
        actual_text = str(actual)
    return actual_text == expected_value or expected_value in actual_text


def _json_contains_key_value_pair_flexible(payload, key, expected_value):
    """Match key/value in JSON supporting exact key, dotted paths, and global-key suffixes."""
    key_text = (key or '').strip()
    if not key_text:
        return False

    # 1) Exact key search anywhere in payload.
    if _json_contains_key_value_pair(payload, key_text, expected_value):
        return True

    # 2) Dotted path lookup (example: foo.bar.baz).
    if '.' in key_text:
        path_parts = [p.strip() for p in key_text.split('.') if p.strip()]
        if path_parts:
            found, value = _json_path_lookup(payload, path_parts)
            if found and _json_value_matches(value, expected_value):
                return True

        # 3) Global data keys may include prefixes (example: api.test001.fieldName).
        #    Fall back to matching by terminal key segment anywhere in response.
        terminal_key = path_parts[-1] if path_parts else ''
        if terminal_key and _json_contains_key_value_pair(payload, terminal_key, expected_value):
            return True

    return False


def _expected_content_mismatch_summary(expected_content, response_body, max_items=8):
    """Return a concise mismatch summary for expected-response validation failures."""
    expected_text = (expected_content or '').strip()
    body_text = (response_body or '').strip()
    if not expected_text:
        return ''

    if _match_expected_content(expected_text, body_text):
        return ''

    expected_json = None
    expected_json_text = expected_text
    if not expected_json_text.startswith('{') and not expected_json_text.startswith('[') and ':' in expected_json_text:
        fragment_text = expected_json_text
        if fragment_text.endswith('}') and fragment_text.count('{') < fragment_text.count('}'):
            fragment_text = fragment_text[:-1].rstrip()
        expected_json_text = '{' + fragment_text + '}'

    try:
        expected_json = json.loads(expected_json_text)
    except (json.JSONDecodeError, TypeError, ValueError):
        if expected_text.startswith('{') and '},' in expected_text:
            try:
                expected_json = json.loads(f'[{expected_text}]')
            except (json.JSONDecodeError, TypeError, ValueError):
                expected_json = None

    try:
        body_json = json.loads(body_text) if body_text else None
    except (json.JSONDecodeError, TypeError, ValueError):
        body_json = None

    if expected_json is None:
        return 'Expected text was not found in response body.'
    if body_json is None:
        return 'Expected JSON could not be matched because response body is not valid JSON.'

    pairs = []

    def collect(node):
        if isinstance(node, dict):
            for key, value in node.items():
                if isinstance(value, (dict, list)):
                    collect(value)
                else:
                    pairs.append((str(key), '' if value is None else str(value)))
            return
        if isinstance(node, list):
            for item in node:
                collect(item)

    collect(expected_json)
    if not pairs:
        return 'Expected response content mismatch.'

    mismatched = []
    for key, value in pairs:
        if not _json_contains_key_value_pair_flexible(body_json, key, value):
            if value:
                mismatched.append(f'{key}={value}')
            else:
                mismatched.append(f'{key}')

    if not mismatched:
        return 'Expected response content mismatch.'

    visible = mismatched[:max_items]
    suffix = '' if len(mismatched) <= max_items else f' (+{len(mismatched) - max_items} more)'
    return 'Missing/mismatched fields: ' + ', '.join(visible) + suffix


def _parse_form_data_rows(raw_form_data):
    """Parse stored form-data JSON rows into a normalized key/value list."""
    if not raw_form_data:
        return []
    try:
        rows = json.loads(raw_form_data)
    except (json.JSONDecodeError, TypeError, ValueError):
        return []

    normalized = []
    for row in rows or []:
        key = str((row or {}).get('key', '')).strip()
        if not key:
            continue
        value = str((row or {}).get('value', '')).strip()
        normalized.append({'key': key, 'value': value})
    return normalized


def _format_expected_key_pairs(rows):
    """Render expected key/value rows as a readable multiline string."""
    if not rows:
        return '-'
    lines = []
    for row in rows:
        key = row.get('key', '')
        value = row.get('value', '')
        if value:
            lines.append(f'{key}={value}')
        else:
            lines.append(key)
    return '\n'.join(lines)


def _form_data_mismatch_summary(form_data_rows, response_body, max_items=8):
    """Return mismatch details for expected key-pair values from Form Data."""
    if not form_data_rows:
        return ''

    match_result = _match_form_data_contains(form_data_rows, response_body)
    if match_result in (True, None):
        return ''

    body_text = response_body or ''
    try:
        body_json = json.loads(body_text) if body_text else None
    except (json.JSONDecodeError, TypeError, ValueError):
        body_json = None

    missing = []
    for row in form_data_rows:
        key = str((row or {}).get('key', '')).strip()
        value = str((row or {}).get('value', '')).strip()
        if not key:
            continue

        if body_json is not None:
            if not _json_contains_key_value_pair_flexible(body_json, key, value):
                missing.append(f'{key}={value}' if value else key)
        else:
            terminal_key = key.split('.')[-1].strip() if key else key
            has_key = (key in body_text) or (terminal_key in body_text)
            has_value = (not value) or (value in body_text)
            if not (has_key and has_value):
                missing.append(f'{key}={value}' if value else key)

    if not missing:
        return 'Expected key-pair values were not matched in response.'

    visible = missing[:max_items]
    suffix = '' if len(missing) <= max_items else f' (+{len(missing) - max_items} more)'
    return 'Missing key-pair values: ' + ', '.join(visible) + suffix


def _match_form_data_contains(form_data_rows, response_body):
    """Validate that response contains populated form-data key/value pairs."""
    pairs = []
    for row in form_data_rows or []:
        key = str((row or {}).get('key', '')).strip()
        value = str((row or {}).get('value', '')).strip()
        if key:
            pairs.append((key, value))

    if not pairs:
        return None

    body_text = response_body or ''
    try:
        body_json = json.loads(body_text) if body_text else None
    except (json.JSONDecodeError, TypeError, ValueError):
        body_json = None

    for key, value in pairs:
        if body_json is not None:
            if not _json_contains_key_value_pair_flexible(body_json, key, value):
                return False
        else:
            terminal_key = key.split('.')[-1].strip() if key else key
            if key not in body_text and terminal_key not in body_text:
                return False
            if value and value not in body_text:
                return False

    return True


def _safe_data_prefix(raw_name):
    """Create a stable prefix used for generated project data keys."""
    cleaned = re.sub(r'[^0-9A-Za-z]+', '_', (raw_name or '').strip()).strip('_').lower()
    return cleaned or 'api_response'


def _flatten_response_json(value, prefix=''):
    """Flatten nested JSON into (field_name, value) tuples for the shared data table."""
    rows = []
    if isinstance(value, dict):
        for key, child in value.items():
            key_str = str(key)
            child_prefix = f"{prefix}.{key_str}" if prefix else key_str
            rows.extend(_flatten_response_json(child, child_prefix))
        return rows

    if isinstance(value, list):
        if not value:
            rows.append((prefix or 'value', '[]'))
            return rows
        for idx, child in enumerate(value):
            child_prefix = f"{prefix}[{idx}]" if prefix else f"[{idx}]"
            rows.extend(_flatten_response_json(child, child_prefix))
        return rows

    if value is None:
        value_text = ''
    elif isinstance(value, bool):
        value_text = 'true' if value else 'false'
    else:
        value_text = str(value)

    rows.append((prefix or 'value', value_text))
    return rows


def _store_response_in_project_data(test_case, response_body, target_folder_name=None):
    """Persist flattened response JSON as global entries in recorder project data."""
    try:
        parsed = json.loads(response_body) if response_body else None
    except (json.JSONDecodeError, TypeError, ValueError):
        return

    if parsed is None:
        return

    testcase_prefix = _safe_data_prefix(test_case.name)
    field_prefix = f"api.{testcase_prefix}"

    explicit_folder_name = (target_folder_name or '').strip()
    default_folder_name = (getattr(test_case, 'project', '') or '').strip()

    folder_names = set()
    if explicit_folder_name:
        folder_names.add(explicit_folder_name)
    elif default_folder_name:
        folder_names.add(default_folder_name)

    # Keep existing folder mappings in sync: if this testcase already has stored
    # keys in other folders (e.g. Project001), overwrite those as well.
    with connection.cursor() as cur:
        cur.execute(
            """
            SELECT DISTINCT TRIM(COALESCE(folder_name, '')) AS folder
              FROM data
             WHERE field_name LIKE %s
               AND TRIM(COALESCE(folder_name, '')) <> ''
            """,
            [field_prefix + '%'],
        )
        for row in cur.fetchall():
            folder = (row[0] or '').strip()
            if folder:
                folder_names.add(folder)

    if not folder_names:
        return

    flattened = _flatten_response_json(parsed, field_prefix)
    if not flattened:
        return

    # Keep a practical cap so very large responses do not flood project data rows.
    flattened = flattened[:2000]

    with connection.cursor() as cur:
        for folder_name in folder_names:
            # Overwrite semantics: remove prior generated keys for this API testcase
            # in each synced folder, regardless of originating record/step.
            cur.execute(
                """
                DELETE FROM data
                 WHERE TRIM(COALESCE(folder_name, '')) = %s
                   AND field_name LIKE %s
                """,
                [
                    folder_name,
                    field_prefix + '%',
                ],
            )

            for field_name, value in flattened:
                cur.execute(
                    """
                    INSERT INTO data (record_id, step_no, field_name, value, folder_name, is_global, created_at)
                    VALUES (%s, 0, %s, %s, %s, TRUE, NOW())
                    """,
                    [
                        '00000000-0000-0000-0000-000000000000',
                        field_name,
                        value,
                        folder_name,
                    ],
                )


def _apply_auth_to_request(auth_type, auth_creds, headers, params):
    """Apply auth credentials to request objects and return (auth, error)."""
    auth = None
    auth_type = (auth_type or 'none').strip().lower()
    auth_creds = auth_creds or {}

    if auth_type in ('none', 'inherit', ''):
        return auth, None

    if auth_type == 'basic':
        auth = (auth_creds.get('username', ''), auth_creds.get('password', ''))
        return auth, None

    if auth_type == 'digest':
        auth = HTTPDigestAuth(auth_creds.get('username', ''), auth_creds.get('password', ''))
        return auth, None

    if auth_type in ('bearer', 'oauth2'):
        token = auth_creds.get('token', '')
        add_to = auth_creds.get('add_to', 'request_headers')
        header_prefix = auth_creds.get('header_prefix', 'Bearer') or 'Bearer'
        query_param_name = auth_creds.get('query_param_name', 'access_token') or 'access_token'
        if token:
            if add_to == 'request_url':
                params[query_param_name] = token
            else:
                headers['Authorization'] = f"{header_prefix} {token}".strip()
        return auth, None

    if auth_type == 'api_key':
        key_name = auth_creds.get('key_name', auth_creds.get('key', 'X-API-Key'))
        key_value = auth_creds.get('key_value', auth_creds.get('value', ''))
        add_to = auth_creds.get('add_to', 'request_headers')
        if add_to == 'request_url':
            params[key_name] = key_value
        elif add_to == 'cookie':
            existing_cookie = headers.get('Cookie', '').strip()
            new_cookie = f"{key_name}={key_value}"
            headers['Cookie'] = f"{existing_cookie}; {new_cookie}".strip('; ') if existing_cookie else new_cookie
        else:
            headers[key_name] = key_value
        return auth, None

    if auth_type == 'oauth1':
        if OAuth1 is None:
            return None, 'OAuth 1.0 requires requests-oauthlib package.'

        add_to = auth_creds.get('add_to', 'request_headers')
        signature_type = {
            'request_headers': 'AUTH_HEADER',
            'request_url': 'QUERY',
            'request_body': 'BODY',
        }.get(add_to, 'AUTH_HEADER')

        auth = OAuth1(
            client_key=auth_creds.get('consumer_key', auth_creds.get('client_key', '')),
            client_secret=auth_creds.get('consumer_secret', auth_creds.get('client_secret', '')),
            resource_owner_key=auth_creds.get('token', auth_creds.get('resource_owner_key', '')),
            resource_owner_secret=auth_creds.get('token_secret', auth_creds.get('resource_owner_secret', '')),
            signature_method=auth_creds.get('signature_method', 'HMAC-SHA1'),
            signature_type=signature_type,
        )
        return auth, None

    if auth_type == 'awsv4':
        if AWS4Auth is None:
            return None, 'AWS Signature requires requests-aws4auth package.'

        access_key = auth_creds.get('access_key', '')
        secret_key = auth_creds.get('secret_key', '')
        region = auth_creds.get('region', '')
        service = auth_creds.get('service', '')
        session_token = auth_creds.get('session_token', '')

        kwargs = {}
        if session_token:
            kwargs['session_token'] = session_token

        auth = AWS4Auth(access_key, secret_key, region, service, **kwargs)
        return auth, None

    if auth_type == 'ntlm':
        if HttpNtlmAuth is None:
            return None, 'NTLM auth requires requests-ntlm package.'

        auth = HttpNtlmAuth(auth_creds.get('username', ''), auth_creds.get('password', ''))
        return auth, None

    return None, f'Unsupported auth type: {auth_type}'


def _safe_json_loads(raw_value, default):
    try:
        return json.loads(raw_value) if raw_value else default
    except (json.JSONDecodeError, TypeError, ValueError):
        return default


def _safe_python_name(raw_name, fallback):
    slug = re.sub(r'[^0-9a-zA-Z_]+', '_', (raw_name or '').strip().lower()).strip('_')
    if not slug:
        slug = fallback
    if slug[0].isdigit():
        slug = f'test_{slug}'
    return slug


def _safe_download_name(raw_name, fallback):
    cleaned = re.sub(r'[^0-9A-Za-z._-]+', '_', (raw_name or '').strip()).strip('_')
    return cleaned or fallback


def _resolve_testcase_runtime_context(tc):
    env = Environment.objects.filter(is_active=True).order_by('id').first()
    if not env:
        env = Environment.objects.order_by('id').first()

    module = None
    module_name = (tc.module or '').strip()
    if module_name:
        module = ApiModule.objects.filter(name__iexact=module_name).first()

    endpoint_template = (tc.endpoint or '').strip()
    url = endpoint_template
    if env:
        env_base_path = urlparse(env.base_url or '').path.strip('/')
        module_base_path = urlparse((module.base_path if module else '') or '').path.strip('/')

        context_root_replacement = ''
        if module_base_path and module_base_path != env_base_path:
            context_root_replacement = '/' + module_base_path

        endpoint_path = endpoint_template.replace('{{context-root}}', context_root_replacement).replace('{{ context-root }}', context_root_replacement)
        endpoint_path = endpoint_path.replace('//', '/')
        endpoint_relative = endpoint_path.lstrip('/')

        if env_base_path and (endpoint_relative == env_base_path or endpoint_relative.startswith(env_base_path + '/')):
            endpoint_relative = endpoint_relative[len(env_base_path):].lstrip('/')

        url = env.base_url.rstrip('/')
        if endpoint_relative:
            url += '/' + endpoint_relative

    path_params = _safe_json_loads(tc.path_params, {})
    for key, value in path_params.items():
        url = url.replace(f'{{{key}}}', str(value))

    headers = _safe_json_loads(tc.headers, {})
    params = _safe_json_loads(tc.query_params, {})
    form_data_rows = _safe_json_loads(tc.form_data, [])

    if module and module.module_auth_type == 'oauth2':
        auth_source = 'module'
        auth_type = 'oauth2'
        auth_creds = {
            'add_to': module.oauth2_add_to or 'request_headers',
            'client_id': module.oauth2_client_id or '',
            'client_secret': module.oauth2_client_secret or '',
            'token_url': module.oauth2_token_url or '',
            'header_prefix': module.oauth2_header_prefix or 'Bearer',
            'token': module.oauth2_current_token or '',
        }
    elif tc.auth_type == 'bearer' and env and env.auth_type == 'oauth2':
        env_oauth_creds = _safe_json_loads(env.auth_credentials, {})
        if env_oauth_creds.get('token_url') and env_oauth_creds.get('client_id') and env_oauth_creds.get('client_secret'):
            auth_source = 'environment'
            auth_type = 'oauth2'
            auth_creds = env_oauth_creds
        else:
            auth_source = 'testcase'
            auth_type = tc.auth_type
            auth_creds = _safe_json_loads(tc.auth_credentials, {})
    elif tc.auth_type != 'inherit':
        auth_source = 'testcase'
        auth_type = tc.auth_type
        auth_creds = _safe_json_loads(tc.auth_credentials, {})
    else:
        auth_source = 'environment'
        auth_type = env.auth_type if env else 'none'
        auth_creds = _safe_json_loads(env.auth_credentials if env else '', {})

    headers_for_request = dict(headers)
    params_for_request = dict(params)
    auth, auth_error = _apply_auth_to_request(auth_type, auth_creds, headers_for_request, params_for_request)

    return {
        'environment': env,
        'module': module,
        'url': url,
        'headers': headers_for_request,
        'params': params_for_request,
        'path_params': path_params,
        'form_data_rows': form_data_rows,
        'auth': auth,
        'auth_type': auth_type,
        'auth_source': auth_source,
        'auth_creds': auth_creds,
        'auth_error': auth_error,
    }


def _request_body_snippet_parts(tc):
    form_data_rows = _safe_json_loads(tc.form_data, [])
    file_rows = [row for row in form_data_rows if row.get('type') == 'file' and row.get('key')]
    text_rows = [row for row in form_data_rows if row.get('type') != 'file' and row.get('key')]
    body_lines = []
    request_parts = []
    notes = []
    extra_imports = []
    body_mode = ''

    if text_rows:
        form_payload = {row.get('key'): row.get('value', '') for row in text_rows}
        body_lines.append(f"    form_data = {pformat(form_payload, width=88)}")
        request_parts.append('data=form_data')
        body_mode = 'form'
    if file_rows:
        files_payload = {row.get('key'): "open('path/to/file', 'rb')" for row in file_rows}
        body_lines.append("    files = {")
        for key, value in files_payload.items():
            body_lines.append(f"        {key!r}: {value},")
        body_lines.append("    }")
        request_parts.append('files=files')
        notes.append('Replace file placeholders with real file paths before running.')
        if not body_mode:
            body_mode = 'files'
    if body_lines:
        return body_lines, request_parts, notes, extra_imports, body_mode

    body_text = (tc.request_body or '').strip()
    if not body_text:
        return body_lines, request_parts, notes, extra_imports, body_mode

    parsed_body = _safe_json_loads(body_text, None)
    if isinstance(parsed_body, (dict, list)):
        extra_imports.append('import json')
        body_lines.append("    payload = json.loads(r'''" + body_text + "''')")
        request_parts.append('json=payload')
        body_mode = 'json'
    else:
        body_lines.append(f"    payload = {body_text!r}")
        request_parts.append('data=payload')
        body_mode = 'raw'
    return body_lines, request_parts, notes, extra_imports, body_mode


def _requests_auth_lines(runtime):
    auth_type = (runtime.get('auth_type') or 'none').strip().lower()
    auth_creds = runtime.get('auth_creds') or {}
    imports = []
    lines = []
    request_parts = []
    notes = []

    if auth_type == 'basic':
        lines.append(f"    auth = ({auth_creds.get('username', '')!r}, {auth_creds.get('password', '')!r})")
        request_parts.append('auth=auth')
    elif auth_type == 'digest':
        imports.append('from requests.auth import HTTPDigestAuth')
        lines.append(f"    auth = HTTPDigestAuth({auth_creds.get('username', '')!r}, {auth_creds.get('password', '')!r})")
        request_parts.append('auth=auth')
    elif auth_type == 'oauth1':
        imports.append('from requests_oauthlib import OAuth1')
        lines.append("    auth = OAuth1(")
        lines.append(f"        {auth_creds.get('client_key', '')!r},")
        lines.append(f"        client_secret={auth_creds.get('client_secret', '')!r},")
        lines.append(f"        resource_owner_key={auth_creds.get('resource_owner_key', '')!r},")
        lines.append(f"        resource_owner_secret={auth_creds.get('resource_owner_secret', '')!r},")
        lines.append("    )")
        request_parts.append('auth=auth')
    elif auth_type == 'awsv4':
        imports.append('from requests_aws4auth import AWS4Auth')
        lines.append("    auth = AWS4Auth(")
        lines.append(f"        {auth_creds.get('access_key', '')!r},")
        lines.append(f"        {auth_creds.get('secret_key', '')!r},")
        lines.append(f"        {auth_creds.get('region', '')!r},")
        lines.append(f"        {auth_creds.get('service', '')!r},")
        lines.append(f"        session_token={auth_creds.get('session_token', '')!r},")
        lines.append("    )")
        request_parts.append('auth=auth')
    elif auth_type == 'ntlm':
        imports.append('from requests_ntlm import HttpNtlmAuth')
        lines.append(f"    auth = HttpNtlmAuth({auth_creds.get('username', '')!r}, {auth_creds.get('password', '')!r})")
        request_parts.append('auth=auth')
    elif auth_type in ('oauth2', 'bearer') and not (runtime.get('headers') or {}).get('Authorization') and not (runtime.get('params') or {}).get('access_token'):
        notes.append('Authentication token is not populated yet. Add a token before running.')

    if runtime.get('auth_error'):
        notes.append(runtime['auth_error'])

    return imports, lines, request_parts, notes


def _generated_response_output_helper_lines(base_name):
    return [
        '',
        '',
        'def _save_response_outputs(data):',
        "    output_dir = Path.cwd() / 'Response'",
        '    output_dir.mkdir(exist_ok=True)',
        f"    file_base = output_dir / '{base_name}_response'",
        "    json_path = file_base.with_suffix('.json')",
        "    csv_path = file_base.with_suffix('.csv')",
        '',
        "    with open(json_path, 'w', encoding='utf-8') as handle:",
        '        json.dump(data, handle, indent=2, ensure_ascii=False)',
        '',
        "    with open(csv_path, 'w', newline='', encoding='utf-8') as handle:",
        '        writer = csv.writer(handle)',
        '        if isinstance(data, list):',
        '            if data and isinstance(data[0], dict):',
        '                fieldnames = sorted({key for item in data if isinstance(item, dict) for key in item.keys()})',
        '                writer.writerow(fieldnames)',
        '                for item in data:',
        "                    writer.writerow([json.dumps(item.get(field, ''), ensure_ascii=False) if isinstance(item.get(field, ''), (dict, list)) else item.get(field, '') for field in fieldnames])",
        '            else:',
        "                writer.writerow(['Value'])",
        '                for item in data:',
        '                    writer.writerow([item])',
        '        elif isinstance(data, dict):',
        "            writer.writerow(['Key', 'Value'])",
        '            for key, value in data.items():',
        '                if isinstance(value, (dict, list)):',
        '                    value = json.dumps(value, ensure_ascii=False)',
        '                writer.writerow([key, value])',
        '        else:',
        "            writer.writerow(['Response'])",
        '            writer.writerow([data])',
        '',
        "    print(f'JSON saved to: {json_path}')",
        "    print(f'CSV saved to: {csv_path}')",
    ]


def _build_requests_snippet(tc):
    runtime = _resolve_testcase_runtime_context(tc)
    imports = ['import csv', 'import json', 'from datetime import datetime', 'from pathlib import Path', 'import requests', 'import urllib3']
    auth_imports, auth_lines, auth_request_parts, auth_notes = _requests_auth_lines(runtime)
    imports.extend(auth_imports)
    body_lines, body_request_parts, body_notes, body_imports, _body_mode = _request_body_snippet_parts(tc)
    imports.extend(body_imports)
    imports = list(dict.fromkeys(imports))

    function_name = _safe_python_name(tc.name, f'testcase_{tc.pk}')
    safe_name = _safe_download_name(tc.name, f'testcase_{tc.pk}')
    lines = imports
    lines.extend(_generated_response_output_helper_lines(safe_name))
    lines.extend(['', '', 'urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)', '', '', f'def {function_name}():', f"    url = {runtime['url']!r}"])

    if runtime.get('headers'):
        lines.append(f"    headers = {pformat(runtime['headers'], width=88)}")
    if runtime.get('params'):
        lines.append(f"    params = {pformat(runtime['params'], width=88)}")

    lines.extend(auth_lines)
    lines.extend(body_lines)

    request_parts = [f"method={tc.http_method!r}", 'url=url']
    if runtime.get('headers'):
        request_parts.append('headers=headers')
    if runtime.get('params'):
        request_parts.append('params=params')
    request_parts.extend(auth_request_parts)
    request_parts.extend(body_request_parts)
    request_parts.append('timeout=30')
    request_parts.append('verify=False')

    lines.append(f"    response = requests.request({', '.join(request_parts)})")
    lines.append('')
    lines.append('    try:')
    lines.append('        data = response.json()')
    lines.append('    except ValueError:')
    lines.append('        data = response.text')
    lines.append("    print(f'Status: {response.status_code}')")
    lines.append("    print('Response:')")
    lines.append('    print(data)')
    lines.append('    _save_response_outputs(data)')
    if tc.expected_status_code:
        lines.append(f"    if response.status_code != {tc.expected_status_code}:")
        lines.append(f"        raise AssertionError(f'Expected status {tc.expected_status_code}, got {{response.status_code}}')")
    else:
        lines.append('    if not response.ok:')
        lines.append("        raise AssertionError(f'Request failed with status {response.status_code}')")
    if tc.expected_response_content:
        lines.append(f"    if {tc.expected_response_content!r} not in response.text:")
        lines.append("        raise AssertionError('Expected response content was not found in response body')")

    notes = []
    if not runtime.get('environment'):
        notes.append('No active environment was found; the snippet uses the testcase endpoint as-is.')
    notes.extend(auth_notes)
    notes.extend(body_notes)
    if notes:
        note_block = ['# Notes:'] + [f'# - {note}' for note in notes] + ['']
        lines = note_block + lines
    lines.extend([
        '',
        '',
        "if __name__ == '__main__':",
        f'    {function_name}()',
    ])
    return '\n'.join(lines).rstrip() + '\n'


def _build_playwright_snippet(tc):
    runtime = _resolve_testcase_runtime_context(tc)
    function_name = _safe_python_name(tc.name, f'testcase_{tc.pk}')
    safe_name = _safe_download_name(tc.name, f'testcase_{tc.pk}')
    body_lines, body_request_parts, body_notes, body_imports, body_mode = _request_body_snippet_parts(tc)
    lines = [
        'import csv',
        'import json',
        'import sys',
        'from datetime import datetime',
        'from pathlib import Path',
        'try:',
        '    from playwright.sync_api import sync_playwright',
        'except ModuleNotFoundError as exc:',
        "    if exc.name == 'playwright':",
        "        raise SystemExit('Playwright is not installed. Run: python -m pip install playwright && python -m playwright install chromium')",
        '    raise',
    ]
    for extra_import in body_imports:
        if extra_import not in lines:
            lines.append(extra_import)
    lines.extend(_generated_response_output_helper_lines(safe_name))
    lines.extend([
        '',
        '',
        f'def {function_name}():',
        f"    url = {runtime['url']!r}",
    ])

    if runtime.get('headers'):
        lines.append(f"    headers = {pformat(runtime['headers'], width=88)}")
    else:
        lines.append('    headers = {}')
    if runtime.get('params'):
        lines.append(f"    params = {pformat(runtime['params'], width=88)}")

    lines.extend(body_lines)
    lines.append('')
    lines.append('    with sync_playwright() as p:')
    lines.append('        browser = p.chromium.launch(headless=True)')
    lines.append('        browser_context = browser.new_context(ignore_https_errors=True)')
    lines.append('        page = browser_context.new_page()')
    lines.append('        request_context = p.request.new_context(extra_http_headers=headers, ignore_https_errors=True)')
    fetch_parts = [f"url", f"method={tc.http_method!r}"]
    if runtime.get('params'):
        fetch_parts.append('params=params')
    for part in body_request_parts:
        if part.startswith('json='):
            if body_mode == 'json':
                fetch_parts.append('data=json.dumps(payload)')
            else:
                fetch_parts.append('data=payload')
        else:
            fetch_parts.append(part)
    lines.append(f"        response = request_context.fetch({', '.join(fetch_parts)})")
    lines.append('')
    lines.append('        try:')
    lines.append('            data = response.json()')
    lines.append('        except Exception:')
    lines.append('            data = response.text()')
    lines.append("        print(f'Status: {response.status}')")
    lines.append("        print('Response:')")
    lines.append('        print(data)')
    lines.append('        _save_response_outputs(data)')
    if tc.expected_status_code:
        lines.append(f"        if response.status != {tc.expected_status_code}:")
        lines.append(f"            raise AssertionError(f'Expected status {tc.expected_status_code}, got {{response.status}}')")
    else:
        lines.append('        if not response.ok:')
        lines.append("            raise AssertionError(f'Request failed with status {response.status}')")
    if tc.expected_response_content:
        lines.append(f"        if {tc.expected_response_content!r} not in response.text():")
        lines.append("            raise AssertionError('Expected response content was not found in response body')")
    lines.append('')
    base_ui_url = ''
    if runtime.get('environment') and runtime['environment'].base_url:
        parsed = urlparse(runtime['environment'].base_url)
        base_ui_url = f"{parsed.scheme}://{parsed.netloc}" if parsed.scheme and parsed.netloc else runtime['environment'].base_url
    lines.append(f"        page.goto({(base_ui_url or runtime['url'])!r})")
    lines.append('        # Add UI assertions or navigation steps here if this API call prepares UI state.')
    lines.append('')
    lines.append('        request_context.dispose()')
    lines.append('        browser_context.close()')
    lines.append('        browser.close()')

    notes = []
    if not runtime.get('environment'):
        notes.append('No active environment was found; the snippet uses the testcase endpoint as-is.')
    if runtime.get('auth_error'):
        notes.append(runtime['auth_error'])
    notes.extend(body_notes)
    if (runtime.get('auth_type') or '').strip().lower() in ('basic', 'digest', 'oauth1', 'awsv4', 'ntlm'):
        notes.append('This Playwright snippet forwards resolved headers. If your auth flow requires browser/session-specific setup, extend the snippet before use.')
    if notes:
        note_block = ['# Notes:'] + [f'# - {note}' for note in notes] + ['']
        lines = note_block + lines
    lines.extend([
        '',
        '',
        "if __name__ == '__main__':",
        f'    {function_name}()',
    ])
    return '\n'.join(lines).rstrip() + '\n'


def _build_generated_snippet(tc, snippet_type):
    snippet_type = (snippet_type or '').strip().lower()
    if snippet_type == 'requests':
        return _build_requests_snippet(tc)
    if snippet_type == 'playwright':
        return _build_playwright_snippet(tc)
    raise ValueError(f'Unsupported snippet type: {snippet_type}')


# ========== Authentication ==========

def login_view(request):
    if request.user.is_authenticated:
        return redirect('api:dashboard')
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            user = authenticate(
                request,
                username=form.cleaned_data['username'],
                password=form.cleaned_data['password']
            )
            if user is not None:
                login(request, user)
                log_action(user, 'api:login', request=request)
                return redirect('api:dashboard')
            else:
                messages.error(request, 'Invalid username or password.')
    else:
        form = LoginForm()
    return render(request, 'api/auth/login.html', {'form': form})


@login_required
def logout_view(request):
    log_action(request.user, 'logout', request=request)
    logout(request)
    return redirect('api:login')


# ========== Dashboard ==========

@login_required
def dashboard(request):
    total_test_cases = TestCase.objects.filter(is_active=True).count()
    total_executions = TestExecution.objects.count()
    passed_tests = TestExecution.objects.filter(result_status='passed').count()
    failed_tests = TestExecution.objects.filter(result_status='failed').count()
    error_tests = TestExecution.objects.filter(result_status='error').count()

    recent_executions = TestExecution.objects.select_related(
        'test_case', 'environment', 'executed_by'
    )[:10]

    # Stats for the last 7 days
    seven_days_ago = timezone.now() - timedelta(days=7)
    daily_stats = []
    for i in range(7):
        day = timezone.now() - timedelta(days=6 - i)
        day_start = day.replace(hour=0, minute=0, second=0, microsecond=0)
        day_end = day.replace(hour=23, minute=59, second=59, microsecond=999999)
        day_executions = TestExecution.objects.filter(
            executed_at__range=(day_start, day_end)
        )
        daily_stats.append({
            'date': day.strftime('%m/%d'),
            'passed': day_executions.filter(result_status='passed').count(),
            'failed': day_executions.filter(result_status='failed').count(),
            'error': day_executions.filter(result_status='error').count(),
        })

    context = {
        'total_test_cases': total_test_cases,
        'total_executions': total_executions,
        'passed_tests': passed_tests,
        'failed_tests': failed_tests,
        'error_tests': error_tests,
        'recent_executions': recent_executions,
        'daily_stats': json.dumps(daily_stats),
    }
    return render(request, 'api/dashboard.html', context)


# ========== Environment Management ==========

@login_required
def environment_list(request):
    environments = Environment.objects.all()
    return render(request, 'api/environments/list.html', {'environments': environments})


@login_required
def environment_detail(request, pk):
    env = get_object_or_404(Environment, pk=pk)

    auth_credentials_pretty = ''
    oauth2_details = None
    if env.auth_credentials:
        try:
            creds = json.loads(env.auth_credentials)
            auth_credentials_pretty = json.dumps(creds, indent=2)
            if env.auth_type == 'oauth2':
                oauth2_details = {
                    'add_to': creds.get('add_to', 'request_headers'),
                    'client_id': creds.get('client_id', ''),
                    'token_url': creds.get('token_url', ''),
                    'header_prefix': creds.get('header_prefix', 'Bearer'),
                    'token': creds.get('token', ''),
                }
        except json.JSONDecodeError:
            auth_credentials_pretty = env.auth_credentials

    return render(request, 'api/environments/detail.html', {
        'env': env,
        'auth_credentials_pretty': auth_credentials_pretty,
        'oauth2_details': oauth2_details,
    })


@login_required
@role_required(['admin', 'tester'])
def environment_test_connection(request, pk):
    """Test connectivity for an environment base URL using its auth settings."""
    if request.method not in ['GET', 'POST']:
        return JsonResponse({'error': 'Method not allowed'}, status=400)

    env = get_object_or_404(Environment, pk=pk)
    url = (env.base_url or '').strip()
    if not url:
        return JsonResponse({'ok': False, 'error': 'Environment base URL is empty.'}, status=400)

    headers = {}
    params = {}
    auth = None

    creds = {}
    if env.auth_credentials:
        try:
            creds = json.loads(env.auth_credentials)
        except json.JSONDecodeError:
            creds = {}

    if env.auth_type == 'oauth2':
        token = creds.get('token', '')
        add_to = creds.get('add_to', 'request_headers')
        header_prefix = creds.get('header_prefix', 'Bearer')

        if env.auth_type == 'oauth2' and not token and creds.get('token_url') and creds.get('client_id') and creds.get('client_secret'):
            refreshed_token, error = fetch_oauth2_access_token(
                creds.get('token_url'),
                creds.get('client_id'),
                creds.get('client_secret'),
            )
            if not error and refreshed_token:
                token = refreshed_token
                creds['token'] = refreshed_token
                env.auth_credentials = json.dumps(creds)
                env.save(update_fields=['auth_credentials', 'updated_at'])

        creds['token'] = token
        creds['add_to'] = add_to
        creds['header_prefix'] = header_prefix

    auth, auth_error = _apply_auth_to_request(env.auth_type, creds, headers, params)
    if auth_error:
        return JsonResponse({'ok': False, 'error': auth_error, 'message': 'Auth configuration error.'}, status=200)

    start_time = time.time()
    try:
        response = requests.get(
            url,
            headers=headers,
            params=params,
            auth=auth,
            timeout=20,
            verify=False,
            allow_redirects=True,
        )
        elapsed_ms = int((time.time() - start_time) * 1000)
        ok = 200 <= response.status_code < 400
        reachable = True

        log_action(
            request.user,
            'environment_update',
            'Environment',
            env.id,
            f'Test connection: status={response.status_code}, time={elapsed_ms}ms',
            request,
        )

        return JsonResponse({
            'ok': ok,
            'reachable': reachable,
            'status_code': response.status_code,
            'response_time_ms': elapsed_ms,
            'final_url': response.url,
            'message': 'Connection successful.' if ok else 'Connection reached server but returned non-success status.',
        })
    except requests.exceptions.RequestException as e:
        elapsed_ms = int((time.time() - start_time) * 1000)
        log_action(
            request.user,
            'environment_update',
            'Environment',
            env.id,
            f'Test connection failed: {str(e)}',
            request,
        )
        return JsonResponse({
            'ok': False,
            'reachable': False,
            'response_time_ms': elapsed_ms,
            'error': str(e),
            'message': 'Connection error.',
        }, status=200)


@login_required
@role_required(['admin', 'tester'])
def environment_create(request):
    if request.method == 'POST':
        form = EnvironmentForm(request.POST)
        if form.is_valid():
            env = form.save(commit=False)
            env.created_by = request.user
            env.save()
            log_action(request.user, 'environment_create', 'Environment', env.id,
                       f'Created environment: {env.name}', request)
            messages.success(request, f'Environment "{env.name}" created successfully.')
            return redirect('api:environment_list')
    else:
        form = EnvironmentForm()
    return render(request, 'api/environments/form.html', {'form': form, 'title': 'Create Environment'})


@login_required
@role_required(['admin', 'tester'])
def environment_edit(request, pk):
    env = Environment.objects.filter(pk=pk).first()
    if not env:
        messages.error(request, 'Environment not found. It may have been deleted.')
        return redirect('api:environment_list')
    if request.method == 'POST':
        form = EnvironmentForm(request.POST, instance=env)
        if form.is_valid():
            form.save()
            log_action(request.user, 'environment_update', 'Environment', env.id,
                       f'Updated environment: {env.name}', request)
            messages.success(request, f'Environment "{env.name}" updated successfully.')
            return redirect('api:environment_list')
    else:
        form = EnvironmentForm(instance=env)
    return render(request, 'api/environments/form.html', {'form': form, 'title': 'Edit Environment'})


@login_required
@role_required(['admin', 'tester'])
def environment_get_new_access_token(request, pk):
    """Fetch a fresh OAuth 2.0 token using values from the edit form."""
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Method not allowed'}, status=400)

    env = get_object_or_404(Environment, pk=pk)

    try:
        payload = json.loads(request.body) if request.body else {}
    except json.JSONDecodeError:
        payload = {}

    add_to = (payload.get('add_to') or 'request_headers').strip() or 'request_headers'
    client_id = (payload.get('client_id') or '').strip()
    client_secret = (payload.get('client_secret') or '').strip()
    token_url = (payload.get('token_url') or '').strip()
    header_prefix = (payload.get('header_prefix') or 'Bearer').strip() or 'Bearer'

    token, error = fetch_oauth2_access_token(token_url, client_id, client_secret)
    if error:
        return JsonResponse({'ok': False, 'error': error, 'message': 'Unable to fetch token.'}, status=200)

    auth_credentials = {
        'add_to': add_to,
        'client_id': client_id,
        'client_secret': client_secret,
        'token_url': token_url,
        'header_prefix': header_prefix,
        'token': token,
    }

    log_action(
        request.user,
        'environment_update',
        'Environment',
        env.id,
        'Fetched new OAuth 2.0 access token from Edit Environment page.',
        request,
    )

    return JsonResponse({
        'ok': True,
        'message': 'New access token fetched successfully.',
        'token': token,
        'auth_credentials': auth_credentials,
    })


@login_required
@role_required(['admin'])
def environment_delete(request, pk):
    env = get_object_or_404(Environment, pk=pk)
    if request.method == 'POST':
        name = env.name
        log_action(request.user, 'environment_delete', 'Environment', env.id,
                   f'Deleted environment: {name}', request)
        env.delete()
        messages.success(request, f'Environment "{name}" deleted successfully.')
    return redirect('api:environment_list')


# ========== Test Case Management ==========

@login_required
def testcase_list(request):
    latest_result_subquery = TestExecution.objects.filter(
        test_case_id=OuterRef('pk')
    ).order_by('-executed_at').values('result_status')[:1]

    test_cases = TestCase.objects.filter(is_active=True).annotate(
        latest_result=Subquery(latest_result_subquery)
    )
    module_filter = request.GET.get('module', '')
    project_filter = request.GET.get('project', '')
    method_filter = request.GET.get('method', '')
    search = request.GET.get('search', '')

    if module_filter:
        test_cases = test_cases.filter(module__icontains=module_filter)
    if project_filter:
        test_cases = test_cases.filter(project__icontains=project_filter)
    if method_filter:
        test_cases = test_cases.filter(http_method=method_filter)
    if search:
        test_cases = test_cases.filter(
            Q(name__icontains=search) | Q(endpoint__icontains=search)
        )

    modules = TestCase.objects.values_list('module', flat=True).distinct().order_by('module')
    projects = TestCase.objects.values_list('project', flat=True).distinct().order_by('project')
    project_folders = Project.objects.all()

    # Build hierarchy: folders with test cases inside
    hierarchy = []
    for folder in project_folders:
        folder_cases_qs = test_cases.filter(project=folder.name).order_by('order', 'name')
        if folder_cases_qs.exists():
            folder_cases = list(folder_cases_qs)
            hierarchy.append({
                'folder_id': folder.id,
                'folder_name': folder.name,
                'test_cases': folder_cases,
            })

    # Add ungrouped test cases (no project)
    ungrouped_qs = test_cases.filter(project='').order_by('order', 'name')
    if ungrouped_qs.exists():
        ungrouped = list(ungrouped_qs)
        hierarchy.insert(0, {
            'folder_id': None,
            'folder_name': 'Ungrouped',
            'test_cases': ungrouped,
        })

    context = {
        'hierarchy': hierarchy,
        'modules': [m for m in modules if m],
        'projects': [p for p in projects if p],
        'environments': Environment.objects.filter(is_active=True),
        'project_form': ProjectCreateForm(),
        'filters': {
            'module': module_filter,
            'project': project_filter,
            'method': method_filter,
            'search': search,
        }
    }
    return render(request, 'api/testcases/list.html', context)


@login_required
def global_test_data_page(request):
    """Show global API test data entries stored in the shared recorder data table."""
    search = (request.GET.get('search') or '').strip()
    project_filter = (request.GET.get('project') or '').strip()

    entries = []
    projects = []

    try:
        with connection.cursor() as cur:
            cur.execute(
                """
                SELECT DISTINCT TRIM(COALESCE(folder_name, '')) AS folder_name
                  FROM data
                 WHERE COALESCE(is_global, FALSE) = TRUE
                   AND field_name LIKE %s
                   AND TRIM(COALESCE(folder_name, '')) <> ''
                 ORDER BY folder_name
                """,
                ['api.%'],
            )
            projects = [row[0] for row in cur.fetchall() if row and row[0]]

            where_clauses = [
                'COALESCE(is_global, FALSE) = TRUE',
                'field_name LIKE %s',
            ]
            params = ['api.%']

            if project_filter:
                where_clauses.append("TRIM(COALESCE(folder_name, '')) = %s")
                params.append(project_filter)

            if search:
                where_clauses.append(
                    "(field_name ILIKE %s OR value ILIKE %s OR TRIM(COALESCE(folder_name, '')) ILIKE %s)"
                )
                like_value = f'%{search}%'
                params.extend([like_value, like_value, like_value])

            query = f"""
                SELECT id,
                       TRIM(COALESCE(folder_name, '')) AS folder_name,
                       COALESCE(field_name, '') AS field_name,
                       COALESCE(value, '') AS value,
                       COALESCE(formula, '') AS formula,
                       COALESCE(is_global, FALSE) AS is_global,
                      category,
                      sub_category,
                      increment_value,
                      COALESCE(increment_frequency, '') AS increment_frequency,
                      decrement_value,
                      COALESCE(decrement_frequency, '') AS decrement_frequency,
                      COALESCE(calculate_on::text, '') AS calculate_on,
                      COALESCE(calculate_mode, '') AS calculate_mode,
                       COALESCE(record_id::text, '') AS record_id,
                       created_at
                  FROM data
                 WHERE {' AND '.join(where_clauses)}
                 ORDER BY TRIM(COALESCE(folder_name, '')), field_name, id
                 LIMIT 5000
            """
            cur.execute(query, params)
            for row in cur.fetchall():
                entries.append({
                    'id': row[0],
                    'folder_name': row[1] or '',
                    'field_name': row[2] or '',
                    'value': row[3] or '',
                    'formula': row[4] or '',
                    'is_global': bool(row[5]),
                    'category': row[6] or '',
                    'sub_category': row[7] or '',
                    'increment_value': row[8],
                    'increment_frequency': row[9] or '',
                    'decrement_value': row[10],
                    'decrement_frequency': row[11] or '',
                    'calculate_on': row[12] or '',
                    'calculate_mode': row[13] or '',
                    'record_id': str(row[14]) if row[14] is not None else '',
                    'created_at': row[15],
                })
    except Exception as exc:
        messages.error(request, f'Unable to load global test data: {exc}')

    return render(
        request,
        'api/testcases/global_data.html',
        {
            'entries': entries,
            'projects': projects,
            'filters': {
                'search': search,
                'project': project_filter,
            },
            'total_entries': len(entries),
        },
    )


@login_required
@role_required(['admin', 'tester'])
def testcase_create(request):
    if request.method == 'POST':
        form = TestCaseForm(request.POST)
        if form.is_valid():
            tc = form.save(commit=False)
            tc.created_by = request.user
            tc.save()
            log_action(request.user, 'api:testcase_create', 'TestCase', tc.id,
                       f'Created test case: {tc.name}', request)
            messages.success(request, f'Test case "{tc.name}" created successfully.')
            return redirect('api:testcase_list')
    else:
        initial = {}
        selected_project = request.GET.get('project', '').strip()
        if selected_project:
            initial['project'] = selected_project
        form = TestCaseForm(initial=initial)
    return render(request, 'api/testcases/form.html', {
        'form': form,
        'title': 'Create Test Case',
        'tc_id': None,
        'environments': Environment.objects.filter(is_active=True),
    })


@login_required
@role_required(['admin', 'tester'])
def bulk_duplicate_testcases(request):
    """Duplicate multiple test cases via AJAX."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST allowed'}, status=400)
    
    try:
        data = json.loads(request.body)
        test_case_ids = data.get('test_case_ids', [])
        
        if not test_case_ids:
            return JsonResponse({'error': 'No test case IDs provided'}, status=400)
        
        test_cases = TestCase.objects.filter(pk__in=test_case_ids, is_active=True)
        duplicated = []
        
        for tc in test_cases:
            new_tc = TestCase.objects.create(
                name=f"{tc.name} (copy)",
                module=tc.module,
                endpoint=tc.endpoint,
                http_method=tc.http_method,
                headers=tc.headers,
                query_params=tc.query_params,
                path_params=tc.path_params,
                request_body=tc.request_body,
                auth_type=tc.auth_type,
                auth_credentials=tc.auth_credentials,
                expected_status_code=tc.expected_status_code,
                expected_response_content=tc.expected_response_content,
                expected_response_time_ms=tc.expected_response_time_ms,
                project=tc.project,
                created_by=request.user,
                is_active=tc.is_active,
            )
            duplicated.append({'id': new_tc.id, 'name': new_tc.name})
            log_action(request.user, 'testcase_bulk_duplicate', 'TestCase', tc.id,
                       f'Bulk duplicated to "{new_tc.name}"', request)
        
        return JsonResponse({
            'success': True,
            'duplicated_count': len(duplicated),
            'message': f'Successfully duplicated {len(duplicated)} test case(s)',
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@role_required(['admin', 'tester'])
def project_create(request):
    if request.method != 'POST':
        return redirect('api:testcase_list')

    form = ProjectCreateForm(request.POST)
    if not form.is_valid():
        messages.error(request, form.errors.get('name', ['Unable to create project folder.'])[0])
        return redirect('api:testcase_list')

    project = Project.objects.create(
        name=form.cleaned_data['name'],
        created_by=request.user,
    )
    messages.success(request, f'Project folder "{project.name}" created.')
    return redirect(f"{reverse('api:testcase_create')}?project={project.name}")


@login_required
@role_required(['admin', 'tester'])
def project_edit(request, pk):
    project = get_object_or_404(Project, pk=pk)
    if request.method == 'POST':
        form = ProjectCreateForm(request.POST)
        if not form.is_valid():
            messages.error(request, form.errors.get('name', ['Unable to update project folder.'])[0])
            return redirect('api:testcase_list')
        
        # Check if new name already exists (case-insensitive)
        existing = Project.objects.filter(
            name__iexact=form.cleaned_data['name']
        ).exclude(pk=pk).exists()
        if existing:
            messages.error(request, f'A project folder with this name already exists.')
            return redirect('api:testcase_list')
        
        old_name = project.name
        project.name = form.cleaned_data['name']
        project.save()
        
        # Update all test cases that reference the old project name
        TestCase.objects.filter(project=old_name).update(project=project.name)
        
        log_action(request.user, 'project_update', 'Project', project.id,
                   f'Updated project folder from "{old_name}" to "{project.name}"', request)
        messages.success(request, f'Project folder renamed to "{project.name}".')
    
    return redirect('api:testcase_list')


@login_required
@role_required(['admin', 'tester'])
def project_delete(request, pk):
    project = get_object_or_404(Project, pk=pk)
    if request.method == 'POST':
        project_name = project.name
        test_case_count = TestCase.objects.filter(project=project_name).count()

        if test_case_count > 0:
            # Keep test cases by moving them to Ungrouped before deleting the folder.
            TestCase.objects.filter(project=project_name).update(project='')
        
        project.delete()
        log_action(request.user, 'project_delete', 'Project', pk,
                   f'Deleted project folder: {project_name}', request)
        if test_case_count > 0:
            messages.success(
                request,
                f'Project folder "{project_name}" deleted. {test_case_count} test case(s) moved to Ungrouped.'
            )
        else:
            messages.success(request, f'Project folder "{project_name}" deleted.')
    
    return redirect('api:testcase_list')


@login_required
@role_required(['admin', 'tester'])
def ungrouped_delete(request):
    """Delete all test cases in the Ungrouped folder (project='')."""
    if request.method == 'POST':
        count = TestCase.objects.filter(project='', is_active=True).count()
        TestCase.objects.filter(project='').update(is_active=False)
        log_action(request.user, 'ungrouped_delete', 'TestCase', None,
                   f'Deleted {count} ungrouped test case(s)', request)
        messages.success(request, f'Ungrouped folder deleted. {count} test case(s) removed.')
    return redirect('api:testcase_list')


@login_required
@role_required(['admin', 'tester'])
def project_duplicate(request, pk):
    """Duplicate a project folder with all its test cases."""
    project = get_object_or_404(Project, pk=pk)
    if request.method == 'POST':
        original_name = project.name
        new_name = f"{original_name} (copy)"
        
        # Check if name already exists
        counter = 1
        while Project.objects.filter(name__iexact=new_name).exists():
            counter += 1
            new_name = f"{original_name} (copy {counter})"
        
        # Create new folder
        new_project = Project.objects.create(
            name=new_name,
            created_by=request.user,
        )
        
        # Duplicate all test cases
        original_cases = TestCase.objects.filter(project=original_name)
        for tc in original_cases:
            TestCase.objects.create(
                name=f"{tc.name} (copy)",
                module=tc.module,
                endpoint=tc.endpoint,
                http_method=tc.http_method,
                headers=tc.headers,
                query_params=tc.query_params,
                path_params=tc.path_params,
                request_body=tc.request_body,
                auth_type=tc.auth_type,
                auth_credentials=tc.auth_credentials,
                expected_status_code=tc.expected_status_code,
                expected_response_content=tc.expected_response_content,
                expected_response_time_ms=tc.expected_response_time_ms,
                project=new_name,
                created_by=request.user,
                is_active=tc.is_active,
            )
        
        log_action(request.user, 'project_duplicate', 'Project', pk,
                   f'Duplicated project folder "{original_name}" to "{new_name}"', request)
        messages.success(request, f'Project folder duplicated as "{new_name}".')
    
    return redirect('api:testcase_list')


@login_required
@role_required(['admin', 'tester'])
def testcase_edit(request, pk):
    tc = get_object_or_404(TestCase, pk=pk)
    if request.method == 'POST':
        form = TestCaseForm(request.POST, instance=tc)
        if form.is_valid():
            form.save()
            log_action(request.user, 'testcase_update', 'TestCase', tc.id,
                       f'Updated test case: {tc.name}', request)
            messages.success(request, f'Test case "{tc.name}" updated successfully.')
            return redirect('api:testcase_list')
    else:
        form = TestCaseForm(instance=tc)
    return render(request, 'api/testcases/form.html', {
        'form': form,
        'title': 'Edit Test Case',
        'tc_id': tc.pk,
        'environments': Environment.objects.filter(is_active=True),
    })


@login_required
def api_global_field_values(request):
    """Return global data field names and their latest values for Form Data autocomplete."""
    project = (request.GET.get('project') or '').strip()
    rows = []

    try:
        with connection.cursor() as cur:
            where = ["COALESCE(is_global, FALSE) = TRUE", "field_name IS NOT NULL", "field_name <> ''"]
            params = []

            if project:
                where.append("(TRIM(COALESCE(folder_name, '')) = %s OR TRIM(COALESCE(folder_name, '')) LIKE %s)")
                params.extend([project, f"{project}/%"])

            cur.execute(
                f"""
                SELECT field_name, COALESCE(value, '') AS value
                  FROM data
                 WHERE {' AND '.join(where)}
                 ORDER BY id DESC
                """,
                params,
            )

            seen = set()
            for field_name, value in cur.fetchall():
                key = (field_name or '').strip()
                if not key or key in seen:
                    continue
                seen.add(key)
                rows.append({'field_name': key, 'value': value or ''})
    except Exception:
        return JsonResponse({'ok': False, 'rows': []}, status=500)

    return JsonResponse({'ok': True, 'rows': rows})


@login_required
@role_required(['admin', 'tester'])
def testcase_duplicate(request, pk):
    tc = get_object_or_404(TestCase, pk=pk)
    tc.pk = None
    tc.name = f"{tc.name} (Copy)"
    tc.created_by = request.user
    tc.save()
    log_action(request.user, 'api:testcase_create', 'TestCase', tc.id,
               f'Duplicated test case: {tc.name}', request)
    messages.success(request, f'Test case duplicated as "{tc.name}".')
    return redirect('api:testcase_edit', pk=tc.pk)


@login_required
@role_required(['admin', 'tester'])
def testcase_delete(request, pk):
    tc = get_object_or_404(TestCase, pk=pk)
    if request.method == 'POST':
        name = tc.name
        log_action(request.user, 'testcase_delete', 'TestCase', tc.id,
                   f'Deleted test case: {name}', request)
        tc.delete()
        messages.success(request, f'Test case "{name}" deleted successfully.')
    return redirect('api:testcase_list')


@login_required
@role_required(['admin', 'tester'])
def testcase_duplicate(request, pk):
    """Duplicate a test case."""
    tc = get_object_or_404(TestCase, pk=pk)
    if request.method == 'POST':
        new_tc = TestCase.objects.create(
            name=f"{tc.name} (copy)",
            module=tc.module,
            endpoint=tc.endpoint,
            http_method=tc.http_method,
            headers=tc.headers,
            query_params=tc.query_params,
            path_params=tc.path_params,
            request_body=tc.request_body,
            auth_type=tc.auth_type,
            auth_credentials=tc.auth_credentials,
            expected_status_code=tc.expected_status_code,
            expected_response_content=tc.expected_response_content,
            expected_response_time_ms=tc.expected_response_time_ms,
            project=tc.project,
            created_by=request.user,
            is_active=tc.is_active,
        )
        log_action(request.user, 'testcase_duplicate', 'TestCase', tc.id,
                   f'Duplicated test case "{tc.name}" to "{new_tc.name}"', request)
        messages.success(request, f'Test case "{tc.name}" duplicated successfully.')
    return redirect('api:testcase_list')


@login_required
@role_required(['admin', 'tester'])
def bulk_execute(request):
    """Execute multiple test cases in serial order matching folder/file sequence."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST allowed'}, status=400)

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    test_case_ids = data.get('test_case_ids', [])
    execution_mode = data.get('mode', 'serial')
    project_folder_name = (data.get('project_folder_name') or data.get('folder_name') or '').strip()

    if not test_case_ids:
        return JsonResponse({'error': 'No test cases selected'}, status=400)

    # Auto-select environment (same logic as execute_test)
    environment = Environment.objects.filter(is_active=True).first()
    if not environment:
        return JsonResponse({'error': 'No active environment found'}, status=400)

    # Fetch test cases and preserve the incoming ID order (folder/file sequence)
    tc_map = {tc.pk: tc for tc in TestCase.objects.filter(pk__in=test_case_ids, is_active=True)}
    ordered_cases = [tc_map[pk] for pk in test_case_ids if pk in tc_map]

    execution_results = []
    for tc in ordered_cases:
        result = _execute_tc(tc, environment, request.user, request, project_folder_name or None)
        execution_results.append({
            'test_case_id': tc.id,
            'test_case_name': tc.name,
            'status': result['result_status'],
            'response_code': result.get('response_status_code'),
            'response_time_ms': result.get('response_time_ms'),
            'error': result.get('error_message') or '',
        })

    passed = sum(1 for r in execution_results if r['status'] == 'passed')
    failed = sum(1 for r in execution_results if r['status'] == 'failed')
    errors = sum(1 for r in execution_results if r['status'] == 'error')

    return JsonResponse({
        'success': True,
        'results': execution_results,
        'summary': {
            'total': len(execution_results),
            'passed': passed,
            'failed': failed,
            'error': errors,
        }
    })


def _execute_tc(tc, env, user, request, target_project_folder_name=None):
    """Full execution of a single TestCase with auth, URL normalisation, and retry logic.
    Returns a result dict (not a JsonResponse) so both execute_test and bulk_execute can use it."""

    module = None
    module_name = (tc.module or '').strip()
    if module_name:
        module = ApiModule.objects.filter(name__iexact=module_name).first()

    # Build normalised URL
    endpoint_template = (tc.endpoint or '').strip()
    env_base_path = urlparse(env.base_url or '').path.strip('/')
    module_base_path = urlparse((module.base_path if module else '') or '').path.strip('/')

    context_root_replacement = ''
    if module_base_path and module_base_path != env_base_path:
        context_root_replacement = '/' + module_base_path

    endpoint_path = endpoint_template.replace('{{context-root}}', context_root_replacement).replace('{{ context-root }}', context_root_replacement)
    endpoint_path = endpoint_path.replace('//', '/')
    endpoint_relative = endpoint_path.lstrip('/')

    if env_base_path and (endpoint_relative == env_base_path or endpoint_relative.startswith(env_base_path + '/')):
        endpoint_relative = endpoint_relative[len(env_base_path):].lstrip('/')

    url = env.base_url.rstrip('/')
    if endpoint_relative:
        url += '/' + endpoint_relative

    # Replace path parameters
    try:
        path_params = json.loads(tc.path_params) if tc.path_params else {}
        for key, value in path_params.items():
            url = url.replace(f'{{{key}}}', str(value))
    except json.JSONDecodeError:
        pass

    # Parse headers and query params
    try:
        headers = json.loads(tc.headers) if tc.headers else {}
    except json.JSONDecodeError:
        headers = {}

    try:
        params = json.loads(tc.query_params) if tc.query_params else {}
    except json.JSONDecodeError:
        params = {}

    # Determine auth source
    if module and module.module_auth_type == 'oauth2':
        auth_source = 'module'
        auth_type = 'oauth2'
        auth_creds = {
            'add_to': module.oauth2_add_to or 'request_headers',
            'client_id': module.oauth2_client_id or '',
            'client_secret': module.oauth2_client_secret or '',
            'token_url': module.oauth2_token_url or '',
            'header_prefix': module.oauth2_header_prefix or 'Bearer',
            'token': module.oauth2_current_token or '',
        }
        auth_creds_raw = json.dumps(auth_creds)
    elif tc.auth_type == 'bearer':
        env_oauth_creds = {}
        if env.auth_type == 'oauth2' and env.auth_credentials:
            try:
                env_oauth_creds = json.loads(env.auth_credentials)
            except json.JSONDecodeError:
                pass
        if env_oauth_creds.get('token_url') and env_oauth_creds.get('client_id') and env_oauth_creds.get('client_secret'):
            auth_source = 'environment'
            auth_type = 'oauth2'
            auth_creds_raw = env.auth_credentials
        else:
            auth_source = 'testcase'
            auth_type = tc.auth_type
            auth_creds_raw = tc.auth_credentials
    elif tc.auth_type != 'inherit':
        auth_source = 'testcase'
        auth_type = tc.auth_type
        auth_creds_raw = tc.auth_credentials
    else:
        auth_source = 'environment'
        auth_type = env.auth_type
        auth_creds_raw = env.auth_credentials

    try:
        auth_creds = json.loads(auth_creds_raw) if auth_creds_raw else {}
    except json.JSONDecodeError:
        auth_creds = {}

    # Apply auth to request
    if auth_type == 'oauth2':
        token = auth_creds.get('token', '')
        add_to = auth_creds.get('add_to', 'request_headers')
        header_prefix = auth_creds.get('header_prefix', 'Bearer')
        if not token and auth_creds.get('token_url') and auth_creds.get('client_id') and auth_creds.get('client_secret'):
            refreshed_token, error = fetch_oauth2_access_token(auth_creds['token_url'], auth_creds['client_id'], auth_creds['client_secret'])
            if not error and refreshed_token:
                token = refreshed_token
                auth_creds['token'] = refreshed_token
                if auth_source == 'module' and module:
                    module.oauth2_current_token = refreshed_token
                    module.oauth2_token_updated_at = timezone.now()
                    module.save(update_fields=['oauth2_current_token', 'oauth2_token_updated_at', 'updated_at'])
                elif auth_source == 'environment':
                    env.auth_credentials = json.dumps(auth_creds)
                    env.save(update_fields=['auth_credentials', 'updated_at'])
                else:
                    tc.auth_credentials = json.dumps(auth_creds)
                    tc.save(update_fields=['auth_credentials', 'updated_at'])
        auth_creds['token'] = token
        auth_creds['add_to'] = add_to
        auth_creds['header_prefix'] = header_prefix

    auth_error = ''
    auth, auth_apply_error = _apply_auth_to_request(auth_type, auth_creds, headers, params)
    if auth_apply_error:
        auth_error = auth_apply_error

    # Build request body / form data
    body = tc.request_body if tc.request_body else None
    form_data_rows = []
    if tc.form_data:
        try:
            form_data_rows = json.loads(tc.form_data)
        except json.JSONDecodeError:
            pass

    req_kwargs = dict(method=tc.http_method, url=url, headers=headers, params=params, auth=auth, timeout=30, verify=False)
    # Form Data rows in API TestLab are used for response contains validation only.
    # Keep sending the configured request body so JSON endpoints are not broken by multipart uploads.
    req_kwargs['data'] = body

    # Execute (with 401-retry token refresh)
    error_message = ''
    response_status = None
    response_headers_str = ''
    response_body = ''
    response_time = None

    try:
        if auth_error:
            raise requests.exceptions.RequestException(auth_error)

        start = time.time()
        resp = requests.request(**req_kwargs)
        response_time = int((time.time() - start) * 1000)
        response_status = resp.status_code
        response_headers_str = json.dumps(dict(resp.headers))
        response_body = resp.text[:50000]

        if response_status == 401 and auth_type in ('oauth2', 'bearer'):
            refresh_creds = None
            refresh_source = None
            if auth_type == 'oauth2' and auth_creds.get('token_url') and auth_creds.get('client_id') and auth_creds.get('client_secret'):
                refresh_source = auth_source
                refresh_creds = {'token_url': auth_creds['token_url'], 'client_id': auth_creds['client_id'], 'client_secret': auth_creds['client_secret'], 'add_to': auth_creds.get('add_to', 'request_headers'), 'header_prefix': auth_creds.get('header_prefix', 'Bearer')}
            elif module and module.module_auth_type == 'oauth2' and module.oauth2_token_url and module.oauth2_client_id:
                refresh_source = 'module'
                refresh_creds = {'token_url': module.oauth2_token_url, 'client_id': module.oauth2_client_id, 'client_secret': module.oauth2_client_secret, 'add_to': module.oauth2_add_to or 'request_headers', 'header_prefix': module.oauth2_header_prefix or 'Bearer'}

            if refresh_creds:
                refreshed_token, refresh_error = fetch_oauth2_access_token(refresh_creds['token_url'], refresh_creds['client_id'], refresh_creds['client_secret'])
                if not refresh_error and refreshed_token:
                    if refresh_source == 'module' and module:
                        module.oauth2_current_token = refreshed_token
                        module.oauth2_token_updated_at = timezone.now()
                        module.save(update_fields=['oauth2_current_token', 'oauth2_token_updated_at', 'updated_at'])
                    add_to = refresh_creds.get('add_to', 'request_headers')
                    header_prefix = refresh_creds.get('header_prefix', 'Bearer')
                    if add_to == 'request_url':
                        params['access_token'] = refreshed_token
                    else:
                        headers['Authorization'] = f"{header_prefix} {refreshed_token}".strip()
                    req_kwargs['headers'] = headers
                    req_kwargs['params'] = params
                    retry_start = time.time()
                    resp = requests.request(**req_kwargs)
                    response_time = int((time.time() - retry_start) * 1000)
                    response_status = resp.status_code
                    response_headers_str = json.dumps(dict(resp.headers))
                    response_body = resp.text[:50000]

    except requests.exceptions.Timeout:
        error_message = 'Request timed out after 30 seconds'
    except requests.exceptions.ConnectionError as e:
        error_message = f'Connection error: {str(e)}'
    except requests.exceptions.RequestException as e:
        error_message = f'Request error: {str(e)}'

    # Validate results
    status_code_match = content_match = time_within_threshold = None
    if not error_message:
        if tc.expected_status_code:
            status_code_match = (response_status == tc.expected_status_code)
        expected_content_match = None
        form_data_contains_match = None

        if tc.expected_response_content:
            expected_content_match = _match_expected_content(tc.expected_response_content, response_body)

        form_data_contains_match = _match_form_data_contains(form_data_rows, response_body)

        if expected_content_match is None:
            content_match = form_data_contains_match
        elif form_data_contains_match is None:
            content_match = expected_content_match
        else:
            content_match = bool(expected_content_match and form_data_contains_match)

        if tc.expected_response_time_ms and response_time is not None:
            time_within_threshold = (response_time <= tc.expected_response_time_ms)
        checks = [v for v in [status_code_match, content_match, time_within_threshold] if v is not None]
        result_status = ('passed' if all(checks) else 'failed') if checks else 'passed'
    else:
        result_status = 'error'

    execution = TestExecution.objects.create(
        test_case=tc,
        environment=env,
        executed_by=user,
        result_status=result_status,
        request_url=url,
        request_headers=json.dumps(headers),
        request_body=body or '',
        response_status_code=response_status,
        response_headers=response_headers_str,
        response_body=response_body,
        response_time_ms=response_time,
        status_code_match=status_code_match,
        content_match=content_match,
        time_within_threshold=time_within_threshold,
        error_message=error_message,
    )

    # Save response files to Response folder
    save_response_files(tc.name, response_body)

    # Always refresh shared project data when we received a response payload,
    # even if validation checks failed (latest response should overwrite prior data).
    if not error_message and response_body:
        try:
            _store_response_in_project_data(tc, response_body, target_project_folder_name)
        except Exception:
            # Keep execution flow resilient if shared data persistence fails.
            pass

    log_action(user, 'testcase_execute', 'TestCase', tc.id,
               f'Executed test case: {tc.name} - Result: {result_status}', request)

    return {
        'execution_id': execution.id,
        'environment_id': env.id,
        'environment_name': env.name,
        'auth_source': auth_source,
        'result_status': result_status,
        'request_url': url,
        'response_status_code': response_status,
        'response_headers': response_headers_str,
        'response_body': response_body,
        'response_time_ms': response_time,
        'status_code_match': status_code_match,
        'content_match': content_match,
        'time_within_threshold': time_within_threshold,
        'error_message': error_message,
        'project_data_folder': (target_project_folder_name or tc.project or ''),
    }


@login_required
def testcase_detail(request, pk):
    tc = get_object_or_404(TestCase, pk=pk)
    executions = tc.executions.all()[:20]
    environments = Environment.objects.filter(is_active=True)
    snippet_runtime = _resolve_testcase_runtime_context(tc)
    expected_key_pairs = _parse_form_data_rows(tc.form_data)
    selected_tab = (request.GET.get('tab') or 'details').strip().lower()
    selected_code_mode = (request.GET.get('code_mode') or '').strip().lower()
    if selected_tab not in {'details', 'executions', 'requests', 'playwright'}:
        selected_tab = 'details'
    if selected_tab in {'requests', 'playwright'}:
        code_view_mode = selected_tab
    elif selected_tab == 'details' and selected_code_mode in {'requests', 'playwright'}:
        code_view_mode = selected_code_mode
    else:
        code_view_mode = ''
    context = {
        'test_case': tc,
        'executions': executions,
        'environments': environments,
        'expected_key_pairs': expected_key_pairs,
        'snippet_environment': snippet_runtime.get('environment'),
        'requests_snippet': _build_requests_snippet(tc),
        'playwright_snippet': _build_playwright_snippet(tc),
        'selected_tab': selected_tab,
        'code_view_mode': code_view_mode,
    }
    return render(request, 'api/testcases/detail.html', context)


@login_required
def testcase_download_requests_py(request, pk):
    tc = get_object_or_404(TestCase, pk=pk)
    response = HttpResponse(_build_requests_snippet(tc), content_type='text/x-python; charset=utf-8')
    safe_name = _safe_download_name(tc.name, f'testcase_{tc.pk}')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_requests.py"'
    return response


@login_required
def testcase_download_playwright_py(request, pk):
    tc = get_object_or_404(TestCase, pk=pk)
    response = HttpResponse(_build_playwright_snippet(tc), content_type='text/x-python; charset=utf-8')
    safe_name = _safe_download_name(tc.name, f'testcase_{tc.pk}')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_playwright.py"'
    return response


@login_required
@role_required(['admin', 'tester'])
def testcase_run_generated_py(request, pk, snippet_type):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    tc = get_object_or_404(TestCase, pk=pk)
    try:
        script_text = _build_generated_snippet(tc, snippet_type)
    except ValueError as exc:
        return JsonResponse({'error': str(exc)}, status=400)

    safe_name = _safe_download_name(tc.name, f'testcase_{tc.pk}')
    file_suffix = '_playwright.py' if snippet_type == 'playwright' else '_requests.py'
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile('w', suffix=file_suffix, prefix=f'{safe_name}_', delete=False, encoding='utf-8') as handle:
            handle.write(script_text)
            temp_path = handle.name

        result = subprocess.run(
            [sys.executable, temp_path],
            capture_output=True,
            text=True,
            cwd=str(Path(__file__).resolve().parent.parent),
            timeout=180,
        )
        return JsonResponse({
            'ok': result.returncode == 0,
            'returncode': result.returncode,
            'stdout': result.stdout,
            'stderr': result.stderr,
            'script_path': temp_path,
            'snippet_type': snippet_type,
        })
    except subprocess.TimeoutExpired as exc:
        return JsonResponse({
            'ok': False,
            'returncode': None,
            'stdout': exc.stdout or '',
            'stderr': exc.stderr or '',
            'error': 'Snippet execution timed out after 180 seconds.',
            'snippet_type': snippet_type,
        }, status=500)
    except Exception as exc:
        return JsonResponse({
            'ok': False,
            'returncode': None,
            'stdout': '',
            'stderr': '',
            'error': str(exc),
            'snippet_type': snippet_type,
        }, status=500)
    finally:
        if temp_path:
            try:
                os.remove(temp_path)
            except OSError:
                pass


@login_required
def testcase_recent_executions(request, pk):
    """AJAX endpoint: latest executions for a single test case."""
    tc = get_object_or_404(TestCase, pk=pk)
    executions = TestExecution.objects.select_related('environment', 'executed_by').filter(test_case=tc)[:20]

    rows = []
    for ex in executions:
        rows.append({
            'id': ex.id,
            'environment': ex.environment.name if ex.environment else '-',
            'result_status': ex.result_status,
            'response_status_code': ex.response_status_code,
            'response_time_ms': ex.response_time_ms,
            'executed_by': ex.executed_by.username if ex.executed_by else '-',
            'executed_at': timezone.localtime(ex.executed_at).strftime('%b %d, %H:%M') if ex.executed_at else '-',
            'detail_url': reverse('api:execution_detail', args=[ex.id]),
        })

    return JsonResponse({'rows': rows})


@login_required
def testcase_transformed_response(request, pk):
    """Return transformed response JSON (flattened key/value pairs) for a test case."""
    tc = get_object_or_404(TestCase, pk=pk)

    execution_id_raw = (request.GET.get('execution_id') or '').strip()
    execution_qs = TestExecution.objects.filter(test_case=tc)
    if execution_id_raw:
        try:
            execution_qs = execution_qs.filter(pk=int(execution_id_raw))
        except ValueError:
            return JsonResponse({'ok': False, 'error': 'execution_id must be numeric.'}, status=400)

    execution = execution_qs.order_by('-executed_at').first()
    if not execution:
        return JsonResponse({'ok': False, 'error': 'No execution found for this test case.'}, status=404)

    if not execution.response_body:
        return JsonResponse({'ok': False, 'error': 'Execution has no response body.', 'execution_id': execution.id}, status=404)

    try:
        parsed = json.loads(execution.response_body)
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({
            'ok': False,
            'error': 'Latest response body is not valid JSON.',
            'execution_id': execution.id,
        }, status=400)

    field_prefix = f"api.{_safe_data_prefix(tc.name)}"
    flattened = _flatten_response_json(parsed, field_prefix)
    if not flattened:
        return JsonResponse({
            'ok': True,
            'execution_id': execution.id,
            'test_case_id': tc.id,
            'test_case_name': tc.name,
            'project_folder': tc.project or '',
            'rows': [],
            'transformed_json': {},
            'transformed_json_pretty': '{}',
        })

    # Cap preview payload for UI responsiveness.
    max_rows = 1000
    preview_rows = flattened[:max_rows]
    transformed_map = {field_name: value for field_name, value in preview_rows}
    rows = [
        {
            'field_name': field_name,
            'value': value,
        }
        for field_name, value in preview_rows
    ]

    return JsonResponse({
        'ok': True,
        'execution_id': execution.id,
        'test_case_id': tc.id,
        'test_case_name': tc.name,
        'project_folder': tc.project or '',
        'truncated': len(flattened) > max_rows,
        'total_rows': len(flattened),
        'rows': rows,
        'transformed_json': transformed_map,
        'transformed_json_pretty': json.dumps(transformed_map, indent=2, ensure_ascii=False),
    })


# ========== API Execution ==========

@login_required
@role_required(['admin', 'tester'])
def execute_test(request):
    """Execute an API test case via AJAX."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    test_case_id = data.get('test_case_id')
    environment_id = data.get('environment_id')
    project_folder_name = (data.get('project_folder_name') or data.get('folder_name') or '').strip()

    if not test_case_id:
        return JsonResponse({'error': 'test_case_id is required'}, status=400)

    tc = get_object_or_404(TestCase, pk=test_case_id)
    if environment_id:
        env = Environment.objects.filter(pk=environment_id).first()
        if not env:
            return JsonResponse({'error': 'Selected environment was not found.'}, status=404)
    else:
        env = Environment.objects.filter(is_active=True).order_by('id').first()
        if not env:
            # Fallback for migrated workspaces where environments exist but none are marked active.
            env = Environment.objects.order_by('id').first()
        if not env:
            return JsonResponse({'error': 'No environment is configured for execution.'}, status=400)

    module = None
    module_name = (tc.module or '').strip()
    if module_name:
        module = ApiModule.objects.filter(name__iexact=module_name).first()

    result = _execute_tc(tc, env, request.user, request, project_folder_name or None)
    return JsonResponse(result)


@login_required
def testcase_latest_results(request):
    """AJAX endpoint: latest result status for a set of test case IDs."""
    ids_raw = (request.GET.get('ids') or '').strip()
    if not ids_raw:
        return JsonResponse({'results': {}})

    ids = []
    for part in ids_raw.split(','):
        part = part.strip()
        if not part:
            continue
        try:
            ids.append(int(part))
        except ValueError:
            continue

    if not ids:
        return JsonResponse({'results': {}})

    executions = TestExecution.objects.filter(test_case_id__in=ids).order_by('test_case_id', '-executed_at')
    latest = {}
    for ex in executions:
        if ex.test_case_id not in latest:
            latest[ex.test_case_id] = ex.result_status

    return JsonResponse({'results': latest})


# ========== Execution History ==========

@login_required
def execution_history(request):
    executions = TestExecution.objects.select_related('test_case', 'environment', 'executed_by')

    # Filters
    status_filter = request.GET.get('status', '')
    env_filter = request.GET.get('environment', '')
    user_filter = request.GET.get('user', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')

    if status_filter:
        executions = executions.filter(result_status=status_filter)
    if env_filter:
        executions = executions.filter(environment_id=env_filter)
    if user_filter:
        executions = executions.filter(executed_by_id=user_filter)
    if date_from:
        executions = executions.filter(executed_at__date__gte=date_from)
    if date_to:
        executions = executions.filter(executed_at__date__lte=date_to)

    environments = Environment.objects.all()
    users = User.objects.filter(is_active=True)

    context = {
        'executions': executions[:100],
        'environments': environments,
        'users': users,
        'filters': {
            'status': status_filter,
            'environment': env_filter,
            'user': user_filter,
            'date_from': date_from,
            'date_to': date_to,
        }
    }
    return render(request, 'api/executions/history.html', context)


@login_required
def execution_detail(request, pk):
    execution = get_object_or_404(
        TestExecution.objects.select_related('test_case', 'environment', 'executed_by'),
        pk=pk
    )
    try:
        parsed_body = json.loads(execution.response_body) if execution.response_body else None
    except (json.JSONDecodeError, TypeError):
        parsed_body = execution.response_body or ''

    if isinstance(parsed_body, (dict, list)):
        response_json_view = json.dumps(parsed_body, indent=2, ensure_ascii=False)
    else:
        response_json_view = parsed_body if parsed_body else ''

    return render(request, 'api/executions/detail.html', {
        'execution': execution,
        'response_json_view': response_json_view,
    })


@login_required
def execution_download_csv(request, pk):
    """Download execution detail as CSV."""
    execution = get_object_or_404(
        TestExecution.objects.select_related('test_case', 'environment', 'executed_by'),
        pk=pk
    )
    safe_name = _safe_download_name(execution.test_case.name, f'testcase_{execution.test_case_id}')
    mismatch_summary = _expected_content_mismatch_summary(
        execution.test_case.expected_response_content,
        execution.response_body,
    )
    expected_key_pairs_rows = _parse_form_data_rows(execution.test_case.form_data)
    expected_key_pairs_text = _format_expected_key_pairs(expected_key_pairs_rows)
    key_pair_mismatch_summary = _form_data_mismatch_summary(
        expected_key_pairs_rows,
        execution.response_body,
    )
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.csv"'
    writer = csv.writer(response)

    writer.writerow(['Field', 'Value'])
    writer.writerow(['Test Case', execution.test_case.name])
    writer.writerow(['HTTP Method', execution.test_case.http_method])
    writer.writerow(['Environment', execution.environment.name if execution.environment else ''])
    writer.writerow(['Executed By', execution.executed_by.username if execution.executed_by else ''])
    writer.writerow(['Executed At', execution.executed_at.strftime('%Y-%m-%d %H:%M:%S')])
    writer.writerow(['Result', execution.result_status.upper()])
    writer.writerow(['Request URL', execution.request_url])
    writer.writerow(['Response Status Code', execution.response_status_code or ''])
    writer.writerow(['Response Time (ms)', execution.response_time_ms or ''])
    writer.writerow(['Expected Response Content', execution.test_case.expected_response_content or ''])
    writer.writerow(['Expected key-pair value', expected_key_pairs_text])
    writer.writerow(['Status Code Match', 'Pass' if execution.status_code_match else ('Fail' if execution.status_code_match is False else 'N/A')])
    writer.writerow(['Content Match', 'Pass' if execution.content_match else ('Fail' if execution.content_match is False else 'N/A')])
    if mismatch_summary:
        writer.writerow(['Mismatch (Expected Response Content)', mismatch_summary])
    if key_pair_mismatch_summary:
        writer.writerow(['Mismatch (Expected key-pair value)', key_pair_mismatch_summary])
    writer.writerow(['Time Threshold', 'Pass' if execution.time_within_threshold else ('Fail' if execution.time_within_threshold is False else 'N/A')])
    writer.writerow(['Error Message', execution.error_message or ''])
    writer.writerow([])
    writer.writerow(['Request Headers', execution.request_headers])
    writer.writerow(['Request Body', execution.request_body])
    writer.writerow(['Response Headers', execution.response_headers])
    writer.writerow(['Response Body', execution.response_body])
    return response


@login_required
def execution_download_docx(request, pk):
    """Download execution detail as Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    execution = get_object_or_404(
        TestExecution.objects.select_related('test_case', 'environment', 'executed_by'),
        pk=pk
    )
    mismatch_summary = _expected_content_mismatch_summary(
        execution.test_case.expected_response_content,
        execution.response_body,
    )
    expected_key_pairs_rows = _parse_form_data_rows(execution.test_case.form_data)
    expected_key_pairs_text = _format_expected_key_pairs(expected_key_pairs_rows)
    key_pair_mismatch_summary = _form_data_mismatch_summary(
        expected_key_pairs_rows,
        execution.response_body,
    )

    doc = Document()
    doc.core_properties.title = f'Execution Report - {execution.test_case.name}'

    heading = doc.add_heading('Execution Report', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Summary', level=1)
    summary = [
        ('Test Case', execution.test_case.name),
        ('HTTP Method', execution.test_case.http_method),
        ('Environment', execution.environment.name if execution.environment else '-'),
        ('Executed By', execution.executed_by.username if execution.executed_by else '-'),
        ('Executed At', execution.executed_at.strftime('%Y-%m-%d %H:%M:%S')),
        ('Result', execution.result_status.upper()),
        ('Request URL', execution.request_url),
        ('Response Status Code', str(execution.response_status_code or '-')),
        ('Response Time (ms)', str(execution.response_time_ms or '-')),
        ('Expected Response Content', execution.test_case.expected_response_content or '-'),
        ('Expected key-pair value', expected_key_pairs_text),
    ]
    table = doc.add_table(rows=len(summary), cols=2)
    table.style = 'Table Grid'
    for i, (field, value) in enumerate(summary):
        table.cell(i, 0).text = field
        table.cell(i, 1).text = str(value)

    doc.add_heading('Validation Results', level=1)
    checks = [
        ('Status Code Match', execution.status_code_match),
        ('Content Match', execution.content_match),
        ('Time Threshold', execution.time_within_threshold),
    ]
    vtable = doc.add_table(rows=len(checks), cols=2)
    vtable.style = 'Table Grid'
    for i, (label, val) in enumerate(checks):
        vtable.cell(i, 0).text = label
        vtable.cell(i, 1).text = 'Pass' if val is True else ('Fail' if val is False else 'N/A')

    if mismatch_summary:
        doc.add_heading('Mismatch Details (Expected Response Content)', level=1)
        p = doc.add_paragraph(mismatch_summary)
        p.runs[0].font.size = Pt(9)

    if key_pair_mismatch_summary:
        doc.add_heading('Mismatch Details (Expected key-pair value)', level=1)
        p = doc.add_paragraph(key_pair_mismatch_summary)
        p.runs[0].font.size = Pt(9)

    if execution.error_message:
        doc.add_heading('Error', level=1)
        doc.add_paragraph(execution.error_message)

    for section_title, content in [
        ('Request Headers', execution.request_headers),
        ('Request Body', execution.request_body),
        ('Response Headers', execution.response_headers),
        ('Response Body', execution.response_body),
    ]:
        if content:
            doc.add_heading(section_title, level=1)
            p = doc.add_paragraph(content)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(8)

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    safe_name = _safe_download_name(execution.test_case.name, f'testcase_{execution.test_case_id}')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response


@login_required
def execution_download_pdf(request, pk):
    """Download execution detail as PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
    import io

    execution = get_object_or_404(
        TestExecution.objects.select_related('test_case', 'environment', 'executed_by'),
        pk=pk
    )
    mismatch_summary = _expected_content_mismatch_summary(
        execution.test_case.expected_response_content,
        execution.response_body,
    )
    expected_key_pairs_rows = _parse_form_data_rows(execution.test_case.form_data)
    expected_key_pairs_text = _format_expected_key_pairs(expected_key_pairs_rows)
    key_pair_mismatch_summary = _form_data_mismatch_summary(
        expected_key_pairs_rows,
        execution.response_body,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    mono = ParagraphStyle('mono', parent=styles['Normal'], fontName='Courier', fontSize=7, leading=10)

    STATUS_COLORS = {'passed': colors.HexColor('#16a34a'), 'failed': colors.HexColor('#dc2626'), 'error': colors.HexColor('#d97706')}
    status_color = STATUS_COLORS.get(execution.result_status, colors.grey)

    title_style = ParagraphStyle('title', parent=styles['Title'], textColor=colors.HexColor('#0f172a'), fontSize=18)
    h1 = ParagraphStyle('h1', parent=styles['Heading1'], textColor=colors.HexColor('#0f766e'), fontSize=12, spaceBefore=12)

    elements = []
    elements.append(Paragraph('Execution Report', title_style))
    elements.append(Spacer(1, 0.3*cm))

    result_text = f'<font color="#{status_color.hexval()[2:] if hasattr(status_color, "hexval") else "000000"}"><b>{execution.result_status.upper()}</b></font>'
    elements.append(Paragraph(f'Result: {execution.result_status.upper()}', ParagraphStyle('res', parent=styles['Normal'], fontSize=11,
        textColor=status_color, spaceAfter=8)))

    elements.append(Paragraph('Summary', h1))
    summary_data = [
        ['Field', 'Value'],
        ['Test Case', execution.test_case.name],
        ['HTTP Method', execution.test_case.http_method],
        ['Environment', execution.environment.name if execution.environment else '-'],
        ['Executed By', execution.executed_by.username if execution.executed_by else '-'],
        ['Executed At', execution.executed_at.strftime('%Y-%m-%d %H:%M:%S')],
        ['Request URL', Paragraph(execution.request_url or '-', mono)],
        ['Response Status', str(execution.response_status_code or '-')],
        ['Response Time', f"{execution.response_time_ms or '-'} ms"],
        ['Expected Response Content', Paragraph((execution.test_case.expected_response_content or '-').replace('\n', '<br/>'), mono)],
        ['Expected key-pair value', Paragraph(expected_key_pairs_text.replace('\n', '<br/>'), mono)],
        ['Error', execution.error_message or '-'],
    ]
    t = Table(summary_data, colWidths=[4.5*cm, 12*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f766e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#f8fafc'), colors.white]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dbe3ec')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(t)

    elements.append(Paragraph('Validation', h1))
    checks_data = [
        ['Check', 'Result'],
        ['Status Code Match', 'Pass' if execution.status_code_match is True else ('Fail' if execution.status_code_match is False else 'N/A')],
        ['Content Match', 'Pass' if execution.content_match is True else ('Fail' if execution.content_match is False else 'N/A')],
        ['Time Threshold', 'Pass' if execution.time_within_threshold is True else ('Fail' if execution.time_within_threshold is False else 'N/A')],
    ]
    if mismatch_summary:
        checks_data.append(['Mismatch (Expected Response Content)', Paragraph(mismatch_summary.replace('\n', '<br/>'), mono)])
    if key_pair_mismatch_summary:
        checks_data.append(['Mismatch (Expected key-pair value)', Paragraph(key_pair_mismatch_summary.replace('\n', '<br/>'), mono)])
    ct = Table(checks_data, colWidths=[8*cm, 8.5*cm])
    ct.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f766e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#f8fafc'), colors.white]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dbe3ec')),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(ct)

    for title, content in [('Request Headers', execution.request_headers), ('Request Body', execution.request_body),
                           ('Response Headers', execution.response_headers), ('Response Body', execution.response_body)]:
        if content:
            elements.append(Paragraph(title, h1))
            elements.append(Paragraph(content.replace('\n', '<br/>'), mono))

    doc.build(elements)
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    safe_name = _safe_download_name(execution.test_case.name, f'testcase_{execution.test_case_id}')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.pdf"'
    return response


@login_required
def execution_download_response_json(request, pk):
    """Download response JSON in JSON format."""
    execution = get_object_or_404(TestExecution, pk=pk)
    try:
        parsed_body = json.loads(execution.response_body) if execution.response_body else {}
    except (json.JSONDecodeError, TypeError):
        parsed_body = {}

    test_case_name = execution.test_case.name.replace(' ', '_').replace('/', '_')
    response = HttpResponse(
        json.dumps(parsed_body, indent=2, ensure_ascii=False),
        content_type='application/json'
    )
    response['Content-Disposition'] = f'attachment; filename="{test_case_name}_response.json"'
    return response


@login_required
def execution_download_response_csv(request, pk):
    """Download response JSON in CSV format."""
    execution = get_object_or_404(TestExecution, pk=pk)
    try:
        parsed_body = json.loads(execution.response_body) if execution.response_body else None
    except (json.JSONDecodeError, TypeError):
        parsed_body = None

    test_case_name = execution.test_case.name.replace(' ', '_').replace('/', '_')
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{test_case_name}_response.csv"'
    writer = csv.writer(response)

    if isinstance(parsed_body, list):
        # If response is a list of objects, write as rows
        if parsed_body and isinstance(parsed_body[0], dict):
            # Get all keys from all objects
            fieldnames = set()
            for item in parsed_body:
                if isinstance(item, dict):
                    fieldnames.update(item.keys())
            fieldnames = sorted(list(fieldnames))
            writer.writerow(fieldnames)
            for item in parsed_body:
                row = [str(item.get(field, '')) for field in fieldnames]
                writer.writerow(row)
        else:
            # List of primitives
            writer.writerow(['Value'])
            for item in parsed_body:
                writer.writerow([str(item)])
    elif isinstance(parsed_body, dict):
        # If response is a single object, write as key-value pairs
        writer.writerow(['Key', 'Value'])
        for key, value in parsed_body.items():
            if isinstance(value, (dict, list)):
                writer.writerow([key, json.dumps(value)])
            else:
                writer.writerow([key, str(value)])
    else:
        # Fallback for other types
        writer.writerow(['Response'])
        writer.writerow([str(parsed_body)])

    return response


def _parse_execution_text_json(raw_text):
    """Parse stored execution text as JSON when possible, else return plain string."""
    text = (raw_text or '').strip()
    if not text:
        return {}
    try:
        return json.loads(text)
    except (json.JSONDecodeError, TypeError, ValueError):
        return text


def _write_jsonish_csv(writer, parsed, fallback_label='Value'):
    """Write dict/list/scalar payload to CSV in a predictable shape."""
    if isinstance(parsed, list):
        if parsed and isinstance(parsed[0], dict):
            fieldnames = set()
            for item in parsed:
                if isinstance(item, dict):
                    fieldnames.update(item.keys())
            fieldnames = sorted(list(fieldnames))
            writer.writerow(fieldnames)
            for item in parsed:
                row = [str(item.get(field, '')) if isinstance(item, dict) else '' for field in fieldnames]
                writer.writerow(row)
            return

        writer.writerow([fallback_label])
        for item in parsed:
            writer.writerow([str(item)])
        return

    if isinstance(parsed, dict):
        writer.writerow(['Key', 'Value'])
        for key, value in parsed.items():
            if isinstance(value, (dict, list)):
                writer.writerow([key, json.dumps(value, ensure_ascii=False)])
            else:
                writer.writerow([key, str(value)])
        return

    writer.writerow([fallback_label])
    writer.writerow([str(parsed)])


@login_required
def execution_download_request_headers_json(request, pk):
    """Download request headers payload in JSON format."""
    execution = get_object_or_404(TestExecution, pk=pk)
    parsed = _parse_execution_text_json(execution.request_headers)

    safe_name = execution.test_case.name.replace(' ', '_').replace('/', '_')
    response = HttpResponse(
        json.dumps(parsed, indent=2, ensure_ascii=False),
        content_type='application/json'
    )
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_request_headers.json"'
    return response


@login_required
def execution_download_request_headers_csv(request, pk):
    """Download request headers payload in CSV format."""
    execution = get_object_or_404(TestExecution, pk=pk)
    parsed = _parse_execution_text_json(execution.request_headers)

    safe_name = execution.test_case.name.replace(' ', '_').replace('/', '_')
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_request_headers.csv"'
    writer = csv.writer(response)
    _write_jsonish_csv(writer, parsed, fallback_label='Header')
    return response


@login_required
def execution_download_request_body_json(request, pk):
    """Download request body payload in JSON format."""
    execution = get_object_or_404(TestExecution, pk=pk)
    parsed = _parse_execution_text_json(execution.request_body)
    if isinstance(parsed, str):
        parsed = {'request_body': parsed}

    safe_name = execution.test_case.name.replace(' ', '_').replace('/', '_')
    response = HttpResponse(
        json.dumps(parsed, indent=2, ensure_ascii=False),
        content_type='application/json'
    )
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_request_body.json"'
    return response


@login_required
def execution_download_request_body_csv(request, pk):
    """Download request body payload in CSV format."""
    execution = get_object_or_404(TestExecution, pk=pk)
    parsed = _parse_execution_text_json(execution.request_body)

    safe_name = execution.test_case.name.replace(' ', '_').replace('/', '_')
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_request_body.csv"'
    writer = csv.writer(response)
    _write_jsonish_csv(writer, parsed, fallback_label='Request Body')
    return response


# ========== User Management ==========

@login_required
@role_required(['admin'])
def user_list(request):
    users = User.objects.select_related('profile').all()
    return render(request, 'api/users/list.html', {'users': users})


@login_required
@role_required(['admin'])
def user_create(request):
    if request.method == 'POST':
        form = UserCreateForm(request.POST)
        if form.is_valid():
            user = User.objects.create_user(
                username=form.cleaned_data['username'],
                email=form.cleaned_data['email'],
                password=form.cleaned_data['password'],
                first_name=form.cleaned_data.get('first_name', ''),
                last_name=form.cleaned_data.get('last_name', ''),
            )
            UserProfile.objects.create(user=user, role=form.cleaned_data['role'])
            log_action(request.user, 'user_create', 'User', user.id,
                       f'Created user: {user.username}', request)
            messages.success(request, f'User "{user.username}" created successfully.')
            return redirect('api:user_list')
    else:
        form = UserCreateForm()
    return render(request, 'api/users/form.html', {'form': form, 'title': 'Create User'})


@login_required
@role_required(['admin'])
def user_edit(request, pk):
    target_user = get_object_or_404(User, pk=pk)
    profile, _ = UserProfile.objects.get_or_create(user=target_user)

    if request.method == 'POST':
        form = UserEditForm(request.POST)
        if form.is_valid():
            target_user.email = form.cleaned_data['email']
            target_user.first_name = form.cleaned_data.get('first_name', '')
            target_user.last_name = form.cleaned_data.get('last_name', '')
            target_user.is_active = form.cleaned_data.get('is_active', True)
            target_user.save()
            profile.role = form.cleaned_data['role']
            profile.save()
            # Handle password change
            new_password = form.cleaned_data.get('new_password')
            if new_password:
                target_user.set_password(new_password)
                target_user.save()
            log_action(request.user, 'user_update', 'User', target_user.id,
                       f'Updated user: {target_user.username}', request)
            messages.success(request, f'User "{target_user.username}" updated successfully.')
            return redirect('api:user_list')
    else:
        form = UserEditForm(initial={
            'email': target_user.email,
            'first_name': target_user.first_name,
            'last_name': target_user.last_name,
            'role': profile.role,
            'is_active': target_user.is_active,
        })
    return render(request, 'api/users/form.html', {
        'form': form, 'title': f'Edit User: {target_user.username}', 'target_user': target_user
    })


@login_required
@role_required(['admin'])
def user_deactivate(request, pk):
    target_user = get_object_or_404(User, pk=pk)
    if request.method == 'POST':
        target_user.is_active = False
        target_user.save()
        log_action(request.user, 'user_deactivate', 'User', target_user.id,
                   f'Deactivated user: {target_user.username}', request)
        messages.success(request, f'User "{target_user.username}" deactivated.')
    return redirect('api:user_list')


@login_required
@role_required(['admin'])
def user_delete(request, pk):
    target_user = get_object_or_404(User, pk=pk)

    if request.method == 'POST':
        if target_user.pk == request.user.pk:
            messages.error(request, 'You cannot delete your own account.')
            return redirect('api:user_list')

        username = target_user.username
        user_id = target_user.id
        target_user.delete()

        log_action(request.user, 'user_deactivate', 'User', user_id,
                   f'Deleted user: {username}', request)
        messages.success(request, f'User "{username}" deleted successfully.')

    return redirect('api:user_list')


@login_required
@role_required(['admin'])
def theme_settings(request):
    theme, _ = ThemeSettings.objects.get_or_create(name='default')

    if request.method == 'POST':
        form = ThemeSettingsForm(request.POST, instance=theme)
        if form.is_valid():
            updated = form.save(commit=False)
            updated.updated_by = request.user
            updated.save()
            log_action(request.user, 'user_update', 'ThemeSettings', updated.id,
                       f'Updated application theme ({updated.theme_mode})', request)
            messages.success(request, 'Theme settings updated successfully.')
            return redirect('api:theme_settings')
    else:
        form = ThemeSettingsForm(instance=theme)

    return render(request, 'api/users/theme.html', {
        'form': form,
        'title': 'Theme Settings',
        'theme': theme,
    })


# ========== Audit Logs ==========

@login_required
@role_required(['admin'])
def audit_log_list(request):
    logs = AuditLog.objects.select_related('user')

    action_filter = request.GET.get('action', '')
    user_filter = request.GET.get('user', '')

    if action_filter:
        logs = logs.filter(action=action_filter)
    if user_filter:
        logs = logs.filter(user_id=user_filter)

    users = User.objects.filter(is_active=True)

    context = {
        'logs': logs[:200],
        'users': users,
        'action_choices': AuditLog.ACTION_CHOICES,
        'filters': {
            'action': action_filter,
            'user': user_filter,
        }
    }
    return render(request, 'api/audit/list.html', context)


# ========== Reports ==========

@login_required
def export_executions_csv(request):
    """Export execution history as CSV."""
    executions = TestExecution.objects.select_related('test_case', 'environment', 'executed_by')

    status_filter = request.GET.get('status', '')
    env_filter = request.GET.get('environment', '')
    if status_filter:
        executions = executions.filter(result_status=status_filter)
    if env_filter:
        executions = executions.filter(environment_id=env_filter)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="test_executions.csv"'

    writer = csv.writer(response)
    writer.writerow([
        'Test Case', 'HTTP Method', 'Endpoint', 'Environment',
        'Executed By', 'Executed At', 'Result', 'Status Code',
        'Response Time (ms)', 'Error'
    ])

    for ex in executions[:1000]:
        writer.writerow([
            ex.test_case.name,
            ex.test_case.http_method,
            ex.test_case.endpoint,
            ex.environment.name if ex.environment else '',
            ex.executed_by.username if ex.executed_by else '',
            ex.executed_at.strftime('%Y-%m-%d %H:%M:%S'),
            ex.result_status,
            ex.response_status_code or '',
            ex.response_time_ms or '',
            ex.error_message,
        ])

    return response


# ========== Module Management ==========

def parse_postman_collection(data):
    """Parse a Postman collection JSON and extract endpoints."""
    endpoints = []

    def extract_items(items):
        for item in items:
            # If item has sub-items (folder), recurse
            if 'item' in item and isinstance(item['item'], list):
                extract_items(item['item'])
            elif 'request' in item:
                req = item['request']
                method = req.get('method', 'GET')
                name = item.get('name', 'Unnamed')
                description = ''
                if isinstance(req.get('description'), str):
                    description = req['description']

                # Extract URL path
                url = req.get('url', {})
                if isinstance(url, str):
                    endpoint_path = url
                elif isinstance(url, dict):
                    raw = url.get('raw', '')
                    path_parts = url.get('path', [])
                    if path_parts:
                        endpoint_path = '/' + '/'.join(path_parts)
                    else:
                        endpoint_path = raw
                else:
                    endpoint_path = ''

                # Extract headers
                headers = {}
                for h in req.get('header', []):
                    if isinstance(h, dict) and h.get('key'):
                        headers[h['key']] = h.get('value', '')

                # Extract body
                body = ''
                if req.get('body'):
                    body_obj = req['body']
                    if body_obj.get('mode') == 'raw':
                        body = body_obj.get('raw', '')

                endpoints.append({
                    'name': name,
                    'method': method,
                    'endpoint_path': endpoint_path,
                    'headers': json.dumps(headers) if headers else '{}',
                    'request_body': body,
                    'description': description,
                })

    items = data.get('item', [])
    extract_items(items)
    return endpoints


def parse_openapi_spec(data):
    """Parse an OpenAPI/Swagger JSON and extract endpoints."""
    endpoints = []
    paths = data.get('paths', {})
    for path, methods in paths.items():
        for method, details in methods.items():
            if method.lower() in ('get', 'post', 'put', 'patch', 'delete'):
                name = details.get('summary') or details.get('operationId') or f"{method.upper()} {path}"
                description = details.get('description', '')
                endpoints.append({
                    'name': name,
                    'method': method.upper(),
                    'endpoint_path': path,
                    'headers': '{}',
                    'request_body': '',
                    'description': description,
                })
    return endpoints


@login_required
def module_list(request):
    modules = ApiModule.objects.annotate(endpoint_count=Count('endpoints'))
    return render(request, 'api/modules/list.html', {'modules': modules})


@login_required
@role_required(['admin', 'tester'])
def module_upload(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('collection_file')
        if not uploaded_file:
            messages.error(request, 'Please select a JSON file to upload.')
            return redirect('api:module_upload')

        if not uploaded_file.name.endswith('.json'):
            messages.error(request, 'Only JSON files are supported.')
            return redirect('api:module_upload')

        try:
            content = uploaded_file.read().decode('utf-8')
            data = json.loads(content)
        except (json.JSONDecodeError, UnicodeDecodeError):
            messages.error(request, 'Invalid JSON file.')
            return redirect('api:module_upload')

        # Determine type: Postman collection or OpenAPI spec
        if 'info' in data and 'item' in data:
            # Postman collection
            collection_name = data.get('info', {}).get('name', uploaded_file.name)
            description = data.get('info', {}).get('description', '')
            endpoints = parse_postman_collection(data)
        elif 'openapi' in data or 'swagger' in data:
            # OpenAPI / Swagger spec
            info = data.get('info', {})
            collection_name = info.get('title', uploaded_file.name)
            description = info.get('description', '')
            endpoints = parse_openapi_spec(data)
        else:
            messages.error(request, 'Unsupported JSON format. Upload a Postman collection or OpenAPI spec.')
            return redirect('api:module_upload')

        if not endpoints:
            messages.error(request, 'No endpoints found in the uploaded file.')
            return redirect('api:module_upload')

        # Create module
        module = ApiModule.objects.create(
            name=collection_name,
            description=description,
            source_file=uploaded_file.name,
            uploaded_by=request.user,
        )

        # Create endpoints
        for ep in endpoints:
            ModuleEndpoint.objects.create(
                module=module,
                name=ep['name'],
                http_method=ep['method'],
                endpoint_path=ep['endpoint_path'],
                headers=ep['headers'],
                request_body=ep['request_body'],
                description=ep['description'],
            )

        messages.success(request, f'Module "{collection_name}" uploaded with {len(endpoints)} endpoints.')
        return redirect('api:module_list')

    return render(request, 'api/modules/upload.html')


@login_required
def module_detail(request, pk):
    module = get_object_or_404(ApiModule, pk=pk)
    search_query = (request.GET.get('q') or '').strip()
    endpoints = module.endpoints.all()
    if search_query:
        endpoints = endpoints.filter(
            Q(name__icontains=search_query)
            | Q(endpoint_path__icontains=search_query)
            | Q(http_method__icontains=search_query)
            | Q(description__icontains=search_query)
        )

    total_endpoints = module.endpoints.count()
    environments = Environment.objects.filter(is_active=True).order_by('name')
    return render(request, 'api/modules/detail.html', {
        'module': module,
        'endpoints': endpoints,
        'environments': environments,
        'search_query': search_query,
        'total_endpoints': total_endpoints,
    })


@login_required
@role_required(['admin'])
def module_delete(request, pk):
    module = get_object_or_404(ApiModule, pk=pk)
    if request.method == 'POST':
        name = module.name
        module.delete()
        messages.success(request, f'Module "{name}" deleted.')
    return redirect('api:module_list')


@login_required
@role_required(['admin', 'tester'])
def module_update_base_path(request, pk):
    module = get_object_or_404(ApiModule, pk=pk)
    if request.method == 'POST':
        module.base_path = request.POST.get('base_path', '').strip()
        module.save()
        messages.success(request, f'Base path updated for "{module.name}".')
    return redirect('api:module_detail', pk=module.pk)


@login_required
@role_required(['admin', 'tester'])
def module_update_oauth(request, pk):
    module = get_object_or_404(ApiModule, pk=pk)
    if request.method != 'POST':
        return redirect('api:module_detail', pk=module.pk)

    module_auth_type = request.POST.get('module_auth_type', 'none').strip()
    module.module_auth_type = module_auth_type

    if module_auth_type == 'oauth2':
        module.oauth2_add_to = request.POST.get('oauth2_add_to', 'request_headers').strip() or 'request_headers'
        module.oauth2_client_id = request.POST.get('oauth2_client_id', '').strip()
        module.oauth2_client_secret = request.POST.get('oauth2_client_secret', '').strip()
        module.oauth2_token_url = request.POST.get('oauth2_token_url', '').strip()
        module.oauth2_header_prefix = request.POST.get('oauth2_header_prefix', 'Bearer').strip() or 'Bearer'

        token, error = fetch_oauth2_access_token(
            module.oauth2_token_url,
            module.oauth2_client_id,
            module.oauth2_client_secret,
        )
        if error:
            messages.error(request, error)
            return redirect('api:module_detail', pk=module.pk)

        module.oauth2_current_token = token
        module.oauth2_token_updated_at = timezone.now()
        module.save()
        messages.success(request, 'OAuth 2.0 token configuration saved and token generated.')
        return redirect('api:module_detail', pk=module.pk)

    # Clear OAuth fields when auth type is set to none
    module.oauth2_add_to = 'request_headers'
    module.oauth2_client_id = ''
    module.oauth2_client_secret = ''
    module.oauth2_token_url = ''
    module.oauth2_current_token = ''
    module.oauth2_header_prefix = 'Bearer'
    module.oauth2_token_updated_at = None
    module.save()
    messages.success(request, 'Module authentication disabled.')
    return redirect('api:module_detail', pk=module.pk)


@login_required
def api_module_endpoints(request, pk):
    """AJAX endpoint: return endpoints for a given module."""
    module = get_object_or_404(ApiModule, pk=pk)
    endpoints = module.endpoints.all().values(
        'id', 'name', 'http_method', 'endpoint_path', 'headers',
        'request_body', 'default_payload', 'expected_responses', 'description'
    )
    return JsonResponse({'endpoints': list(endpoints)})


@login_required
def api_module_auth(request, pk):
    """AJAX endpoint: return module auth config for test case prefill."""
    module = get_object_or_404(ApiModule, pk=pk)
    if module.module_auth_type != 'oauth2':
        return JsonResponse({'auth_type': 'none'})

    token = module.oauth2_current_token
    token_fetched_recently = module.oauth2_token_updated_at and (
        timezone.now() - module.oauth2_token_updated_at <= timedelta(minutes=50)
    )
    if not token or not token_fetched_recently:
        refreshed_token, error = fetch_oauth2_access_token(
            module.oauth2_token_url,
            module.oauth2_client_id,
            module.oauth2_client_secret,
        )
        if not error:
            token = refreshed_token
            module.oauth2_current_token = refreshed_token
            module.oauth2_token_updated_at = timezone.now()
            module.save(update_fields=['oauth2_current_token', 'oauth2_token_updated_at', 'updated_at'])

    return JsonResponse({
        'auth_type': module.module_auth_type,
        'oauth2_add_to': module.oauth2_add_to,
        'oauth2_header_prefix': module.oauth2_header_prefix,
        'oauth2_client_id': module.oauth2_client_id,
        'oauth2_token_url': module.oauth2_token_url,
        'oauth2_current_token': token,
    })


@login_required
@role_required(['admin', 'tester'])
def reorder_projects(request):
    """AJAX endpoint: persist new project folder order."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    try:
        data = json.loads(request.body)
        ids = data.get('ids', [])
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    for position, pk in enumerate(ids):
        Project.objects.filter(pk=pk).update(order=position)
    return JsonResponse({'ok': True})


@login_required
@role_required(['admin', 'tester'])
def reorder_testcases(request):
    """AJAX endpoint: persist new test case order within a folder."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    try:
        data = json.loads(request.body)
        ids = data.get('ids', [])
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    for position, pk in enumerate(ids):
        TestCase.objects.filter(pk=pk).update(order=position)
    return JsonResponse({'ok': True})


@login_required
def api_modules_list(request):
    """AJAX endpoint: return all modules for dropdown population."""
    modules = ApiModule.objects.all().values('id', 'name', 'base_path')
    return JsonResponse({'modules': list(modules)})


@login_required
def api_environment_details(request, pk):
    """AJAX endpoint: return environment details for module auto-population."""
    env = get_object_or_404(Environment, pk=pk)

    creds = {}
    if env.auth_credentials:
        try:
            creds = json.loads(env.auth_credentials)
        except json.JSONDecodeError:
            creds = {}

    return JsonResponse({
        'id': env.id,
        'name': env.name,
        'base_path': env.base_url,
        'auth_type': env.auth_type,
        'auth_type_display': env.get_auth_type_display(),
        'oauth2_add_to': creds.get('add_to', 'request_headers'),
        'oauth2_client_id': creds.get('client_id', ''),
        'oauth2_token_url': creds.get('token_url', ''),
        'oauth2_header_prefix': creds.get('header_prefix', 'Bearer'),
        'oauth2_current_token': creds.get('token', ''),
        'updated_at': timezone.localtime(env.updated_at).strftime('%b %d, %Y %H:%M') if env.updated_at else '',
    })


@login_required
@role_required(['admin', 'tester'])
def endpoint_edit(request, pk):
    """Edit an endpoint's headers, payload and expected responses."""
    endpoint = get_object_or_404(ModuleEndpoint, pk=pk)

    if request.method == 'POST':
        headers_raw      = request.POST.get('headers', '').strip()
        default_payload  = request.POST.get('default_payload', '').strip()
        expected_responses = request.POST.get('expected_responses', '').strip()

        # Validate JSON fields
        for label, value in [('Headers', headers_raw), ('Default Payload', default_payload), ('Expected Responses', expected_responses)]:
            if value:
                try:
                    json.loads(value)
                except json.JSONDecodeError:
                    messages.error(request, f'{label} is not valid JSON.')
                    return render(request, 'api/modules/endpoint_edit.html', {'endpoint': endpoint})

        endpoint.headers           = headers_raw or '{}'
        endpoint.default_payload   = default_payload
        endpoint.expected_responses = expected_responses or '{}'
        endpoint.save()
        messages.success(request, f'Endpoint "{endpoint.name}" updated.')
        return redirect('api:module_detail', pk=endpoint.module.pk)

    return render(request, 'api/modules/endpoint_edit.html', {'endpoint': endpoint})


@login_required
@role_required(['admin', 'tester'])
def endpoint_upload_payload(request, pk):
    """Upload a JSON file as the default payload for an endpoint."""
    endpoint = get_object_or_404(ModuleEndpoint, pk=pk)

    if request.method == 'POST':
        uploaded_file = request.FILES.get('payload_file')
        if not uploaded_file:
            messages.error(request, 'Please select a JSON file.')
            return redirect('api:endpoint_edit', pk=endpoint.pk)

        try:
            content = uploaded_file.read().decode('utf-8')
            json.loads(content)  # Validate
            endpoint.default_payload = content
            endpoint.save()
            messages.success(request, f'Payload uploaded for "{endpoint.name}".')
        except (json.JSONDecodeError, UnicodeDecodeError):
            messages.error(request, 'Invalid JSON file.')

    return redirect('api:endpoint_edit', pk=endpoint.pk)

